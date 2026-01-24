import streamlit as st
import pandas as pd
import pdfplumber
import re
from datetime import datetime
import time
import base64
import hashlib
import sqlite3
import os
import io
import sys
import subprocess
import random
import string
from PIL import Image, ImageEnhance
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import ImageReader
from reportlab.lib.colors import HexColor
from typing import Any, List, Optional, Union, Literal, overload, Dict
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from lunardate import LunarDate

# --- QUAN TR·ªåNG: C·∫§U H√åNH TRANG PH·∫¢I ·ªû ƒê·∫¶U TI√äN ---
st.set_page_config(
    page_title="Qu·∫£n L√Ω H√≥a ƒê∆°n Pro", 
    page_icon="üå∏", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- IMPORTS KH√ÅC (Kh√¥ng d√πng auto_install n·ªØa) ---
import gspread
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
import cv2
import numpy as np

# --- OCR CONFIGURATION ---
try:
    import pytesseract
    # C·∫•u h√¨nh ƒë∆∞·ªùng d·∫´n Tesseract
    if sys.platform.startswith('win'):
        # ƒê∆∞·ªùng d·∫´n cho Windows (Local)
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    else:
        # Ch·∫°y tr√™n Streamlit Cloud (Linux) - Kh√¥ng c·∫ßn set path, n√≥ t·ª± t√¨m
        # N·∫øu c·∫ßn thi·∫øt: pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'
        pass
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    pytesseract = None

# --- EXCEL & DOCX LIBS ---
import openpyxl
import xlsxwriter
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
HAS_OPENPYXL = True
HAS_XLSXWRITER = True
HAS_CV = True
HAS_DOCX = True

# --- C·∫§U H√åNH GOOGLE (ƒê√£ c·∫≠p nh·∫≠t theo th√¥ng tin c·ªßa b·∫°n) ---
SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive"
]

# T√™n file ch√¨a kh√≥a (H√£y ƒë·ªïi t√™n file b·∫°n t·∫£i v·ªÅ th√†nh t√™n n√†y)
SERVICE_ACCOUNT_FILE = 'service_account.json'

# ID Google Drive (N∆°i l∆∞u ·∫£nh/pdf)
# Link: https://drive.google.com/drive/folders/1PMCKIUirYwbacu0evnRyuF0xSq-bQtBv?usp=drive_link
DRIVE_FOLDER_ID = '1PMCKIUirYwbacu0evnRyuF0xSq-bQtBv'

# ID Google Sheet (L·∫•y t·ª´ link b·∫°n g·ª≠i)
# Link: https://docs.google.com/spreadsheets/d/1coeIPogjKEJSKv1hW1dFBrSAwF6V7c-tkVCZPuPQjoc/edit?gid=0#gid=0
SPREADSHEET_ID = '1coeIPogjKEJSKv1hW1dFBrSAwF6V7c-tkVCZPuPQjoc'

def get_gspread_client():
    # Ki·ªÉm tra xem ƒëang ch·∫°y tr√™n Cloud (d√πng secrets) hay Local (d√πng file json)
    if "gcp_service_account" in st.secrets:
        creds_dict = dict(st.secrets["gcp_service_account"])
        creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPES)
    else:
        creds = Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    
    client = gspread.authorize(creds)
    return client

def get_drive_service():
    if "gcp_service_account" in st.secrets:
        creds_dict = dict(st.secrets["gcp_service_account"])
        creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPES)
    else:
        creds = Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    
    service = build('drive', 'v3', credentials=creds)
    return service

# --- C√ÅC H√ÄM X·ª¨ L√ù D·ªÆ LI·ªÜU M·ªöI (Thay th·∫ø SQL) ---

def load_table(table_name):
    """ƒê·ªçc d·ªØ li·ªáu t·ª´ Local SQLite (Thay th·∫ø Google Sheet)"""
    conn = get_connection()
    try:
        df = pd.read_sql_query(f"SELECT * FROM {table_name}", conn)
        return df
    except Exception as e:
        print(f"L·ªói ƒë·ªçc b·∫£ng {table_name}: {e}")
        return pd.DataFrame()

def add_row_to_table(table_name, row_dict):
    """Th√™m d√≤ng m·ªõi v√†o Local SQLite V√Ä Google Sheet"""
    # 1. Ghi v√†o SQLite (Local)
    conn = get_connection()
    c = conn.cursor()
    success = False
    try:
        columns = ', '.join(row_dict.keys())
        placeholders = ', '.join(['?'] * len(row_dict))
        sql = f"INSERT INTO {table_name} ({columns}) VALUES ({placeholders})"
        c.execute(sql, list(row_dict.values()))
        conn.commit()
        success = True
    except Exception as e:
        st.error(f"L·ªói ghi d·ªØ li·ªáu v√†o {table_name}: {e}")
        return False

    # 2. Ghi v√†o Google Sheet (Cloud)
    if success:
        try:
            gc = get_gspread_client()
            sh = gc.open_by_key(SPREADSHEET_ID)
            try:
                wks = sh.worksheet(table_name)
            except:
                wks = sh.add_worksheet(title=table_name, rows=100, cols=20)
            
            # X·ª≠ l√Ω header v√† map d·ªØ li·ªáu
            existing = wks.get_all_values()
            if not existing:
                headers = list(row_dict.keys())
                wks.append_row(headers)
            else:
                headers = existing[0]
            
            row_values = []
            for h in headers:
                val = row_dict.get(h, "")
                if val is None: val = ""
                row_values.append(val)
                
            wks.append_row(row_values)
        except Exception as e:
            # [DEBUG] Thay ƒë·ªïi ƒë·ªÉ hi·ªÉn th·ªã l·ªói chi ti·∫øt h∆°n
            st.error(f"‚ö†Ô∏è L·ªñI ƒê·ªíNG B·ªò GOOGLE SHEET (ƒê√£ l∆∞u v√†o m√°y nh∆∞ng kh√¥ng ƒë·∫©y l√™n cloud ƒë∆∞·ª£c)")
            st.exception(e)
            
    return success

def upload_to_drive(file_obj, file_name, mimetype=None):
    """Upload file l√™n Google Drive"""
    try:
        service = get_drive_service()
        file_metadata = {'name': file_name, 'parents': [DRIVE_FOLDER_ID]}
        
        if not mimetype and hasattr(file_obj, 'type'):
            mimetype = file_obj.type
            
        media = MediaIoBaseUpload(file_obj, mimetype=mimetype or 'application/octet-stream', resumable=True)
        file = service.files().create(body=file_metadata, media_body=media, fields='id, webViewLink').execute()
        return file.get('webViewLink')
    except Exception as e:
        st.warning(f"‚ö†Ô∏è L·ªói upload Drive: {e}")
        return None

def sync_all_data_to_gsheet():
    """ƒê·ªçc t·∫•t c·∫£ d·ªØ li·ªáu t·ª´ SQLite v√† ghi ƒë√® l√™n Google Sheet."""
    TABLES_TO_SYNC = [
        'users', 'invoices', 'projects', 'project_links', 'company_info', 
        'flight_tickets', 'flight_groups', 'flight_group_links', 
        'service_bookings', 'customers', 'tours', 'tour_items', 'ocr_learning',
        'transaction_history'
    ]

    try:
        gc = get_gspread_client()
        sh = gc.open_by_key(SPREADSHEET_ID)
        conn = get_connection()

        st.info(f"B·∫Øt ƒë·∫ßu ƒë·ªìng b·ªô {len(TABLES_TO_SYNC)} b·∫£ng...")
        status_placeholder = st.empty()
        progress_bar = st.progress(0)
        
        for i, table_name in enumerate(TABLES_TO_SYNC):
            status_placeholder.info(f"ƒêang x·ª≠ l√Ω b·∫£ng: **{table_name}**...")
            
            try:
                df = pd.read_sql_query(f"SELECT * FROM {table_name}", conn)
            except Exception:
                st.warning(f"B·∫£ng '{table_name}' kh√¥ng c√≥ trong DB, b·ªè qua.")
                progress_bar.progress((i + 1) / len(TABLES_TO_SYNC))
                continue

            try:
                wks = sh.worksheet(table_name)
                wks.clear()
            except gspread.WorksheetNotFound:
                wks = sh.add_worksheet(title=table_name, rows=1, cols=20)

            if not df.empty:
                df = df.astype(str).replace({'nan': '', 'NaT': ''})
                # [FIX] Truncate cells that are too long for Google Sheets API to prevent 400 error
                df = df.map(lambda x: x[:49999] if isinstance(x, str) and len(x) >= 50000 else x)
                data_to_upload = [df.columns.tolist()] + df.values.tolist()
                wks.update(data_to_upload, 'A1')
                st.toast(f"‚úÖ ƒê·ªìng b·ªô '{table_name}' ({len(df)} d√≤ng) OK.")
            else:
                st.toast(f"‚ÑπÔ∏è B·∫£ng '{table_name}' r·ªóng, ƒë√£ d·ªçn d·∫πp tr√™n cloud.")

            progress_bar.progress((i + 1) / len(TABLES_TO_SYNC))

        status_placeholder.empty()
        st.success("üéâ ƒê·ªìng b·ªô to√†n b·ªô d·ªØ li·ªáu ho√†n t·∫•t!")
    except Exception as e:
        st.error("‚ùå L·ªói nghi√™m tr·ªçng khi ƒë·ªìng b·ªô:")
        st.exception(e)
        st.info("üí° G·ª£i √Ω: H√£y ch·∫Øc ch·∫Øn r·∫±ng email c·ªßa t√†i kho·∫£n d·ªãch v·ª• (`client_email` trong file .json) ƒë√£ ƒë∆∞·ª£c c·∫•p quy·ªÅn 'Editor' (Ng∆∞·ªùi ch·ªânh s·ª≠a) cho file Google Sheet n√†y.")

# --- QU·∫¢N L√ù SESSION STATE ---
if "logged_in" not in st.session_state: st.session_state.logged_in = False
if "user_info" not in st.session_state: st.session_state.user_info = None
if "db_initialized" not in st.session_state: st.session_state.db_initialized = False

# Bi·∫øn l∆∞u tr·ªØ
if "ready_pdf_bytes" not in st.session_state: st.session_state.ready_pdf_bytes = None
if "ready_file_name" not in st.session_state: st.session_state.ready_file_name = None
if "uploader_key" not in st.session_state: st.session_state.uploader_key = 0
if "pdf_data" not in st.session_state: st.session_state.pdf_data = None
if "edit_lock" not in st.session_state: st.session_state.edit_lock = True
if "local_edit_count" not in st.session_state: st.session_state.local_edit_count = 0
if "current_doc_type" not in st.session_state: st.session_state.current_doc_type = "H√≥a ƒë∆°n"
if "invoice_view_page" not in st.session_state: st.session_state.invoice_view_page = 0

# Bi·∫øn ri√™ng cho Edit Mode
if "unc_edit_mode" not in st.session_state: st.session_state.unc_edit_mode = False
if "est_edit_mode" not in st.session_state: st.session_state.est_edit_mode = False
if "current_tour_id_est" not in st.session_state: st.session_state.current_tour_id_est = None
if "est_editor_key" not in st.session_state: st.session_state.est_editor_key = 0

# Initialize tab variables to avoid Pylance undefined errors
tab_est = tab_act = tab_rpt = None

# FIX L·ªñI OUT T√ÄI KHO·∫¢N
UPLOAD_FOLDER = ".uploaded_invoices"
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

DB_FILE = "invoice_app.db"

# ==========================================
# 2. X·ª¨ L√ù DATABASE (SQLite)
# ==========================================
@st.cache_resource
def get_connection():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def migrate_db_columns():
    conn = get_connection()
    c = conn.cursor()
    # Th√™m c√°c c·ªôt n·∫øu ch∆∞a c√≥ cho H√≥a ƒë∆°n/D·ª± √°n c≈©
    try: c.execute("ALTER TABLE invoices ADD COLUMN request_edit INTEGER DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE flight_tickets ADD COLUMN airline TEXT")
    except: pass
    try: c.execute("ALTER TABLE projects ADD COLUMN pending_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE projects ADD COLUMN type TEXT DEFAULT 'NORMAL'")
    except: pass
    try: c.execute("ALTER TABLE tour_items ADD COLUMN category TEXT")
    except: pass
    try: c.execute("ALTER TABLE tour_items ADD COLUMN times REAL DEFAULT 1")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN pending_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN request_delete INTEGER DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN request_edit_act INTEGER DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN tour_code TEXT")
    except: pass
    try: c.execute("ALTER TABLE invoices ADD COLUMN cost_code TEXT")
    except: pass
    
    # --- C·∫≠p nh·∫≠t cho B√†n Giao Tour (M·ªõi) ---
    try: c.execute("ALTER TABLE tours ADD COLUMN pickup_location TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN pickup_time TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN flight_code TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN driver_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN driver_phone TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN car_plate TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN car_type TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN itinerary_summary TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN guide_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN guide_phone TEXT")
    except: pass
    try: c.execute("CREATE TABLE IF NOT EXISTS ocr_learning (keyword TEXT UNIQUE, weight INTEGER DEFAULT 1)")
    except: pass

    # --- B·∫£ng ƒêi·ªÉm Tham Quan (M·ªõi) ---
    try: c.execute('''CREATE TABLE IF NOT EXISTS tour_sightseeings (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        name TEXT,
        address TEXT,
        quantity INTEGER,
        note TEXT
    )''')
    except: pass
    
    # --- C·∫≠p nh·∫≠t c·ªôt T√†i ch√≠nh cho KS/NH (M·ªõi) ---
    try: c.execute("ALTER TABLE tour_hotels ADD COLUMN total_amount REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tour_hotels ADD COLUMN deposit REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tour_restaurants ADD COLUMN total_amount REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tour_restaurants ADD COLUMN deposit REAL DEFAULT 0")
    except: pass
    # --- C·∫≠p nh·∫≠t c·ªôt Ng√†y v√† T√†i ch√≠nh cho Nh√† h√†ng/Tham quan (M·ªõi nh·∫•t) ---
    try: c.execute("ALTER TABLE tour_restaurants ADD COLUMN date TEXT")
    except: pass
    try: c.execute("ALTER TABLE tour_sightseeings ADD COLUMN date TEXT")
    except: pass
    try: c.execute("ALTER TABLE tour_sightseeings ADD COLUMN total_amount REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tour_sightseeings ADD COLUMN deposit REAL DEFAULT 0")
    except: pass

    # --- C·∫≠p nh·∫≠t cho Payment Reminders (M·ªõi) ---
    try: c.execute("ALTER TABLE payment_reminders ADD COLUMN cc_email TEXT")
    except: pass
    try: c.execute("ALTER TABLE payment_reminders ADD COLUMN sender_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE payment_reminders ADD COLUMN bank_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE payment_reminders ADD COLUMN bank_account TEXT")
    except: pass
    try: c.execute("ALTER TABLE payment_reminders ADD COLUMN bank_holder TEXT")
    except: pass

    # --- B·∫£ng L·ªãch Tr√¨nh Tour (M·ªõi) ---
    try: c.execute('''CREATE TABLE IF NOT EXISTS tour_itineraries (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        day_index INTEGER,
        content TEXT
    )''')
    except: pass

    # --- B·∫£ng Chi Ph√≠ Ph√°t Sinh (M·ªõi) ---
    try: c.execute('''CREATE TABLE IF NOT EXISTS tour_incurred_costs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        name TEXT,
        unit TEXT,
        quantity REAL,
        price REAL,
        total_amount REAL,
        deposit REAL DEFAULT 0,
        note TEXT
    )''')
    except: pass

    # --- B·∫£ng Booking D·ªãch V·ª• (M·ªõi) ---
    try: c.execute('''CREATE TABLE IF NOT EXISTS service_bookings (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        code TEXT UNIQUE,
        name TEXT,
        created_at TEXT,
        status TEXT DEFAULT 'active'
    )''')
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN type TEXT")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN details TEXT")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN customer_info TEXT")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN net_price REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN tax_percent REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN selling_price REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN profit REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN sale_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN hotel_code TEXT")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN room_type TEXT")
    except: pass
    try: c.execute("ALTER TABLE service_bookings ADD COLUMN guest_list TEXT")
    except: pass

    # --- B·∫£ng Kh√°ch H√†ng (M·ªõi) ---
    try: c.execute('''CREATE TABLE IF NOT EXISTS customers (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        phone TEXT,
        email TEXT,
        address TEXT,
        notes TEXT,
        created_at TEXT
    )''')
    except: pass
    try: c.execute("ALTER TABLE customers ADD COLUMN sale_name TEXT")
    except: pass

    # --- C·∫≠p nh·∫≠t c·ªôt m·ªõi cho Tour (Gi√° ch·ªët, Gi√° tr·∫ª em, Gi√° tr·ªã h·ª£p ƒë·ªìng) ---
    try: c.execute("ALTER TABLE tours ADD COLUMN final_tour_price REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN child_price REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN contract_value REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN final_qty REAL DEFAULT 0")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN child_qty REAL DEFAULT 0")
    except: pass

    # --- C·∫≠p nh·∫≠t th√¥ng tin kh√°ch h√†ng cho Tour ---
    try: c.execute("ALTER TABLE tours ADD COLUMN customer_name TEXT")
    except: pass
    try: c.execute("ALTER TABLE tours ADD COLUMN customer_phone TEXT")
    except: pass

    # --- C·∫≠p nh·∫≠t cho ph·∫ßn Danh s√°ch & D·ªãch v·ª• (M·ªõi) ---
    try: c.execute("ALTER TABLE tours ADD COLUMN handover_checklist TEXT")
    except: pass

    c.execute('''CREATE TABLE IF NOT EXISTS tour_guests (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        name TEXT,
        dob TEXT,
        hometown TEXT,
        cccd TEXT,
        type TEXT
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS tour_hotels (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        hotel_name TEXT,
        address TEXT,
        phone TEXT,
        total_rooms TEXT,
        room_type TEXT
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS tour_restaurants (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        meal_name TEXT,
        restaurant_name TEXT,
        address TEXT,
        phone TEXT,
        menu TEXT
    )''')

    # --- C·∫≠p nh·∫≠t m√£ tour cho d·ªØ li·ªáu c≈© ---
    try:
        old_tours = c.execute("SELECT id FROM tours WHERE tour_code IS NULL OR tour_code = ''").fetchall()
        for t in old_tours:
            code = ''.join(random.choices(string.ascii_uppercase, k=5))
            c.execute("UPDATE tours SET tour_code=? WHERE id=?", (code, t['id'])) # type: ignore
    except: pass
    
    # --- C·∫≠p nh·∫≠t d·ªØ li·ªáu c≈© ƒë·ªÉ hi·ªán th·ªã d·ª± √°n ---
    try: 
        c.execute("UPDATE projects SET type='NORMAL' WHERE type IS NULL OR type=''")
    except: pass

    # --- FIX QUAN TR·ªåNG: ƒê·∫¢M B·∫¢O B·∫¢NG TOURS T·ªíN T·∫†I KHI C·∫¨P NH·∫¨T ---
    # Ph·∫ßn n√†y gi√∫p t·∫°o b·∫£ng ngay c·∫£ khi DB ƒë√£ t·ªìn t·∫°i t·ª´ tr∆∞·ªõc
    c.execute('''CREATE TABLE IF NOT EXISTS tours (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_name TEXT,
        sale_name TEXT,
        start_date TEXT,
        end_date TEXT,
        guest_count INTEGER,
        created_at TEXT,
        est_profit_percent REAL DEFAULT 10.0,
        est_tax_percent REAL DEFAULT 8.0,
        status TEXT DEFAULT 'running'
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS tour_items (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        item_type TEXT, 
        category TEXT,
        description TEXT,
        unit TEXT,
        quantity REAL,
        times REAL DEFAULT 1,
        unit_price REAL,
        total_amount REAL
    )''')
    
    # --- B·∫£ng C√¥ng N·ª£ (M·ªõi) ---
    try: c.execute('''CREATE TABLE IF NOT EXISTS transaction_history (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ref_code TEXT,
        type TEXT,
        amount REAL,
        payment_method TEXT,
        note TEXT,
        created_at TEXT
    )''')
    except: pass

    conn.commit()

def init_db():
    conn = get_connection()
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users (id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT UNIQUE, password TEXT, role TEXT, status TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS invoices (
        id INTEGER PRIMARY KEY AUTOINCREMENT, type TEXT, date TEXT, invoice_number TEXT, invoice_symbol TEXT, 
        seller_name TEXT, buyer_name TEXT, pre_tax_amount REAL, tax_amount REAL, total_amount REAL, 
        file_name TEXT, status TEXT, edit_count INTEGER, created_at TEXT, memo TEXT, file_path TEXT, request_edit INTEGER DEFAULT 0
    )''')
    # Th√™m c·ªôt pending_name v√† type v√†o b·∫£ng projects
    c.execute('''CREATE TABLE IF NOT EXISTS projects (
        id INTEGER PRIMARY KEY AUTOINCREMENT, 
        project_name TEXT, 
        created_at TEXT,
        pending_name TEXT,
        type TEXT DEFAULT 'NORMAL'
    )''')
    c.execute('''CREATE TABLE IF NOT EXISTS project_links (id INTEGER PRIMARY KEY AUTOINCREMENT, project_id INTEGER, invoice_id INTEGER)''')
    c.execute('''CREATE TABLE IF NOT EXISTS company_info (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT, address TEXT, phone TEXT, logo_base64 TEXT)''')
    
    # B·∫£ng V√© m√°y bay
    c.execute('''CREATE TABLE IF NOT EXISTS flight_tickets (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ticket_code TEXT,
        flight_date TEXT,
        route TEXT,
        passenger_names TEXT,
        file_path TEXT,
        created_at TEXT,
        airline TEXT
    )''')
    
    # B·∫£ng ƒêo√†n bay (C≈© - Gi·ªØ nguy√™n ƒë·ªÉ t∆∞∆°ng th√≠ch)
    c.execute('''CREATE TABLE IF NOT EXISTS flight_groups (id INTEGER PRIMARY KEY AUTOINCREMENT, group_name TEXT, created_at TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS flight_group_links (id INTEGER PRIMARY KEY AUTOINCREMENT, group_id INTEGER, ticket_id INTEGER)''')

    # --- B·∫¢NG BOOKING D·ªäCH V·ª§ ---
    c.execute('''CREATE TABLE IF NOT EXISTS service_bookings (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        code TEXT UNIQUE,
        name TEXT,
        created_at TEXT,
        status TEXT DEFAULT 'active'
    )''')

    # --- B·∫¢NG QU·∫¢N L√ù TOUR  ---
    c.execute('''CREATE TABLE IF NOT EXISTS tours (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_name TEXT,
        sale_name TEXT,
        start_date TEXT,
        end_date TEXT,
        guest_count INTEGER,
        created_at TEXT,
        est_profit_percent REAL DEFAULT 10.0,
        est_tax_percent REAL DEFAULT 8.0,
        status TEXT DEFAULT 'running'
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS tour_items (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        tour_id INTEGER,
        item_type TEXT, 
        category TEXT,
        description TEXT,
        unit TEXT,
        quantity REAL,
        times REAL DEFAULT 1,
        unit_price REAL,
        total_amount REAL
    )''')
    # item_type: 'EST' (D·ª± to√°n), 'ACT' (Quy·∫øt to√°n)

    # --- B·∫£ng C√¥ng N·ª£ (M·ªõi) ---
    c.execute('''CREATE TABLE IF NOT EXISTS transaction_history (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ref_code TEXT,
        type TEXT,
        amount REAL,
        payment_method TEXT,
        note TEXT,
        created_at TEXT
    )''')

    # --- TH√äM V√ÄO TRONG H√ÄM init_db() ---
    c.execute('''CREATE TABLE IF NOT EXISTS payment_reminders (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        ref_code TEXT,       -- M√£ Booking/Tour
        ref_name TEXT,       -- T√™n kh√°ch/Tour
        amount REAL,         -- S·ªë ti·ªÅn c·∫ßn thu
        due_date TEXT,       -- Ng√†y h·∫πn th√¥ng b√°o l·∫°i
        receiver_email TEXT, -- Email ng∆∞·ªùi nh·∫≠n (N·ªôi b·ªô ho·∫∑c Kh√°ch)
        content TEXT,        -- N·ªôi dung nh·∫Øc
        status TEXT,         -- 'sent_1': ƒê√£ g·ª≠i l·∫ßn 1, 'sent_2': ƒê√£ g·ª≠i l·∫ßn 2 (ho√†n t·∫•t)
        created_at TEXT
    )''')

    c.execute("SELECT * FROM users WHERE username = 'admin'")
    if not c.fetchone():
        admin_pw = hashlib.sha256("admin123".encode()).hexdigest()
        c.execute("INSERT INTO users (username, password, role, status) VALUES (?, ?, ?, ?)", ('admin', admin_pw, 'admin', 'approved'))
    
    c.execute("SELECT * FROM company_info WHERE id = 1")
    if not c.fetchone():
        c.execute("INSERT INTO company_info (name, address, phone, logo_base64) VALUES (?, ?, ?, ?)", ('T√™n C√¥ng Ty C·ªßa B·∫°n', 'ƒê·ªãa ch·ªâ...', '090...', ''))

    conn.commit()

if not st.session_state.db_initialized:
    init_db()
    st.session_state.db_initialized = True

# Lu√¥n ch·∫°y migration ƒë·ªÉ ƒë·∫£m b·∫£o c·ªôt m·ªõi ƒë∆∞·ª£c th√™m v√†o (Fix l·ªói Admin kh√¥ng nh·∫≠n y√™u c·∫ßu)
migrate_db_columns()

# --- C√ÅC H√ÄM H·ªñ TR·ª¢ ---
@overload
def run_query(query: str, params: Any = ..., fetch_one: Literal[False] = ..., commit: Literal[False] = ...) -> List[Any]: ...

@overload
def run_query(query: str, params: Any, fetch_one: Literal[True], commit: Literal[False] = ...) -> Any: ...

@overload
def run_query(query: str, *, fetch_one: Literal[True], commit: Literal[False] = ...) -> Any: ...

@overload
def run_query(query: str, params: Any = ..., fetch_one: Any = ..., *, commit: Literal[True]) -> bool: ...

def run_query(query, params=(), fetch_one=False, commit=False):
    conn = get_connection()
    c = conn.cursor()
    try:
        c.execute(query, params)
        if commit:
            conn.commit()
            return True
        if fetch_one:
            return c.fetchone()
        return c.fetchall()
    except Exception as e:
        print(f"L·ªói truy v·∫•n DB: {e}")
        if commit: return False
        if fetch_one: return None
        return []

def run_query_many(query, data):
    """Th·ª±c thi nhi·ªÅu c√¢u l·ªánh (th∆∞·ªùng l√† INSERT) c√πng l√∫c."""
    conn = get_connection()
    c = conn.cursor()
    try:
        c.executemany(query, data)
        conn.commit()
        return True
    except Exception as e:
        print(f"L·ªói truy v·∫•n DB (many): {e}")
        return False

def save_customer_check(name, phone, sale_name=None):
    """L∆∞u kh√°ch h√†ng m·ªõi n·∫øu ch∆∞a t·ªìn t·∫°i"""
    if not name: return
    try:
        exist = run_query("SELECT id FROM customers WHERE name=?", (name,), fetch_one=True)
        if not exist:
            data = {'name': name, 'phone': phone, 'created_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
            if sale_name:
                data['sale_name'] = sale_name
            add_row_to_table('customers', data)
    except: pass

def hash_pass(password):
    return hashlib.sha256(str.encode(password)).hexdigest()

def save_file_local(file_bytes, original_name):
    try:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        clean_name = re.sub(r'[\\/*?:"<>|]', "", original_name)
        if not clean_name.lower().endswith('.pdf'):
            clean_name = os.path.splitext(clean_name)[0] + ".pdf"
            
        final_name = f"{ts}_{clean_name}"
        file_path = os.path.join(UPLOAD_FOLDER, final_name)
        
        with open(file_path, "wb") as f:
            f.write(file_bytes)
                
        return file_path, final_name
    except: return None, None

def format_vnd(amount):
    if amount is None: return "0"
    try: return "{:,.0f}".format(float(amount)).replace(",", ".")
    except: return "0"

@st.cache_data
def get_company_data():
    row = run_query("SELECT * FROM company_info WHERE id = 1", fetch_one=True)
    if isinstance(row, sqlite3.Row):
        return {'name': row['name'], 'address': row['address'], 'phone': row['phone'], 'logo_b64_str': row['logo_base64']}
    return {'name': 'Company', 'address': '...', 'phone': '...', 'logo_b64_str': ''}

def update_company_info(name, address, phone, logo_bytes=None):
    b64_str = base64.b64encode(logo_bytes).decode('utf-8') if logo_bytes else ""
    if not logo_bytes:
        old = run_query("SELECT logo_base64 FROM company_info WHERE id = 1", fetch_one=True)
        if isinstance(old, sqlite3.Row): b64_str = old['logo_base64'] # type: ignore
    run_query("UPDATE company_info SET name=?, address=?, phone=?, logo_base64=? WHERE id=1", (name, address, phone, b64_str), commit=True)
    get_company_data.clear()# type: ignore

# --- H√ÄM G·ª¨I EMAIL ---
def send_email_notification(to_email, subject, body_html, cc_emails=None):
    """H√†m g·ª≠i email qua SMTP Gmail"""
    try:
        # L·∫•y c·∫•u h√¨nh t·ª´ secrets.toml
        email_sender = st.secrets["email"]["sender"]
        email_password = st.secrets["email"]["password"]
        
        msg = MIMEMultipart()
        msg['From'] = f"Bali Tourist System <{email_sender}>"
        msg['To'] = to_email
        msg['Subject'] = subject
        if cc_emails:
            msg['Cc'] = cc_emails
        msg.attach(MIMEText(body_html, 'html'))

        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(email_sender, email_password)
        
        # X·ª≠ l√Ω danh s√°ch ng∆∞·ªùi nh·∫≠n (To + Cc)
        recipients = [to_email]
        if cc_emails:
            if isinstance(cc_emails, str):
                recipients.extend([e.strip() for e in cc_emails.split(',') if e.strip()])
            elif isinstance(cc_emails, list):
                recipients.extend(cc_emails)
        
        server.send_message(msg, to_addrs=recipients)
        server.quit()
        return True, "ƒê√£ g·ª≠i mail th√†nh c√¥ng!"
    except Exception as e:
        return False, f"L·ªói g·ª≠i mail: {str(e)}"

# --- H√ÄM T·ª∞ ƒê·ªòNG QU√âT & G·ª¨I L·∫¶N 2 ---
def check_and_send_due_reminders():
    """Ki·ªÉm tra c√°c l·ªãch h·∫πn ƒë·∫øn ng√†y h√¥m nay ƒë·ªÉ g·ª≠i email l·∫ßn 2"""
    now_str = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    # T√¨m c√°c nh·∫Øc h·∫πn c√≥ ng√†y = h√¥m nay (ho·∫∑c qu√° kh·ª©) V√Ä m·ªõi ch·ªâ g·ª≠i l·∫ßn 1 ('sent_1')
    reminders = run_query("SELECT * FROM payment_reminders WHERE status='sent_1' AND due_date <= ?", (now_str,))
    
    count = 0
    if reminders:
        for row in reminders:
            r = dict(row)
            # G·ª≠i email l·∫ßn 2
            cc = r.get('cc_email', '')
            sender = r.get('sender_name', 'Bali Tourist System')
            
            # [FIX] Format ng√†y hi·ªÉn th·ªã trong mail (DD/MM/YYYY HH:MM)
            try:
                d_obj = datetime.strptime(r['due_date'], '%Y-%m-%d %H:%M:%S')
                date_display = d_obj.strftime('%H:%M %d/%m/%Y')
            except:
                try:
                    d_obj = datetime.strptime(r['due_date'], '%Y-%m-%d')
                    date_display = d_obj.strftime('%d/%m/%Y')
                except:
                    date_display = r['due_date']

            # [NEW] Bank Info for Automated Email
            bank_info_html = ""
            if r.get('bank_name') and r.get('bank_account'):
                bank_info_html = f"""
                <div style="background-color: #f8f9fa; padding: 15px; border-radius: 5px; margin: 15px 0;">
                    <h4 style="margin-top: 0;">üè¶ TH√îNG TIN CHUY·ªÇN KHO·∫¢N</h4>
                    <p><strong>Ng√¢n h√†ng:</strong> {r['bank_name']}</p>
                    <p><strong>S·ªë t√†i kho·∫£n:</strong> {r['bank_account']}</p>
                    <p><strong>Ch·ªß t√†i kho·∫£n:</strong> {r.get('bank_holder', '')}</p>
                </div>
                """

            subject = f"üîî [NH·∫ÆC H·∫∏N L·∫¶N 2] Thanh to√°n cho m√£ {r['ref_code']}"
            content = f"""
            <h3>üîî NH·∫ÆC H·∫∏N THANH TO√ÅN (L·∫¶N 2)</h3>
            <p>H·ªá th·ªëng t·ª± ƒë·ªông nh·∫Øc b·∫°n v·ªÅ kho·∫£n thanh to√°n ƒë√£ ƒë·∫øn h·∫πn:</p>
            <ul>
                <li><strong>M√£ h·ªì s∆°:</strong> {r['ref_code']}</li>
                <li><strong>T√™n:</strong> {r['ref_name']}</li>
                <li><strong>S·ªë ti·ªÅn:</strong> {format_vnd(r['amount'])} VND</li>
                <li><strong>N·ªôi dung:</strong> {r['content']}</li>
                <li><strong>Ng√†y h·∫πn:</strong> {date_display}</li>
            </ul>
            {bank_info_html}
            <p>Vui l√≤ng ki·ªÉm tra v√† x·ª≠ l√Ω.</p>
            <p>Tr√¢n tr·ªçng,<br>{sender}</p>
            """
            success, msg = send_email_notification(r['receiver_email'], subject, content, cc_emails=cc)
            if success:
                # C·∫≠p nh·∫≠t tr·∫°ng th√°i th√†nh sent_2 (ƒë√£ xong)
                run_query("UPDATE payment_reminders SET status='sent_2' WHERE id=?", (r['id'],), commit=True)
                count += 1
    return count

# --- H√ÄM H·ªñ TR·ª¢ L·ªäCH √ÇM/D∆Ø∆†NG ---
def convert_solar_to_lunar(solar_date):
    """Chuy·ªÉn D∆∞∆°ng l·ªãch -> √Çm l·ªãch"""
    try:
        ld = LunarDate.fromSolarDate(solar_date.year, solar_date.month, solar_date.day)
        return f"{ld.day:02d}/{ld.month:02d}/{ld.year} (√Çm l·ªãch)"
    except:
        return "Kh√¥ng x√°c ƒë·ªãnh"

def convert_lunar_to_solar(day, month, year, is_leap=False):
    """Chuy·ªÉn √Çm l·ªãch -> D∆∞∆°ng l·ªãch"""
    try:
        sd = LunarDate(year, month, day, is_leap).toSolarDate()
        return sd # Tr·∫£ v·ªÅ object date
    except ValueError:
        return None

def get_tour_financials(tour_id, tour_info):
    """
    T√≠nh to√°n doanh thu v√† chi ph√≠ cho m·ªôt tour.
    """
    # L·∫•y t·ªïng chi ph√≠ quy·∫øt to√°n (ACT) t·ª´ b·∫£ng k√™
    act_items = run_query("SELECT SUM(total_amount) as total FROM tour_items WHERE tour_id=? AND item_type='ACT'", (tour_id,), fetch_one=True)
    act_cost_items = act_items['total'] if act_items and act_items['total'] else 0

    # L·∫•y t·ªïng chi ph√≠ t·ª´ h√≥a ƒë∆°n ƒë·∫ßu v√†o li√™n k·∫øt v·ªõi tour (kh√¥ng t√≠nh UNC)
    inv_items = run_query("SELECT SUM(total_amount) as total FROM invoices WHERE cost_code=? AND status='active' AND type='IN' AND invoice_number NOT LIKE '%UNC%'", (tour_info['tour_code'],), fetch_one=True)
    inv_cost = inv_items['total'] if inv_items and inv_items['total'] else 0

    cost = (act_cost_items or 0) + (inv_cost or 0)

    # L·∫•y t·ªïng chi ph√≠ d·ª± to√°n (EST) ƒë·ªÉ t√≠nh doanh thu n·∫øu c·∫ßn
    est_items = run_query("SELECT SUM(total_amount) as total FROM tour_items WHERE tour_id=? AND item_type='EST'", (tour_id,), fetch_one=True)
    est_cost = est_items['total'] if est_items and est_items['total'] else 0

    # T√≠nh doanh thu d·ª±a tr√™n gi√° ch·ªët
    t_dict = dict(tour_info)
    final_price = float(t_dict.get('final_tour_price', 0) or 0)
    child_price = float(t_dict.get('child_price', 0) or 0)
    final_qty = float(t_dict.get('final_qty', 0) or 0)
    child_qty = float(t_dict.get('child_qty', 0) or 0)
    if final_qty == 0: final_qty = float(t_dict.get('guest_count', 1))
    
    revenue = (final_price * final_qty) + (child_price * child_qty)

    # N·∫øu chi ph√≠ quy·∫øt to√°n ch∆∞a c√≥, d√πng t·∫°m chi ph√≠ d·ª± to√°n
    if cost == 0 and est_cost > 0:
        cost = est_cost

    return revenue, cost
# ==========================================
# 3. CSS & GIAO DI·ªÜN HI·ªÜN ƒê·∫†I
# ==========================================
comp = get_company_data()
st.markdown("""<style>
/* --- BASE & ANIMATION --- */
@keyframes fadeIn { 0% { opacity: 0; transform: translateY(10px); } 100% { opacity: 1; transform: translateY(0); } }
.stApp {
    background-color: #f8f9fa;
    font-family: 'Inter', 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;
    animation: fadeIn 0.5s ease-in-out;
}

/* --- TYPOGRAPHY & LABELS --- */
h1, h2, h3, h4, h5, h6 { color: #2c3e50; }
div[data-testid="stMarkdownContainer"] p { font-weight: 400; white-space: normal; word-break: break-word; }
.company-info-text p, .report-card p { white-space: normal !important; }

/* --- MODERN INPUTS --- */
.stTextInput input, .stNumberInput input, .stSelectbox div[data-baseweb="select"], .stTextArea textarea, .stDateInput input {
    border-radius: 10px !important;
    border: 1px solid #e0e0e0 !important;
    padding: 10px 12px !important;
    background-color: #ffffff !important;
    transition: all 0.3s;
    font-size: 0.95rem;
}
.stTextInput input:focus, .stNumberInput input:focus, .stTextArea textarea:focus, .stDateInput input:focus {
    border-color: #56ab2f !important;
    box-shadow: 0 4px 12px rgba(86, 171, 47, 0.15) !important;
}

/* --- BUTTONS --- */
.stButton button {
    border-radius: 12px !important;
    font-weight: 600;
    font-size: 1rem;
    padding: 0.6rem 1.2rem !important;
    border: none !important;
    box-shadow: 0 4px 6px rgba(0,0,0,0.05);
    transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    white-space: normal !important;
    height: auto !important;
    min-height: 2.5rem;
}
.stButton button:hover {
    transform: translateY(-2px);
    box-shadow: 0 8px 15px rgba(0,0,0,0.1);
}
.stButton button[kind="primary"] {
    background: linear-gradient(90deg, #56ab2f 0%, #a8e063 100%);
    color: white;
}
.stButton button[kind="secondary"] {
    background-color: #f1f3f5;
    color: #333;
}

/* --- COMPANY HEADER --- */
.company-header-container {
    display: flex; align-items: center; justify-content: center; gap: 30px;
    padding: 25px 40px; background: rgba(255, 255, 255, 0.8);
    backdrop-filter: blur(10px); border-radius: 20px;
    box-shadow: 0 8px 32px rgba(0,0,0,0.05); margin-bottom: 30px;
    border: 1px solid rgba(255,255,255,0.3); flex-wrap: nowrap !important;
}
.company-logo-img { height: 70px; width: auto; object-fit: contain; flex-shrink: 0; }
.company-info-text { text-align: left; flex: 1; display: flex; flex-direction: column; justify-content: center; white-space: normal; }
.company-info-text h1 { margin: 0; font-size: 1.8rem; color: #2e7d32; font-weight: 800; line-height: 1.2; }
.company-info-text p { margin: 5px 0 0 0; color: #555; font-size: 0.9rem; font-weight: 500; display: flex; align-items: center; gap: 10px; }

/* --- CARD STYLES --- */
.report-card, .login-container {
    background-color: white; border: none; border-radius: 20px;
    padding: 25px; margin-bottom: 25px;
    box-shadow: 0 10px 30px rgba(0,0,0,0.04);
    transition: all 0.3s ease;
}
.report-card:hover { transform: translateY(-5px); box-shadow: 0 20px 40px rgba(0,0,0,0.08); }

/* --- MONEY BOX --- */
.money-box {
    background: linear-gradient(135deg, #00b09b, #96c93d) !important;
    color: #ffffff !important; padding: 25px; border-radius: 20px;
    box-shadow: 0 15px 30px -5px rgba(0, 176, 155, 0.3);
    font-size: clamp(1.2rem, 3vw, 2.5rem); font-weight: 800;
    text-align: center; margin: 1.5rem 0; width: 100%;
    text-shadow: 0 2px 4px rgba(0,0,0,0.1); letter-spacing: 1px;
    white-space: normal; word-wrap: break-word;
    transition: transform 0.3s ease;
}
.money-box:hover { transform: scale(1.02); }

/* --- MODERN TABS --- */
div[data-baseweb="tab-list"] { border-bottom: 2px solid #e0e0e0; }
button[data-baseweb="tab"] {
    background-color: transparent !important; border-bottom: 2px solid transparent !important;
    padding-bottom: 10px !important; margin-bottom: -2px !important; transition: all 0.3s !important;
}
button[data-baseweb="tab"]:hover { background-color: #f1f3f5 !important; }
button[aria-selected="true"] {
    border-bottom-color: #56ab2f !important; font-weight: 600; color: #56ab2f !important;
}

/* --- ENHANCED EXPANDER --- */
div[data-testid="stExpander"] {
    border: 1px solid #e0e0e0 !important; border-radius: 15px !important;
    overflow: hidden; box-shadow: none !important; background-color: #fff;
}
div[data-testid="stExpander"] > details > summary {
    font-weight: 600; font-size: 1.05rem; background-color: #fafafa;
    padding: 0.75rem 1rem !important;
}
div[data-testid="stExpander"] > details > summary:hover { background-color: #f1f3f5; }

/* --- DATA EDITOR --- */
div[data-testid="stDataEditor"] {
    border-radius: 15px; overflow: hidden;
    border: 1px solid #f0f0f0; box-shadow: 0 4px 12px rgba(0,0,0,0.03);
}

/* --- FINANCE SUMMARY CARDS --- */
.finance-summary-card {
    background-color: #ffffff; border: 1px solid #e9ecef; border-radius: 15px;
    padding: 20px; margin-top: 15px;
}
.finance-summary-card .row {
    display: flex; justify-content: space-between; align-items: center;
    padding: 8px 0; border-bottom: 1px solid #f1f3f5;
}
.finance-summary-card .row:last-child { border-bottom: none; }
.finance-summary-card .row span { color: #495057; }
.finance-summary-card .row b { color: #212529; }
.finance-summary-card .total-row {
    font-size: 1.2em; font-weight: bold; color: #2e7d32; padding-top: 15px;
}
.finance-summary-card .pax-price {
    text-align: right; font-size: 0.9em; color: #6c757d; margin-top: 5px;
}
.profit-summary-card {
    background-color: #e3f2fd; padding: 20px; border-radius: 15px;
    text-align: center; border: 1px solid #90caf9; margin-top: 10px;
}
.profit-summary-card h3 {
    margin: 0; color: #1565c0; font-size: 1.1rem; font-weight: 600;
}
.profit-summary-card .formula {
    font-size: 1.8em; font-weight: bold; color: #1e88e5; margin-top: 10px;
}
.profit-summary-card .formula .result { color: #d32f2f; }

/* --- RESPONSIVE --- */
@media only screen and (max-width: 600px) {
    .company-header-container { flex-direction: column; text-align: center; gap: 10px; flex-wrap: wrap !important; }
    .company-info-text { text-align: center; }
    .company-info-text p { justify-content: center; }
}
</style>""", unsafe_allow_html=True)

def convert_image_to_pdf(image_file):
    try:
        img = Image.open(image_file)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        img_width, img_height = img.size
        pdf_buffer = io.BytesIO()
        c = canvas.Canvas(pdf_buffer, pagesize=(img_width, img_height))
        temp_img_path = f"temp_img_{int(time.time())}.jpg"
        img.save(temp_img_path)
        c.drawImage(temp_img_path, 0, 0, img_width, img_height)
        c.save()
        if os.path.exists(temp_img_path): os.remove(temp_img_path)
        pdf_buffer.seek(0)
        return pdf_buffer.getvalue()
    except Exception as e:
        return None

# --- H√ÄM OCR ---
def perform_ocr(image_input, lang='vie'):
    """
    Th·ª±c hi·ªán OCR tr√™n ·∫£nh v·ªõi c√°c b∆∞·ªõc ti·ªÅn x·ª≠ l√Ω n√¢ng cao s·ª≠ d·ª•ng OpenCV ƒë·ªÉ c·∫£i thi·ªán ƒë·ªô ch√≠nh x√°c.
    """
    # Check for dependencies and provide clear feedback.
    # This also helps static analysis tools like Pylance understand that `np` and `cv2` are not None below.
    if not HAS_OCR or pytesseract is None:
        st.toast("‚ö†Ô∏è Tesseract OCR ch∆∞a ƒë∆∞·ª£c c√†i ƒë·∫∑t.", icon="üö®")
        return ""
    if not HAS_CV or np is None or cv2 is None:
        st.toast("‚ö†Ô∏è OpenCV ho·∫∑c Numpy ch∆∞a ƒë∆∞·ª£c c√†i ƒë·∫∑t.", icon="üö®")
        return ""
    try:
        # 1. Load ·∫£nh t·ª´ input (c√≥ th·ªÉ l√† file stream ho·∫∑c ƒë·ªëi t∆∞·ª£ng PIL)
        if isinstance(image_input, Image.Image):
            img = image_input
        else:
            image_input.seek(0)
            img = Image.open(image_input)

        # 2. Chuy·ªÉn ƒë·ªïi sang ƒë·ªãnh d·∫°ng OpenCV
        # Chuy·ªÉn sang ·∫£nh x√°m (grayscale) v√† numpy array ƒë·ªÉ x·ª≠ l√Ω
        img_np = np.array(img.convert('L'))

        # 3. TƒÉng k√≠ch th∆∞·ªõc ·∫£nh (Upscaling)
        # OCR ho·∫°t ƒë·ªông t·ªët h∆°n v·ªõi ·∫£nh c√≥ DPI cao (kho·∫£ng 300). Vi·ªác upscale ·∫£nh nh·ªè gi√∫p nh·∫≠n di·ªán k√Ω t·ª± t·ªët h∆°n.
        h, w = img_np.shape
        if w < 2000:
            scale = 2000 / w
            new_w, new_h = int(w * scale), int(h * scale)
            # S·ª≠ d·ª•ng Lanczos interpolation cho k·∫øt qu·∫£ s·∫Øc n√©t khi ph√≥ng to
            img_np = cv2.resize(img_np, (new_w, new_h), interpolation=cv2.INTER_LANCZOS4)

        # 4. Gi·∫£m nhi·ªÖu (Noise Reduction)
        # S·ª≠ d·ª•ng Median Blur hi·ªáu qu·∫£ ƒë·ªÉ lo·∫°i b·ªè nhi·ªÖu "mu·ªëi ti√™u" (salt-and-pepper noise) m√† kh√¥ng l√†m m·ªù c√°c c·∫°nh qu√° nhi·ªÅu.
        img_np = cv2.medianBlur(img_np, 3)

        # 5. Binarization th√¥ng minh (Adaptive Thresholding)
        # ƒê√¢y l√† b∆∞·ªõc quan tr·ªçng nh·∫•t, thay th·∫ø cho vi·ªác tƒÉng contrast v√† d√πng ng∆∞·ª°ng c·ªë ƒë·ªãnh.
        # N√≥ t·ª± ƒë·ªông t√≠nh to√°n ng∆∞·ª°ng cho c√°c v√πng ·∫£nh nh·ªè, r·∫•t hi·ªáu qu·∫£ v·ªõi ·∫£nh c√≥ ƒëi·ªÅu ki·ªán s√°ng kh√¥ng ƒë·ªìng ƒë·ªÅu.
        img_processed = cv2.adaptiveThreshold(
            img_np,
            255,  # Gi√° tr·ªã t·ªëi ƒëa cho pixel
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,  # Ph∆∞∆°ng ph√°p t√≠nh ng∆∞·ª°ng d·ª±a tr√™n v√πng l√¢n c·∫≠n theo ph√¢n ph·ªëi Gaussian
            cv2.THRESH_BINARY, # Chuy·ªÉn ·∫£nh th√†nh ƒëen v√† tr·∫Øng
            15,  # K√≠ch th∆∞·ªõc v√πng l√¢n c·∫≠n (block size), n√™n l√† s·ªë l·∫ª
            4    # H·∫±ng s·ªë C, m·ªôt gi√° tr·ªã ƒë∆∞·ª£c tr·ª´ ƒëi t·ª´ gi√° tr·ªã trung b√¨nh t√≠nh ƒë∆∞·ª£c
        )

        # 6. C·∫•u h√¨nh Tesseract ƒë·ªÉ c√≥ k·∫øt qu·∫£ t·ªët nh·∫•t
        # --psm 4: Gi·∫£ ƒë·ªãnh vƒÉn b·∫£n l√† m·ªôt c·ªôt duy nh·∫•t v·ªõi k√≠ch th∆∞·ªõc thay ƒë·ªïi (t·ªët cho h√≥a ƒë∆°n, UNC).
        # --oem 3: S·ª≠ d·ª•ng engine m·∫∑c ƒë·ªãnh (k·∫øt h·ª£p Legacy v√† LSTM), th∆∞·ªùng cho k·∫øt qu·∫£ ·ªïn ƒë·ªãnh.
        config = '--psm 4 --oem 3'
        text = pytesseract.image_to_string(img_processed, lang='vie+eng', config=config) if pytesseract else ""
        return text
    except Exception as e:
        print(f"OCR Error: {e}")
        return ""

def extract_money_smart(line):
    cleaned = re.sub(r'[^\d.,]', '', line) 
    potential_numbers = []
    raw_digits = re.findall(r'\d+', cleaned)
    for rd in raw_digits:
        if len(rd) > 8 and str(rd).startswith('0'): continue
        if len(rd) >= 4: potential_numbers.append(float(rd))
    matches = re.findall(r'\d[\d.,\s]*\d', line) 
    for m in matches:
        s = m.replace('VND', '').replace('ƒë', '').replace(' ', '').strip()
        if len(s) > 8 and s.startswith('0'): continue
        try:
            val = 0.0
            if ',' in s and '.' not in s: val = float(s.replace(',', ''))
            elif '.' in s and ',' not in s: val = float(s.replace('.', ''))
            elif ',' in s and '.' in s:
                last_dot = s.rfind('.')
                last_comma = s.rfind(',')
                if last_dot > last_comma: val = float(s.replace(',', '')) 
                else: val = float(s.replace('.', '').replace(',', '.'))
            else: val = float(s)
            if (val > 2030 or val < 1900) and val > 1000:
                potential_numbers.append(val)
        except: pass
    return potential_numbers

def extract_numbers_from_line_basic(line):
    clean_line = line.replace("-", "").replace("VND", "").replace("ƒë", "").strip()
    raw_integers = re.findall(r'(?<!\d)\d{4,}(?!\d)', clean_line)
    results = []
    for n in raw_integers:
        try:
            val = float(n)
            if not (1990 <= val <= 2030): results.append(val)
        except: pass
    return results

# --- X·ª¨ L√ù H√ìA ƒê∆†N & UNC (LOGIC C≈®) ---
def extract_data_smart(file_obj, is_image, doc_type="H√≥a ƒë∆°n"):
    text_content = ""
    msg = None
    try:
        if is_image:
            if HAS_OCR:
                # G·ªçi h√†m OCR ƒë√£ s·ª≠a ƒë·ªïi
                text_content = perform_ocr(file_obj)
                if not text_content.strip(): msg = "Hic, ·∫£nh m·ªù qu√° ho·∫∑c kh√¥ng t√¨m th·∫•y ch·ªØ s·ªë n√†o üò≠."
            else: msg = "‚ö†Ô∏è T√¨nh y√™u ∆°i, m√°y ch∆∞a c√†i Tesseract OCR n√™n kh√¥ng ƒë·ªçc ƒë∆∞·ª£c ·∫£nh n√®."
        else:
            # X·ª≠ l√Ω PDF (C·∫£ text v√† scan)
            file_obj.seek(0)
            with pdfplumber.open(file_obj) as pdf:
                for page in pdf.pages: 
                    extracted = page.extract_text()
                    if extracted and len(extracted.strip()) > 10: 
                        text_content += extracted + "\n"
                    else:
                        if HAS_OCR:
                            im = page.to_image(resolution=300).original
                            text_content += perform_ocr(im) + "\n"
            
            if not text_content.strip(): 
                if not HAS_OCR: msg = "‚ö†Ô∏è File PDF n√†y l√† ·∫£nh scan, c·∫ßn c√†i Tesseract OCR ƒë·ªÉ ƒë·ªçc."
                else: msg = "‚ö†Ô∏è File tr·∫Øng tinh ho·∫∑c kh√¥ng ƒë·ªçc ƒë∆∞·ª£c n·ªôi dung."

    except Exception as e: return None, f"L·ªói x√≠u xiu: {str(e)}"
    
    info = {"date": "", "seller": "", "buyer": "", "inv_num": "", "inv_sym": "", "pre_tax": 0.0, "tax": 0.0, "total": 0.0, "content": ""}
    if not text_content: return info, msg

    lines = text_content.split('\n')
    all_found_numbers = set()

    # --- T√åM NG√ÄY TH√ÅNG ---
    m_date = re.search(r'(?:Ng√†y|ng√†y)\s+(\d{1,2})\s+(?:th√°ng|Th√°ng|[/.-])\s+(\d{1,2})\s+(?:nƒÉm|NƒÉm|[/.-])\s+(\d{4})', text_content)
    if m_date: 
        try: info["date"] = f"{int(m_date.group(1)):02d}/{int(m_date.group(2)):02d}/{m_date.group(3)}"
        except: pass
    else:
        m_date_alt = re.search(r'(\d{2}/\d{2}/\d{4})', text_content)
        if m_date_alt: info["date"] = m_date_alt.group(1)

    # --- LOGIC X·ª¨ L√ù S·ªê TI·ªÄN ---
    if doc_type == "H√≥a ƒë∆°n":
        # ... (Gi·ªØ nguy√™n logic H√≥a ƒë∆°n c≈© c·ªßa b·∫°n ·ªü ƒë√¢y n·∫øu c·∫ßn, ho·∫∑c d√πng ƒëo·∫°n d∆∞·ªõi ƒë√¢y)
        m_no = re.search(r'(?:S·ªë h√≥a ƒë∆°n|S·ªë Hƒê|S·ªë|No)[:\s\.]*(\d{1,8})\b', text_content, re.IGNORECASE)
        if m_no: info["inv_num"] = m_no.group(1).zfill(7)
        m_sym = re.search(r'(?:K√Ω hi·ªáu|M·∫´u s·ªë|Serial)[:\s\.]*([A-Z0-9]{1,2}[A-Z0-9/-]{3,10})', text_content, re.IGNORECASE)
        if m_sym: info["inv_sym"] = m_sym.group(1)
        
        for line in lines:
            line_l = line.lower()
            nums = extract_money_smart(line)
            for n in nums: all_found_numbers.add(n)
            if not nums: continue
            val = max(nums)
            if any(kw in line_l for kw in ["thanh to√°n", "t·ªïng c·ªông", "c·ªông ti·ªÅn h√†ng"]): info["total"] = val
            elif any(kw in line_l for kw in ["ti·ªÅn h√†ng", "th√†nh ti·ªÅn", "tr∆∞·ªõc thu·∫ø"]): info["pre_tax"] = val
            elif "thu·∫ø" in line_l and "su·∫•t" not in line_l: info["tax"] = val
        
        if info["total"] == 0 and all_found_numbers: info["total"] = max(all_found_numbers)
        if info["pre_tax"] == 0: info["pre_tax"] = round(info["total"] / 1.08)
        if info["tax"] == 0: info["tax"] = info["total"] - info["pre_tax"]
        
        # T√¨m Buyer/Seller cho H√≥a ƒë∆°n
        for line in lines[:35]:
            l_c = line.strip()
            if re.search(r'^(ƒê∆°n v·ªã b√°n|Ng∆∞·ªùi b√°n|B√™n A|Nh√† cung c·∫•p)', l_c, re.IGNORECASE): 
                parts = l_c.split(':')
                if len(parts) > 1: info["seller"] = parts[-1].strip()
            elif re.search(r'^(ƒê∆°n v·ªã mua|Ng∆∞·ªùi mua|Kh√°ch h√†ng|B√™n B)', l_c, re.IGNORECASE): 
                parts = l_c.split(':')
                if len(parts) > 1: info["buyer"] = parts[-1].strip()

    else: # === UNC (N√ÇNG C·∫§P LOGIC) ===
        candidates_total = []
        BLOCK_KEYWORDS = ['s·ªë d∆∞', 'balance', 'ph√≠', 'fee', 'charge', 'vat', 'tax', 'ƒëi·ªán tho·∫°i', 'tel', 'fax', 'mst', 'm√£ s·ªë thu·∫ø', 'l·ªá ph√≠', 'so du', 'le phi']
        CONFIRM_KEYWORDS = ['s·ªë ti·ªÅn', 'amount', 'thanh to√°n', 'chuy·ªÉn kho·∫£n', 'transaction', 'gi√° tr·ªã', 'total', 'c·ªông', 'money', 'so tien', 'chuyen khoan', 'gia tri']
        
        # --- LOAD T·ª™ KH√ìA ƒê√É H·ªåC T·ª™ DB ---
        learned_kws = run_query("SELECT keyword FROM ocr_learning")
        if learned_kws:
            CONFIRM_KEYWORDS.extend([r['keyword'] for r in learned_kws]) # type: ignore
            
        CURRENCY_KEYWORDS = ['vnd', 'ƒë', 'vnƒë', 'usd']
        prev_line_score_boost = 0
        fallback_numbers = []

        for i, line in enumerate(lines):
            line_l = line.lower()
            
            is_label_line = False
            if any(kw in line_l for kw in CONFIRM_KEYWORDS):
                nums_in_line = extract_money_smart(line)
                if not nums_in_line: 
                    prev_line_score_boost = 15 
                    is_label_line = True
            
            if is_label_line: continue

            nums = extract_money_smart(line)
            if not nums: 
                prev_line_score_boost = 0
                continue
            
            max_val = max(nums)
            if max_val < 1000: 
                prev_line_score_boost = 0
                continue 
            
            is_blocked = any(bad in line_l for bad in BLOCK_KEYWORDS)
            if not is_blocked:
                fallback_numbers.append(max_val)
            
            score = 0
            score += prev_line_score_boost
            prev_line_score_boost = 0 
            
            if any(kw in line_l for kw in CONFIRM_KEYWORDS): score += 10
            if any(kw in line_l for kw in CURRENCY_KEYWORDS): score += 5
            if is_blocked and not any(good in line_l for good in CONFIRM_KEYWORDS):
                score -= 20
            if 't√†i kho·∫£n' in line_l or 'account' in line_l or 'stk' in line_l: score -= 5

            val_str = "{:,.0f}".format(max_val) # 10,000,000
            val_str_dot = val_str.replace(",", ".") # 10.000.000
            
            if val_str in line or val_str_dot in line:
                score += 3
            elif max_val > 100000000: 
                score -= 3

            if score > -10: candidates_total.append((max_val, score))
        
        if candidates_total:
            candidates_total.sort(key=lambda x: (x[1], x[0]), reverse=True)
            info["total"] = candidates_total[0][0]
        elif fallback_numbers:
            info["total"] = max(fallback_numbers)
            
        info["pre_tax"] = info["total"]
        
        for line in lines:
            if re.search(r'(?:n·ªôi dung|di·ªÖn gi·∫£i|l√Ω do|remarks|narrative|description|message)', line, re.IGNORECASE):
                parts = re.split(r'[:\.\-]', line, 1)
                if len(parts) > 1: info["content"] = parts[1].strip()
                else: info["content"] = line.strip()
                break

        for i, line in enumerate(lines):
            line_clean = line.strip()
            if re.search(r'(?:ng∆∞·ªùi h∆∞·ªüng|ƒë∆°n v·ªã th·ª• h∆∞·ªüng|t√†i kho·∫£n nh·∫≠n|t√™n ng∆∞·ªùi nh·∫≠n|b√™n nh·∫≠n|beneficiary)', line_clean, re.IGNORECASE):
                parts = line_clean.split(':')
                if len(parts) > 1 and len(parts[-1].strip()) > 3:
                    info["seller"] = parts[-1].strip()
                    break
                elif i + 1 < len(lines):
                    info["seller"] = lines[i+1].strip()
                    break

    info["raw_text"] = text_content
    return info, msg

def create_handover_docx(tour_info, guests, hotels, restaurants, sightseeings, checklist_str):
    if not HAS_DOCX: return None
    
    # L·∫•y th√¥ng tin c√¥ng ty
    comp_data = get_company_data()

    doc = Document() # type: ignore
    
    # Styles
    style = doc.styles['Normal']
    font = style.font # type: ignore
    font.name = 'Times New Roman'
    font.size = Pt(11) # type: ignore
    
    # --- HEADER (COMPANY INFO) ---
    header_table = doc.add_table(rows=1, cols=1)
    header_table.autofit = False
    header_table.allow_autofit = False # type: ignore
    cell = header_table.cell(0, 0)
    cell.width = Inches(6.5) # type: ignore
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER # type: ignore
    
    run_comp = p.add_run(f"{comp_data['name'].upper()}\n")
    run_comp.bold = True
    run_comp.font.size = Pt(14) # type: ignore
    run_comp.font.color.rgb = None # Black
    
    p.add_run(f"ƒê·ªãa ch·ªâ: {comp_data['address']}\n")
    p.add_run(f"M√£ S·ªë Thu·∫ø: {comp_data['phone']}")
    
    doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER # type: ignore
    doc.add_paragraph()
    
    # --- TITLE ---
    p_title = doc.add_heading('H·ªí S∆† B√ÄN GIAO ƒêO√ÄN', 0)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER # type: ignore
    p_title.style.font.name = 'Times New Roman' # type: ignore
    p_title.style.font.size = Pt(16) # type: ignore
    p_title.style.font.bold = True # type: ignore
    p_title.style.font.color.rgb = None # Black # type: ignore
    
    doc.add_paragraph()
    
    # --- I. TH√îNG TIN CHUNG ---
    doc.add_heading('I. TH√îNG TIN CHUNG', level=1)
    
    table_info = doc.add_table(rows=0, cols=4)
    table_info.style = 'Table Grid'
    table_info.autofit = True
    
    def add_kv(k1, v1, k2, v2):
        row = table_info.add_row()
        c = row.cells
        c[0].text = k1
        c[0].paragraphs[0].runs[0].bold = True
        c[1].text = str(v1)
        c[2].text = k2
        c[2].paragraphs[0].runs[0].bold = True
        c[3].text = str(v2)

    cust_info = f"{tour_info.get('customer_name','')} - {tour_info.get('customer_phone','')}"
    
    add_kv("T√™n ƒëo√†n:", tour_info['tour_name'], "M√£ ƒëo√†n:", tour_info['tour_code'])
    add_kv("Ng√†y ƒëi:", tour_info['start_date'], "Ng√†y v·ªÅ:", tour_info['end_date'])
    add_kv("S·ªë l∆∞·ª£ng kh√°ch:", str(tour_info['guest_count']), "Sales:", tour_info['sale_name'])
    
    # Row for Customer
    r = table_info.add_row()
    r.cells[0].text = "Kh√°ch h√†ng:"
    r.cells[0].paragraphs[0].runs[0].bold = True
    r.cells[1].merge(r.cells[3])
    r.cells[1].text = cust_info
    
    doc.add_paragraph()
    
    # --- II. DANH S√ÅCH ƒêO√ÄN ---
    doc.add_heading('II. DANH S√ÅCH ƒêO√ÄN', level=1)
    if guests:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['STT', 'H·ªç v√† t√™n', 'Ng√†y sinh', 'S·ªë CCCD', 'Ph√¢n lo·∫°i']
        for i, h in enumerate(headers): 
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # type: ignore
            
        for i, g in enumerate(guests):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i + 1)
            row_cells[1].text = str(g['name'])
            row_cells[2].text = str(g['dob'])
            row_cells[3].text = str(g['cccd'])
            row_cells[4].text = str(g['type'])
    else: doc.add_paragraph("(Ch∆∞a c√≥ danh s√°ch ƒëo√†n)")
    doc.add_paragraph()
        
    # --- III. KH√ÅCH S·∫†N ---
    doc.add_heading('III. TH√îNG TIN L∆ØU TR√ö', level=1)
    if hotels:
        table_h = doc.add_table(rows=1, cols=7)
        table_h.style = 'Table Grid'
        hdr = table_h.rows[0].cells
        for i, h in enumerate(['T√™n Kh√°ch s·∫°n', 'ƒê·ªãa ch·ªâ & Li√™n h·ªá', 'T·ªïng ph√≤ng', 'Lo·∫°i ph√≤ng', 'T·ªïng ti·ªÅn', 'ƒê√£ c·ªçc', 'C√≤n l·∫°i']):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            
        for h in hotels:
            total = float(h.get('total_amount', 0) or 0)
            dep = float(h.get('deposit', 0) or 0)
            rem = total - dep
            row = table_h.add_row().cells
            row[0].text = f"üè® {h['hotel_name']}"
            row[1].text = f"{h['address']}\nSƒêT: {h['phone']}"
            row[2].text = str(h['total_rooms'])
            row[3].text = str(h['room_type'])
            row[4].text = "{:,.0f}".format(total)
            row[5].text = "{:,.0f}".format(dep)
            row[6].text = "{:,.0f}".format(rem)
    else: doc.add_paragraph("(Ch∆∞a c√≥ th√¥ng tin kh√°ch s·∫°n)")
    doc.add_paragraph()

    # --- IV. NH√Ä H√ÄNG ---
    doc.add_heading('IV. ·∫®M TH·ª∞C & TH·ª∞C ƒê∆†N', level=1)
    if restaurants:
        table_r = doc.add_table(rows=1, cols=7)
        table_r.style = 'Table Grid'
        hdr = table_r.rows[0].cells
        for i, h in enumerate(['B·ªØa ƒÉn', 'Nh√† h√†ng', 'Li√™n h·ªá', 'Th·ª±c ƒë∆°n', 'T·ªïng ti·ªÅn', 'ƒê√£ c·ªçc', 'C√≤n l·∫°i']):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            
        for r in restaurants:
            total = float(r.get('total_amount', 0) or 0)
            dep = float(r.get('deposit', 0) or 0)
            rem = total - dep
            row = table_r.add_row().cells
            row[0].text = f"üçΩÔ∏è {r['meal_name']}"
            row[1].text = str(r['restaurant_name'])
            row[2].text = f"{r['address']}\nSƒêT: {r['phone']}"
            row[3].text = str(r['menu'])
            row[4].text = "{:,.0f}".format(total)
            row[5].text = "{:,.0f}".format(dep)
            row[6].text = "{:,.0f}".format(rem)
    else: doc.add_paragraph("(Ch∆∞a c√≥ th√¥ng tin nh√† h√†ng)")
    doc.add_paragraph()

    # --- V. ƒêI·ªÇM THAM QUAN ---
    doc.add_heading('V. ƒêI·ªÇM THAM QUAN', level=1)
    if sightseeings:
        table_s = doc.add_table(rows=1, cols=4)
        table_s.style = 'Table Grid'
        hdr = table_s.rows[0].cells
        for i, h in enumerate(['T√™n ƒë·ªãa ƒëi·ªÉm', 'ƒê·ªãa ch·ªâ', 'S·ªë l∆∞·ª£ng', 'L∆∞u √Ω']):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            
        for s in sightseeings:
            row = table_s.add_row().cells
            row[0].text = f"üìç {s['name']}"
            row[1].text = str(s['address'])
            row[2].text = str(s['quantity'])
            row[3].text = str(s['note'])
    else: doc.add_paragraph("(Ch∆∞a c√≥ th√¥ng tin ƒëi·ªÉm tham quan)")
    doc.add_paragraph()
        
    # --- VI. CHECKLIST ---
    doc.add_heading('VI. CHECKLIST B√ÄN GIAO', level=1)
    checked_items = checklist_str.split(',') if checklist_str else []
    all_items = ["Ch∆∞∆°ng tr√¨nh ƒë√≥ng m·ªôc", "Danh s√°ch ƒë√≥ng m·ªôc", "B·∫£o hi·ªÉm du l·ªãch", "Th·ª±c ƒë∆°n ƒë√≥ng m·ªôc", "V√© m√°y bay", "X√°c nh·∫≠n khu du l·ªãch/nh√† h√†ng (N·∫øu c√≥)", "H·ª£p ƒë·ªìng h∆∞·ªõng d·∫´n"]
    
    table_c = doc.add_table(rows=0, cols=2)
    for item in all_items:
        mark = "‚òë" if item in checked_items else "‚òê"
        row = table_c.add_row()
        row.cells[0].text = mark
        row.cells[0].width = Pt(20) # type: ignore
        row.cells[1].text = item
        
    # Footer
    doc.add_paragraph("\n")
    p_foot = doc.add_paragraph(f"Ng√†y xu·∫•t h·ªì s∆°: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT # type: ignore
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- H√ÄM ƒê·ªåC S·ªê TI·ªÄN B·∫∞NG CH·ªÆ (VIETNAMESE) ---
def read_money_vietnamese(amount):
    if amount == 0: return "Kh√¥ng ƒë·ªìng"
    
    digits = ["kh√¥ng", "m·ªôt", "hai", "ba", "b·ªën", "nƒÉm", "s√°u", "b·∫£y", "t√°m", "ch√≠n"]
    units = ["", "ngh√¨n", "tri·ªáu", "t·ª∑"]
    
    def read_group(n):
        res = ""
        h = n // 100
        t = (n % 100) // 10
        u = n % 10
        
        if h > 0:
            res += digits[h] + " trƒÉm "
        elif n > 0: # C√≥ h√†ng ch·ª•c ho·∫∑c ƒë∆°n v·ªã nh∆∞ng h√†ng trƒÉm = 0 (x·ª≠ l√Ω ·ªü loop ch√≠nh t·ªët h∆°n, ƒë√¢y l√† logic ƒë∆°n gi·∫£n)
            pass 
            
        if t > 1:
            res += digits[t] + " m∆∞∆°i "
            if u == 1: res += "m·ªët "
            elif u == 5: res += "lƒÉm "
            elif u > 0: res += digits[u] + " "
        elif t == 1:
            res += "m∆∞·ªùi "
            if u == 1: res += "m·ªôt "
            elif u == 5: res += "lƒÉm "
            elif u > 0: res += digits[u] + " "
        else: # t = 0
            if h > 0 and u > 0: res += "l·∫ª "
            if u > 0: res += digits[u] + " "
        return res

    s_num = "{:.0f}".format(amount)
    groups = []
    while len(s_num) > 0:
        groups.append(int(s_num[-3:]))
        s_num = s_num[:-3]
    
    ret = ""
    for i, g in enumerate(groups):
        if g > 0:
            s_g = read_group(g)
            # X·ª≠ l√Ω s·ªë 0 trƒÉm
            if i < len(groups) - 1 and g < 100 and g > 0: 
                s_g = "kh√¥ng trƒÉm " + s_g
                
            ret = s_g + units[i] + " " + ret
            
    ret = ret.strip()
    # Capitalize first letter
    if ret:
        ret = ret[0].upper() + ret[1:]
    
    return ret + " ƒë·ªìng"

def create_voucher_pdf(voucher_data):
    """T·∫°o file PDF phi·∫øu thu/chi ƒë·∫πp, c√≥ logo v√† m√†u s·∫Øc"""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # C·∫•u h√¨nh Font
    font_name = 'Helvetica' # Fallback
    font_name_bold = 'Helvetica-Bold'
    try:
        font_path = r"C:\Windows\Fonts\times.ttf"
        font_path_bd = r"C:\Windows\Fonts\timesbd.ttf"
        if os.path.exists(font_path) and os.path.exists(font_path_bd):
            pdfmetrics.registerFont(TTFont('TimesNewRoman', font_path))
            pdfmetrics.registerFont(TTFont('TimesNewRoman-Bold', font_path_bd))
            font_name = 'TimesNewRoman'
            font_name_bold = 'TimesNewRoman-Bold'
        else:
            font_path = r"C:\Windows\Fonts\arial.ttf"
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('Arial', font_path))
                font_name = 'Arial'
                font_name_bold = 'Arial'
    except: pass

    comp = get_company_data()
    
    # M√†u s·∫Øc ch·ªß ƒë·∫°o
    primary_color = "#2E7D32" if voucher_data['type'] == 'THU' else "#C62828" # Xanh cho Thu, ƒê·ªè cho Chi
    text_color = "#212121"
    
    # --- HEADER ---
    # Logo
    logo_height = 60
    header_y = height - 50
    header_x_text = 50
    
    if comp['logo_b64_str']:
        try:
            logo_data = base64.b64decode(comp['logo_b64_str'])
            image_stream = io.BytesIO(logo_data)
            img_reader = ImageReader(image_stream)
            # T√≠nh t·ª∑ l·ªá ·∫£nh
            iw, ih = img_reader.getSize()
            aspect = iw / float(ih)
            draw_w = logo_height * aspect
            
            c.drawImage(img_reader, 50, header_y - logo_height, width=draw_w, height=logo_height, mask='auto')
            header_x_text = 50 + draw_w + 20
        except: pass

    # --- LOGO CH√åM (WATERMARK) ---
    if comp['logo_b64_str']:
        try:
            c.saveState()
            logo_data = base64.b64decode(comp['logo_b64_str'])
            image_stream = io.BytesIO(logo_data)
            img_reader = ImageReader(image_stream)
            iw, ih = img_reader.getSize()
            aspect = iw / float(ih)
            wm_width = 300
            wm_height = wm_width / aspect
            c.setFillAlpha(0.1) # ƒê·ªô m·ªù 10%
            # V·∫Ω ch√≠nh gi·ªØa trang
            c.drawImage(img_reader, (width - wm_width)/2, (height - wm_height)/2, width=wm_width, height=wm_height, mask='auto')
            c.restoreState()
        except: pass

    # Th√¥ng tin c√¥ng ty
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_name_bold, 16)
    c.drawString(header_x_text, header_y - 15, comp['name'].upper())
    
    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 10)
    c.drawString(header_x_text, header_y - 35, f"ƒêC: {comp['address']}")
    c.drawString(header_x_text, header_y - 50, f"MST: {comp['phone']}")
    
    # ƒê∆∞·ªùng k·∫ª trang tr√≠
    c.setStrokeColor(HexColor(primary_color))
    c.setLineWidth(2)
    c.line(50, header_y - 70, width - 50, header_y - 70)
    
    # --- TI√äU ƒê·ªÄ ---
    title = "PHI·∫æU THU TI·ªÄN" if voucher_data['type'] == 'THU' else "PHI·∫æU CHI TI·ªÄN"
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_name_bold, 24)
    c.drawCentredString(width/2, height - 150, title)
    
    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 11)
    c.drawCentredString(width/2, height - 170, f"Ng√†y: {voucher_data['date']}")
    
    # --- N·ªòI DUNG ---
    # L·∫•y t√™n kh√°ch h√†ng n·∫øu c√≥
    person_name = ""
    ref_code = voucher_data.get('ref_code', '')
    if ref_code:
        try:
            # Th·ª≠ t√¨m trong Tours
            t = run_query("SELECT customer_name FROM tours WHERE tour_code=?", (ref_code,), fetch_one=True)
            if t and t['customer_name']: person_name = t['customer_name']
            else:
                # Th·ª≠ t√¨m trong Bookings
                b = run_query("SELECT customer_info FROM service_bookings WHERE code=?", (ref_code,), fetch_one=True)
                if b and b['customer_info']: person_name = b['customer_info'].split(' - ')[0]
        except: pass

    # --- T√çNH TO√ÅN T√ÄI CH√çNH (M·ªöI) ---
    contract_val = 0.0
    total_paid = 0.0
    remaining = 0.0
    
    if ref_code:
        # 1. L·∫•y gi√° tr·ªã h·ª£p ƒë·ªìng
        # Th·ª≠ t√¨m Tour
        t_info = run_query("SELECT * FROM tours WHERE tour_code=?", (ref_code,), fetch_one=True)
        if t_info:
            t_dict = dict(t_info)
            final_price = float(t_dict.get('final_tour_price', 0) or 0)
            child_price = float(t_dict.get('child_price', 0) or 0)
            final_qty = float(t_dict.get('final_qty', 0) or 0)
            child_qty = float(t_dict.get('child_qty', 0) or 0)
            if final_qty == 0: final_qty = float(t_dict.get('guest_count', 1))
            contract_val = (final_price * final_qty) + (child_price * child_qty)
        else:
            # Th·ª≠ t√¨m Booking
            b_info = run_query("SELECT selling_price FROM service_bookings WHERE code=?", (ref_code,), fetch_one=True)
            if b_info:
                contract_val = float(b_info['selling_price'] or 0)
        
        # 2. L·∫•y t·ªïng ƒë√£ thu (Bao g·ªìm c·∫£ phi·∫øu v·ª´a t·∫°o n·∫øu ƒë√£ l∆∞u DB)
        txns = run_query("SELECT type, amount FROM transaction_history WHERE ref_code=?", (ref_code,))
        if txns:
            paid_sum = sum(r['amount'] for r in txns if r['type'] == 'THU')
            refund_sum = sum(r['amount'] for r in txns if r['type'] == 'CHI')
            total_paid = paid_sum - refund_sum
            
        remaining = contract_val - total_paid

    # --- V·∫º PDF ---
    y = height - 220
    x_label = 70
    x_val = 200
    line_height = 30
    
    # V·∫Ω khung n·ªÅn m·ªù
    bg_color = "#E8F5E9" if voucher_data['type'] == 'THU' else "#FFEBEE"
    c.setFillColor(HexColor(bg_color))
    # TƒÉng chi·ªÅu cao khung ƒë·ªÉ ch·ª©a th√™m th√¥ng tin (210 -> 330)
    c.roundRect(50, y - 310, width - 100, 330, 10, fill=1, stroke=0)
    
    c.setFillColor(HexColor(text_color))
    
    def draw_line_content(label, value, y_pos, is_money=False):
        c.setFont(font_name, 12)
        c.drawString(x_label, y_pos, label)
        
        if is_money:
            c.setFont(font_name_bold, 14)
            c.setFillColor(HexColor(primary_color))
            c.drawString(x_val, y_pos, value)
            c.setFillColor(HexColor(text_color)) # Reset
        else:
            c.setFont(font_name, 12)
            if value:
                c.drawString(x_val, y_pos, value)
            else:
                # V·∫Ω d√≤ng ch·∫•m
                c.setStrokeColor(HexColor("#BDBDBD"))
                c.setLineWidth(1)
                c.setDash(1, 3)
                c.line(x_val, y_pos - 3, width - 70, y_pos - 3)
                c.setDash([])

    label_person = "Ng∆∞·ªùi n·ªôp ti·ªÅn:" if voucher_data['type'] == 'THU' else "Ng∆∞·ªùi nh·∫≠n ti·ªÅn:"
    draw_line_content(label_person, person_name, y); y -= line_height
    draw_line_content("ƒê·ªãa ch·ªâ/SƒêT:", "", y); y -= line_height
    draw_line_content("L√Ω do:", f"{voucher_data['note']} (M√£: {voucher_data['ref_code']})", y); y -= line_height
    draw_line_content("S·ªë ti·ªÅn:", f"{format_vnd(voucher_data['amount'])} VND", y, is_money=True); y -= line_height
    draw_line_content("B·∫±ng ch·ªØ:", read_money_vietnamese(voucher_data['amount']), y); y -= line_height
    
    # --- C√ÅC D√íNG M·ªöI ---
    draw_line_content("T·ªïng gi√° tr·ªã Hƒê:", f"{format_vnd(contract_val)} VND", y); y -= line_height
    draw_line_content("ƒê√£ thanh to√°n:", f"{format_vnd(total_paid)} VND", y); y -= line_height
    draw_line_content("C√≤n l·∫°i:", f"{format_vnd(remaining)} VND", y); y -= line_height
    draw_line_content("Ng∆∞·ªùi xu·∫•t phi·∫øu:", voucher_data.get('issuer', ''), y); y -= line_height

    draw_line_content("K√®m theo:", "", y); y -= line_height
    
    # --- CH·ªÆ K√ù ---
    y_sig = y - 40
    sigs = ["Gi√°m ƒë·ªëc", "K·∫ø to√°n tr∆∞·ªüng", "Ng∆∞·ªùi l·∫≠p phi·∫øu", "Ng∆∞·ªùi n·ªôp/nh·∫≠n"]
    x_positions = [50, 180, 310, 440]
    for i, sig in enumerate(sigs):
        c.setFont(font_name, 11)
        c.setFillColor(HexColor(text_color))
        c.drawCentredString(x_positions[i] + 40, y_sig, sig)
        c.setFont(font_name, 9)
        c.setFillColor(HexColor("#757575"))
        c.drawCentredString(x_positions[i] + 40, y_sig - 15, "(K√Ω, h·ªç t√™n)")
        
    c.save()
    buffer.seek(0)
    return buffer

def create_booking_cfm_pdf(booking_info, company_info, lang='en'):
    """T·∫°o file PDF Booking Confirmation (CFM)"""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    # --- HELPER: WATERMARK & NEW PAGE ---
    def draw_watermark():
        if company_info['logo_b64_str']:
            try:
                c.saveState()
                logo_data = base64.b64decode(company_info['logo_b64_str'])
                image_stream = io.BytesIO(logo_data)
                img_reader = ImageReader(image_stream)
                iw, ih = img_reader.getSize()
                aspect = iw / float(ih)
                wm_width = 400
                wm_height = wm_width / aspect
                c.setFillAlpha(0.08) # ƒê·ªô m·ªù 8%
                c.drawImage(img_reader, (width - wm_width)/2, (height - wm_height)/2, width=wm_width, height=wm_height, mask='auto')
                c.restoreState()
            except: pass

    draw_watermark()

    # --- C·∫§U H√åNH FONT (T∆∞∆°ng t·ª± create_voucher_pdf) ---
    font_name = 'Helvetica'
    font_bold = 'Helvetica-Bold'
    try:
        font_path = r"C:\Windows\Fonts\times.ttf"
        font_path_bd = r"C:\Windows\Fonts\timesbd.ttf"
        if os.path.exists(font_path) and os.path.exists(font_path_bd):
            pdfmetrics.registerFont(TTFont('TimesNewRoman', font_path))
            pdfmetrics.registerFont(TTFont('TimesNewRoman-Bold', font_path_bd))
            font_name = 'TimesNewRoman'
            font_bold = 'TimesNewRoman-Bold'
        else:
            # Fallback cho Arial n·∫øu kh√¥ng c√≥ Times
            font_path_arial = r"C:\Windows\Fonts\arial.ttf"
            if os.path.exists(font_path_arial):
                pdfmetrics.registerFont(TTFont('Arial', font_path_arial))
                font_name = 'Arial'
                font_bold = 'Arial' # Arial th∆∞·ªùng kh√¥ng t√°ch file bold r√µ r√†ng trong code ƒë∆°n gi·∫£n
    except: pass

    # --- M√ÄU S·∫ÆC ---
    primary_color = "#1B5E20" # Xanh ƒë·∫≠m th∆∞∆°ng hi·ªáu
    text_color = "#212121"
    line_color = "#BDBDBD"

    # --- T·ª™ ƒêI·ªÇN NG√îN NG·ªÆ ---
    labels = {
        'en': {
            'add': 'Add:', 'tax': 'Tax Code:',
            'title': 'BOOKING CONFIRMATION',
            'greeting': f"A warm greeting from {company_info['name']}!",
            'gen_info': 'I. GENERAL INFORMATION',
            'attn': 'Attention to:', 'bk_code': 'Booking Code:', 'svc_date': 'Service Date:',
            'date_created': 'Date Created:', 'checkout': 'Check-out:', 'created_by': 'Created By:',
            'status': 'Status:', 'hotel_code': 'Hotel Code:',
            'confirmed': 'Confirmed', 'cancelled': 'Cancelled',
            'svc_details': 'II. SERVICE DETAILS',
            'svc_details_cont': 'II. SERVICE DETAILS (Cont.)',
            'tbl_name': 'SERVICE NAME', 'tbl_det': 'DETAILS / ROOM INFO', 'tbl_note': 'NOTE',
            'guest_list': 'III. GUEST LIST',
            'guest_list_cont': 'III. GUEST LIST (Cont.)',
            'included': 'INCLUDED SERVICES',
            'inc_1': '- Tax and Service charges.',
            'inc_2': '- 24/7 Support from our team.',
            'confirmed_by': 'CONFIRMED BY',
            'signed': '[SIGNED]',
            'page': 'Page'
        },
        'vi': {
            'add': 'ƒêC:', 'tax': 'MST:',
            'title': 'X√ÅC NH·∫¨N ƒê·∫∂T D·ªäCH V·ª§',
            'greeting': f"L·ªùi ch√†o tr√¢n tr·ªçng t·ª´ {company_info['name']}!",
            'gen_info': 'I. TH√îNG TIN CHUNG',
            'attn': 'K√≠nh g·ª≠i:', 'bk_code': 'M√£ ƒë·∫∑t ch·ªó:', 'svc_date': 'Ng√†y s·ª≠ d·ª•ng:',
            'date_created': 'Ng√†y t·∫°o:', 'checkout': 'Ng√†y tr·∫£ ph√≤ng:', 'created_by': 'Ng∆∞·ªùi t·∫°o:',
            'status': 'Tr·∫°ng th√°i:', 'hotel_code': 'M√£ kh√°ch s·∫°n:',
            'confirmed': 'ƒê√£ x√°c nh·∫≠n', 'cancelled': 'ƒê√£ h·ªßy',
            'svc_details': 'II. CHI TI·∫æT D·ªäCH V·ª§',
            'svc_details_cont': 'II. CHI TI·∫æT D·ªäCH V·ª§ (Ti·∫øp)',
            'tbl_name': 'T√äN D·ªäCH V·ª§', 'tbl_det': 'CHI TI·∫æT / TH√îNG TIN PH√íNG', 'tbl_note': 'GHI CH√ö',
            'guest_list': 'III. DANH S√ÅCH KH√ÅCH',
            'guest_list_cont': 'III. DANH S√ÅCH KH√ÅCH (Ti·∫øp)',
            'included': 'D·ªäCH V·ª§ BAO G·ªíM',
            'inc_1': '- Thu·∫ø v√† ph√≠ ph·ª•c v·ª•.',
            'inc_2': '- H·ªó tr·ª£ 24/7 t·ª´ ƒë·ªôi ng≈© c·ªßa ch√∫ng t√¥i.',
            'confirmed_by': 'X√ÅC NH·∫¨N B·ªûI',
            'signed': '[ƒê√É K√ù]',
            'page': 'Trang'
        }
    }
    txt = labels[lang]

    # --- HEADER ---
    y = height - 50
    # Logo
    if company_info['logo_b64_str']:
        try:
            logo_data = base64.b64decode(company_info['logo_b64_str'])
            image_stream = io.BytesIO(logo_data)
            img_reader = ImageReader(image_stream)
            iw, ih = img_reader.getSize()
            aspect = iw / float(ih)
            logo_h = 85 # Logo to h∆°n
            logo_w = logo_h * aspect
            c.drawImage(img_reader, 40, y - logo_h, width=logo_w, height=logo_h, mask='auto')
        except: pass

    # X·ª≠ l√Ω ƒë·ªãa ch·ªâ (D·ªãch s∆° b·ªô n·∫øu l√† ti·∫øng Anh)
    # X·ª≠ l√Ω ƒë·ªãa ch·ªâ v√† t√™n c√¥ng ty (D·ªãch s∆° b·ªô n·∫øu l√† ti·∫øng Anh)
    comp_addr = company_info['address']
    comp_name = company_info['name']
    
    if lang == 'en':
        # [UPDATED] Hardcoded English details
        comp_name = "BALI TOURIST TRAVEL COMPANY LIMITED"
        comp_addr = "No. 46 Nguyen Oanh, Hanh Thong Ward, Ho Chi Minh City, Vietnam"
        txt['greeting'] = f"A warm greeting from {comp_name}!"

    # Th√¥ng tin c√¥ng ty (CƒÉn ph·∫£i)
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_bold, 18)
    c.drawRightString(width - 40, y - 25, comp_name.upper())
    
    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 10)
    c.drawRightString(width - 40, y - 45, f"{txt['add']} {comp_addr}")
    c.drawRightString(width - 40, y - 60, f"{txt['tax']} {company_info['phone']}")
    
    y -= 100
    c.setStrokeColor(HexColor(primary_color))
    c.setLineWidth(2)
    c.line(40, y, width - 40, y)
    
    # --- TITLE ---
    y -= 40
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_bold, 20)
    c.drawCentredString(width/2, y, txt['title'])
    
    y -= 25
    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 11)
    c.drawCentredString(width/2, y, txt['greeting'])
    
    # --- X·ª¨ L√ù D·ªÆ LI·ªÜU BOOKING ---
    # Parse Customer Info
    cust_raw = booking_info.get('customer_info', '')
    cust_name = cust_raw.split(' - ')[0] if ' - ' in cust_raw else cust_raw
    
    # Parse Dates from Details
    details = booking_info.get('details', '')
    dates = re.findall(r'\d{1,2}[/-]\d{1,2}[/-]\d{4}', details)
    check_in = dates[0] if len(dates) > 0 else booking_info.get('created_at', '')
    check_out = dates[1] if len(dates) > 1 else "N/A"
    
    # --- PH·∫¶N 1: TH√îNG TIN CHUNG ---
    y -= 40
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_bold, 12)
    c.drawString(40, y, txt['gen_info'])
    y -= 20
    
    # T√≠nh to√°n chi·ªÅu cao khung th√¥ng tin
    info_lines = 4
    if booking_info.get('type') == 'HOTEL' and booking_info.get('hotel_code'):
        info_lines = 5
    box_height = info_lines * 25 + 10

    # V·∫Ω khung th√¥ng tin
    c.setStrokeColor(HexColor(line_color))
    c.setLineWidth(1)
    c.rect(40, y - box_height, width - 80, box_height, fill=0)
    
    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 11)
    
    # D√≤ng 1: Attention to (Ri√™ng 1 d√≤ng)
    row_y = y - 25
    c.drawString(50, row_y, txt['attn'])
    c.setFont(font_bold, 11); c.drawString(130, row_y, cust_name); c.setFont(font_name, 11)
    
    # D√≤ng 2
    row_y -= 25
    c.drawString(50, row_y, txt['bk_code'])
    c.setFont(font_bold, 11); c.drawString(130, row_y, booking_info['code']); c.setFont(font_name, 11)
    
    c.drawString(300, row_y, txt['svc_date'])
    c.drawString(400, row_y, check_in)
    
    # D√≤ng 3
    row_y -= 25
    c.drawString(50, row_y, txt['date_created'])
    c.drawString(130, row_y, booking_info.get('created_at', ''))
    
    if check_out != "N/A":
        c.drawString(300, row_y, txt['checkout'])
        c.drawString(400, row_y, check_out)
        
    # D√≤ng 4
    row_y -= 25
    c.drawString(50, row_y, txt['created_by'])
    c.drawString(130, row_y, booking_info.get('sale_name', ''))

    c.drawString(300, row_y, txt['status'])
    status_txt = txt['confirmed'] if booking_info.get('status') != 'deleted' else txt['cancelled']
    c.setFillColor(HexColor("#2E7D32" if status_txt == "Confirmed" else "#C62828"))
    c.setFont(font_bold, 11)
    c.drawString(400, row_y, status_txt)
    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 11)

    # D√≤ng 5 (N·∫øu c√≥ Hotel Code)
    if info_lines == 5:
        row_y -= 25
        c.drawString(50, row_y, txt['hotel_code'])
        c.setFont(font_bold, 11)
        c.drawString(130, row_y, booking_info.get('hotel_code', ''))
        c.setFont(font_name, 11)

    # --- PH·∫¶N 2: CHI TI·∫æT D·ªäCH V·ª§ ---
    y -= (box_height + 40)
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_bold, 12)
    c.drawString(40, y, txt['svc_details'])
    y -= 25
    
    # Header B·∫£ng
    c.setFillColor(HexColor("#E8F5E9"))
    c.rect(40, y - 5, width - 80, 20, fill=1, stroke=0) # Header BG
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_bold, 10)
    c.drawString(50, y, txt['tbl_name'])
    c.drawString(250, y, txt['tbl_det'])
    c.drawString(450, y, txt['tbl_note'])
    
    y -= 20
    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 10)
    
    # N·ªôi dung b·∫£ng (X·ª≠ l√Ω Combo t√°ch d√≤ng)
    items = []
    
    # [HELPER] Translate content if English
    def translate_content(text):
        if lang != 'en' or not text: return text
        replacements = {
            "Ng√†y:": "Date:", "SL:": "Qty:", "L∆∞u tr√∫:": "Stay:",
            "Xe ": "Car ", "V√©:": "Ticket:", "M√°y bay": "Flight", 
            "T√†u h·ªèa": "Train", "Du thuy·ªÅn": "Cruise", "Cabin:": "Cabin:",
            "ph√≤ng": "rooms", "ƒë√™m": "nights", "kh√°ch": "pax",
            "[KS]": "[Hotel]", "[XE]": "[Car]", "[BAY]": "[Flight]", 
            "[TAU]": "[Train]", "[THUYEN]": "[Cruise]", "[CB]": "[Combo]"
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
        return text

    if booking_info.get('type') == 'COMBO':
        # T√°ch c√°c item trong combo (ngƒÉn c√°ch b·ªüi | ho·∫∑c d√≤ng m·ªõi)
        raw_items = re.split(r'[|\n]', details)
        for item in raw_items:
            if item.strip(): 
                display_item = translate_content(item.strip())
                items.append((display_item, ""))
    else:
        # Translate basic keywords for non-hotel types
        details_display = translate_content(details)
        name_display = translate_content(booking_info['name'])
        items.append((name_display, details_display))
        
    # [NEW] X·ª≠ l√Ω hi·ªÉn th·ªã chi ti·∫øt cho Booking Kh√°ch s·∫°n (Hotel Code, Room Type, Guest List)
    if booking_info.get('type') == 'HOTEL':
        # Override items list to show detailed info
        r_type = booking_info.get('room_type', '')
        g_list = booking_info.get('guest_list', '')
        
        # Format l·∫°i ph·∫ßn Details
        # details c≈©: "L∆∞u tr√∫: 01/01 - 03/01 (2 ƒë√™m, 1 ph√≤ng)"
        # Regex extract numbers
        new_details = details
        if lang == 'en':
            match = re.search(r'\((\d+)\s+ƒë√™m,\s+(\d+)\s+ph√≤ng\)', details)
            if match:
                nights = match.group(1)
                rooms = match.group(2)
                new_details = f"{nights} nights, {rooms} rooms"
            else:
                new_details = translate_content(details)
                
            if r_type: new_details += f"\nRoom Type: {r_type}"
        else:
            if r_type: new_details += f"\nLo·∫°i ph√≤ng: {r_type}"
        
        # Format ph·∫ßn Note ho·∫∑c th√™m v√†o Details
        note_part = "" # Guest list moved to separate section
        
        name_display = translate_content(booking_info['name'])
        items = [(name_display, new_details, note_part)]

    for item in items:
        # T·ª± ƒë·ªông xu·ªëng d√≤ng n·∫øu text qu√° d√†i (Logic ƒë∆°n gi·∫£n)
        if len(item) == 3:
            name, det, note = item
        else:
            name, det = item # type: ignore
            note = ""
            
        # V·∫Ω Name
        c.drawString(50, y, name[:45] + "..." if len(name)>45 else name)
        
        # V·∫Ω Details (Multi-line support basic)
        det_lines = det.split('\n')
        dy = y
        for line in det_lines:
            c.drawString(250, dy, line[:50] + "..." if len(line)>50 else line)
            dy -= 12
            
        # V·∫Ω Note (Guest List)
        note_lines = note.split('\n')
        ny = y
        for line in note_lines:
            c.drawString(450, ny, line[:40] + "..." if len(line)>40 else line)
            ny -= 12
            
        # T√≠nh to√°n y ti·∫øp theo d·ª±a tr√™n s·ªë d√≤ng nhi·ªÅu nh·∫•t
        max_lines = max(len(det_lines), len(note_lines), 1)
        row_height = max(25, max_lines * 12 + 10)
        
        # [NEW] Ki·ªÉm tra ng·∫Øt trang
        if y - row_height < 50:
            c.setFillColor(HexColor(text_color))
            c.setFont(font_name, 9)
            c.drawCentredString(width / 2, 15, f"Page {c.getPageNumber()}")
            c.showPage()
            y = height - 50
            draw_watermark()
            
            # V·∫Ω l·∫°i Header b·∫£ng
            c.setFillColor(HexColor(primary_color))
            c.setFont(font_bold, 12)
            c.drawString(40, y, txt['svc_details_cont'])
            y -= 25
            c.setFillColor(HexColor("#E8F5E9"))
            c.rect(40, y - 5, width - 80, 20, fill=1, stroke=0)
            c.setFillColor(HexColor(primary_color))
            c.setFont(font_bold, 10)
            c.drawString(50, y, txt['tbl_name'])
            c.drawString(250, y, txt['tbl_det'])
            c.drawString(450, y, txt['tbl_note'])
            y -= 20
            c.setFillColor(HexColor(text_color))
            c.setFont(font_name, 10)
        
        # K·∫ª d√≤ng d∆∞·ªõi
        c.setStrokeColor(HexColor("#EEEEEE"))
        line_y = y - row_height + 15
        c.line(40, line_y, width - 40, line_y)
        y -= row_height

    # --- PH·∫¶N 3: GUEST LIST (N·∫æU C√ì) ---
    g_list_content = booking_info.get('guest_list', '')
    next_section_idx = 3
    
    if g_list_content:
        if y < 100:
            c.setFillColor(HexColor(text_color))
            c.setFont(font_name, 9)
            c.drawCentredString(width / 2, 15, f"Page {c.getPageNumber()}")
            c.showPage()
            y = height - 50
            draw_watermark()
            
        y -= 20
        c.setFillColor(HexColor(primary_color))
        c.setFont(font_bold, 12)
        c.drawString(40, y, txt['guest_list'])
        y -= 20
        c.setFillColor(HexColor(text_color))
        c.setFont(font_name, 10)
        
        for line in g_list_content.split('\n'):
            if y < 50:
                c.setFillColor(HexColor(text_color))
                c.setFont(font_name, 9)
                c.drawCentredString(width / 2, 15, f"Page {c.getPageNumber()}")
                c.showPage()
                y = height - 50
                draw_watermark()
                c.setFillColor(HexColor(primary_color))
                c.setFont(font_bold, 12)
                c.drawString(40, y, txt['guest_list_cont'])
                y -= 20
                c.setFillColor(HexColor(text_color))
                c.setFont(font_name, 10)
                
            c.drawString(50, y, line)
            y -= 15
        next_section_idx = 4

    # --- PH·∫¶N 4: INCLUDED & FOOTER ---
    if y < 150:
        c.setFillColor(HexColor(text_color))
        c.setFont(font_name, 9)
        c.drawCentredString(width / 2, 15, f"Page {c.getPageNumber()}")
        c.showPage()
        y = height - 50
        draw_watermark()
        
    y -= 20
    c.setFillColor(HexColor(primary_color))
    c.setFont(font_bold, 12)
    roman_num = "IV" if next_section_idx == 4 else "III"
    c.drawString(40, y, f"{roman_num}. {txt['included']}")
    y -= 20
    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 10)
    c.drawString(50, y, txt['inc_1'])
    c.drawString(50, y - 15, txt['inc_2'])
    

    # Signature
    y -= 45
    c.setFont(font_bold, 11)
    c.drawCentredString(width - 120, y, txt['confirmed_by'])
    c.setFont(font_name, 10)
    c.drawCentredString(width - 120, y - 15, comp_name)
    
    # D·∫•u m·ªôc gi·∫£ l·∫≠p (Text)
    c.setFillColor(HexColor("#C62828"))
    c.setFont(font_bold, 14)
    c.saveState()
    c.translate(width - 120, y - 50)
    c.rotate(15)
    c.drawCentredString(0, 0, txt['signed'])
    c.restoreState()

    c.setFillColor(HexColor(text_color))
    c.setFont(font_name, 9)
    c.drawCentredString(width / 2, 15, f"{txt['page']} {c.getPageNumber()}")

    c.save()
    buffer.seek(0)
    return buffer

# ==========================================
# 4. GIAO DI·ªÜN & LOGIC MODULES
# ==========================================

def render_notification_calendar():
    st.title("üìÖ L·ªãch Th√¥ng B√°o & Nh·∫Øc Thanh To√°n")
    
    # --- T·ª∞ ƒê·ªòNG CH·∫†Y KI·ªÇM TRA G·ª¨I L·∫¶N 2 ---
    if "auto_check_done" not in st.session_state:
        sent_count = check_and_send_due_reminders()
        if sent_count > 0:
            st.toast(f"üöÄ H·ªá th·ªëng v·ª´a t·ª± ƒë·ªông g·ª≠i {sent_count} email nh·∫Øc h·∫πn ƒë·∫øn h·∫°n!", icon="‚úÖ")
        st.session_state.auto_check_done = True

    # Chia layout
    col_cal, col_form = st.columns([1, 1.5])

    # --- C·ªòT TR√ÅI: DANH S√ÅCH & C√îNG C·ª§ L·ªäCH ---
    with col_cal:
        # === 1. C√îNG C·ª§ CHUY·ªÇN ƒê·ªîI L·ªäCH (M·ªöI) ===
        with st.expander("‚òØÔ∏è C√¥ng c·ª• Chuy·ªÉn ƒë·ªïi √Çm / D∆∞∆°ng", expanded=False):
            st.caption("Tra c·ª©u nhanh ng√†y √Çm/D∆∞∆°ng l·ªãch")
            cv_mode = st.radio("Ch·∫ø ƒë·ªô:", ["D∆∞∆°ng ‚û° √Çm", "√Çm ‚û° D∆∞∆°ng"], horizontal=True, label_visibility="collapsed")
            
            if cv_mode == "D∆∞∆°ng ‚û° √Çm":
                d_in = st.date_input("Ch·ªçn ng√†y D∆∞∆°ng:", datetime.now(), format="DD/MM/YYYY")
                if d_in:
                    lunar_txt = convert_solar_to_lunar(d_in)
                    st.success(f"üóìÔ∏è **{lunar_txt}**")
            else:
                c_d, c_m, c_y = st.columns(3)
                l_day = c_d.number_input("Ng√†y", 1, 30, 1)
                l_month = c_m.number_input("Th√°ng", 1, 12, 1)
                l_year = c_y.number_input("NƒÉm", 2024, 2030, datetime.now().year)
                is_leap = st.checkbox("Th√°ng nhu·∫≠n")
                
                if st.button("Tra c·ª©u D∆∞∆°ng l·ªãch"):
                    res_date = convert_lunar_to_solar(l_day, l_month, l_year, is_leap)
                    if res_date:
                        st.success(f"‚òÄÔ∏è Ng√†y D∆∞∆°ng: **{res_date.strftime('%d/%m/%Y')}**")
                        weekday_map = {0:"Th·ª© Hai", 1:"Th·ª© Ba", 2:"Th·ª© T∆∞", 3:"Th·ª© NƒÉm", 4:"Th·ª© S√°u", 5:"Th·ª© B·∫£y", 6:"Ch·ªß Nh·∫≠t"}
                        st.caption(f"({weekday_map[res_date.weekday()]})")
                    else:
                        st.error("Ng√†y √¢m l·ªãch kh√¥ng h·ª£p l·ªá!")

        st.divider()
        
        # === 2. DANH S√ÅCH L·ªäCH H·∫∏N ===
        st.markdown("### üóìÔ∏è L·ªãch s·∫Øp t·ªõi")
        with st.container(border=True):
            # L·∫•y danh s√°ch nh·∫Øc h·∫πn
            upcoming = run_query("SELECT * FROM payment_reminders WHERE status != 'sent_2' ORDER BY due_date ASC")
            
            if upcoming:
                for item in upcoming:
                    try:
                        d_obj = datetime.strptime(item['due_date'], '%Y-%m-%d %H:%M:%S')
                    except:
                        d_obj = datetime.strptime(item['due_date'], '%Y-%m-%d')
                        
                    days_left = (d_obj.date() - datetime.now().date()).days
                    
                    # Format ng√†y th√°ng nƒÉm
                    date_display = d_obj.strftime('%H:%M %d/%m/%Y')
                    lunar_display = convert_solar_to_lunar(d_obj).replace(" (√Çm l·ªãch)", "")
                    
                    color = "orange" if days_left == 0 else "green" if days_left > 0 else "red"
                    icon = "üîî" if days_left == 0 else "üìÖ"
                    
                    with st.expander(f"{icon} {date_display} (√Çm: {lunar_display}) | {item['ref_code']}"):
                        st.write(f"**N·ªôi dung:** {item['content']}")
                        st.write(f"**Ng∆∞·ªùi nh·∫≠n:** {item['receiver_email']}")
                        st.write(f"**S·ªë ti·ªÅn:** {format_vnd(item['amount'])} VND")
                        
                        status_txt = "Ch·ªù g·ª≠i L·∫ßn 1" if item['status'] == 'pending' else "ƒê√£ g·ª≠i L·∫ßn 1, ch·ªù L·∫ßn 2"
                        st.caption(f"Tr·∫°ng th√°i: {status_txt}")
                        
                        if st.button("üóëÔ∏è X√≥a", key=f"del_cal_{item['id']}"):
                            run_query("DELETE FROM payment_reminders WHERE id=?", (item['id'],), commit=True)
                            st.rerun()
            else:
                st.info("Kh√¥ng c√≥ l·ªãch nh·∫Øc n√†o s·∫Øp t·ªõi.")

    # --- C·ªòT PH·∫¢I: FORM T·∫†O (GI·ªÆ NGUY√äN) ---
    with col_form:
        st.markdown("### ‚úçÔ∏è T·∫°o y√™u c·∫ßu thanh to√°n m·ªõi")
        with st.container(border=True):
            # 1. L·∫•y d·ªØ li·ªáu Booking/Tour ƒë·ªÉ li√™n k·∫øt
            # [UPDATED] Ph√¢n quy·ªÅn xem Booking/Tour
            user_info = st.session_state.get("user_info", {})
            u_role = user_info.get('role')
            u_name = user_info.get('name')
            
            tour_q = "SELECT tour_code, tour_name FROM tours WHERE status='running'"
            tour_p = []
            bk_q = "SELECT code, name FROM service_bookings WHERE status='active'"
            bk_p = []
            
            if u_role not in ['admin', 'admin_f1']:
                tour_q += " AND sale_name=?"
                tour_p.append(u_name)
                bk_q += " AND sale_name=?"
                bk_p.append(u_name)
                
            tours = run_query(tour_q, tuple(tour_p))
            bookings = run_query(bk_q, tuple(bk_p))
            
            opts = ["-- Ch·ªçn m√£ li√™n k·∫øt --"]
            if tours: opts += [f"TOUR | {t['tour_code']} | {t['tour_name']}" for t in tours]
            if bookings: opts += [f"BOOK | {b['code']} | {b['name']}" for b in bookings]
            
            sel_ref = st.selectbox("Li√™n k·∫øt v·ªõi Booking/Tour:", opts, key="notif_ref")
            
            # T·ª± ƒë·ªông ƒëi·ªÅn th√¥ng tin n·∫øu ch·ªçn m√£
            ref_code = ""
            ref_name = ""
            if sel_ref != "-- Ch·ªçn m√£ li√™n k·∫øt --":
                parts = sel_ref.split(" | ")
                ref_code = parts[1]
                ref_name = parts[2]

            c1, c2 = st.columns(2)
            
            # [C·∫¨P NH·∫¨T] Nh·∫≠p s·ªë ti·ªÅn c√≥ ƒë·ªãnh d·∫°ng VND
            if "req_amount_val" not in st.session_state: st.session_state.req_amount_val = ""
            def fmt_req_amount():
                val = st.session_state.req_amount_val
                try:
                    v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                    st.session_state.req_amount_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                except: pass

            amount_input = c1.text_input("S·ªë ti·ªÅn y√™u c·∫ßu:", key="req_amount_val", on_change=fmt_req_amount, help="Nh·∫≠p s·ªë ti·ªÅn (VD: 1000000)")
            try: amount = float(amount_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
            except: amount = 0.0
            
            # [C·∫¨P NH·∫¨T] Ch·ªçn ng√†y v√† gi·ªù v·ªõi format DD/MM/YYYY
            with c2:
                c_d, c_t = st.columns(2)
                due_date = c_d.date_input("Ng√†y h·∫πn (L·∫ßn 2):", min_value=datetime.now(), format="DD/MM/YYYY", key="notif_date")
                due_time = c_t.time_input("Gi·ªù h·∫πn:", value=datetime.now().time(), key="notif_time")
                due_datetime = datetime.combine(due_date, due_time)
            
            # Hi·ªÉn th·ªã ng√†y √Çm l·ªãch t∆∞∆°ng ·ª©ng ngay d∆∞·ªõi ƒë·ªÉ ti·ªán theo d√µi
            if due_date:
                st.caption(f"üóìÔ∏è T∆∞∆°ng ·ª©ng √Çm l·ªãch: {convert_solar_to_lunar(due_date)}")
            
            # [C·∫¨P NH·∫¨T] Th√¥ng tin ng∆∞·ªùi g·ª≠i v√† CC
            current_user_name = st.session_state.user_info.get('name', '')
            sender_name = st.text_input("Ng∆∞·ªùi g·ª≠i (Hi·ªÉn th·ªã trong mail):", value=current_user_name, disabled=True)
            cc_email = st.text_input("CC Email (c√°ch nhau d·∫•u ph·∫©y):", placeholder="boss@gmail.com, ketoan@gmail.com", key="notif_cc")
            
            # Email m·∫∑c ƒë·ªãnh l·∫•y t·ª´ secrets
            def_email = ""
            try: def_email = st.secrets["email"].get("receiver_default", "")
            except: pass
            
            receiver = st.text_input("Email ng∆∞·ªùi nh·∫≠n th√¥ng b√°o:", value=def_email, help="Email c·ªßa K·∫ø to√°n ho·∫∑c Kh√°ch h√†ng", key="notif_receiver")
            content = st.text_area("N·ªôi dung y√™u c·∫ßu thanh to√°n:", height=100, placeholder="VD: Y√™u c·∫ßu thanh to√°n ƒë·ª£t 1 cho ƒëo√†n...", key="notif_content")
            
            # [NEW] Bank Info Inputs
            st.markdown("##### üè¶ Th√¥ng tin chuy·ªÉn kho·∫£n")
            c_b1, c_b2 = st.columns(2)
            bank_name = c_b1.text_input("T√™n Ng√¢n H√†ng", placeholder="VD: Techcombank", key="notif_bank_name")
            bank_acc = c_b2.text_input("S·ªë T√†i Kho·∫£n", placeholder="VD: 1903...", key="notif_bank_acc")
            bank_holder = st.text_input("Ch·ªß T√†i Kho·∫£n", placeholder="VD: NGUYEN VAN A", key="notif_bank_holder")

            st.info("‚ÑπÔ∏è **C∆° ch·∫ø:** Khi b·∫•m n√∫t d∆∞·ªõi, h·ªá th·ªëng s·∫Ω **G·ª¨I NGAY 1 EMAIL** cho ng∆∞·ªùi nh·∫≠n. ƒê·∫øn ng√†y h·∫πn ·ªü tr√™n, h·ªá th·ªëng s·∫Ω **G·ª¨I TI·∫æP 1 EMAIL N·ªÆA**.")

            if st.button("üöÄ L∆∞u & G·ª≠i th√¥ng b√°o ngay", type="primary", use_container_width=True):
                if ref_code and receiver and content:
                    with st.spinner("ƒêang g·ª≠i email l·∫ßn 1..."):
                        # 1. G·ª≠i Email L·∫ßn 1 Ngay l·∫≠p t·ª©c
                        bank_html = ""
                        if bank_name and bank_acc:
                            bank_html = f"""
                            <div style="background-color: #f8f9fa; padding: 15px; border-radius: 5px; margin: 15px 0;">
                                <h4 style="margin-top: 0;">üè¶ TH√îNG TIN CHUY·ªÇN KHO·∫¢N</h4>
                                <p><strong>Ng√¢n h√†ng:</strong> {bank_name}</p>
                                <p><strong>S·ªë t√†i kho·∫£n:</strong> {bank_acc}</p>
                                <p><strong>Ch·ªß t√†i kho·∫£n:</strong> {bank_holder}</p>
                            </div>
                            """

                        subj = f"üì¢ [TH√îNG B√ÅO] Y√™u c·∫ßu thanh to√°n - {ref_code}"
                        html_body = f"""
                        <h3>üì¢ Y√äU C·∫¶U THANH TO√ÅN (L·∫¶N 1)</h3>
                        <p>K√≠nh g·ª≠i,</p>
                        <p>Ch√∫ng t√¥i g·ª≠i th√¥ng b√°o thanh to√°n cho d·ªãch v·ª• <strong>{ref_name}</strong> (M√£: {ref_code}).</p>
                        <p><strong>S·ªë ti·ªÅn:</strong> {format_vnd(amount)} VND</p>
                        <p><strong>N·ªôi dung:</strong> {content}</p>
                        {bank_html}
                        <p>H·ªá th·ªëng s·∫Ω g·ª≠i nh·∫Øc nh·ªü l·∫°i v√†o l√∫c: <strong>{due_datetime.strftime('%H:%M %d/%m/%Y')}</strong>.</p>
                        <hr>
                        <p>Tr√¢n tr·ªçng,<br><strong>{sender_name}</strong><br><small>Bali Tourist Automated System</small></p>
                        """
                        
                        ok, msg = send_email_notification(receiver, subj, html_body, cc_emails=cc_email)
                        if ok:
                            # 2. L∆∞u v√†o DB ƒë·ªÉ h·∫πn gi·ªù g·ª≠i l·∫ßn 2
                            run_query("""INSERT INTO payment_reminders 
                                (ref_code, ref_name, amount, due_date, receiver_email, content, status, created_at, cc_email, sender_name, bank_name, bank_account, bank_holder)
                                VALUES (?, ?, ?, ?, ?, ?, 'sent_1', ?, ?, ?, ?, ?, ?)""", 
                                (ref_code, ref_name, amount, due_datetime.strftime('%Y-%m-%d %H:%M:%S'), receiver, content, datetime.now().strftime('%Y-%m-%d %H:%M:%S'), cc_email, sender_name, bank_name, bank_acc, bank_holder),
                                commit=True)
                            
                            # [FIX] Reset form fields
                            keys_to_reset = ["req_amount_val", "notif_receiver", "notif_content", "notif_bank_name", "notif_bank_acc", "notif_bank_holder", "notif_cc", "notif_date", "notif_time", "notif_ref"]
                            for k in keys_to_reset:
                                if k in st.session_state: del st.session_state[k]

                            st.success(f"‚úÖ ƒê√£ g·ª≠i email l·∫ßn 1 v√† l√™n l·ªãch nh·∫Øc l·∫ßn 2 v√†o ng√†y {due_date.strftime('%d/%m/%Y')}!"); time.sleep(1); st.rerun()
                        else: st.error(msg)
                else:
                    st.warning("Vui l√≤ng ch·ªçn M√£ li√™n k·∫øt, nh·∫≠p Email v√† N·ªôi dung.")

def render_dashboard():
    st.title("üè† Trang Ch·ªß - T·ªïng Quan Kinh Doanh")
    
    # User context
    user_info = st.session_state.get("user_info", {})
    role = user_info.get('role')
    username = user_info.get('name')
    
    # Time context
    now = datetime.now()
    current_month = now.month
    current_year = now.year
    
    st.markdown(f"### üìÖ S·ªë li·ªáu th√°ng {current_month}/{current_year}")
    
    # Data fetching
    # 1. Tours
    tour_query = "SELECT * FROM tours WHERE status != 'deleted'"
    tour_params = []
    if role == 'sale':
        tour_query += " AND sale_name=?"
        tour_params.append(username)
    tours = run_query(tour_query, tuple(tour_params))
    
    # 2. Bookings
    bk_query = "SELECT * FROM service_bookings WHERE status != 'deleted'"
    bk_params = []
    if role == 'sale':
        bk_query += " AND sale_name=?"
        bk_params.append(username)
    bookings = run_query(bk_query, tuple(bk_params))
    
    # 3. Costs (for tours)
    all_items = run_query("SELECT tour_id, item_type, total_amount FROM tour_items")
    items_map = {}
    if all_items:
        for item in all_items:
            tid = item['tour_id']
            itype = item['item_type']
            amt = item['total_amount'] or 0
            if tid not in items_map: items_map[tid] = {'EST': 0, 'ACT': 0}
            items_map[tid][itype] += amt

    # Processing
    total_tour_rev = 0
    total_tour_profit = 0
    count_tours = 0
    tours_in_month = []
    
    total_bk_rev = 0
    total_bk_profit = 0
    count_bks = 0
    bks_in_month = []
    
    # Process Tours
    if tours:
        for t in tours:
            t = dict(t)
            try:
                s_date = datetime.strptime(t['start_date'], '%d/%m/%Y')
                if s_date.month == current_month and s_date.year == current_year:
                    count_tours += 1
                    
                    final_price = float(t.get('final_tour_price', 0) or 0)
                    child_price = float(t.get('child_price', 0) or 0)
                    final_qty = float(t.get('final_qty', 0) or 0)
                    child_qty = float(t.get('child_qty', 0) or 0)
                    if final_qty == 0: final_qty = float(t.get('guest_count', 1))
                    
                    rev = (final_price * final_qty) + (child_price * child_qty)
                    
                    costs = items_map.get(t['id'], {'EST': 0, 'ACT': 0})
                    est_cost = costs['EST']; act_cost = costs['ACT']
                    
                    if rev == 0:
                        p_pct = t.get('est_profit_percent', 0) or 0
                        t_pct = t.get('est_tax_percent', 0) or 0
                        profit_est_val = est_cost * (p_pct/100)
                        rev = (est_cost + profit_est_val) * (1 + t_pct/100)
                    
                    t_pct = t.get('est_tax_percent', 0) or 0
                    net_rev = rev / (1 + t_pct/100) if (1 + t_pct/100) != 0 else rev
                    prof = net_rev - act_cost
                    
                    total_tour_rev += rev; total_tour_profit += prof
                    
                    t_display = dict(t); t_display['revenue'] = rev; t_display['profit'] = prof
                    tours_in_month.append(t_display)
            except: pass

    # Process Bookings
    if bookings:
        for b in bookings:
            try:
                c_date = datetime.strptime(str(b['created_at']).split(' ')[0], '%Y-%m-%d')
                if c_date.month == current_month and c_date.year == current_year:
                    count_bks += 1
                    rev = float(b['selling_price'] or 0); prof = float(b['profit'] or 0)
                    total_bk_rev += rev; total_bk_profit += prof
                    b_display = dict(b); b_display['revenue'] = rev; b_display['profit'] = prof
                    bks_in_month.append(b_display)
            except: pass

    # Display Metrics
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("T·ªïng Doanh Thu", format_vnd(total_tour_rev + total_bk_rev) + " VND")
    m2.metric("T·ªïng L·ª£i Nhu·∫≠n", format_vnd(total_tour_profit + total_bk_profit) + " VND")
    m3.metric("S·ªë l∆∞·ª£ng Tour", count_tours)
    m4.metric("S·ªë l∆∞·ª£ng Booking", count_bks)
    
    st.divider()
    c_left, c_right = st.columns(2)
    with c_left:
        st.subheader("üì¶ Tour trong th√°ng")
        if tours_in_month:
            df_t = pd.DataFrame(tours_in_month)[['start_date', 'tour_name', 'revenue', 'profit']]
            df_t['revenue'] = df_t['revenue'].apply(lambda x: format_vnd(x) + " VND")
            df_t['profit'] = df_t['profit'].apply(lambda x: format_vnd(x) + " VND")
            st.dataframe(df_t, column_config={"start_date": "Ng√†y ƒëi", "tour_name": "T√™n ƒëo√†n", "revenue": "Doanh thu", "profit": "L·ª£i nhu·∫≠n (TT)"}, use_container_width=True, hide_index=True)
        else: st.info("Kh√¥ng c√≥ tour n√†o.")
            
    with c_right:
        st.subheader("üîñ Booking trong th√°ng")
        if bks_in_month:
            df_b = pd.DataFrame(bks_in_month)[['created_at', 'name', 'revenue', 'profit']]
            df_b['revenue'] = df_b['revenue'].apply(lambda x: format_vnd(x) + " VND")
            df_b['profit'] = df_b['profit'].apply(lambda x: format_vnd(x) + " VND")
            st.dataframe(df_b, column_config={"created_at": "Ng√†y t·∫°o", "name": "T√™n d·ªãch v·ª•", "revenue": "Doanh thu", "profit": "L·ª£i nhu·∫≠n"}, use_container_width=True, hide_index=True)
        else: st.info("Kh√¥ng c√≥ booking n√†o.")

# ==========================================
# 4. GIAO DI·ªÜN & LOGIC MODULES
# ==========================================

def render_login_page(comp):
    col_a, col_b, col_c = st.columns([1, 2, 1])
    with col_b:
        st.write("")
        if comp['logo_b64_str']:
            st.markdown(f'''
            <div class="company-header-container">
                <img src="data:image/png;base64,{comp["logo_b64_str"]}" class="company-logo-img">
                <div class="company-info-text">
                    <h1>{comp['name']}</h1>
                    <p>üìç {comp['address']}</p>
                    <p>MST: {comp['phone']}</p>
                </div>
            </div>
            ''', unsafe_allow_html=True)
        else:
            st.markdown(f"""<div style="text-align:center; margin-top:20px;"><h1 style="color:#28a745 !important;">{comp['name']}</h1><p>üìç {comp['address']}<br>MST: {comp['phone']}</p></div>""", unsafe_allow_html=True)
        
        tab_login, tab_reg = st.tabs(["üîê ƒêƒÉng nh·∫≠p", "üìù ƒêƒÉng k√Ω"])
        with tab_login:
            with st.container(border=True):
                with st.form("login"):
                    u = st.text_input("T√†i kho·∫£n"); p = st.text_input("M·∫≠t kh·∫©u", type="password")
                    if st.form_submit_button("ƒêƒÇNG NH·∫¨P", width="stretch"):
                        pw_hash = hash_pass(p)
                        
                        # [CODE M·ªöI] ƒê·ªçc t·ª´ Google Sheet thay v√¨ SQL
                        df_users = load_table('users') 
                        
                        # Ki·ªÉm tra user
                        if not df_users.empty:
                            # L·ªçc user tr√πng username v√† password
                            mask = (df_users['username'] == u) & (df_users['password'] == pw_hash) # type: ignore
                            user_found = df_users.loc[mask]
                            
                            if not user_found.empty and user_found.iloc[0]['status'] == 'approved': # type: ignore
                                st.session_state.logged_in = True
                                st.session_state.user_info = {
                                    "name": user_found.iloc[0]['username'],  # type: ignore
                                    "role": user_found.iloc[0]['role'] # type: ignore
                                }
                                st.rerun()
                            else:
                                st.error("Sai th√¥ng tin ho·∫∑c t√†i kho·∫£n ch∆∞a duy·ªát!")
                        else:
                            st.error("Kh√¥ng k·∫øt n·ªëi ƒë∆∞·ª£c danh s√°ch ng∆∞·ªùi d√πng!")
        with tab_reg:
            with st.container(border=True):
                with st.form("reg"):
                    nu = st.text_input("T√†i kho·∫£n m·ªõi"); np = st.text_input("M·∫≠t kh·∫©u", type="password")
                    if st.form_submit_button("ƒêƒÇNG K√ù", width="stretch"):
                        try:
                            add_row_to_table('users', {'username': nu, 'password': hash_pass(np), 'role': 'user', 'status': 'pending'})
                            st.success("ƒê√£ g·ª≠i y√™u c·∫ßu! Ch·ªù x√≠u nha ü•∞")
                        except: st.error("T√™n n√†y c√≥ ng∆∞·ªùi d√πng r·ªìi n√®!")

def render_admin_notifications():
    st.divider()
    st.markdown("### üîî Trung T√¢m Th√¥ng B√°o & Ph√™ Duy·ªát")
    
    # --- L·∫§Y D·ªÆ LI·ªÜU C·∫¶N DUY·ªÜT ---
    pending_projs = run_query("SELECT * FROM projects WHERE pending_name IS NOT NULL AND pending_name != ''")
    pending_tours = run_query("SELECT * FROM tours WHERE pending_name IS NOT NULL AND pending_name != ''")
    del_tours = run_query("SELECT * FROM tours WHERE request_delete=1")
    req_edit_tours = run_query("SELECT * FROM tours WHERE request_edit_act=1")
    pending_users = run_query("SELECT * FROM users WHERE role='user' AND status='pending'")
    req_invoices = run_query("SELECT * FROM invoices WHERE request_edit=1 AND status='active'")
    
    has_requests = False

    # 1. DUY·ªÜT ƒê·ªîI T√äN D·ª∞ √ÅN
    if pending_projs:
        has_requests = True
        st.markdown(f"#### üìù ƒê·ªïi t√™n D·ª± √°n ({len(pending_projs)})")
        for p in pending_projs:
            with st.container(border=True):
                st.markdown(f"**D·ª± √°n:** `{p['project_name']}` ‚û° <span style='color:green'><b>`{p['pending_name']}`</b></span>", unsafe_allow_html=True) # type: ignore
                c_app, c_rej = st.columns(2)
                if c_app.button("‚úî Duy·ªát", key=f"app_ren_{p['id']}", type="primary"): # type: ignore
                    run_query("UPDATE projects SET project_name=?, pending_name=NULL WHERE id=?", (p['pending_name'], p['id']), commit=True) # type: ignore
                    st.rerun()
                if c_rej.button("‚úñ H·ªßy", key=f"rej_ren_{p['id']}"): # type: ignore
                    run_query("UPDATE projects SET pending_name=NULL WHERE id=?", (p['id'],), commit=True) # type: ignore
                    st.rerun()

    # 2. DUY·ªÜT ƒê·ªîI T√äN TOUR
    if pending_tours:
        has_requests = True
        st.markdown(f"#### üì¶ ƒê·ªïi t√™n Tour ({len(pending_tours)})")
        for t in pending_tours:
            with st.container(border=True):
                st.markdown(f"**Tour:** `{t['tour_name']}` ‚û° <span style='color:green'><b>`{t['pending_name']}`</b></span>", unsafe_allow_html=True) # type: ignore
                c_app, c_rej = st.columns(2)
                if c_app.button("‚úî Duy·ªát", key=f"app_ren_t_{t['id']}", type="primary"): # type: ignore
                    run_query("UPDATE tours SET tour_name=?, pending_name=NULL WHERE id=?", (t['pending_name'], t['id']), commit=True) # type: ignore
                    st.rerun()
                if c_rej.button("‚úñ H·ªßy", key=f"rej_ren_t_{t['id']}"): # type: ignore
                    run_query("UPDATE tours SET pending_name=NULL WHERE id=?", (t['id'],), commit=True) # type: ignore
                    st.rerun()

    # 3. DUY·ªÜT X√ìA TOUR
    if del_tours:
        has_requests = True
        st.markdown(f"#### <span style='color:red;'>üóëÔ∏è X√≥a Tour ({len(del_tours)})</span>", unsafe_allow_html=True)
        for t in del_tours:
            with st.container(border=True):
                st.markdown(f"‚ùå Y√™u c·∫ßu x√≥a Tour: **{t['tour_name']}**") # type: ignore
                c_app, c_rej = st.columns(2)
                if c_app.button("‚úî Duy·ªát x√≥a", key=f"app_del_t_{t['id']}", type="primary"): # type: ignore
                    run_query("UPDATE tours SET request_delete=2 WHERE id=?", (t['id'],), commit=True) # type: ignore
                    st.success("ƒê√£ duy·ªát! Ch·ªù ng∆∞·ªùi d√πng x√°c nh·∫≠n."); time.sleep(1); st.rerun()
                if c_rej.button("‚úñ T·ª´ ch·ªëi", key=f"rej_del_t_{t['id']}"): # type: ignore
                    run_query("UPDATE tours SET request_delete=0 WHERE id=?", (t['id'],), commit=True) # type: ignore
                    st.rerun()

    # 4. DUY·ªÜT S·ª¨A QUY·∫æT TO√ÅN (M·ªöI)
    if req_edit_tours:
        has_requests = True
        st.markdown(f"#### üí∏ S·ª≠a Quy·∫øt to√°n ({len(req_edit_tours)})")
        for t in req_edit_tours:
            with st.container(border=True):
                st.write(f"Tour: **{t['tour_name']}**") # type: ignore
                c1, c2 = st.columns(2)
                if c1.button("‚úî Duy·ªát", key=f"app_edit_act_{t['id']}"): # type: ignore
                    run_query("UPDATE tours SET request_edit_act=2 WHERE id=?", (t['id'],), commit=True); st.rerun() # type: ignore
                if c2.button("‚úñ T·ª´ ch·ªëi", key=f"rej_edit_act_{t['id']}"): # type: ignore
                    run_query("UPDATE tours SET request_edit_act=0 WHERE id=?", (t['id'],), commit=True); st.rerun() # type: ignore

    # 5. DUY·ªÜT USER
    if pending_users:
        has_requests = True
        st.markdown(f"#### üë§ ƒêƒÉng k√Ω m·ªõi ({len(pending_users)})")
        for u in pending_users:
            with st.container(border=True):
                st.write(f"User: **{u['username']}**") # type: ignore
                c1, c2 = st.columns(2)
                if c1.button("‚úî Duy·ªát", key=f"app_user_{u['id']}"): # type: ignore
                    run_query("UPDATE users SET status='approved' WHERE id=?", (u['id'],), commit=True) # type: ignore
                    st.rerun()
                if c2.button("‚úñ X√≥a", key=f"del_user_{u['id']}"): # type: ignore
                    run_query("DELETE FROM users WHERE id=?", (u['id'],), commit=True) # type: ignore
                    st.rerun()

    # 6. DUY·ªÜT S·ª¨A GI√Å H√ìA ƒê∆†N
    if req_invoices:
        has_requests = True
        st.markdown(f"#### üí∞ S·ª≠a gi√° H√≥a ƒë∆°n ({len(req_invoices)})")
        for r in req_invoices:
            with st.container(border=True):
                st.info(f"Hƒê: {r['invoice_number']} | Ti·ªÅn: {format_vnd(r['total_amount'])}") # type: ignore
                c1, c2 = st.columns(2)
                if c1.button("‚úî Duy·ªát", key=f"app_inv_{r['id']}"): # type: ignore
                    run_query("UPDATE invoices SET edit_count=0, request_edit=0 WHERE id=?", (r['id'],), commit=True) # type: ignore
                    st.success("ƒê√£ duy·ªát!"); time.sleep(0.5); st.rerun()
                if c2.button("‚úñ T·ª´ ch·ªëi", key=f"rej_inv_{r['id']}"): # type: ignore
                    run_query("UPDATE invoices SET request_edit=0 WHERE id=?", (r['id'],), commit=True) # type: ignore
                    st.rerun()

    if not has_requests:
        st.success("‚úÖ Hi·ªán kh√¥ng c√≥ y√™u c·∫ßu n√†o c·∫ßn duy·ªát.")

def render_admin_panel(comp):
    with st.expander("‚öôÔ∏è Admin Panel", expanded=False):
        st.caption("C·∫≠p nh·∫≠t th√¥ng tin C√¥ng ty")
        with st.form("comp_update"):
            cn = st.text_input("T√™n", value=comp['name'])
            ca = st.text_input("ƒê·ªãa ch·ªâ", value=comp['address'])
            cp = st.text_input("M√£ S·ªë Thu·∫ø", value=comp['phone'])
            ul = st.file_uploader("Logo", type=['png','jpg'])
            if st.form_submit_button("L∆∞u"):
                update_company_info(cn, ca, cp, ul.read() if ul else None)
                st.success("Xong!"); time.sleep(0.5); st.rerun()
        
        # Ch·ªâ admin ch√≠nh m·ªõi th·∫•y m·ª•c x√≥a
        if (st.session_state.user_info or {}).get('role') == 'admin':
            st.divider()
            st.markdown("##### üóëÔ∏è Qu·∫£n l√Ω d·ªØ li·ªáu")
            
            c1, c2 = st.columns(2)
            with c1:
                if st.button("X√≥a H√≥a ƒê∆°n", use_container_width=True, help="X√≥a TO√ÄN B·ªò d·ªØ li·ªáu H√≥a ƒë∆°n & UNC"):
                    run_query("DELETE FROM invoices", commit=True)
                    run_query("DELETE FROM sqlite_sequence WHERE name='invoices'", commit=True)
                    if os.path.exists(UPLOAD_FOLDER):
                        for f in os.listdir(UPLOAD_FOLDER):
                            if "UNC" not in f and "converted" not in f: 
                                    try: os.remove(os.path.join(UPLOAD_FOLDER, f))
                                    except: pass
                    st.toast("ƒê√£ x√≥a s·∫°ch H√≥a ƒê∆°n!"); time.sleep(1); st.rerun()
                
                if st.button("X√≥a Tour", use_container_width=True, help="X√≥a TO√ÄN B·ªò d·ªØ li·ªáu Tour (D·ª± to√°n v√† Quy·∫øt to√°n)"):
                    run_query("DELETE FROM tours", commit=True)
                    run_query("DELETE FROM tour_items", commit=True)
                    run_query("DELETE FROM sqlite_sequence WHERE name='tours'", commit=True)
                    run_query("DELETE FROM sqlite_sequence WHERE name='tour_items'", commit=True)
                    st.toast("ƒê√£ x√≥a s·∫°ch d·ªØ li·ªáu Tour!"); time.sleep(1); st.rerun()
            
            with c2:
                if st.button("X√≥a Booking", use_container_width=True, help="X√≥a TO√ÄN B·ªò d·ªØ li·ªáu Booking d·ªãch v·ª•"):
                    run_query("DELETE FROM service_bookings", commit=True)
                    run_query("DELETE FROM sqlite_sequence WHERE name='service_bookings'", commit=True)
                    st.toast("ƒê√£ x√≥a s·∫°ch Booking!"); time.sleep(1); st.rerun()
                
                if st.button("X√≥a Kh√°ch H√†ng", use_container_width=True, help="X√≥a TO√ÄN B·ªò d·ªØ li·ªáu Kh√°ch h√†ng"):
                    run_query("DELETE FROM customers", commit=True); run_query("DELETE FROM sqlite_sequence WHERE name='customers'", commit=True)
                    st.toast("ƒê√£ x√≥a s·∫°ch Kh√°ch h√†ng!"); time.sleep(1); st.rerun()

            with st.popover("üí• X√ìA TO√ÄN B·ªò D·ªÆ LI·ªÜU üí•", use_container_width=True):
                st.error("C·∫¢NH B√ÅO C·ª∞C K·ª≤ NGUY HI·ªÇM!")
                st.warning("H√†nh ƒë·ªông n√†y s·∫Ω **X√ìA S·∫†CH TO√ÄN B·ªò** d·ªØ li·ªáu kinh doanh (H√≥a ƒë∆°n, Tour, Booking, Kh√°ch h√†ng...). D·ªØ li·ªáu ng∆∞·ªùi d√πng v√† th√¥ng tin c√¥ng ty s·∫Ω ƒë∆∞·ª£c gi·ªØ l·∫°i. H√†nh ƒë·ªông n√†y kh√¥ng th·ªÉ ho√†n t√°c.")
                st.warning("Ch·ªâ th·ª±c hi·ªán khi b·∫°n mu·ªën b·∫Øt ƒë·∫ßu l·∫°i t·ª´ ƒë·∫ßu. B·∫°n c√≥ ch·∫Øc ch·∫Øn kh√¥ng?")
                if st.button("C√ì, T√îI HI·ªÇU R·ª¶I RO V√Ä MU·ªêN X√ìA T·∫§T C·∫¢", type="primary"):
                    TABLES_TO_DELETE = [
                        'invoices', 'projects', 'project_links', 'service_bookings', 
                        'customers', 'tours', 'tour_items', 'ocr_learning',
                        'transaction_history',
                        'flight_tickets', 'flight_groups', 'flight_group_links'
                    ]
                    with st.spinner("ƒêang d·ªçn d·∫πp h·ªá th·ªëng..."):
                        for table in TABLES_TO_DELETE:
                            run_query(f"DELETE FROM {table}", commit=True)
                            run_query(f"DELETE FROM sqlite_sequence WHERE name='{table}'", commit=True)
                        if os.path.exists(UPLOAD_FOLDER):
                            for f in os.listdir(UPLOAD_FOLDER):
                                try: os.remove(os.path.join(UPLOAD_FOLDER, f))
                                except: pass
                    st.success("ƒê√£ x√≥a to√†n b·ªô d·ªØ li·ªáu kinh doanh v√† c√°c file ƒë√£ upload!")
                    time.sleep(2); st.rerun()

        with st.popover("üîÑ ƒê·ªìng b·ªô l√™n Google Sheet", use_container_width=True):
            st.warning("‚ö†Ô∏è H√†nh ƒë·ªông n√†y s·∫Ω **ghi ƒë√® to√†n b·ªô** d·ªØ li·ªáu tr√™n Google Sheet b·∫±ng d·ªØ li·ªáu hi·ªán t·∫°i tr√™n m√°y c·ªßa b·∫°n. B·∫°n c√≥ ch·∫Øc ch·∫Øn kh√¥ng?")
            if st.button("C√≥, t√¥i mu·ªën ƒë·ªìng b·ªô ngay", type="primary"):
                sync_all_data_to_gsheet()

def render_sidebar(comp):
    with st.sidebar:
        if comp['logo_b64_str']: st.markdown(f'<div style="text-align:center; margin-bottom:20px;"><img src="data:image/png;base64,{comp["logo_b64_str"]}" width="120" style="border-radius:10px;"></div>', unsafe_allow_html=True)
        
        user_info = st.session_state.get("user_info")
        if user_info and isinstance(user_info, dict):
            st.success(f"Xin ch√†o **{user_info.get('name', 'User')}** üëã")
        else:
            st.session_state.logged_in = False
            st.rerun()
        
        st.markdown("### üóÇÔ∏è Ph√¢n H·ªá Qu·∫£n L√Ω")
        module = st.selectbox("Ch·ªçn ch·ª©c nƒÉng:", ["üè† Trang Ch·ªß", "üìÖ L·ªãch Th√¥ng B√°o", "üîñ Qu·∫£n L√Ω Booking", "üí∞ Ki·ªÉm So√°t Chi Ph√≠", "üí≥ Qu·∫£n L√Ω C√¥ng N·ª£", "üì¶ Qu·∫£n L√Ω Tour ", "ü§ù Qu·∫£n L√Ω Kh√°ch H√†ng", "üë• Qu·∫£n L√Ω Nh√¢n S·ª±", "üîç Tra c·ª©u th√¥ng tin"], label_visibility="collapsed")
        
        menu = None
        if module == "üí∞ Ki·ªÉm So√°t Chi Ph√≠":
            menu = st.radio("Menu", ["1. Nh·∫≠p H√≥a ƒê∆°n", "2. B√°o C√°o T·ªïng H·ª£p"])
        
        if st.session_state.user_info and st.session_state.user_info.get('role') in ['admin', 'admin_f1']:
            render_admin_notifications()

        st.divider()

        if st.session_state.user_info and st.session_state.user_info.get('role') in ['admin', 'admin_f1']:
            render_admin_panel(comp)

        if st.button("ƒêƒÉng xu·∫•t", use_container_width=True):
            st.session_state.logged_in = False
            st.rerun()
        with st.popover("üîê ƒê·ªïi m·∫≠t kh·∫©u", use_container_width=True):
            st.markdown("##### C·∫≠p nh·∫≠t m·∫≠t kh·∫©u")
            with st.form("change_pass"):
                op = st.text_input("M·∫≠t kh·∫©u hi·ªán t·∫°i", type="password")
                new_p = st.text_input("M·∫≠t kh·∫©u m·ªõi", type="password")
                cp = st.text_input("X√°c nh·∫≠n m·∫≠t kh·∫©u m·ªõi", type="password")
                if st.form_submit_button("L∆∞u thay ƒë·ªïi"):
                    c_user = (st.session_state.user_info or {}).get('name', '')
                    db_u = run_query("SELECT * FROM users WHERE username=?", (c_user,), fetch_one=True)
                    if isinstance(db_u, sqlite3.Row) and db_u['password'] == hash_pass(op): # type: ignore
                        if new_p and new_p == cp:
                            run_query("UPDATE users SET password=? WHERE username=?", (hash_pass(new_p), c_user), commit=True)
                            st.success("ƒê·ªïi m·∫≠t kh·∫©u th√†nh c√¥ng! ƒêƒÉng nh·∫≠p l·∫°i nh√©.")
                            time.sleep(1)
                            st.session_state.logged_in = False
                            st.rerun()
                        else:
                            st.error("M·∫≠t kh·∫©u m·ªõi kh√¥ng kh·ªõp!")
                    else:
                        st.error("M·∫≠t kh·∫©u c≈© sai r·ªìi!")

        # --- KI·ªÇM TRA K·∫æT N·ªêI GOOGLE (DEBUG) ---
        st.divider() # type: ignore
        with st.expander("üîå Ki·ªÉm tra k·∫øt n·ªëi Google"):
            if st.button("Test K·∫øt N·ªëi Ngay", use_container_width=True):
                try:
                    with st.spinner("ƒêang k·∫øt n·ªëi Google API..."):
                        gc = get_gspread_client()
                        sh = gc.open_by_key(SPREADSHEET_ID)
                        st.success(f"‚úÖ Sheet OK: {sh.title}")
                        drive = get_drive_service()
                        st.success(f"‚úÖ Drive OK (ID: ...{DRIVE_FOLDER_ID[-5:]})")
                except Exception as e:
                    st.error(f"‚ùå L·ªói: {str(e)}")
                    st.info("üí° G·ª£i √Ω: Ki·ªÉm tra file service_account.json ho·∫∑c quy·ªÅn chia s·∫ª c·ªßa Sheet/Folder.")
    return module, menu

# --- H√ÄM HI·ªÇN TH·ªä SO S√ÅNH CHI PH√ç (UNC vs H√ìA ƒê∆†N) ---
def render_cost_comparison(code):
    # L·∫•y t·∫•t c·∫£ h√≥a ƒë∆°n/UNC theo m√£
    docs = run_query("SELECT * FROM invoices WHERE cost_code=? AND status='active'", (code,))
    if not docs:
        st.info("Ch∆∞a c√≥ ch·ª©ng t·ª´ n√†o li√™n k·∫øt.")
        return 0

    df = pd.DataFrame([dict(r) for r in docs])
    
    # L·ªçc chi ph√≠ ƒë·∫ßu v√†o (IN)
    df_in = df.loc[df['type'] == 'IN'].copy() # type: ignore
    if df_in.empty:
        st.info("Ch∆∞a c√≥ chi ph√≠ ƒë·∫ßu v√†o.")
        return 0

    # T√°ch H√≥a ƒë∆°n v√† UNC (D·ª±a v√†o s·ªë h√≥a ƒë∆°n c√≥ ch·ª©a 'UNC' hay kh√¥ng)
    df_in['Is_UNC'] = df_in['invoice_number'].astype(str).str.contains("UNC", case=False, na=False) # type: ignore
    
    df_bills = df_in.loc[~df_in['Is_UNC']]
    df_uncs = df_in.loc[df_in['Is_UNC']]
    
    total_bills = df_bills['total_amount'].sum()
    total_uncs = df_uncs['total_amount'].sum()
    
    # Hi·ªÉn th·ªã so s√°nh
    c1, c2, c3 = st.columns(3)
    c1.metric("T·ªïng H√≥a ƒê∆°n (Chi ph√≠)", format_vnd(total_bills), help="T·ªïng gi√° tr·ªã c√°c h√≥a ƒë∆°n ƒë·∫ßu v√†o (Kh√¥ng t√≠nh UNC)")
    c2.metric("T·ªïng UNC (ƒê√£ chi)", format_vnd(total_uncs), help="T·ªïng s·ªë ti·ªÅn ƒë√£ chuy·ªÉn kho·∫£n (UNC)")
    
    diff = total_uncs - total_bills
    if diff == 0:
        c3.success("‚úÖ ƒê√£ kh·ªõp")
    elif diff > 0:
        c3.warning(f"‚ö†Ô∏è UNC d∆∞: {format_vnd(diff)}")
    else:
        c3.error(f"‚ö†Ô∏è Thi·∫øu UNC: {format_vnd(abs(diff))}")
        
    # B·∫£ng chi ti·∫øt
    t1, t2 = st.tabs(["üìÑ Danh s√°ch H√≥a ƒê∆°n", "üí∏ Danh s√°ch UNC"])
    with t1:
        st.dataframe(df_bills[['date', 'invoice_number', 'seller_name', 'total_amount', 'memo']], 
                     column_config={"total_amount": st.column_config.NumberColumn("S·ªë ti·ªÅn", format="%d")}, use_container_width=True, hide_index=True)
    with t2:
        st.dataframe(df_uncs[['date', 'invoice_number', 'seller_name', 'total_amount', 'memo']], 
                     column_config={"total_amount": st.column_config.NumberColumn("S·ªë ti·ªÅn", format="%d")}, use_container_width=True, hide_index=True)
        
    return total_bills

def render_cost_control(menu):
    if menu == "1. Nh·∫≠p H√≥a ƒê∆°n":
        # 1. Logic Nh·∫≠p UNC m·∫∑c ƒë·ªãnh l√† ƒê·∫ßu v√†o (Nh∆∞ng Type IN)
        doc_type = st.radio("üìÇ Lo·∫°i ch·ª©ng t·ª´", ["·ª¶y nhi·ªám chi ", "H√≥a ƒë∆°n"], horizontal=True, index=1 if st.session_state.current_doc_type == "H√≥a ƒë∆°n" else 0)
        
        if doc_type != st.session_state.current_doc_type:
            st.session_state.current_doc_type = doc_type
            st.session_state.pdf_data = None
            st.session_state.ready_pdf_bytes = None
            st.session_state.ready_file_name = None
            st.session_state.uploader_key += 1
            st.rerun()

        uploaded_file = st.file_uploader(f"Upload {doc_type} (PDF/·∫¢nh)", type=["pdf", "png", "jpg", "jpeg"], key=f"up_{st.session_state.uploader_key}")
        
        if uploaded_file and st.session_state.ready_file_name != uploaded_file.name:
            st.session_state.ready_pdf_bytes = None
            st.session_state.ready_file_name = uploaded_file.name
            st.session_state.pdf_data = None
            st.session_state.invoice_view_page = 0
        
        is_ready_to_analyze = False
        is_pdf_origin = False
        
        if uploaded_file:
            file_type = uploaded_file.type
            is_pdf_origin = "pdf" in file_type
            is_ready_to_analyze = True

            c_view, c_action = st.columns([1, 1])
            with c_view:
                if is_pdf_origin:
                    st.info("üìÑ File PDF G·ªëc")
                    pdf_img = None
                    total_pages = 0
                    try:
                        uploaded_file.seek(0)
                        with pdfplumber.open(uploaded_file) as pdf:
                            total_pages = len(pdf.pages)
                            if st.session_state.invoice_view_page >= total_pages: st.session_state.invoice_view_page = 0
                            pdf_img = pdf.pages[st.session_state.invoice_view_page].to_image(resolution=200).original
                    except: pass
                    
                    if total_pages > 0:
                        if total_pages > 1:
                            c_p, c_n = st.columns(2)
                            if c_p.button("‚¨Ö Tr∆∞·ªõc", key="btn_inv_prev", use_container_width=True): st.session_state.invoice_view_page = max(0, st.session_state.invoice_view_page - 1); st.rerun()
                            if c_n.button("Sau ‚û°", key="btn_inv_next", use_container_width=True): st.session_state.invoice_view_page = min(total_pages - 1, st.session_state.invoice_view_page + 1); st.rerun()
                        if pdf_img:
                            st.image(pdf_img, caption=f"Trang {st.session_state.invoice_view_page+1}/{total_pages}", width="stretch")
                else:
                    st.info("üñºÔ∏è File ·∫¢nh")
                    st.image(uploaded_file, caption="·∫¢nh g·ªëc", width="stretch")
                    
            with c_action:
                if not is_pdf_origin and st.session_state.ready_pdf_bytes is None:
                    st.info("üëâ B·∫°n ƒëang d√πng File ·∫¢nh. H·ªá th·ªëng s·∫Ω d√πng OCR ƒë·ªÉ qu√©t.")
                    if st.button("üîÑ CHUY·ªÇN ƒê·ªîI SANG PDF (ƒê·ªÇ L∆ØU TR·ªÆ)", type="secondary", width="stretch"):
                        with st.spinner("ƒêang chuy·ªÉn ƒë·ªïi..."):
                            uploaded_file.seek(0)
                            converted_bytes = convert_image_to_pdf(uploaded_file)
                            if converted_bytes:
                                st.session_state.ready_pdf_bytes = converted_bytes
                                st.success("ƒê√£ convert xong!")
                                time.sleep(0.5)
                                st.rerun()

                if is_ready_to_analyze:
                    if st.button(f"üîç QU√âT TH√îNG TIN ({doc_type})", type="primary", width="stretch"):
                        file_to_scan = None
                        is_img_input = not is_pdf_origin
                        if is_img_input: file_to_scan = uploaded_file 
                        else: file_to_scan = uploaded_file 
                        
                        if file_to_scan:
                            file_to_scan.seek(0)
                            data, msg = extract_data_smart(file_to_scan, is_img_input, doc_type)
                            if msg: st.warning(msg)
                            if data is None: st.error("L·ªói h·ªá th·ªëng khi ƒë·ªçc file.")
                            else:
                                data['file_name'] = uploaded_file.name
                                st.session_state.pdf_data = data
                                st.session_state.edit_lock = True
                                st.session_state.local_edit_count = 0
                                
                                if not HAS_OCR and is_img_input:
                                    st.error("‚ùå M√°y ch∆∞a c√†i Tesseract OCR. Kh√¥ng th·ªÉ ƒë·ªçc s·ªë t·ª´ ·∫£nh ƒë√¢u √°!")
                                
                                # --- 1. KH√îI PH·ª§C TH√îNG B√ÅO KH·ªöP TI·ªÄN ---
                                if doc_type == "H√≥a ƒë∆°n":
                                    diff = abs(data['total'] - (data['pre_tax'] + data['tax']))
                                    if diff < 10: st.success(f"‚úÖ Chu·∫©n men! T·ªïng: {format_vnd(data['total'])}")
                                    else: st.warning(f"‚ö†Ô∏è L·ªách ti·ªÅn: {format_vnd(diff)} (T·ªïng != Ti·ªÅn h√†ng + Thu·∫ø)")
                                else:
                                    st.success(f"‚úÖ ƒê√£ qu√©t UNC! S·ªë ti·ªÅn: {format_vnd(data['total'])}")
                                st.rerun()

                if st.session_state.pdf_data:
                    d = st.session_state.pdf_data
                    st.divider()
                    
                    # --- LOGIC M√É CHI PH√ç (COST CODE) - MOVED OUTSIDE FORM ---
                    # L·∫•y danh s√°ch Tour ƒëang ch·∫°y ƒë·ªÉ ch·ªçn
                    user_info_cost = st.session_state.get("user_info", {})
                    user_role_cost = user_info_cost.get('role')
                    user_name_cost = user_info_cost.get('name')
                    tour_query = "SELECT tour_name, tour_code FROM tours WHERE status='running'"
                    tour_params = []
                    if user_role_cost == 'sale' and user_name_cost:
                        tour_query += " AND sale_name=?"
                        tour_params.append(user_name_cost)
                    active_tours = run_query(tour_query, tuple(tour_params))
                    tour_choices = {f"[{t['tour_code']}] {t['tour_name']}": t['tour_code'] for t in active_tours} if active_tours else {} # type: ignore
                    tour_choices = {f"üì¶ TOUR: [{t['tour_code']}] {t['tour_name']}": t['tour_code'] for t in active_tours} if active_tours else {} # type: ignore
                    
                    # L·∫•y danh s√°ch c√°c m√£ Cost Code ƒë√£ t·ªìn t·∫°i (t·ª´ UNC ho·∫∑c H√≥a ƒë∆°n tr∆∞·ªõc ƒë√≥) ƒë·ªÉ H√≥a ƒë∆°n ch·ªçn l·∫°i
                    existing_codes_query = run_query("SELECT DISTINCT cost_code FROM invoices WHERE cost_code IS NOT NULL AND cost_code != ''")
                    existing_codes = [r['cost_code'] for r in existing_codes_query] if existing_codes_query else [] # type: ignore
                    
                    # L·∫•y danh s√°ch Booking D·ªãch V·ª• (L·ªçc theo sale n·∫øu c·∫ßn)
                    bk_query = "SELECT name, code FROM service_bookings WHERE status='active'"
                    bk_params = []
                    if user_role_cost == 'sale' and user_name_cost:
                        bk_query += " AND sale_name=?"
                        bk_params.append(user_name_cost)
                    active_bookings = run_query(bk_query, tuple(bk_params))
                    booking_choices = {f"üîñ BOOKING: [{b['code']}] {b['name']}": b['code'] for b in active_bookings} if active_bookings else {} # type: ignore

                    selected_cost_code = ""
                    new_bk_name = None
                    new_bk_code = None
                    
                    st.markdown("##### üîñ Ph√¢n lo·∫°i & Li√™n k·∫øt chi ph√≠")
                    with st.container(border=True):
                        if doc_type == "·ª¶y nhi·ªám chi ":
                            st.info("üîñ Ph√¢n lo·∫°i chi ph√≠")
                            # Logic m·ªõi: Lu√¥n y√™u c·∫ßu ch·ªçn M√£ (Tour ho·∫∑c Booking)
                            link_type = st.radio("Li√™n k·∫øt v·ªõi:", ["Tour", "Booking D·ªãch V·ª•"], horizontal=True)
                            
                            if link_type == "Tour":
                                if tour_choices:
                                    sel_t = st.selectbox("Ch·ªçn Tour:", list(tour_choices.keys()))
                                    selected_cost_code = tour_choices[sel_t]
                                else:
                                    st.warning("Ch∆∞a c√≥ Tour n√†o ƒëang ch·∫°y.")
                            else:
                                # Booking D·ªãch V·ª•
                                bk_action = st.radio("Thao t√°c:", ["Ch·ªçn Booking c√≥ s·∫µn", "‚ûï T·∫°o Booking m·ªõi"], horizontal=True, label_visibility="collapsed")
                                
                                if bk_action == "Ch·ªçn Booking c√≥ s·∫µn":
                                    if booking_choices:
                                        sel_b = st.selectbox("Ch·ªçn Booking:", list(booking_choices.keys()))
                                        selected_cost_code = booking_choices[sel_b]
                                    else:
                                        st.warning("Ch∆∞a c√≥ Tour n√†o ƒëang ch·∫°y.")
                                        st.warning("Ch∆∞a c√≥ Booking n√†o.")
                                else:
                                    # T·ª± t·∫°o m√£ Booking l·∫ª
                                    if "gen_booking_code" not in st.session_state:
                                        st.session_state.gen_booking_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                                    # T·∫°o m·ªõi Booking D·ªãch V·ª• ngay t·∫°i ƒë√¢y
                                    c_new_b1, c_new_b2 = st.columns([1, 2])
                                    if "new_bk_code" not in st.session_state:
                                        st.session_state.new_bk_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                                    
                                    c_gen1, c_gen2 = st.columns([1, 3])
                                    c_gen1.text_input("M√£ Booking:", value=st.session_state.gen_booking_code, disabled=True)
                                    if c_gen2.button("üîÑ T·∫°o m√£ kh√°c"):
                                        st.session_state.gen_booking_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                                        st.rerun()
                                    selected_cost_code = st.session_state.gen_booking_code
                                    new_bk_code = c_new_b1.text_input("M√£ Booking (T·ª± ƒë·ªông)", value=st.session_state.new_bk_code, disabled=True)
                                    new_bk_name = c_new_b2.text_input("T√™n Booking / D·ªãch v·ª•", placeholder="VD: Kh√°ch l·∫ª A, V√© m√°y bay B...")
                                
                        else: # H√≥a ƒë∆°n
                            st.info("üîó Li√™n k·∫øt chi ph√≠")
                            inv_opt = st.radio("Ngu·ªìn g·ªëc:", ["Theo m√£ UNC/Booking/Tour", "Kh√¥ng c√≥ UNC (T·ª± t·∫°o m√£)"], horizontal=True)
                            if inv_opt == "Theo m√£ UNC/Booking/Tour":
                                # G·ªôp c·∫£ m√£ Tour v√† m√£ Booking l·∫ª ƒë√£ c√≥
                                all_avail_codes = sorted(list(set(list(tour_choices.values()) + existing_codes)))
                                if all_avail_codes:
                                    selected_cost_code = st.selectbox("Ch·ªçn M√£ li√™n k·∫øt:", all_avail_codes)
                                else:
                                    st.warning("Ch∆∞a c√≥ m√£ n√†o ƒë·ªÉ li√™n k·∫øt.")
                            else:
                                if "gen_inv_code" not in st.session_state:
                                    st.session_state.gen_inv_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                                st.text_input("M√£ chi ph√≠ m·ªõi:", value=st.session_state.gen_inv_code, disabled=True)
                                selected_cost_code = st.session_state.gen_inv_code
                                st.caption("Vui l√≤ng nh·∫≠p t√™n ƒë·ªÉ t·∫°o m√£.")
                    
                    # Initialize variables to avoid unbound errors
                    txn_content = ""; seller = ""; buyer = ""

                    with st.form("inv_form"):
                        # M·∫∑c ƒë·ªãnh UNC l√† ƒê·∫ßu v√†o
                        default_idx = 0 
                        
                        # --- PH·∫¶N 1: TH√îNG TIN CHUNG ---
                        st.markdown("##### üìù Th√¥ng tin chung")
                        with st.container(border=True):
                            st.text_input("M√£ chi ph√≠ / Booking:", value=selected_cost_code, disabled=True)
                            st.divider()
                            
                            typ = st.radio("Lo·∫°i", ["ƒê·∫ßu v√†o", "ƒê·∫ßu ra"], horizontal=True, index=default_idx)
                            drive_link = st.text_input("üîó Link Drive (T√πy ch·ªçn)")
                            
                            c1, c2 = st.columns(2)
                            if doc_type == "H√≥a ƒë∆°n":
                                memo = st.text_input("G·ª£i nh·ªõ (Memo)", value=d.get('file_name',''))
                                date = st.text_input("Ng√†y", value=d['date'])
                                num = c1.text_input("S·ªë h√≥a ƒë∆°n", value=d['inv_num'])
                                sym = c2.text_input("K√Ω hi·ªáu/M·∫´u s·ªë", value=d['inv_sym'])
                            else:
                                memo = c1.text_input("G·ª£i nh·ªõ (T√™n file)", value=d.get('file_name', ''))
                                date = c2.text_input("Ng√†y chuy·ªÉn kho·∫£n", value=d['date'])
                                content_val = d.get('content', '')
                                txn_content = st.text_area("N·ªôi dung chuy·ªÉn kho·∫£n (OCR)", value=content_val, height=70)
                                num = ""; sym = ""; buyer = "" 
                        
                        # --- PH·∫¶N 2: B√äN MUA / B√ÅN ---
                        if doc_type == "H√≥a ƒë∆°n" or doc_type == "·ª¶y nhi·ªám chi ":
                            st.markdown("##### ü§ù ƒê·ªëi t∆∞·ª£ng")
                            with st.container(border=True):
                                if doc_type == "H√≥a ƒë∆°n":
                                    seller = st.text_input("B√™n B√°n", value=d['seller'])
                                    buyer = st.text_input("B√™n Mua", value=d['buyer'])
                                else:
                                    seller = st.text_input("ƒê∆°n v·ªã nh·∫≠n ti·ªÅn", value=d['seller'])
                        
                        # --- PH·∫¶N 3: T√ÄI CH√çNH ---
                        st.markdown("##### üí∞ T√†i ch√≠nh")
                        with st.container(border=True):
                            if doc_type == "H√≥a ƒë∆°n":
                                pre = st.number_input("Ti·ªÅn h√†ng", value=float(d['pre_tax']), disabled=st.session_state.edit_lock, format="%.0f")
                                tax = st.number_input("VAT", value=float(d['tax']), disabled=st.session_state.edit_lock, format="%.0f")
                                total = pre + tax
                            else:
                                st.caption("(V·ªõi UNC, ch·ªâ c·∫ßn nh·∫≠p S·ªë ti·ªÅn ƒë√£ chuy·ªÉn nha)")
                                pre = 0; tax = 0
                                total = st.number_input("S·ªë ti·ªÅn ƒë√£ chuy·ªÉn", value=float(d['total']), disabled=st.session_state.edit_lock, format="%.0f")

                            is_locked_admin = False
                            # 3. & 5. LOGIC DUY·ªÜT:
                            
                            if st.session_state.local_edit_count == 2:
                                st.markdown('<div style="background:#fff3cd; color:orange; padding:10px; border-radius:5px; margin-bottom:10px;">‚ö†Ô∏è <b>L∆∞u √Ω:</b> N·∫øu ch·ªânh s·ª≠a l·∫ßn 3 ph·∫£i g·ª≠i admin duy·ªát.</div>', unsafe_allow_html=True)
                            elif st.session_state.local_edit_count >= 3 and st.session_state.local_edit_count < 5:
                                is_locked_admin = True
                                st.markdown(f'<div style="background:#ffeef7; color:red; padding:10px; border-radius:5px; margin-bottom:10px;">üîí <b>Ch·∫ø ƒë·ªô duy·ªát:</b> B·∫°n ƒëang s·ª≠a l·∫ßn {st.session_state.local_edit_count}. C·∫ßn Admin duy·ªát.</div>', unsafe_allow_html=True)
                            elif st.session_state.local_edit_count >= 5:
                                st.error("‚õî ƒê√£ qu√° s·ªë l·∫ßn ch·ªânh s·ª≠a cho ph√©p (5 l·∫ßn).")

                            # 6. HI·ªÇN TH·ªä TI·ªÄN 1 H√ÄNG (CSS .money-box ƒë√£ x·ª≠ l√Ω)
                            st.write("") 
                            st.markdown(f'<div class="money-box">{format_vnd(total)}</div>', unsafe_allow_html=True)
                            
                            b1, b2 = st.columns(2)
                            
                            if st.session_state.local_edit_count < 5:
                                if b1.form_submit_button("‚úèÔ∏è S·ª≠a gi√°"):
                                    st.session_state.edit_lock = False
                                    st.rerun()
                            
                            if not st.session_state.edit_lock and b2.form_submit_button("‚úÖ Ch·ªët gi√°"):
                                new_pre = pre if doc_type == "H√≥a ƒë∆°n" else total
                                st.session_state.pdf_data.update({'pre_tax': new_pre, 'tax': tax, 'total': total})
                                st.session_state.edit_lock = True
                                st.session_state.local_edit_count += 1
                                st.rerun()

                        # N√∫t L∆∞u / G·ª≠i Duy·ªát
                        if is_locked_admin:
                            btn_label = "üöÄ G·ª¨I ADMIN DUY·ªÜT"
                        elif st.session_state.local_edit_count >= 5:
                            btn_label = "‚õî ƒê√É KH√ìA"
                        else:
                            btn_label = "üíæ L∆ØU CH·ª®NG T·ª™"
                        
                        if st.form_submit_button(btn_label, type="primary", width="stretch", disabled=(st.session_state.local_edit_count >= 5)):
                            if doc_type == "H√≥a ƒë∆°n" and (not date or not num): st.error("∆† k√¨a, thi·∫øu ng√†y ho·∫∑c s·ªë h√≥a ƒë∆°n r·ªìi!")
                            elif doc_type == "·ª¶y nhi·ªám chi " and not date: st.error("Thi·∫øu ng√†y chuy·ªÉn kho·∫£n r·ªìi n√®!")
                            elif not st.session_state.edit_lock: st.warning("B·∫•m 'Ch·ªët gi√°' tr∆∞·ªõc khi l∆∞u nha!")
                            else:
                                # --- CHU·∫®N B·ªä D·ªÆ LI·ªÜU ---
                                t = 'OUT' if "ƒê·∫ßu ra" in typ else 'IN'
                                save_memo = memo
                                save_num = num
                                
                                if doc_type == "·ª¶y nhi·ªám chi ":
                                    save_memo = f"[UNC] {memo} - {txn_content}"
                                    if not save_num: save_num = f"UNC-{datetime.now().strftime('%y%m%d%H%M')}"

                                # --- T·∫†O T√äN FILE ---
                                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                                clean_name = re.sub(r'[\\/*?:"<>|]', "", uploaded_file.name)
                                final_name = f"{ts}_{clean_name}"
                                if st.session_state.ready_pdf_bytes and not final_name.lower().endswith('.pdf'):
                                    final_name = os.path.splitext(final_name)[0] + ".pdf"

                                # [CODE M·ªöI] 
                                # 1. Upload file l√™n Drive (ƒê√£ t·∫Øt theo y√™u c·∫ßu - Ch·ªâ l∆∞u d·ªØ li·ªáu)
                                drive_link = ""
                                # if uploaded_file:
                                #     # X·ª≠ l√Ω file upload (n·∫øu l√† ·∫£nh ƒë√£ convert sang PDF th√¨ d√πng bytes)
                                #     if st.session_state.ready_pdf_bytes:
                                #         file_obj = io.BytesIO(st.session_state.ready_pdf_bytes)
                                #         drive_link = upload_to_drive(file_obj, final_name, mimetype='application/pdf')
                                #     else:
                                #         drive_link = upload_to_drive(uploaded_file, final_name)
                                
                                # 2. Chu·∫©n b·ªã d·ªØ li·ªáu ƒë·ªÉ l∆∞u
                                new_invoice = {
                                    'type': t, 
                                    'date': date,
                                    'invoice_number': save_num,
                                    'invoice_symbol': sym,
                                    'seller_name': seller,
                                    'buyer_name': buyer,
                                    'pre_tax_amount': pre,
                                    'tax_amount': tax,
                                    'total_amount': total,
                                    'file_name': final_name,
                                    'status': 'active',
                                    'created_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                    'memo': save_memo,
                                    'file_path': drive_link, 
                                    'cost_code': selected_cost_code,
                                    'edit_count': st.session_state.local_edit_count,
                                    'request_edit': 1 if is_locked_admin else 0
                                }
                                
                                # 3. Ghi v√†o Sheet 'invoices'
                                if add_row_to_table('invoices', new_invoice):
                                    st.success("ƒê√£ l∆∞u th√†nh c√¥ng l√™n Cloud! üéâ")
                                    
                                    # Reset state
                                    time.sleep(1)
                                    st.session_state.pdf_data = None
                                    st.session_state.uploader_key += 1
                                    st.session_state.ready_pdf_bytes = None
                                    st.session_state.ready_file_name = None
                                    st.session_state.local_edit_count = 0
                                    if "gen_booking_code" in st.session_state: del st.session_state.gen_booking_code
                                    if "gen_inv_code" in st.session_state: del st.session_state.gen_inv_code
                                    if "new_bk_code" in st.session_state: del st.session_state.new_bk_code
                                    if "pending_booking_create" in st.session_state: del st.session_state.pending_booking_create
                                    st.rerun()

        st.divider()
        # --- 4. L·ªäCH S·ª¨ NH·∫¨P LI·ªÜU (HI·ªÜN T·∫§T C·∫¢ NH∆ØNG C√ì NOTE) ---
        with st.expander("L·ªãch s·ª≠ nh·∫≠p li·ªáu", expanded=True):
            rows = run_query("SELECT id, type, invoice_number, total_amount, status, memo, request_edit, edit_count, cost_code FROM invoices ORDER BY id DESC LIMIT 20")
            if rows:
                df = pd.DataFrame([dict(r) for r in rows])
                df['Ch·ªçn'] = False 
                
                def get_status_note(row): # type: ignore
                    if row['status'] == 'deleted': # type: ignore
                        return "‚ùå ƒê√£ x√≥a"
                    note = ""
                    if row['request_edit'] == 1: # type: ignore
                        note += "‚è≥ Ch·ªù duy·ªát"
                    if row['edit_count'] > 0: # type: ignore
                        if note: note += " | "
                        note += f"‚úèÔ∏è S·ª≠a {row['edit_count']} l·∫ßn" # type: ignore
                    
                    if not note:
                        return "‚úÖ Ho·∫°t ƒë·ªông"
                    return note.strip(" | ")
                
                df['Tr·∫°ng th√°i'] = df.apply(get_status_note, axis=1)
                
                df = df[['Ch·ªçn', 'id', 'cost_code', 'type', 'invoice_number', 'total_amount', 'Tr·∫°ng th√°i', 'memo']]
                df.columns = ['Ch·ªçn', 'ID', 'M√£ Chi Ph√≠', 'Lo·∫°i', 'S·ªë Hƒê', 'T·ªïng Ti·ªÅn', 'Tr·∫°ng th√°i', 'Ghi ch√∫']
                
                df['T·ªïng Ti·ªÅn'] = df['T·ªïng Ti·ªÅn'].apply(format_vnd)

                edited_df = st.data_editor(
                    df,
                    column_config={
                        "Ch·ªçn": st.column_config.CheckboxColumn(required=True),
                        "ID": st.column_config.NumberColumn(disabled=True),
                        "M√£ Chi Ph√≠": st.column_config.TextColumn(disabled=True),
                        "Lo·∫°i": st.column_config.TextColumn(disabled=True),
                        "S·ªë Hƒê": st.column_config.TextColumn(disabled=True),
                        "T·ªïng Ti·ªÅn": st.column_config.TextColumn(disabled=True),
                        "Tr·∫°ng th√°i": st.column_config.TextColumn(disabled=True),
                        "Ghi ch√∫": st.column_config.TextColumn(disabled=True),
                    },
                    hide_index=True,
                    use_container_width=True
                )

                if st.button("üóëÔ∏è X√≥a c√°c m·ª•c ƒë√£ ch·ªçn", type="primary"):
                    selected_ids = edited_df[edited_df['Ch·ªçn']]['ID'].tolist()
                    if selected_ids:
                        for i in selected_ids:
                            run_query("UPDATE invoices SET status='deleted' WHERE id=?", (i,), commit=True)
                        st.success(f"ƒê√£ x√≥a {len(selected_ids)} h√≥a ƒë∆°n!")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.warning("B·∫°n ch∆∞a ch·ªçn m·ª•c n√†o c·∫£.")
            else:
                st.info("Ch∆∞a c√≥ h√≥a ƒë∆°n n√†o.")
    elif menu == "2. B√°o C√°o T·ªïng H·ª£p":
        st.title("üìä B√°o C√°o T√†i Ch√≠nh")

        all_financial_records = []
        with st.spinner("ƒêang t·ªïng h·ª£p d·ªØ li·ªáu t·ª´ t·∫•t c·∫£ c√°c ph√¢n h·ªá..."):
            # --- OPTIMIZED DATA FETCHING ---
            # L·ªçc booking theo sale n·∫øu c·∫ßn
            user_info_rpt = st.session_state.get("user_info", {})
            user_role_rpt = user_info_rpt.get('role')
            user_name_rpt = user_info_rpt.get('name')

            # 1. Fetch all base data in a few queries
            tour_rpt_query = "SELECT * FROM tours WHERE status != 'deleted'"
            tour_rpt_params = []
            if user_role_rpt == 'sale' and user_name_rpt:
                tour_rpt_query += " AND sale_name=?"
                tour_rpt_params.append(user_name_rpt)
            all_tours = run_query(tour_rpt_query, tuple(tour_rpt_params))
            
            bk_rpt_query = "SELECT * FROM service_bookings WHERE status != 'deleted'"
            bk_rpt_params = []
            if user_role_rpt == 'sale' and user_name_rpt:
                bk_rpt_query += " AND sale_name=?"
                bk_rpt_params.append(user_name_rpt)
            all_bookings = run_query(bk_rpt_query, tuple(bk_rpt_params))

            all_linked_invoices = run_query("SELECT cost_code, type, invoice_number, total_amount FROM invoices WHERE status='active' AND request_edit=0 AND cost_code IS NOT NULL AND cost_code != ''")
            # [NEW] Fetch all transactions for debt calculation
            all_transactions = run_query("SELECT ref_code, type, amount FROM transaction_history")

            # 2. Process data in memory using dictionaries for fast lookups
            invoice_costs_by_code = {}
            for inv in all_linked_invoices:
                code = inv['cost_code']
                if code not in invoice_costs_by_code:
                    invoice_costs_by_code[code] = {'IN_INV': 0, 'IN_UNC': 0}
                if inv['type'] == 'IN':
                    is_unc = 'UNC' in (inv.get('invoice_number') or '') # type: ignore
                    if is_unc:
                        invoice_costs_by_code[code]['IN_UNC'] += inv['total_amount'] # type: ignore
                    else:
                        invoice_costs_by_code[code]['IN_INV'] += inv['total_amount'] # type: ignore
            
            # [NEW] Process transactions to get paid amounts
            paid_amounts = {}
            if all_transactions:
                df_txns = pd.DataFrame([dict(r) for r in all_transactions])
                if not df_txns.empty:
                    df_thu = df_txns[df_txns['type'] == 'THU'].groupby('ref_code')['amount'].sum()
                    df_chi = df_txns[df_txns['type'] == 'CHI'].groupby('ref_code')['amount'].sum() # CHI means refund
                    paid_amounts = (df_thu.subtract(df_chi, fill_value=0)).to_dict()

            # --- Process Tours ---
            if all_tours:
                for tour_row in all_tours:
                    tour = dict(tour_row)
                    # [NEW] Add status to record
                    tour_status = tour.get('status', 'running')
                    revenue, cost = get_tour_financials(tour['id'], tour)
                    if revenue > 0: all_financial_records.append({'date_str': tour['start_date'], 'name': tour['tour_name'], 'code': tour['tour_code'], 'category': 'Tour', 'type': 'thu', 'amount': revenue, 'status': tour_status}) # type: ignore
                    if cost > 0: all_financial_records.append({'date_str': tour['start_date'], 'name': tour['tour_name'], 'code': tour['tour_code'], 'category': 'Tour', 'type': 'chi', 'amount': cost, 'status': tour_status}) # type: ignore

            # --- Process Service Bookings ---
            if all_bookings:
                for booking_row in all_bookings:
                    booking = dict(booking_row)
                    
                    # [FIX] Chuy·ªÉn ƒë·ªïi ƒë·ªãnh d·∫°ng ng√†y YYYY-MM-DD sang DD/MM/YYYY ƒë·ªÉ ƒë·ªìng b·ªô
                    try:
                        booking_date_obj = datetime.strptime(str(booking['created_at']).split(" ")[0], '%Y-%m-%d')
                        booking_date_str = booking_date_obj.strftime('%d/%m/%Y')
                    except:
                        booking_date_str = booking['created_at']
                    # [NEW] Add status to record
                    booking_status = booking.get('status', 'active')

                    if booking.get('selling_price', 0) > 0:
                        all_financial_records.append({'date_str': booking_date_str, 'name': booking['name'], 'code': booking['code'], 'category': 'Booking D·ªãch V·ª•', 'type': 'thu', 'amount': booking['selling_price'], 'status': booking_status}) # type: ignore
                    
                    # [FIX] Ch·ªâ t√≠nh chi ph√≠ t·ª´ h√≥a ƒë∆°n (IN_INV), kh√¥ng t√≠nh UNC ƒë·ªÉ tr√°nh double-count.
                    # UNC l√† thanh to√°n cho chi ph√≠, kh√¥ng ph·∫£i b·∫£n th√¢n chi ph√≠.
                    total_cost_booking = invoice_costs_by_code.get(booking['code'], {}).get('IN_INV', 0)
                    if total_cost_booking == 0 and booking.get('net_price', 0) > 0:
                        total_cost_booking = booking['net_price'] # type: ignore
                    if total_cost_booking > 0:
                        all_financial_records.append({'date_str': booking_date_str, 'name': booking['name'], 'code': booking['code'], 'category': 'Booking D·ªãch V·ª•', 'type': 'chi', 'amount': total_cost_booking, 'status': booking_status}) # type: ignore

            # --- Process old Projects & Unlinked Invoices (These queries are already efficient) ---
            project_invoices = run_query("SELECT p.project_name, i.type, i.total_amount, i.date, p.id as project_id FROM projects p JOIN project_links l ON p.id = l.project_id JOIN invoices i ON l.invoice_id = i.id WHERE i.status = 'active' AND i.request_edit = 0")
            if project_invoices:
                for inv in project_invoices:
                    all_financial_records.append({'date_str': inv['date'], 'name': inv['project_name'], 'code': f"PROJ_{inv['project_id']}", 'category': 'D·ª± √°n (c≈©)', 'type': 'thu' if inv['type'] == 'OUT' else 'chi', 'amount': inv['total_amount'], 'status': 'N/A'}) # type: ignore

            unlinked_invoices = run_query("SELECT * FROM invoices i WHERE i.status = 'active' AND i.request_edit = 0 AND (i.cost_code IS NULL OR i.cost_code = '') AND NOT EXISTS (SELECT 1 FROM project_links pl WHERE pl.invoice_id = i.id)")
            if unlinked_invoices:
                for inv in unlinked_invoices:
                    all_financial_records.append({'date_str': inv['date'], 'name': inv['memo'] or inv['seller_name'] or 'Chi ph√≠ chung', 'code': f"INV_{inv['id']}", 'category': 'Chi ph√≠ chung', 'type': 'thu' if inv['type'] == 'OUT' else 'chi', 'amount': inv['total_amount'], 'status': 'N/A'}) # type: ignore

        if not all_financial_records:
            st.info("Ch∆∞a c√≥ d·ªØ li·ªáu t√†i ch√≠nh ƒë·ªÉ b√°o c√°o.")
        else:
            df = pd.DataFrame(all_financial_records)
            df['date'] = pd.to_datetime(df['date_str'], errors='coerce', dayfirst=True)
            df['status'] = df['status'].fillna('N/A') # ƒê·∫£m b·∫£o c·ªôt status kh√¥ng c√≥ gi√° tr·ªã null
            df = df.dropna(subset=['date'])

            # Explicitly create a DatetimeIndex to help Pylance with type inference
            dt_index = pd.DatetimeIndex(df['date'])
            df['year'] = dt_index.year
            df['quarter'] = dt_index.quarter
            df['month_year'] = dt_index.to_period('M').astype(str)
            df['quarter_year'] = df.apply(lambda row: f"Q{row['quarter']}/{row['year']}", axis=1)

            st.markdown("####  L·ªçc b√°o c√°o")
            c1, c2, c3 = st.columns(3)
            filter_type = c1.selectbox("L·ªçc theo th·ªùi gian:", ["Th√°ng", "Qu√Ω", "NƒÉm"])
            
            options = []
            period_col = ''
            if filter_type == "Th√°ng":
                options = sorted(df['month_year'].unique(), reverse=True)
                period_col = 'month_year'
            elif filter_type == "Qu√Ω":
                options = sorted(df['quarter_year'].unique(), reverse=True)
                period_col = 'quarter_year'
            elif filter_type == "NƒÉm":
                options = sorted(df['year'].unique(), reverse=True)
                period_col = 'year'
                
            selected_period = c2.selectbox(f"Ch·ªçn k·ª≥:", ["T·∫•t c·∫£"] + options)

            # [NEW] Th√™m b·ªô l·ªçc tr·∫°ng th√°i
            status_map = {
                "T·∫•t c·∫£ tr·∫°ng th√°i": None,
                "ƒêang ch·∫°y / Ho·∫°t ƒë·ªông": ['running', 'active'],
                "ƒê√£ ho√†n th√†nh": ['completed']
            }
            selected_status_label = c3.selectbox("L·ªçc theo tr·∫°ng th√°i:", list(status_map.keys()))
            selected_statuses = status_map[selected_status_label]

            # √Åp d·ª•ng c√°c b·ªô l·ªçc
            df_filtered = df.copy()
            if selected_period != "T·∫•t c·∫£":
                df_filtered = df_filtered[df_filtered[period_col] == selected_period]
            
            if selected_statuses:
                # Ch·ªâ l·ªçc c√°c m·ª•c c√≥ tr·∫°ng th√°i (Tour/Booking), gi·ªØ l·∫°i c√°c m·ª•c kh√°c (Chi ph√≠ chung...)
                mask = df_filtered['status'].isin(selected_statuses) | (df_filtered['status'] == 'N/A')
                df_filtered = df_filtered[mask]

            if not df_filtered.empty:
                agg = df_filtered.pivot_table(index=['category', 'name', 'code'], columns='type', values='amount', aggfunc='sum').fillna(0)
                agg = agg.reset_index()
                
                if 'thu' not in agg.columns: agg['thu'] = 0
                if 'chi' not in agg.columns: agg['chi'] = 0
                agg['l·ª£i nhu·∫≠n'] = agg['thu'] - agg['chi']
                
                total_thu = agg['thu'].sum()
                total_chi = agg['chi'].sum()
                total_loi_nhuan = agg['l·ª£i nhu·∫≠n'].sum()
                
                m1, m2, m3 = st.columns(3)
                m1.metric(f"T·ªïng Thu ({selected_period})", format_vnd(total_thu))
                m2.metric(f"T·ªïng Chi ({selected_period})", format_vnd(total_chi))
                m3.metric(f"L·ª£i Nhu·∫≠n ({selected_period})", format_vnd(total_loi_nhuan), delta=format_vnd(total_loi_nhuan) if total_loi_nhuan != 0 else None)

                st.divider()
                
                st.markdown("#### Chi ti·∫øt theo h·∫°ng m·ª•c")
                # Sort categories by total profit
                category_profit = agg.groupby('category')['l·ª£i nhu·∫≠n'].sum().sort_values(ascending=False)
                
                for category in category_profit.index:
                    group = agg[agg['category'] == category]
                    with st.expander(f"üìÇ {category} (L·ª£i nhu·∫≠n: {format_vnd(group['l·ª£i nhu·∫≠n'].sum())})", expanded=True):
                        group = group.sort_values('l·ª£i nhu·∫≠n', ascending=False)
                        for _, r in group.iterrows():
                            # --- [NEW] Debt calculation & display ---
                            debt_html = ""
                            # Only calculate for Tours and Bookings which have revenue
                            if r['category'] in ['Tour', 'Booking D·ªãch V·ª•'] and r['thu'] > 0:
                                code = r['code']
                                revenue = r['thu']
                                paid = paid_amounts.get(code, 0.0)
                                remaining = revenue - paid
                                
                                if remaining <= 0.1: # Use a small threshold for float comparison
                                    debt_html = f'''<div style="margin-top: 8px; font-size: 0.9em; text-align: right;">
                                        <span style="color: #2e7d32; font-weight: bold;">‚úÖ ƒê√£ thanh to√°n ƒë·ªß</span>
                                    </div>'''
                                else:
                                    debt_html = f'''<div style="margin-top: 8px; font-size: 0.9em; text-align: right;">
                                        <span style="color: #c62828; font-weight: bold;">C√≤n ph·∫£i thu: {format_vnd(remaining)}</span>
                                    </div>'''
                            # --- End of new code ---

                            st.markdown(f"""
                            <div class="report-card" style="padding: 15px; margin-bottom: 10px; border-left: 5px solid {'#28a745' if r['l·ª£i nhu·∫≠n']>=0 else '#e53935'};">
                                <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom: 8px;">
                                    <h5 style="margin:0; padding-right: 10px;">{r['name']}</h5>
                                    <span style="font-size: 0.8em; color: #6c757d; background-color: #f1f3f5; padding: 2px 6px; border-radius: 5px; white-space: nowrap;">CODE: {r['code']}</span>
                                </div>
                                <div style="display:flex; justify-content:space-between; font-size: 0.95em; border-bottom: 1px solid #f1f3f5; padding-bottom: 8px;">
                                    <span>Thu: <b>{format_vnd(r['thu'])}</b></span>
                                    <span>Chi: <b>{format_vnd(r['chi'])}</b></span>
                                    <span style="font-weight: bold; color:{'#1B5E20' if r['l·ª£i nhu·∫≠n']>=0 else '#c62828'}">L√£i: {format_vnd(r['l·ª£i nhu·∫≠n'])}</span>
                                </div>
                                {debt_html}
                            </div>
                            """, unsafe_allow_html=True)
            else:
                st.info(f"Kh√¥ng c√≥ d·ªØ li·ªáu cho k·ª≥ b√°o c√°o '{selected_period}'.")

def render_debt_management():
    st.title("üí≥ Qu·∫£n L√Ω C√¥ng N·ª£")
    st.caption("Theo d√µi v√† t·ªïng h·ª£p c√°c kho·∫£n ph·∫£i thu t·ª´ kh√°ch h√†ng.")

    tab_lookup, tab_summary = st.tabs(["Tra c·ª©u theo M√£", "T·ªïng h·ª£p C√¥ng n·ª£"])

    with tab_lookup:
        st.subheader("Tra c·ª©u c√¥ng n·ª£ theo M√£ Tour / Booking")
        
        # --- L·∫§Y D·ªÆ LI·ªÜU ƒê·ªÇ T√åM KI·∫æM (CH·ªà HI·ªÜN C√ÅC M√É C√íN N·ª¢) ---
        with st.spinner("ƒêang t·∫£i danh s√°ch c√≤n n·ª£..."):
            # 1. L·∫•y t·∫•t c·∫£ giao d·ªãch v√† t√≠nh to√°n s·ªë ti·ªÅn ƒë√£ tr·∫£ cho m·ªói m√£
            all_txns_cn = run_query("SELECT ref_code, type, amount FROM transaction_history")
            paid_amounts_cn = {}
            if all_txns_cn:
                df_txns_cn = pd.DataFrame([dict(r) for r in all_txns_cn])
                if not df_txns_cn.empty:
                    df_thu_cn = df_txns_cn[df_txns_cn['type'] == 'THU'].groupby('ref_code')['amount'].sum()
                    df_chi_cn = df_txns_cn[df_txns_cn['type'] == 'CHI'].groupby('ref_code')['amount'].sum()
                    paid_amounts_cn = (df_thu_cn.subtract(df_chi_cn, fill_value=0)).to_dict()

            # 2. L·∫•y t·∫•t c·∫£ tour v√† booking (l·ªçc theo sale n·∫øu c·∫ßn)
            user_info_cn = st.session_state.get("user_info", {})
            user_role_cn = user_info_cn.get('role')
            user_name_cn = user_info_cn.get('name')

            # [FIX] L·∫•y t·∫•t c·∫£ tour/booking ch∆∞a b·ªã x√≥a (bao g·ªìm c·∫£ m·ª•c ƒë√£ ho√†n th√†nh) ƒë·ªÉ ki·ªÉm tra c√¥ng n·ª£
            tour_cn_query = "SELECT * FROM tours WHERE COALESCE(status, 'running') NOT IN ('deleted')"
            tour_cn_params = []
            if user_role_cn == 'sale' and user_name_cn:
                tour_cn_query += " AND sale_name=?"
                tour_cn_params.append(user_name_cn)
            all_tours_cn = run_query(tour_cn_query, tuple(tour_cn_params))

            bk_cn_query = "SELECT * FROM service_bookings WHERE COALESCE(status, 'active') NOT IN ('deleted')"
            bk_cn_params = []
            if user_role_cn == 'sale' and user_name_cn:
                bk_cn_query += " AND sale_name=?"
                bk_cn_params.append(user_name_cn)
            all_bookings_cn = run_query(bk_cn_query, tuple(bk_cn_params))

            search_options = {"": "-- Ch·ªçn m√£ ƒë·ªÉ theo d√µi --"}

            # 3. X·ª≠ l√Ω Tours: Ch·ªâ th√™m v√†o danh s√°ch n·∫øu ch∆∞a thu ƒë·ªß
            if all_tours_cn:
                for t_row in all_tours_cn:
                    tour = dict(t_row)
                    # T√≠nh gi√° tr·ªã h·ª£p ƒë·ªìng
                    final_price = float(tour.get('final_tour_price', 0) or 0)
                    child_price = float(tour.get('child_price', 0) or 0)
                    final_qty = float(tour.get('final_qty', 0) or 0)
                    child_qty = float(tour.get('child_qty', 0) or 0)
                    if final_qty == 0: final_qty = float(tour.get('guest_count', 1))
                    contract_value = (final_price * final_qty) + (child_price * child_qty)
                    
                    paid = paid_amounts_cn.get(tour['tour_code'], 0.0)
                    
                    if contract_value > 0 and contract_value - paid > 0.1:
                        search_options[f"üì¶ TOUR: [{tour['tour_code']}] {tour['tour_name']}"] = tour['tour_code']

            # 4. X·ª≠ l√Ω Bookings: Ch·ªâ th√™m v√†o danh s√°ch n·∫øu ch∆∞a thu ƒë·ªß
            if all_bookings_cn:
                for b_row in all_bookings_cn:
                    booking = dict(b_row)
                    contract_value = float(booking.get('selling_price', 0) or 0)
                    paid = paid_amounts_cn.get(booking['code'], 0.0)
                    if contract_value > 0 and contract_value - paid > 0.1:
                        search_options[f"üîñ BOOKING: [{booking['code']}] {booking['name']}"] = booking['code']

        # --- GIAO DI·ªÜN CH√çNH ---
        col1, col2 = st.columns([1, 2])
        remaining = 0.0

        with col1:
            st.markdown("#### üîç Ch·ªçn ƒë·ªëi t∆∞·ª£ng")
            selected_label = st.selectbox("T√¨m theo M√£ Tour / Booking (ch·ªâ hi·ªán m√£ c√≤n n·ª£):", list(search_options.keys()), label_visibility="collapsed")
            selected_code = search_options.get(selected_label)

            if selected_code:
                # Reset tr·∫°ng th√°i phi·∫øu v·ª´a t·∫°o n·∫øu chuy·ªÉn m√£ kh√°c
                if "last_voucher_code" not in st.session_state or st.session_state.last_voucher_code != selected_code:
                    if "last_voucher" in st.session_state: del st.session_state.last_voucher
                    if "last_voucher_pdf" in st.session_state: del st.session_state.last_voucher_pdf
                    st.session_state.last_voucher_code = selected_code

                st.markdown("---")
                st.markdown("#### üìä T·ªïng quan c√¥ng n·ª£")

                contract_value = 0.0
                # X√°c ƒë·ªãnh gi√° tr·ªã h·ª£p ƒë·ªìng
                if "TOUR" in selected_label:
                    tour_info = run_query("SELECT * FROM tours WHERE tour_code=?", (selected_code,), fetch_one=True)
                    if tour_info:
                        t_dict = dict(tour_info)
                        final_price = float(t_dict.get('final_tour_price', 0) or 0)
                        child_price = float(t_dict.get('child_price', 0) or 0)
                        final_qty = float(t_dict.get('final_qty', 0) or 0)
                        child_qty = float(t_dict.get('child_qty', 0) or 0)
                        if final_qty == 0: final_qty = float(t_dict.get('guest_count', 1))
                        contract_value = (final_price * final_qty) + (child_price * child_qty)
                elif "BOOKING" in selected_label:
                    booking_info = run_query("SELECT selling_price FROM service_bookings WHERE code=?", (selected_code,), fetch_one=True)
                    if booking_info:
                        contract_value = float(booking_info['selling_price'] or 0)

                # L·∫•y t·ªïng ƒë√£ thu
                paid_data = run_query("SELECT SUM(amount) as total FROM transaction_history WHERE ref_code=? AND type='THU'", (selected_code,), fetch_one=True)
                total_paid = paid_data['total'] if paid_data and paid_data['total'] else 0.0

                # L·∫•y t·ªïng ƒë√£ chi (ho√†n ti·ªÅn)
                refund_data = run_query("SELECT SUM(amount) as total FROM transaction_history WHERE ref_code=? AND type='CHI'", (selected_code,), fetch_one=True)
                total_refund = refund_data['total'] if refund_data and refund_data['total'] else 0.0
                
                actual_paid = total_paid - total_refund
                
                remaining = contract_value - actual_paid

                with st.container(border=True):
                    st.metric("Gi√° tr·ªã H·ª£p ƒë·ªìng/Booking", format_vnd(contract_value))
                    st.metric("ƒê√£ thu th·ª±c t·∫ø", format_vnd(actual_paid))
                    delta_color = "inverse" if remaining > 0 else "off"
                    st.metric("C√≤n ph·∫£i thu", format_vnd(remaining), delta=f"-{format_vnd(remaining)}" if remaining > 0 else "‚úÖ ƒê√£ thu ƒë·ªß", delta_color=delta_color)

        with col2:
            if selected_code:
                tab_add, tab_history = st.tabs(["‚ûï T·∫°o Phi·∫øu Thu/Chi", "üìú L·ªãch s·ª≠ giao d·ªãch"])

                with tab_add:
                    st.markdown("##### T·∫°o phi·∫øu m·ªõi")
                    
                    # Hi·ªÉn th·ªã n√∫t t·∫£i phi·∫øu v·ª´a t·∫°o (n·∫øu c√≥)
                    if "last_voucher" in st.session_state and st.session_state.last_voucher.get('ref_code') == selected_code:
                        lv = st.session_state.last_voucher
                        st.success("‚úÖ ƒê√£ l∆∞u phi·∫øu th√†nh c√¥ng!")
                        
                        # S·ª≠ d·ª•ng cache PDF n·∫øu c√≥ ƒë·ªÉ tr√°nh t·∫°o l·∫°i li√™n t·ª•c
                        if "last_voucher_pdf" in st.session_state:
                            pdf_data = st.session_state.last_voucher_pdf
                        else:
                            pdf_data = create_voucher_pdf(lv)
                            st.session_state.last_voucher_pdf = pdf_data

                        st.download_button(
                            label=f"üì• T·∫£i Phi·∫øu {lv['type']} (PDF)",
                            data=pdf_data,
                            file_name=f"Phieu_{lv['type']}_{lv['date'].replace('/','')}.pdf",
                            mime="application/pdf",
                            type="primary"
                        )
                        st.divider()

                    # Form nh·∫≠p li·ªáu (Lu√¥n hi·ªÉn th·ªã)
                    k_amt = f"txn_amt_{selected_code}"
                    k_note = f"txn_note_{selected_code}"
                    
                    if k_amt not in st.session_state:
                        if remaining >= 1:
                            st.session_state[k_amt] = "{:,.0f}".format(remaining).replace(",", ".") + " VND"
                        else:
                            st.session_state[k_amt] = ""
                    
                    def fmt_txn_amt_dynamic(key_name):
                        if key_name in st.session_state:
                            val = st.session_state[key_name]
                            try:
                                clean_val = val.replace('.', '').replace(',', '').replace(' VND', '').strip()
                                if clean_val:
                                    v_float = float(clean_val)
                                    st.session_state[key_name] = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                            except: pass

                    c1, c2 = st.columns(2)
                    txn_type = c1.radio("Lo·∫°i phi·∫øu", ["THU", "CHI (Ho√†n ti·ªÅn)"], horizontal=True, key=f"txn_type_{selected_code}")
                    
                    txn_amount_input = c2.text_input("S·ªë ti·ªÅn", key=k_amt, on_change=fmt_txn_amt_dynamic, args=(k_amt,), help="Nh·∫≠p s·ªë ti·ªÅn (VD: 1.000.000)")
                    
                    try:
                        txn_amount = float(txn_amount_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                    except: txn_amount = 0.0
                    
                    c3, c4 = st.columns(2)
                    txn_method = c3.selectbox("H√¨nh th·ª©c", ["Chuy·ªÉn kho·∫£n", "Ti·ªÅn m·∫∑t"], key=f"txn_method_{selected_code}")
                    txn_note = c4.text_input("N·ªôi dung", placeholder="VD: C·ªçc l·∫ßn 1, Thanh to√°n...", key=k_note)
                    
                    btn_label = "üíæ T·∫°o Phi·∫øu Thu" if txn_type == "THU" else "üíæ T·∫°o Phi·∫øu Chi"
                    if st.button(btn_label, type="primary", use_container_width=True, key=f"btn_save_txn_{selected_code}"):
                        if txn_amount > 0 and txn_note:
                            now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                            run_query(
                                "INSERT INTO transaction_history (ref_code, type, amount, payment_method, note, created_at) VALUES (?, ?, ?, ?, ?, ?)",
                                (selected_code, txn_type, txn_amount, txn_method, txn_note, now_str),
                                commit=True
                            )
                            
                            # L∆∞u th√¥ng tin ƒë·ªÉ t·∫°o PDF
                            v_data = {
                                'ref_code': selected_code,
                                'type': txn_type,
                                'amount': txn_amount,
                                'method': txn_method,
                                'note': txn_note,
                                'date': datetime.now().strftime("%d/%m/%Y"),
                                'issuer': st.session_state.user_info.get('name', '')
                            }
                            st.session_state.last_voucher = v_data
                            
                            # T·∫°o PDF ngay v√† cache l·∫°i
                            pdf_bytes = create_voucher_pdf(v_data)
                            st.session_state.last_voucher_pdf = pdf_bytes
                            
                            if k_amt in st.session_state: del st.session_state[k_amt]
                            if k_note in st.session_state: del st.session_state[k_note]
                            st.rerun()
                        else:
                            st.warning("Vui l√≤ng nh·∫≠p s·ªë ti·ªÅn v√† n·ªôi dung.")

                with tab_history:
                    st.markdown("##### L·ªãch s·ª≠ c√°c l·∫ßn thanh to√°n")
                    history = run_query("SELECT * FROM transaction_history WHERE ref_code=? ORDER BY id DESC", (selected_code,))
                    
                    if history:
                        # Hi·ªÉn th·ªã d·∫°ng b·∫£ng cho g·ªçn
                        df_hist = pd.DataFrame([dict(r) for r in history])
                        
                        # Format d·ªØ li·ªáu hi·ªÉn th·ªã
                        df_display = df_hist.copy()
                        df_display['created_at'] = pd.to_datetime(df_display['created_at'], errors='coerce').dt.strftime('%d/%m/%Y')
                        df_display['amount'] = df_display['amount'].apply(lambda x: format_vnd(x))
                        
                        df_display = df_display.rename(columns={
                            'created_at': 'Ng√†y',
                            'type': 'Lo·∫°i',
                            'amount': 'S·ªë ti·ªÅn',
                            'payment_method': 'H√¨nh th·ª©c',
                            'note': 'N·ªôi dung',
                            'id': 'ID'
                        })
                        
                        st.dataframe(
                            df_display[['ID', 'Ng√†y', 'Lo·∫°i', 'S·ªë ti·ªÅn', 'H√¨nh th·ª©c', 'N·ªôi dung']],
                            use_container_width=True,
                            hide_index=True
                        )
                        
                        st.divider()
                        st.markdown("###### üõ†Ô∏è Thao t√°c (T·∫£i phi·∫øu / X√≥a)")
                        
                        # T·∫°o danh s√°ch l·ª±a ch·ªçn
                        txn_options = {}
                        for r in history:
                            try: d_lbl = datetime.strptime(r['created_at'], "%Y-%m-%d %H:%M:%S").strftime("%d/%m/%Y")
                            except: d_lbl = r['created_at']
                            label = f"#{r['id']} | {d_lbl} | {r['type']} | {format_vnd(r['amount'])}"
                            txn_options[label] = r

                        selected_txn_label = st.selectbox("Ch·ªçn giao d·ªãch:", ["-- Ch·ªçn giao d·ªãch --"] + list(txn_options.keys()))
                        
                        if selected_txn_label and selected_txn_label != "-- Ch·ªçn giao d·ªãch --":
                            txn = txn_options[selected_txn_label]
                            
                            # Ch·ªâ t·∫°o PDF khi ƒë√£ ch·ªçn (T·ªëi ∆∞u hi·ªáu nƒÉng)
                            try: d_str = datetime.strptime(txn['created_at'], "%Y-%m-%d %H:%M:%S").strftime("%d/%m/%Y")
                            except: d_str = txn['created_at']
                            
                            v_data = {
                                'ref_code': selected_code,
                                'type': txn['type'],
                                'amount': txn['amount'],
                                'method': txn['payment_method'],
                                'note': txn['note'],
                                'date': d_str,
                                'issuer': st.session_state.user_info.get('name', '')
                            }
                            pdf_bytes = create_voucher_pdf(v_data)
                            
                            c_dl, c_del = st.columns([1, 1])
                            with c_dl:
                                st.download_button(
                                    label="üì• T·∫£i Phi·∫øu (PDF)",
                                    data=pdf_bytes,
                                    file_name=f"Phieu_{txn['type']}_{txn['id']}.pdf",
                                    mime="application/pdf",
                                    key=f"dl_hist_btn_{txn['id']}",
                                    use_container_width=True,
                                    type="primary"
                                )
                            
                            with c_del:
                                if st.button("üóëÔ∏è X√≥a giao d·ªãch n√†y", key=f"del_hist_btn_{txn['id']}", use_container_width=True):
                                    run_query("DELETE FROM transaction_history WHERE id=?", (txn['id'],), commit=True)
                                    st.success("ƒê√£ x√≥a!")
                                    time.sleep(0.5)
                                    st.rerun()
                    else:
                        st.info("Ch∆∞a c√≥ l·ªãch s·ª≠ giao d·ªãch cho m√£ n√†y.")
            else:
                st.info("üëÜ Vui l√≤ng ch·ªçn m·ªôt M√£ Tour ho·∫∑c M√£ Booking ƒë·ªÉ xem c√¥ng n·ª£.")

    with tab_summary:
        st.subheader("T·ªïng h·ª£p c√°c kho·∫£n ph·∫£i thu")
        with st.spinner("ƒêang t√≠nh to√°n c√¥ng n·ª£..."):
            # 1. L·∫•y t·∫•t c·∫£ giao d·ªãch v√† t√≠nh to√°n s·ªë ti·ªÅn ƒë√£ tr·∫£ cho m·ªói m√£
            all_txns = run_query("SELECT ref_code, type, amount FROM transaction_history")
            paid_amounts = {}
            if all_txns:
                df_txns = pd.DataFrame([dict(r) for r in all_txns])
                if not df_txns.empty:
                    df_thu = df_txns[df_txns['type'] == 'THU'].groupby('ref_code')['amount'].sum()
                    df_chi = df_txns[df_txns['type'] == 'CHI'].groupby('ref_code')['amount'].sum()
                    paid_amounts = (df_thu.subtract(df_chi, fill_value=0)).to_dict()

            debt_records = []

            # 2. L·∫•y t·∫•t c·∫£ tour ƒëang ho·∫°t ƒë·ªông v√† t√≠nh c√¥ng n·ª£
            user_info_debt = st.session_state.get("user_info", {})
            user_role_debt = user_info_debt.get('role')
            user_name_debt = user_info_debt.get('name')
            # [FIX] L·∫•y t·∫•t c·∫£ tour ch∆∞a b·ªã x√≥a (bao g·ªìm c·∫£ tour ƒë√£ ho√†n th√†nh) ƒë·ªÉ t·ªïng h·ª£p c√¥ng n·ª£
            tour_debt_query = "SELECT * FROM tours WHERE COALESCE(status, 'running') NOT IN ('deleted')"
            tour_debt_params = []
            if user_role_debt == 'sale' and user_name_debt:
                tour_debt_query += " AND sale_name=?"
                tour_debt_params.append(user_name_debt)
            active_tours = run_query(tour_debt_query, tuple(tour_debt_params))
            if active_tours:
                for tour_row in active_tours:
                    tour = dict(tour_row)
                    final_price = float(tour.get('final_tour_price', 0) or 0)
                    child_price = float(tour.get('child_price', 0) or 0)
                    final_qty = float(tour.get('final_qty', 0) or 0)
                    child_qty = float(tour.get('child_qty', 0) or 0)
                    if final_qty == 0: final_qty = float(tour.get('guest_count', 1))
                    contract_value = (final_price * final_qty) + (child_price * child_qty)
 
                    if contract_value > 0:
                        paid = paid_amounts.get(tour['tour_code'], 0.0)
                        remaining = contract_value - paid
                        if remaining > 0.1:
                            debt_records.append({'customer_name': tour.get('customer_name', 'N/A'), 'ref_name': tour['tour_name'], 'ref_code': tour['tour_code'], 'type': 'Tour', 'contract_value': contract_value, 'paid': paid, 'remaining': remaining})
 
            # 3. L·∫•y t·∫•t c·∫£ booking l·∫ª ƒëang ho·∫°t ƒë·ªông v√† t√≠nh c√¥ng n·ª£
            # [FIX] L·∫•y t·∫•t c·∫£ booking ch∆∞a b·ªã x√≥a (bao g·ªìm c·∫£ booking ƒë√£ ho√†n th√†nh) ƒë·ªÉ t·ªïng h·ª£p c√¥ng n·ª£
            bk_debt_query = "SELECT * FROM service_bookings WHERE COALESCE(status, 'active') NOT IN ('deleted')"
            bk_debt_params = []
            if user_role_debt == 'sale' and user_name_debt:
                bk_debt_query += " AND sale_name=?"
                bk_debt_params.append(user_name_debt)
            active_bookings = run_query(bk_debt_query, tuple(bk_debt_params))
            if active_bookings:
                for booking_row in active_bookings:
                    booking = dict(booking_row)
                    contract_value = float(booking.get('selling_price', 0) or 0)
 
                    if contract_value > 0:
                        paid = paid_amounts.get(booking['code'], 0.0)
                        remaining = contract_value - paid
                        if remaining > 0.1:
                            customer_info = booking.get('customer_info', 'N/A')
                            customer_name = customer_info.split(' - ')[0] if ' - ' in customer_info else customer_info
                            debt_records.append({'customer_name': customer_name, 'ref_name': booking['name'], 'ref_code': booking['code'], 'type': 'Booking', 'contract_value': contract_value, 'paid': paid, 'remaining': remaining})
 
            # 4. Hi·ªÉn th·ªã k·∫øt qu·∫£
            if not debt_records:
                st.success("üéâ Kh√¥ng c√≥ c√¥ng n·ª£ n√†o c·∫ßn thu.")
            else:
                df_debt = pd.DataFrame(debt_records)
                total_debt = df_debt['remaining'].sum()
                
                st.metric("T·ªîNG S·ªê TI·ªÄN C·∫¶N THU", format_vnd(total_debt))
                
                st.divider()
                st.markdown("#### Danh s√°ch kh√°ch h√†ng ƒëang n·ª£")
                
                customer_debt = df_debt.groupby('customer_name')['remaining'].sum().reset_index().sort_values('remaining', ascending=False)
                customer_debt.columns = ['Kh√°ch h√†ng', 'T·ªïng n·ª£']
                
                st.dataframe(customer_debt, column_config={"T·ªïng n·ª£": st.column_config.NumberColumn(format="%d VND")}, use_container_width=True, hide_index=True)
                
                st.divider()
                st.markdown("#### Chi ti·∫øt c√°c kho·∫£n n·ª£")
                st.dataframe(
                    df_debt.sort_values('remaining', ascending=False),
                    column_config={ 'customer_name': 'Kh√°ch h√†ng', 'ref_name': 'T√™n Tour/Booking', 'ref_code': 'M√£', 'type': 'Lo·∫°i', 'contract_value': st.column_config.NumberColumn("Gi√° tr·ªã Hƒê", format="%d VND"), 'paid': st.column_config.NumberColumn("ƒê√£ thu", format="%d VND"), 'remaining': st.column_config.NumberColumn("C√≤n l·∫°i", format="%d VND"), },
                    use_container_width=True, hide_index=True
                )

                # --- T√çNH NƒÇNG XU·∫§T EXCEL C√îNG N·ª¢ ---
                st.write("")
                if "debt_xls_data" not in st.session_state: st.session_state.debt_xls_data = None
                
                if st.button("üìä T·∫°o file Excel b√°o c√°o"):
                    buffer_debt = io.BytesIO()
                    try:
                        with pd.ExcelWriter(buffer_debt, engine='xlsxwriter') as writer:
                            workbook: Any = writer.book
                            worksheet = workbook.add_worksheet('CongNo')
                            
                            # Formats
                            fmt_title = workbook.add_format({'bold': True, 'font_size': 16, 'align': 'center', 'valign': 'vcenter', 'font_color': '#B71C1C', 'font_name': 'Times New Roman'})
                            fmt_header = workbook.add_format({'bold': True, 'bg_color': '#FFEBEE', 'border': 1, 'align': 'center', 'valign': 'vcenter', 'font_color': '#B71C1C', 'text_wrap': True, 'font_name': 'Times New Roman'})
                            fmt_text = workbook.add_format({'border': 1, 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                            fmt_money = workbook.add_format({'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'font_name': 'Times New Roman'})
                            fmt_comp = workbook.add_format({'bold': True, 'font_size': 12, 'font_color': '#1B5E20', 'font_name': 'Times New Roman'})
                            
                            # Company Info
                            comp_data = get_company_data()
                            worksheet.write('A1', comp_data['name'], fmt_comp)
                            worksheet.write('A2', f"ƒêC: {comp_data['address']}")
                            worksheet.write('A3', f"MST: {comp_data['phone']}")
                            
                            # Title
                            worksheet.merge_range('A5:G5', "B√ÅO C√ÅO C√îNG N·ª¢ KH√ÅCH H√ÄNG", fmt_title)
                            worksheet.write('A6', f"Ng√†y xu·∫•t: {datetime.now().strftime('%d/%m/%Y')}")
                            
                            # Headers
                            headers = ['Kh√°ch h√†ng', 'T√™n Tour/Booking', 'M√£', 'Lo·∫°i', 'Gi√° tr·ªã Hƒê', 'ƒê√£ thu', 'C√≤n l·∫°i']
                            for i, h in enumerate(headers):
                                worksheet.write(7, i, h, fmt_header)
                            
                            # Data
                            df_export = df_debt.sort_values(['customer_name', 'remaining'], ascending=[True, False])
                            
                            row = 8
                            for _, r in df_export.iterrows():
                                worksheet.write(row, 0, r['customer_name'], fmt_text)
                                worksheet.write(row, 1, r['ref_name'], fmt_text)
                                worksheet.write(row, 2, r['ref_code'], fmt_text)
                                worksheet.write(row, 3, r['type'], fmt_text)
                                worksheet.write(row, 4, r['contract_value'], fmt_money)
                                worksheet.write(row, 5, r['paid'], fmt_money)
                                worksheet.write(row, 6, r['remaining'], fmt_money)
                                row += 1
                                
                            # Total row
                            fmt_total = workbook.add_format({'bold': True, 'bg_color': '#FFCDD2', 'border': 1, 'num_format': '#,##0', 'align': 'right', 'font_name': 'Times New Roman'})
                            worksheet.merge_range(row, 0, row, 3, "T·ªîNG C·ªòNG", fmt_total)
                            worksheet.write(row, 4, df_export['contract_value'].sum(), fmt_total)
                            worksheet.write(row, 5, df_export['paid'].sum(), fmt_total)
                            worksheet.write(row, 6, df_export['remaining'].sum(), fmt_total)
                            
                            # Column widths
                            worksheet.set_column('A:A', 25)
                            worksheet.set_column('B:B', 35)
                            worksheet.set_column('C:D', 15)
                            worksheet.set_column('E:G', 18)

                        st.session_state.debt_xls_data = buffer_debt.getvalue()
                        st.rerun()
                    except Exception as e:
                        st.error(f"L·ªói t·∫°o file Excel: {e}")

                if st.session_state.debt_xls_data:
                    st.download_button(
                        label="üì• T·∫£i B√°o C√°o C√¥ng N·ª£ (Excel)",
                        data=st.session_state.debt_xls_data,
                        file_name=f"BaoCao_CongNo_{datetime.now().strftime('%d%m%Y')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        type="primary"
                    )

def render_booking_management():
    st.title("üîñ Qu·∫£n L√Ω Booking")
    st.caption("Qu·∫£n l√Ω c√°c booking l·∫ª, booking d·ªãch v·ª• (Kh√¥ng ph·∫£i Tour tr·ªçn g√≥i)")
    
    # L·∫•y th√¥ng tin user hi·ªán t·∫°i ƒë·ªÉ g√°n cho booking v√† l·ªçc d·ªØ li·ªáu
    current_user_info = st.session_state.get("user_info", {})
    current_user_name = current_user_info.get('name', 'N/A')
    current_user_role = current_user_info.get('role')

    # --- 2. T√ÅCH LI√äN K·∫æT RA 2 PH·∫¶N RI√äNG BI·ªÜT ---
    tab1, tab2, tab3 = st.tabs(["‚ú® T·∫°o Booking", "üîó Chi ti·∫øt Booking", "üìú L·ªãch s·ª≠ Booking"])
    
    # ---------------- TAB 1: T·∫†O BOOKING ----------------
    with tab1:
        with st.container(border=True):
            st.markdown("### ‚ûï T·∫°o Booking M·ªõi")
            
            # --- G·ª¢I √ù KH√ÅCH H√ÄNG ---
            cust_query = "SELECT * FROM customers ORDER BY id DESC"
            cust_params = []
            if current_user_role == 'sale' and current_user_name:
                cust_query = "SELECT * FROM customers WHERE sale_name=? ORDER BY id DESC"
                cust_params.append(current_user_name)
            customers = run_query(cust_query, tuple(cust_params))
            cust_opts = ["-- Kh√°ch m·ªõi --"]
            if customers:
                cust_opts += [f"{c['name']} | {c['phone']}" for c in customers]
            sel_cust = st.selectbox("üîç Ch·ªçn kh√°ch h√†ng c≈© (G·ª£i √Ω):", cust_opts, key="bk_cust_suggest")
            
            pre_name, pre_phone = "", ""
            if sel_cust and sel_cust != "-- Kh√°ch m·ªõi --":
                parts = sel_cust.split(" | ")
                pre_name = parts[0]
                pre_phone = parts[1] if len(parts) > 1 else ""
            
            # Ch·ªçn lo·∫°i d·ªãch v·ª•
            bk_type = st.radio("Ch·ªçn lo·∫°i d·ªãch v·ª•:", ["üè® Kh√°ch s·∫°n", "üöå V·∫≠n chuy·ªÉn", "üß© Combo / ƒêa d·ªãch v·ª•", "üîñ Kh√°c"], horizontal=True)
            st.divider()

            if bk_type == "üè® Kh√°ch s·∫°n":
                st.markdown("##### üè® Th√¥ng tin l∆∞u tr√∫ & T√†i ch√≠nh")
                
                # [NEW] Move dates out to calculate nights
                c_date, c_room = st.columns([2, 1])
                dates = c_date.date_input("Th·ªùi gian l∆∞u tr√∫", value=[], help="Ch·ªçn ng√†y nh·∫≠n v√† tr·∫£ ph√≤ng", format="DD/MM/YYYY")
                room_count = c_room.number_input("S·ªë l∆∞·ª£ng ph√≤ng", min_value=1, step=1, value=1)
                
                nights = 1
                if len(dates) == 2:
                    nights = (dates[1] - dates[0]).days
                    if nights < 1: nights = 1
                    st.caption(f"Th·ªùi gian: {dates[0].strftime('%d/%m')} - {dates[1].strftime('%d/%m')} ({nights} ƒë√™m) x {room_count} ph√≤ng")
                elif len(dates) == 1:
                    st.caption("Vui l√≤ng ch·ªçn ng√†y tr·∫£ ph√≤ng.")
                else:
                    st.caption("Vui l√≤ng ch·ªçn ng√†y nh·∫≠n v√† tr·∫£ ph√≤ng.")

                # Financials
                c1, c2, c3 = st.columns(3)
                
                # [CODE M·ªöI] X·ª≠ l√Ω nh·∫≠p ti·ªÅn c√≥ ƒë·ªãnh d·∫°ng
                if "bk_hotel_net_val" not in st.session_state: st.session_state.bk_hotel_net_val = ""
                if "bk_hotel_sell_val" not in st.session_state: st.session_state.bk_hotel_sell_val = ""

                def fmt_hotel_net():
                    val = st.session_state.bk_hotel_net_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_hotel_net_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass
                
                def fmt_hotel_sell():
                    val = st.session_state.bk_hotel_sell_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_hotel_sell_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass

                net_price_input = c1.text_input("Gi√° n√©t / ƒë√™m / ph√≤ng", key="bk_hotel_net_val", on_change=fmt_hotel_net, help="Nh·∫≠p s·ªë ti·ªÅn (VD: 1000000)")
                selling_price_input = c2.text_input("Gi√° b√°n / ƒë√™m / ph√≤ng", key="bk_hotel_sell_val", on_change=fmt_hotel_sell, help="Nh·∫≠p s·ªë ti·ªÅn (VD: 1500000)")
                
                try: net_price_unit = float(net_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: net_price_unit = 0.0
                
                try: selling_price_unit = float(selling_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: selling_price_unit = 0.0
                
                # Tax logic
                tax_option = st.radio("Gi√° n√©t ƒë√£ bao g·ªìm thu·∫ø?", ["ƒê√£ bao g·ªìm thu·∫ø", "Ch∆∞a bao g·ªìm thu·∫ø"], horizontal=True)
                tax_percent = 0.0
                if tax_option == "Ch∆∞a bao g·ªìm thu·∫ø":
                    tax_percent = st.number_input("Nh·∫≠p % Thu·∫ø", min_value=0.0, max_value=100.0, step=0.5, format="%.1f")
                
                # Calculations
                net_price_unit_incl_tax = net_price_unit * (1 + tax_percent / 100)
                total_net = net_price_unit_incl_tax * nights * room_count
                total_sell = selling_price_unit * nights * room_count
                total_profit = total_sell - total_net
                
                # Display Table
                st.markdown("###### üìä B·∫£ng t√≠nh chi ti·∫øt")
                calc_df = pd.DataFrame({
                    "Lo·∫°i": ["Gi√° N√©t (V·ªën)", "Gi√° B√°n (Doanh thu)"],
                    "ƒê∆°n gi√°": [format_vnd(net_price_unit_incl_tax) + " VND", format_vnd(selling_price_unit) + " VND"],
                    "S·ªë l∆∞·ª£ng": [f"{nights} ƒë√™m x {room_count} ph√≤ng", f"{nights} ƒë√™m x {room_count} ph√≤ng"],
                    "Th√†nh ti·ªÅn": [format_vnd(total_net) + " VND", format_vnd(total_sell) + " VND"]
                })
                st.dataframe(calc_df, use_container_width=True, hide_index=True)
                
                st.markdown(f"""<div style="background-color: #e8f5e9; padding: 10px; border-radius: 5px; border: 1px solid #c8e6c9; text-align: center;">
                    <span style="color: #2e7d32; font-weight: bold; font-size: 1.1em;">L·ª¢I NHU·∫¨N D·ª∞ KI·∫æN: {format_vnd(total_profit)} VND</span>
                </div>""", unsafe_allow_html=True)

                st.divider()
                st.text_input("Sales ph·ª• tr√°ch", value=current_user_name, disabled=True)
                with st.form("bk_hotel", clear_on_submit=True):
                    h_name = st.text_input("T√™n Kh√°ch s·∫°n", placeholder="VD: M∆∞·ªùng Thanh Luxury")
                    
                    # [NEW] Th√™m c√°c tr∆∞·ªùng th√¥ng tin chi ti·∫øt
                    c_h1, c_h2 = st.columns(2)
                    hotel_code = c_h1.text_input("M√£ code kh√°ch s·∫°n (Booking ID)", placeholder="VD: 12345678")
                    room_type = c_h2.text_area("H·∫°ng ph√≤ng", placeholder="VD: 2 Deluxe, 1 Suite (Xu·ªëng d√≤ng n·∫øu nhi·ªÅu h·∫°ng)", height=68)
                    guest_list = st.text_area("Danh s√°ch kh√°ch l∆∞u tr√∫", placeholder="VD: Nguyen Van A, Tran Thi B...", height=100)
                    
                    # Dates are already outside
                    
                    c_cust_n, c_cust_p = st.columns(2)
                    cust_name = c_cust_n.text_input("T√™n kh√°ch h√†ng (*)", value=pre_name, placeholder="Nh·∫≠p t√™n kh√°ch")
                    cust_phone = c_cust_p.text_input("S·ªë ƒëi·ªán tho·∫°i", value=pre_phone, placeholder="Nh·∫≠p SƒêT (T√πy ch·ªçn)")

                    new_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                    st.caption(f"M√£ Booking d·ª± ki·∫øn: {new_code}")
                    if st.form_submit_button("T·∫°o Booking Kh√°ch s·∫°n", type="primary"):
                        if h_name and len(dates) == 2 and cust_name:
                            cust_info = f"{cust_name} - {cust_phone}" if cust_phone else cust_name
                            d_range = f"{dates[0].strftime('%d/%m/%Y')} - {dates[1].strftime('%d/%m/%Y')} ({nights} ƒë√™m, {room_count} ph√≤ng)"
                            save_customer_check(cust_name, cust_phone, current_user_name)

                            add_row_to_table('service_bookings', {
                                'code': new_code, 'name': f"[KS] {h_name}", 'created_at': datetime.now().strftime("%Y-%m-%d"),
                                'type': 'HOTEL', 'details': f"L∆∞u tr√∫: {d_range}", 'customer_info': cust_info,
                                'net_price': total_net, # Storing TOTAL
                                'tax_percent': tax_percent,
                                'selling_price': total_sell, # Storing TOTAL
                                'profit': total_profit,
                                'sale_name': current_user_name,
                                'hotel_code': hotel_code,
                                'room_type': room_type,
                                'guest_list': guest_list
                            })
                            # Clear inputs
                            if "bk_hotel_net_val" in st.session_state: del st.session_state.bk_hotel_net_val
                            if "bk_hotel_sell_val" in st.session_state: del st.session_state.bk_hotel_sell_val
                            st.success("ƒê√£ t·∫°o!"); time.sleep(0.5); st.rerun()
                        else: st.warning("Vui l√≤ng nh·∫≠p t√™n kh√°ch s·∫°n, t√™n kh√°ch h√†ng v√† ch·ªçn ƒë·ªß ng√†y ƒëi/v·ªÅ.")

            elif bk_type == "üöå V·∫≠n chuy·ªÉn":
                trans_type = st.radio("Lo·∫°i ph∆∞∆°ng ti·ªán:", ["Xe (√î t√¥)", "M√°y bay", "T√†u h·ªèa", "Du thuy·ªÅn"], horizontal=True)
                
                st.divider()
                st.markdown("##### üí∞ Th√¥ng tin t√†i ch√≠nh")
                
                c_qty, c_net, c_sell = st.columns(3)
                qty = c_qty.number_input("S·ªë l∆∞·ª£ng (V√©/Kh√°ch)", min_value=1, value=1, key="trans_qty")
                
                # [CODE M·ªöI] X·ª≠ l√Ω nh·∫≠p ti·ªÅn c√≥ ƒë·ªãnh d·∫°ng cho V·∫≠n chuy·ªÉn
                if "bk_trans_net_val" not in st.session_state: st.session_state.bk_trans_net_val = ""
                if "bk_trans_sell_val" not in st.session_state: st.session_state.bk_trans_sell_val = ""

                def fmt_trans_net():
                    val = st.session_state.bk_trans_net_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_trans_net_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass
                
                def fmt_trans_sell():
                    val = st.session_state.bk_trans_sell_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_trans_sell_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass

                net_price_input = c_net.text_input("Gi√° n√©t / v√©", key="bk_trans_net_val", on_change=fmt_trans_net, help="Nh·∫≠p s·ªë ti·ªÅn (VD: 1000000)")
                selling_price_input = c_sell.text_input("Gi√° b√°n / v√©", key="bk_trans_sell_val", on_change=fmt_trans_sell, help="Nh·∫≠p s·ªë ti·ªÅn (VD: 1500000)")
                
                try: net_price_unit = float(net_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: net_price_unit = 0.0
                
                try: selling_price_unit = float(selling_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: selling_price_unit = 0.0
                
                tax_option = st.radio("Gi√° n√©t ƒë√£ bao g·ªìm thu·∫ø?", ["ƒê√£ bao g·ªìm thu·∫ø", "Ch∆∞a bao g·ªìm thu·∫ø"], horizontal=True, key="trans_tax_opt")
                tax_percent = 0.0
                
                if tax_option == "Ch∆∞a bao g·ªìm thu·∫ø":
                    tax_percent = st.number_input("Nh·∫≠p % Thu·∫ø", min_value=0.0, max_value=100.0, step=0.5, format="%.1f", key="trans_tax_pct")
                
                # Calculations
                net_price_unit_incl_tax = net_price_unit * (1 + tax_percent / 100)
                total_net = net_price_unit_incl_tax * qty
                total_sell = selling_price_unit * qty
                profit = total_sell - total_net

                # Display Table
                st.markdown("###### üìä B·∫£ng t√≠nh chi ti·∫øt")
                calc_df = pd.DataFrame({
                    "Lo·∫°i": ["Gi√° N√©t (V·ªën)", "Gi√° B√°n (Doanh thu)"],
                    "ƒê∆°n gi√°": [format_vnd(net_price_unit_incl_tax) + " VND", format_vnd(selling_price_unit) + " VND"],
                    "S·ªë l∆∞·ª£ng": [qty, qty],
                    "Th√†nh ti·ªÅn": [format_vnd(total_net) + " VND", format_vnd(total_sell) + " VND"]
                })
                st.dataframe(calc_df, use_container_width=True, hide_index=True)

                st.markdown(f"""<div style="background-color: #e8f5e9; padding: 10px; border-radius: 5px; border: 1px solid #c8e6c9; text-align: center;">
                    <span style="color: #2e7d32; font-weight: bold; font-size: 1.1em;">L·ª¢I NHU·∫¨N D·ª∞ KI·∫æN: {format_vnd(profit)} VND</span>
                </div>""", unsafe_allow_html=True)

                st.divider()
                st.text_input("Sales ph·ª• tr√°ch", value=current_user_name, disabled=True, key="trans_sale")
                with st.form("bk_trans", clear_on_submit=True):
                    details = ""
                    bk_name = ""
                    is_valid = False

                    if trans_type == "Xe (√î t√¥)":
                        c1, c2 = st.columns(2)
                        route_from = c1.text_input("ƒêi·ªÉm ƒëi")
                        route_to = c2.text_input("ƒêi·ªÉm ƒë·∫øn")
                        c3, c4, c5 = st.columns(3)
                        car_type = c3.selectbox("Lo·∫°i xe", ["4S", "7S", "16S", "29S", "35S", "45S"])
                        car_no = c4.text_input("Bi·ªÉn s·ªë / M√£ xe")
                        t_date = c5.date_input("Ng√†y ƒëi", format="DD/MM/YYYY")
                        
                        if route_from and route_to:
                            is_valid = True
                            bk_name = f"[XE] {route_from} - {route_to}"
                            details = f"Xe {car_type}: {car_no} | Ng√†y: {t_date.strftime('%d/%m/%Y')} | SL: {qty}"

                    elif trans_type == "M√°y bay":
                        c1, c2 = st.columns(2)
                        ticket_code = c1.text_input("M√£ v√© / S·ªë hi·ªáu")
                        flight_date = c2.date_input("Ng√†y bay", format="DD/MM/YYYY")
                        flight_route = st.text_input("H√†nh tr√¨nh / H√£ng bay (T√πy ch·ªçn)", placeholder="VD: VN123 HAN-SGN")
                        
                        if ticket_code:
                            is_valid = True
                            desc = flight_route if flight_route else ticket_code
                            bk_name = f"[BAY] {desc}"
                            details = f"V√©: {ticket_code} | Ng√†y: {flight_date.strftime('%d/%m/%Y')} | SL: {qty}"

                    elif trans_type == "T√†u h·ªèa":
                        c1, c2 = st.columns(2)
                        ticket_code = c1.text_input("M√£ v√© / Toa / Gh·∫ø")
                        train_date = c2.date_input("Ng√†y ƒëi", format="DD/MM/YYYY")
                        train_route = st.text_input("Ga ƒëi - Ga ƒë·∫øn (T√πy ch·ªçn)", placeholder="VD: H√† N·ªôi - Vinh")
                        
                        if ticket_code:
                            is_valid = True
                            desc = train_route if train_route else ticket_code
                            bk_name = f"[TAU] {desc}"
                            details = f"V√©: {ticket_code} | Ng√†y: {train_date.strftime('%d/%m/%Y')} | SL: {qty}"

                    elif trans_type == "Du thuy·ªÅn":
                        c1, c2 = st.columns(2)
                        cruise_name = c1.text_input("T√™n du thuy·ªÅn / Tuy·∫øn")
                        cruise_date = c2.date_input("Ng√†y ƒëi", format="DD/MM/YYYY")
                        cabin_type = st.text_input("Lo·∫°i Cabin / Ghi ch√∫", placeholder="VD: Junior Suite, Balcony...")
                        
                        if cruise_name:
                            is_valid = True
                            bk_name = f"[THUYEN] {cruise_name}"
                            details = f"Cabin: {cabin_type} | Ng√†y: {cruise_date.strftime('%d/%m/%Y')} | SL: {qty}"

                    st.divider()
                    c_cust_n, c_cust_p = st.columns(2)
                    cust_name = c_cust_n.text_input("T√™n kh√°ch h√†ng (*)", value=pre_name, placeholder="Nh·∫≠p t√™n kh√°ch")
                    cust_phone = c_cust_p.text_input("S·ªë ƒëi·ªán tho·∫°i", value=pre_phone, placeholder="Nh·∫≠p SƒêT (T√πy ch·ªçn)")

                    new_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                    st.caption(f"M√£ Booking d·ª± ki·∫øn: {new_code}")
                    if st.form_submit_button("T·∫°o Booking V·∫≠n chuy·ªÉn", type="primary"):
                        if is_valid and cust_name:
                            cust_info = f"{cust_name} - {cust_phone}" if cust_phone else cust_name
                            save_customer_check(cust_name, cust_phone, current_user_name)
                            add_row_to_table('service_bookings', {
                                'code': new_code, 'name': bk_name, 'created_at': datetime.now().strftime("%Y-%m-%d"),
                                'type': 'TRANS', 'details': details, 'customer_info': cust_info,
                                'net_price': total_net,
                                'tax_percent': tax_percent,
                                'selling_price': total_sell, 'profit': profit,
                                'sale_name': current_user_name
                            })
                            # Clear inputs
                            if "bk_trans_net_val" in st.session_state: del st.session_state.bk_trans_net_val
                            if "bk_trans_sell_val" in st.session_state: del st.session_state.bk_trans_sell_val
                            st.success("ƒê√£ t·∫°o!"); time.sleep(0.5); st.rerun()
                        else: st.warning("Vui l√≤ng nh·∫≠p ƒë·ªß th√¥ng tin (H√†nh tr√¨nh/M√£ v√© v√† T√™n kh√°ch).")

            elif bk_type == "üß© Combo / ƒêa d·ªãch v·ª•":
                if "combo_list" not in st.session_state: st.session_state.combo_list = []
                c_add, c_list = st.columns([1, 1.5])
                with c_add:
                    st.markdown("##### Th√™m d·ªãch v·ª• con")
                    sub_type = st.selectbox("Lo·∫°i", ["Kh√°ch s·∫°n", "V·∫≠n chuy·ªÉn", "Kh√°c"], key="cb_sub")
                    if sub_type == "Kh√°ch s·∫°n":
                        sh_n = st.text_input("T√™n KS", key="cb_h_n")
                        c_qty, c_date = st.columns([1, 2])
                        sh_qty = c_qty.number_input("S·ªë l∆∞·ª£ng ph√≤ng", min_value=1, value=1, key="cb_h_q")
                        sh_d = c_date.date_input("Ng√†y ·ªü", [], key="cb_h_d", format="DD/MM/YYYY")
                        if st.button("Th√™m KS") and sh_n and len(sh_d)==2:
                            st.session_state.combo_list.append(f"üè® {sh_n} - {sh_qty} ph√≤ng ({sh_d[0].strftime('%d/%m')} - {sh_d[1].strftime('%d/%m')})"); st.rerun()
                    elif sub_type == "V·∫≠n chuy·ªÉn":
                        tr_mode = st.selectbox("Lo·∫°i ph∆∞∆°ng ti·ªán", ["Xe", "M√°y bay", "T√†u h·ªèa", "Du thuy·ªÅn"], key="cb_tr_mode")
                        st_r = st.text_input("H√†nh tr√¨nh / M√£ v√© / T√™n t√†u", key="cb_t_r")
                        st_d = st.date_input("Ng√†y", key="cb_t_d", format="DD/MM/YYYY")
                        
                        icon_map = {"Xe": "üöå", "M√°y bay": "‚úàÔ∏è", "T√†u h·ªèa": "üöÜ", "Du thuy·ªÅn": "üö¢"}
                        
                        if st.button("Th√™m V·∫≠n chuy·ªÉn") and st_r:
                            st.session_state.combo_list.append(f"{icon_map[tr_mode]} {st_r} ({st_d.strftime('%d/%m')})"); st.rerun()
                    else:
                        so_n = st.text_input("T√™n d·ªãch v·ª•", key="cb_o_n")
                        if st.button("Th√™m DV") and so_n:
                            st.session_state.combo_list.append(f"üîñ {so_n}"); st.rerun()
                with c_list:
                    st.markdown("##### Danh s√°ch ƒë√£ th√™m")
                    # [FIX] D√πng list() ƒë·ªÉ √©p ki·ªÉu r√µ r√†ng, Pylance s·∫Ω hi·ªÉu ƒë√¢y l√† danh s√°ch l·∫∑p ƒë∆∞·ª£c
                    safe_combo_list = list(st.session_state.get("combo_list", []))
                    for i, item in enumerate(safe_combo_list): st.text(f"{i+1}. {item}")
                    
                    if st.session_state.get("combo_list") and st.button("X√≥a h·∫øt", type="secondary"): 
                        st.session_state.combo_list = []
                        st.rerun()
                
                st.divider()
                st.markdown("##### üí∞ Th√¥ng tin t√†i ch√≠nh")
                
                c_qty, c_net, c_sell = st.columns(3)
                qty = c_qty.number_input("S·ªë l∆∞·ª£ng (Combo/Pax)", min_value=1, value=1, key="combo_qty")
                
                # [CODE M·ªöI] X·ª≠ l√Ω nh·∫≠p ti·ªÅn c√≥ ƒë·ªãnh d·∫°ng cho Combo
                if "bk_combo_net_val" not in st.session_state: st.session_state.bk_combo_net_val = ""
                if "bk_combo_sell_val" not in st.session_state: st.session_state.bk_combo_sell_val = ""

                def fmt_combo_net():
                    val = st.session_state.bk_combo_net_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_combo_net_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass
                
                def fmt_combo_sell():
                    val = st.session_state.bk_combo_sell_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_combo_sell_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass

                net_price_input = c_net.text_input("Gi√° n√©t / combo", key="bk_combo_net_val", on_change=fmt_combo_net, help="Nh·∫≠p s·ªë ti·ªÅn (VD: 1000000)")
                selling_price_input = c_sell.text_input("Gi√° b√°n / combo", key="bk_combo_sell_val", on_change=fmt_combo_sell, help="Nh·∫≠p s·ªë ti·ªÅn (VD: 1500000)")
                
                try: net_price_unit = float(net_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: net_price_unit = 0.0
                
                try: selling_price_unit = float(selling_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: selling_price_unit = 0.0
                
                tax_option = st.radio("Gi√° n√©t ƒë√£ bao g·ªìm thu·∫ø?", ["ƒê√£ bao g·ªìm thu·∫ø", "Ch∆∞a bao g·ªìm thu·∫ø"], horizontal=True, key="combo_tax_opt")
                tax_percent = 0.0
                
                if tax_option == "Ch∆∞a bao g·ªìm thu·∫ø":
                    tax_percent = st.number_input("Nh·∫≠p % Thu·∫ø", min_value=0.0, max_value=100.0, step=0.5, format="%.1f", key="combo_tax_pct")
                
                # Calculations
                net_price_unit_incl_tax = net_price_unit * (1 + tax_percent / 100)
                total_net = net_price_unit_incl_tax * qty
                total_sell = selling_price_unit * qty
                profit = total_sell - total_net

                # Display Table
                st.markdown("###### üìä B·∫£ng t√≠nh chi ti·∫øt")
                calc_df = pd.DataFrame({
                    "Lo·∫°i": ["Gi√° N√©t (V·ªën)", "Gi√° B√°n (Doanh thu)"],
                    "ƒê∆°n gi√°": [format_vnd(net_price_unit_incl_tax) + " VND", format_vnd(selling_price_unit) + " VND"],
                    "S·ªë l∆∞·ª£ng": [qty, qty],
                    "Th√†nh ti·ªÅn": [format_vnd(total_net) + " VND", format_vnd(total_sell) + " VND"]
                })
                st.dataframe(calc_df, use_container_width=True, hide_index=True)

                st.markdown(f"""<div style="background-color: #e8f5e9; padding: 10px; border-radius: 5px; border: 1px solid #c8e6c9; text-align: center;">
                    <span style="color: #2e7d32; font-weight: bold; font-size: 1.1em;">L·ª¢I NHU·∫¨N D·ª∞ KI·∫æN: {format_vnd(profit)} VND</span>
                </div>""", unsafe_allow_html=True)

                st.divider()
                st.text_input("Sales ph·ª• tr√°ch", value=current_user_name, disabled=True, key="combo_sale")
                with st.form("bk_combo", clear_on_submit=True):
                    combo_name = st.text_input("T√™n Combo / G√≥i", placeholder="VD: Combo ƒê√† N·∫µng 3N2ƒê")
                    # [NEW] Th√™m danh s√°ch kh√°ch cho Combo
                    guest_list_cb = st.text_area("Danh s√°ch kh√°ch (ƒêo√†n)", placeholder="Nh·∫≠p danh s√°ch kh√°ch h√†ng...", height=100)
                    
                    c_cust_n, c_cust_p = st.columns(2)
                    cust_name = c_cust_n.text_input("T√™n kh√°ch h√†ng (*)", value=pre_name, placeholder="Nh·∫≠p t√™n kh√°ch")
                    cust_phone = c_cust_p.text_input("S·ªë ƒëi·ªán tho·∫°i", value=pre_phone, placeholder="Nh·∫≠p SƒêT (T√πy ch·ªçn)")

                    new_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                    if st.form_submit_button("L∆∞u Combo", type="primary"):
                        if combo_name and st.session_state.combo_list and cust_name:
                            cust_info = f"{cust_name} - {cust_phone}" if cust_phone else cust_name
                            save_customer_check(cust_name, cust_phone, current_user_name)
                            details_str = " | ".join(st.session_state.combo_list) + f" (SL: {qty})"
                            add_row_to_table('service_bookings', {
                                'code': new_code, 'name': f"[CB] {combo_name}", 'created_at': datetime.now().strftime("%Y-%m-%d"),
                                'type': 'COMBO', 'details': details_str, 'customer_info': cust_info,
                                'net_price': total_net,
                                'tax_percent': tax_percent,
                                'selling_price': total_sell, 'profit': profit,
                                'sale_name': current_user_name,
                                'guest_list': guest_list_cb
                            })
                            # Clear inputs
                            if "bk_combo_net_val" in st.session_state: del st.session_state.bk_combo_net_val
                            if "bk_combo_sell_val" in st.session_state: del st.session_state.bk_combo_sell_val
                            st.session_state.combo_list = []; st.success("ƒê√£ t·∫°o!"); time.sleep(0.5); st.rerun()
                        else: st.warning("C·∫ßn t√™n Combo, t√™n kh√°ch h√†ng v√† √≠t nh·∫•t 1 d·ªãch v·ª•.")

            else:
                st.markdown("##### üí∞ Th√¥ng tin t√†i ch√≠nh")
                c_qty, c_net, c_sell = st.columns(3)
                qty = c_qty.number_input("S·ªë l∆∞·ª£ng", min_value=1, value=1, key="other_qty")
                
                # [CODE M·ªöI] X·ª≠ l√Ω nh·∫≠p ti·ªÅn c√≥ ƒë·ªãnh d·∫°ng cho Kh√°c
                if "bk_other_net_val" not in st.session_state: st.session_state.bk_other_net_val = ""
                if "bk_other_sell_val" not in st.session_state: st.session_state.bk_other_sell_val = ""

                def fmt_other_net():
                    val = st.session_state.bk_other_net_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_other_net_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass
                
                def fmt_other_sell():
                    val = st.session_state.bk_other_sell_val
                    try:
                        v_float = float(val.replace('.', '').replace(',', '').replace(' VND', '').strip())
                        st.session_state.bk_other_sell_val = "{:,.0f}".format(v_float).replace(",", ".") + " VND"
                    except: pass

                net_price_input = c_net.text_input("Gi√° n√©t / ƒë∆°n v·ªã", key="bk_other_net_val", on_change=fmt_other_net, help="Nh·∫≠p s·ªë ti·ªÅn (VD: 1000000)")
                selling_price_input = c_sell.text_input("Gi√° b√°n / ƒë∆°n v·ªã", key="bk_other_sell_val", on_change=fmt_other_sell, help="Nh·∫≠p s·ªë ti·ªÅn (VD: 1500000)")
                
                try: net_price_unit = float(net_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: net_price_unit = 0.0
                
                try: selling_price_unit = float(selling_price_input.replace('.', '').replace(',', '').replace(' VND', '').strip())
                except: selling_price_unit = 0.0
                
                tax_option = st.radio("Gi√° n√©t ƒë√£ bao g·ªìm thu·∫ø?", ["ƒê√£ bao g·ªìm thu·∫ø", "Ch∆∞a bao g·ªìm thu·∫ø"], horizontal=True, key="other_tax_opt")
                tax_percent = 0.0
                
                if tax_option == "Ch∆∞a bao g·ªìm thu·∫ø":
                    tax_percent = st.number_input("Nh·∫≠p % Thu·∫ø", min_value=0.0, max_value=100.0, step=0.5, format="%.1f", key="other_tax_pct")
                
                # Calculations
                net_price_unit_incl_tax = net_price_unit * (1 + tax_percent / 100)
                total_net = net_price_unit_incl_tax * qty
                total_sell = selling_price_unit * qty
                profit = total_sell - total_net

                # Display Table
                st.markdown("###### üìä B·∫£ng t√≠nh chi ti·∫øt")
                calc_df = pd.DataFrame({
                    "Lo·∫°i": ["Gi√° N√©t (V·ªën)", "Gi√° B√°n (Doanh thu)"],
                    "ƒê∆°n gi√°": [format_vnd(net_price_unit_incl_tax) + " VND", format_vnd(selling_price_unit) + " VND"],
                    "S·ªë l∆∞·ª£ng": [qty, qty],
                    "Th√†nh ti·ªÅn": [format_vnd(total_net) + " VND", format_vnd(total_sell) + " VND"]
                })
                st.dataframe(calc_df, use_container_width=True, hide_index=True)

                st.markdown(f"""<div style="background-color: #e8f5e9; padding: 10px; border-radius: 5px; border: 1px solid #c8e6c9; text-align: center;">
                    <span style="color: #2e7d32; font-weight: bold; font-size: 1.1em;">L·ª¢I NHU·∫¨N D·ª∞ KI·∫æN: {format_vnd(profit)} VND</span>
                </div>""", unsafe_allow_html=True)

                st.divider()
                st.text_input("Sales ph·ª• tr√°ch", value=current_user_name, disabled=True, key="other_sale")
                with st.form("bk_other", clear_on_submit=True):
                    new_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                    c1, c2 = st.columns([1, 3])
                    c1.text_input("M√£ (Auto)", value=new_code, disabled=True)
                    new_name = c2.text_input("T√™n Booking / D·ªãch v·ª•")
                    
                    c_cust_n, c_cust_p = st.columns(2)
                    cust_name = c_cust_n.text_input("T√™n kh√°ch h√†ng (*)", value=pre_name, placeholder="Nh·∫≠p t√™n kh√°ch")
                    cust_phone = c_cust_p.text_input("S·ªë ƒëi·ªán tho·∫°i", value=pre_phone, placeholder="Nh·∫≠p SƒêT (T√πy ch·ªçn)")

                    if st.form_submit_button("T·∫°o"):
                        if new_name and cust_name:
                            cust_info = f"{cust_name} - {cust_phone}" if cust_phone else cust_name
                            save_customer_check(cust_name, cust_phone, current_user_name)
                            add_row_to_table('service_bookings', {
                                'code': new_code, 'name': new_name, 'created_at': datetime.now().strftime("%Y-%m-%d"),
                                'type': 'OTHER', 'customer_info': cust_info, 'details': f"SL: {qty}",
                                'net_price': total_net,
                                'tax_percent': tax_percent,
                                'selling_price': total_sell, 'profit': profit,
                                'sale_name': current_user_name
                            })
                            # Clear inputs
                            if "bk_other_net_val" in st.session_state: del st.session_state.bk_other_net_val
                            if "bk_other_sell_val" in st.session_state: del st.session_state.bk_other_sell_val
                            st.success("ƒê√£ t·∫°o!"); time.sleep(0.5); st.rerun()
                        else: st.warning("Vui l√≤ng nh·∫≠p t√™n d·ªãch v·ª• v√† t√™n kh√°ch h√†ng.")

    # ---------------- TAB 2: KH·ªöP UNC & H√ìA ƒê∆†N (D·ª∞ √ÅN UNC) ----------------
    with tab2:
        st.subheader("üîó Chi ti·∫øt Booking")
        # --- L·ªçc danh s√°ch booking theo sale ---
        bk_query = "SELECT * FROM service_bookings WHERE status='active'"
        bk_params = []
        if current_user_role == 'sale' and current_user_name:
            bk_query += " AND sale_name=?"
            bk_params.append(current_user_name)
        bk_query += " ORDER BY id DESC"
        bookings = run_query(bk_query, tuple(bk_params))
        
        if bookings:
            bk_map = {f"[{b['code']}] {b['name']}": b['code'] for b in bookings} # type: ignore
            selected_bk_label = st.selectbox("Ch·ªçn Booking ƒë·ªÉ xem chi ti·∫øt:", list(bk_map.keys()))
            
            if selected_bk_label:
                code = bk_map[selected_bk_label] # type: ignore
                
                bk_info = run_query("SELECT * FROM service_bookings WHERE code=?", (code,), fetch_one=True)
                st.divider()
                st.markdown(f"### üìä Chi ti·∫øt: {selected_bk_label}")
                if isinstance(bk_info, sqlite3.Row):
                    st.markdown("##### üí∞ T·ªïng quan t√†i ch√≠nh")
                    fin1, fin2, fin3 = st.columns(3)
                    net_p = bk_info['net_price'] or 0 # type: ignore
                    sell_p = bk_info['selling_price'] or 0 # type: ignore
                    prof_p = bk_info['profit'] or 0 # type: ignore
                    fin1.metric("Gi√° n√©t (ƒë√£ g·ªìm thu·∫ø)", format_vnd(net_p))
                    fin2.metric("Gi√° b√°n", format_vnd(sell_p))
                    fin3.metric("L·ª£i nhu·∫≠n", format_vnd(prof_p))

                    if bk_info['customer_info']:
                        st.markdown(f"**üë§ Kh√°ch h√†ng:** {bk_info['customer_info']}")
                    if bk_info['details']:
                        st.info(f"‚ÑπÔ∏è **Th√¥ng tin:** {bk_info['details']}")
                
                # G·ªçi h√†m hi·ªÉn th·ªã so s√°nh
                render_cost_comparison(code)
                
                # --- N√öT T·∫¢I BOOKING CONFIRMATION (M·ªöI) ---
                st.write("")
                c_lang, c_dl_btn = st.columns([1, 2])
                sel_lang = c_lang.radio("Ng√¥n ng·ªØ PDF:", ["Ti·∫øng Vi·ªát", "English"], horizontal=True)
                lang_code = 'vi' if sel_lang == "Ti·∫øng Vi·ªát" else 'en'
                
                comp_data_cfm = get_company_data()
                pdf_cfm = create_booking_cfm_pdf(dict(bk_info), comp_data_cfm, lang=lang_code)
                c_dl_btn.download_button(
                    label="üì• T·∫£i Booking Confirmation (PDF)",
                    data=pdf_cfm,
                    file_name=f"Booking_CFM_{code}.pdf",
                    mime="application/pdf",
                    type="secondary"
                )
                
                st.divider()
                # N√∫t ho√†n t·∫•t & x√≥a booking
                c_complete, c_delete = st.columns(2)
                if c_complete.button("‚úÖ Ho√†n t·∫•t Booking", type="primary", use_container_width=True):
                    run_query("UPDATE service_bookings SET status='completed' WHERE code=?", (code,), commit=True)
                    st.success("ƒê√£ ho√†n t·∫•t! Booking ƒë√£ ƒë∆∞·ª£c chuy·ªÉn sang tab L·ªãch s·ª≠."); time.sleep(1); st.rerun()

                if c_delete.button("üóëÔ∏è X√≥a Booking n√†y", use_container_width=True):
                    run_query("UPDATE service_bookings SET status='deleted' WHERE code=?", (code,), commit=True)
                    st.success("ƒê√£ x√≥a!"); time.sleep(0.5); st.rerun()
        else:
            st.info("Ch∆∞a c√≥ booking n√†o.")

    # ---------------- TAB 3: L·ªäCH S·ª¨ BOOKING ----------------
    with tab3:
        st.subheader("üìú L·ªãch s·ª≠ Booking ƒë√£ ho√†n t·∫•t")
        # --- L·ªçc danh s√°ch booking theo sale ---
        hist_bk_query = "SELECT * FROM service_bookings WHERE status='completed'"
        hist_bk_params = []
        if current_user_role == 'sale' and current_user_name:
            hist_bk_query += " AND sale_name=?"
            hist_bk_params.append(current_user_name)
        hist_bk_query += " ORDER BY id DESC"
        history_bk = run_query(hist_bk_query, tuple(hist_bk_params))
        if history_bk:
            df_hist = pd.DataFrame([dict(r) for r in history_bk])
            st.dataframe(
                df_hist[['code', 'name', 'created_at', 'type', 'customer_info', 'details', 'net_price', 'selling_price', 'profit']],
                column_config={
                    "code": "M√£ Booking",
                    "name": "T√™n Booking",
                    "created_at": "Ng√†y t·∫°o",
                    "type": "Lo·∫°i",
                    "customer_info": "Kh√°ch h√†ng",
                    "details": "Chi ti·∫øt",
                    "net_price": st.column_config.NumberColumn("Gi√° n√©t", format="%d"),
                    "selling_price": st.column_config.NumberColumn("Gi√° b√°n", format="%d"),
                    "profit": st.column_config.NumberColumn("L·ª£i nhu·∫≠n", format="%d"),
                },
                use_container_width=True,
                hide_index=True
            )
        else:
            st.info("Ch∆∞a c√≥ booking n√†o ho√†n t·∫•t.")

def render_tour_management():
    st.title("üì¶ Qu·∫£n L√Ω Tour ")
    
    # S·ª≠ d·ª•ng Tabs theo y√™u c·∫ßu
    tab_est, tab_list_srv, tab_act, tab_hist, tab_rpt = st.tabs(["üìù D·ª± To√°n Chi Ph√≠", "üìã Danh s√°ch & D·ªãch v·ª•", "üí∏ Quy·∫øt To√°n Tour", "üìú L·ªãch s·ª≠ Tour", "üìà T·ªïng H·ª£p L·ª£i Nhu·∫≠n"])
    
    # L·∫•y th√¥ng tin user hi·ªán t·∫°i ƒë·ªÉ l·ªçc
    current_user_info_tour = st.session_state.get("user_info", {})
    current_user_name_tour = current_user_info_tour.get('name', 'N/A')
    current_user_role_tour = current_user_info_tour.get('role')

    # L·∫•y danh s√°ch Tour cho Selectbox d√πng chung
    all_tours_query = "SELECT * FROM tours ORDER BY id DESC"
    all_tours_params = []
    if current_user_role_tour == 'sale' and current_user_name_tour:
        all_tours_query = "SELECT * FROM tours WHERE sale_name=? ORDER BY id DESC"
        all_tours_params.append(current_user_name_tour)
    all_tours = run_query(all_tours_query, tuple(all_tours_params))
    running_tours = [t for t in all_tours if t['status'] == 'running']
    tour_options = {f"[{t['tour_code']}] {t['tour_name']} ({t['start_date']})": t['id'] for t in running_tours} if running_tours else {} # type: ignore
    
    # ---------------- TAB 1: D·ª∞ TO√ÅN CHI PH√ç ----------------
    with tab_est:
        with st.expander("‚ûï T·∫°o Th√¥ng Tin ƒêo√†n M·ªõi", expanded=False):
            # --- G·ª¢I √ù KH√ÅCH H√ÄNG ---
            cust_query_t = "SELECT * FROM customers ORDER BY id DESC"
            cust_params_t = []
            if current_user_role_tour == 'sale' and current_user_name_tour:
                cust_query_t = "SELECT * FROM customers WHERE sale_name=? ORDER BY id DESC"
                cust_params_t.append(current_user_name_tour)
            customers = run_query(cust_query_t, tuple(cust_params_t))
            cust_opts_t = ["-- Kh√°ch m·ªõi --"] + [f"{c['name']} | {c['phone']}" for c in customers] if customers else ["-- Kh√°ch m·ªõi --"] # type: ignore
            sel_cust_t = st.selectbox("üîç G·ª£i √Ω kh√°ch h√†ng:", cust_opts_t, key="tour_cust_suggest")
            
            t_pre_name, t_pre_phone = "", ""
            if sel_cust_t and sel_cust_t != "-- Kh√°ch m·ªõi --":
                parts = sel_cust_t.split(" | ")
                t_pre_name = parts[0]
                t_pre_phone = parts[1] if len(parts) > 1 else ""

            with st.form("create_tour_form", clear_on_submit=True):
                c1, c2 = st.columns(2)
                t_name = c1.text_input("T√™n ƒêo√†n")
                t_sale = c2.text_input("Sales ph·ª• tr√°ch", value=current_user_name_tour, disabled=True)
                c_cust1, c_cust2 = st.columns(2)
                t_cust_name = c_cust1.text_input("T√™n Kh√°ch / ƒê·∫°i di·ªán", value=t_pre_name)
                t_cust_phone = c_cust2.text_input("SƒêT Kh√°ch", value=t_pre_phone)
                c3, c4, c5 = st.columns(3)
                t_start = c3.date_input("Ng√†y ƒëi", format="DD/MM/YYYY")
                t_end = c4.date_input("Ng√†y v·ªÅ", format="DD/MM/YYYY")
                t_pax = c5.number_input("S·ªë l∆∞·ª£ng kh√°ch", min_value=1, step=1)
                
                if st.form_submit_button("T·∫°o ƒêo√†n"):
                    if t_name:
                        save_customer_check(t_cust_name, t_cust_phone, current_user_name_tour)
                        new_tour_code = ''.join(random.choices(string.ascii_uppercase, k=5))
                        add_row_to_table('tours', {
                            'tour_name': t_name, 'sale_name': current_user_name_tour, 'start_date': t_start.strftime('%d/%m/%Y'),
                            'end_date': t_end.strftime('%d/%m/%Y'), 'guest_count': t_pax, 'created_at': datetime.now().strftime('%Y-%m-%d'),
                            'tour_code': new_tour_code, 'customer_name': t_cust_name, 'customer_phone': t_cust_phone
                        })
                        st.success(f"ƒê√£ t·∫°o ƒëo√†n m·ªõi! M√£ tour: {new_tour_code}. H√£y ch·ªçn ·ªü danh s√°ch b√™n d∆∞·ªõi ƒë·ªÉ l√†m d·ª± to√°n.")
                        time.sleep(1)
                        st.rerun()
                    else:
                        st.error("Vui l√≤ng nh·∫≠p t√™n ƒëo√†n.")

        st.divider()
        st.subheader("B·∫£ng T√≠nh D·ª± To√°n")
        
        selected_tour_label = st.selectbox("Ch·ªçn ƒêo√†n ƒë·ªÉ l√†m d·ª± to√°n:", list(tour_options.keys()) if tour_options else [], key="sel_tour_est")
        
        if selected_tour_label:
            tour_id = tour_options[selected_tour_label] # type: ignore
            tour_info = next((t for t in all_tours if t['id'] == tour_id), None)
            if not tour_info:
                st.error("Kh√¥ng t√¨m th·∫•y th√¥ng tin tour.")
                st.stop()
            assert tour_info is not None

            # --- TOOLBAR: S·ª¨A / X√ìA TOUR ---
            c_ren, c_del = st.columns(2)
            with c_ren:
                with st.popover("‚úèÔ∏è S·ª≠a th√¥ng tin", use_container_width=True):
                    with st.form(f"edit_tour_{tour_id}"):
                        en_n = st.text_input("T√™n ƒêo√†n", value=tour_info['tour_name']) # type: ignore
                        en_s = st.text_input("Sales", value=tour_info['sale_name']) # type: ignore
                        en_p = st.number_input("S·ªë kh√°ch", value=tour_info['guest_count'], min_value=1) # type: ignore
                        if st.form_submit_button("L∆∞u thay ƒë·ªïi"):
                            if en_n != tour_info['tour_name']: # type: ignore
                                run_query("UPDATE tours SET pending_name=?, sale_name=?, guest_count=? WHERE id=?", (en_n, en_s, en_p, tour_id), commit=True)
                                st.success("ƒê√£ c·∫≠p nh·∫≠t th√¥ng tin & G·ª≠i y√™u c·∫ßu ƒë·ªïi t√™n (Ch·ªù Admin duy·ªát)!"); time.sleep(0.5); st.rerun()
                            else:
                                run_query("UPDATE tours SET sale_name=?, guest_count=? WHERE id=?", (en_s, en_p, tour_id), commit=True)
                                st.success("ƒê√£ c·∫≠p nh·∫≠t!"); time.sleep(0.5); st.rerun()
            with c_del:
                req_status = tour_info['request_delete'] # type: ignore
                if req_status == 0:
                    with st.popover("üóëÔ∏è Y√™u c·∫ßu x√≥a", use_container_width=True):
                        st.warning(f"G·ª≠i y√™u c·∫ßu x√≥a ƒëo√†n: {tour_info['tour_name']}?") # type: ignore
                        if st.button("G·ª≠i y√™u c·∫ßu", type="primary", use_container_width=True, key=f"req_del_t_{tour_id}"):
                            run_query("UPDATE tours SET request_delete=1 WHERE id=?", (tour_id,), commit=True)
                            st.success("ƒê√£ g·ª≠i y√™u c·∫ßu x√≥a (Ch·ªù Admin duy·ªát)!"); time.sleep(0.5); st.rerun()
                elif req_status == 1:
                    st.warning("‚è≥ ƒêang ch·ªù Admin duy·ªát x√≥a...")
                    if st.button("H·ªßy y√™u c·∫ßu", key=f"cancel_req_{tour_id}", use_container_width=True): # type: ignore
                        run_query("UPDATE tours SET request_delete=0 WHERE id=?", (tour_id,), commit=True)
                        st.rerun()
                elif req_status == 2:
                    st.success("‚úÖ Admin ƒë√£ duy·ªát x√≥a!")
                    c_conf, c_can = st.columns(2)
                    if c_conf.button("üóëÔ∏è X√≥a ngay", type="primary", key=f"final_del_{tour_id}"): # type: ignore
                        run_query("DELETE FROM tours WHERE id=?", (tour_id,), commit=True)
                        run_query("DELETE FROM tour_items WHERE tour_id=?", (tour_id,), commit=True)
                        st.success("ƒê√£ x√≥a vƒ©nh vi·ªÖn!"); time.sleep(0.5); st.rerun()
                    if c_can.button("H·ªßy x√≥a", key=f"keep_tour_{tour_id}"):
                        run_query("UPDATE tours SET request_delete=0 WHERE id=?", (tour_id,), commit=True)
                        st.rerun()

            # Reset edit mode when changing tour
            if st.session_state.current_tour_id_est != tour_id:
                st.session_state.est_edit_mode = False
                st.session_state.current_tour_id_est = tour_id
                if "est_df_temp" in st.session_state: del st.session_state.est_df_temp
                st.session_state.est_editor_key += 1
            
            # --- IMPORT EXCEL (M·ªöI - D·ª∞ TO√ÅN) ---
            with st.expander("üì• Nh·∫≠p d·ªØ li·ªáu t·ª´ Excel (Import)", expanded=False):
                st.caption("üí° File Excel c·∫ßn c√≥ d√≤ng ti√™u ƒë·ªÅ: **H·∫°ng m·ª•c, Di·ªÖn gi·∫£i, ƒê∆°n v·ªã, ƒê∆°n gi√°, S·ªë l∆∞·ª£ng, S·ªë l·∫ßn**")
                
                # Widget upload file
                uploaded_est_file = st.file_uploader("Ch·ªçn file Excel d·ª± to√°n", type=["xlsx", "xls"], key="up_est_tool")
                
                if uploaded_est_file:
                    if st.button("üöÄ ƒê·ªçc file & ƒêi·ªÅn v√†o b·∫£ng", type="primary"):
                        try:
                            # 1. ƒê·ªçc file Excel (T√¨m d√≤ng ti√™u ƒë·ªÅ t·ª± ƒë·ªông)
                            uploaded_est_file.seek(0)
                            df_raw = pd.read_excel(uploaded_est_file, header=None)
                            
                            header_idx = 0
                            detect_kws = ['h·∫°ng m·ª•c', 't√™n h√†ng', 'di·ªÖn gi·∫£i', 'ƒë∆°n gi√°', 's·ªë l∆∞·ª£ng', 'th√†nh ti·ªÅn', 'item', 'price', 'qty', 'ƒëvt']
                            
                            # Qu√©t 15 d√≤ng ƒë·∫ßu ƒë·ªÉ t√¨m d√≤ng ch·ª©a nhi·ªÅu t·ª´ kh√≥a nh·∫•t
                            for i in range(min(15, len(df_raw))):
                                row_vals = [str(x).lower() for x in df_raw.iloc[i].tolist()]
                                if sum(1 for kw in detect_kws if any(kw in val for val in row_vals)) >= 2:
                                    header_idx = i
                                    break
                            
                            uploaded_est_file.seek(0)
                            df_in = pd.read_excel(uploaded_est_file, header=header_idx)
                            
                            # 2. Chu·∫©n h√≥a t√™n c·ªôt
                            # Chuy·ªÉn h·∫øt v·ªÅ ch·ªØ th∆∞·ªùng ƒë·ªÉ so s√°nh
                            df_in.columns = [str(c).lower().strip() for c in df_in.columns]
                            
                            # ƒê·ªãnh nghƒ©a c√°c t·ª´ kh√≥a (Aliases) cho t·ª´ng c·ªôt DB - ∆Øu ti√™n t·ª´ tr√°i sang ph·∫£i
                            col_aliases = {
                                'category': ['h·∫°ng m·ª•c', 'hang muc', 't√™n h√†ng', 'ten hang', 't√™n d·ªãch v·ª•', 'ten dich vu', 'n·ªôi dung', 'noi dung', 'item'],
                                'description': ['di·ªÖn gi·∫£i', 'dien giai', 'chi ti·∫øt', 'chi tiet', 'ghi ch√∫', 'ghi chu', 'm√¥ t·∫£', 'mo ta', 'description', 'desc'],
                                'unit': ['ƒë∆°n v·ªã', 'don vi', 'ƒëvt', 'dvt', 'unit', 'uom'],
                                'quantity': ['s·ªë l∆∞·ª£ng', 'so luong', 'sl', 'qty', 'quantity', 'vol'],
                                'unit_price': ['ƒë∆°n gi√°', 'don gia', 'gi√°', 'gia', 'price', 'unit_price', 'unit price'],
                                'times': ['s·ªë l·∫ßn', 'so lan', 'l·∫ßn', 'lan', 'times']
                            }
                            
                            # X√°c ƒë·ªãnh c·ªôt n√†o trong Excel map v√†o c·ªôt n√†o trong DB
                            final_col_map = {}
                            for db_col, aliases in col_aliases.items():
                                for alias in aliases:
                                    if alias in df_in.columns:
                                        final_col_map[db_col] = alias
                                        break
                            
                            new_data = []
                            if not final_col_map:
                                st.warning("‚ö†Ô∏è Kh√¥ng t√¨m th·∫•y c√°c c·ªôt th√¥ng tin c·∫ßn thi·∫øt (H·∫°ng m·ª•c, ƒê∆°n gi√°...). Vui l√≤ng ki·ªÉm tra t√™n c·ªôt trong file Excel.")
                            else:
                                for _, row in df_in.iterrows():
                                    item = {}
                                    for db_col, xls_col in final_col_map.items():
                                        val = row[xls_col]
                                        if pd.isna(val):
                                            val = 0 if db_col in ['quantity', 'unit_price', 'times'] else ""
                                        item[db_col] = val
                                    
                                    # Default values
                                    if 'category' not in item: item['category'] = ""
                                    if 'description' not in item: item['description'] = ""
                                    if 'unit' not in item: item['unit'] = ""
                                    
                                    # Safe numeric conversion
                                    def safe_float(v):
                                        try: return float(v)
                                        except: return 0.0
                                    
                                    item['quantity'] = safe_float(item.get('quantity', 1))
                                    item['unit_price'] = safe_float(item.get('unit_price', 0))
                                    item['times'] = safe_float(item.get('times', 1))
                                    if item['times'] == 0: item['times'] = 1
                                    
                                    if str(item['category']).strip() or str(item['description']).strip():
                                        new_data.append(item)

                            if new_data:
                                # 3. C·∫≠p nh·∫≠t v√†o Session State (Hi·ªÉn th·ªã l√™n m√†n h√¨nh)
                                st.session_state.est_df_temp = pd.DataFrame(new_data)
                                st.session_state.est_edit_mode = True # B·∫≠t ch·∫ø ƒë·ªô s·ª≠a ƒë·ªÉ hi·ªán n√∫t L∆∞u
                                st.success(f"ƒê√£ ƒë·ªçc th√†nh c√¥ng {len(new_data)} d√≤ng! Vui l√≤ng ki·ªÉm tra b·∫£ng b√™n d∆∞·ªõi v√† b·∫•m L∆ØU.")
                                time.sleep(1)
                                st.rerun()
                            else:
                                st.warning(f"Kh√¥ng ƒë·ªçc ƒë∆∞·ª£c d·ªØ li·ªáu! (ƒê√£ th·ª≠ d√≤ng {header_idx+1} l√†m ti√™u ƒë·ªÅ). Vui l√≤ng ki·ªÉm tra t√™n c·ªôt.")
                                
                        except Exception as e:
                            st.error(f"L·ªói khi ƒë·ªçc file: {str(e)}")

            # --- Fetch Items (EST) ---
            if "est_df_temp" not in st.session_state:
                existing_items = run_query("SELECT * FROM tour_items WHERE tour_id=? AND item_type='EST'", (tour_id,))
                if existing_items:
                    df_est = pd.DataFrame([dict(r) for r in existing_items])
                    if 'times' not in df_est.columns: df_est['times'] = 1.0
                    df_est = df_est[['category', 'description', 'unit', 'unit_price', 'quantity', 'times']]
                else:
                    df_est = pd.DataFrame([
                        {"category": "V·∫≠n chuy·ªÉn", "description": "Xe 16 ch·ªó", "unit": "Xe", "unit_price": 0, "quantity": 1, "times": 1},
                        {"category": "L∆∞u tr√∫", "description": "Kh√°ch s·∫°n 3 sao", "unit": "Ph√≤ng", "unit_price": 0, "quantity": 1, "times": 1},
                        {"category": "ƒÇn u·ªëng", "description": "B·ªØa tr∆∞a ng√†y 1", "unit": "Su·∫•t", "unit_price": 0, "quantity": 1, "times": 1},
                    ])
                st.session_state.est_df_temp = df_est

            # Prepare Display Data (T·∫°o b·∫£n sao ƒë·ªÉ hi·ªÉn th·ªã format ƒë·∫πp)
            df_display = st.session_state.est_df_temp.copy()
            # [NEW] T·∫°o s·ªë th·ª© t·ª± (STT) t·ª± ƒë·ªông
            df_display.index = pd.RangeIndex(start=1, stop=len(df_display) + 1)
            
            # [MODIFIED] T√≠nh Gi√°/Pax v√† ·∫©n c·ªôt Times
            guest_cnt = tour_info['guest_count'] if tour_info['guest_count'] else 1 # type: ignore
            df_display['total_val'] = df_display['quantity'] * df_display['unit_price'] * df_display['times']
            df_display['price_per_pax'] = df_display['total_val'] / guest_cnt
            
            df_display['price_per_pax'] = df_display['price_per_pax'].apply(lambda x: format_vnd(x) + " VND")
            df_display['total_display'] = df_display['total_val'].apply(lambda x: format_vnd(x) + " VND")
            df_display['unit_price'] = df_display['unit_price'].apply(lambda x: format_vnd(x) + " VND") # type: ignore

            st.markdown(f"**ƒêo√†n:** {tour_info['tour_name']} (M√£: {tour_info['tour_code']}) | **Pax:** {tour_info['guest_count']}")
            
            is_disabled = not st.session_state.est_edit_mode

            # --- DATA EDITOR ---
            edited_est = st.data_editor(
                df_display,
                disabled=is_disabled,
                num_rows="dynamic",
                column_config={
                    "_index": st.column_config.NumberColumn("STT", disabled=True),
                    "category": st.column_config.TextColumn("H·∫°ng m·ª•c chi ph√≠", required=False),
                    "description": st.column_config.TextColumn("Di·ªÖn gi·∫£i"),
                    "unit": st.column_config.TextColumn("ƒê∆°n v·ªã"),
                    "unit_price": st.column_config.TextColumn("ƒê∆°n gi√° (VND)", required=False),
                    "quantity": st.column_config.NumberColumn("S·ªë l∆∞·ª£ng", min_value=0),
                    "times": st.column_config.NumberColumn("S·ªë l·∫ßn", min_value=1),
                    "price_per_pax": st.column_config.TextColumn("Gi√°/Pax", disabled=True),
                    "total_display": st.column_config.TextColumn("T·ªïng chi ph√≠", disabled=True),
                    "total_val": st.column_config.NumberColumn("Hidden", disabled=True),
                },
                column_order=("category", "description", "unit", "unit_price", "quantity", "times", "price_per_pax", "total_display"),
                use_container_width=True,
                hide_index=False,
                key=f"editor_est_{st.session_state.est_editor_key}"
            )
            
            # --- AUTO-UPDATE CALCULATION ---
            if st.session_state.est_edit_mode:
                # T·ª± ƒë·ªông c·∫≠p nh·∫≠t khi d·ªØ li·ªáu thay ƒë·ªïi
                df_new = edited_est.copy()
                
                def clean_vnd_auto(x):
                    if isinstance(x, str):
                        return float(x.replace('.', '').replace(',', '').replace(' VND', '').strip())
                    return float(x) if x else 0.0
                
                df_new['unit_price'] = df_new['unit_price'].apply(clean_vnd_auto)
                df_new['quantity'] = pd.to_numeric(df_new['quantity'], errors='coerce').fillna(0)
                if 'times' not in df_new.columns: df_new['times'] = 1
                df_new['times'] = pd.to_numeric(df_new['times'], errors='coerce').fillna(1)
                
                # So s√°nh v·ªõi d·ªØ li·ªáu c≈©
                cols_check = ['category', 'description', 'unit', 'unit_price', 'quantity', 'times']
                df_old = st.session_state.est_df_temp.copy()
                if 'times' not in df_old.columns: df_old['times'] = 1
                
                # Reset index v√† fillna ƒë·ªÉ so s√°nh
                df_new_check = df_new[cols_check].reset_index(drop=True).fillna(0)
                df_old_check = df_old[cols_check].reset_index(drop=True).fillna(0)
                
                has_changes = False
                if len(df_new_check) != len(df_old_check): has_changes = True
                elif not df_new_check.equals(df_old_check): has_changes = True
                
                if has_changes:
                    st.session_state.est_df_temp = df_new[cols_check]
                    st.rerun()

            # --- T√çNH TO√ÅN REAL-TIME ---
            total_cost = 0
            if not edited_est.empty:
                # [FIX] Handle case where a cell is None, which becomes the string 'None' after astype(str)
                cleaned_prices_est = edited_est['unit_price'].astype(str).str.replace('.', '', regex=False).str.replace(' VND', '', regex=False).str.strip()
                p_price = cleaned_prices_est.apply(lambda x: float(x) if x and x.lower() != 'none' else 0.0)
                t_times = edited_est['times'].fillna(1) # type: ignore
                total_cost = (edited_est['quantity'] * p_price * t_times).sum()
            
            st.divider()
            
            # --- PH·∫¶N T√çNH L·ª¢I NHU·∫¨N & THU·∫æ (Y√äU C·∫¶U 2: S·∫Øp x·∫øp h√†ng ngang) ---
            c_cost, c_profit, c_tax = st.columns(3)
            
            with c_cost:
                st.metric("T·ªïng Chi Ph√≠ D·ª± To√°n", format_vnd(total_cost) + " VND") # type: ignore
            with c_profit:
                p_percent = st.number_input("L·ª£i Nhu·∫≠n Mong Mu·ªën (%)", value=float(tour_info['est_profit_percent']), step=0.5, key="p_pct", disabled=is_disabled) # type: ignore
            with c_tax:
                t_percent = st.number_input("Thu·∫ø VAT ƒê·∫ßu Ra (%)", value=float(tour_info['est_tax_percent']), step=1.0, key="t_pct", disabled=is_disabled) # type: ignore
            
            # C√¥ng th·ª©c: Gi√° B√°n = Chi Ph√≠ + L·ª£i Nhu·∫≠n + Thu·∫ø
            # L·ª£i nhu·∫≠n = Chi Ph√≠ * %
            # Thu·∫ø = (Chi Ph√≠ + L·ª£i Nhu·∫≠n) * %
            profit_amt = total_cost * (p_percent / 100)
            base_price = total_cost + profit_amt
            tax_amt = base_price * (t_percent / 100)
            final_price = base_price + tax_amt

            st.markdown(f"""<div class="finance-summary-card">
                <div class="row"><span>Ti·ªÅn L·ª£i Nhu·∫≠n ({p_percent}%):</span> <b>{format_vnd(profit_amt)} VND</b></div>
                <div class="row"><span>Ti·ªÅn Thu·∫ø ({t_percent}%):</span> <b>{format_vnd(tax_amt)} VND</b></div>
                <div class="row total-row"><span>T·ªîNG GI√Å B√ÅN D·ª∞ KI·∫æN:</span> <b>{format_vnd(final_price)} VND</b></div>
                <div class="pax-price">(Gi√° trung b√¨nh/kh√°ch: {format_vnd(final_price/tour_info['guest_count'] if tour_info['guest_count'] else 1)} VND)</div>
            </div>
            """, unsafe_allow_html=True)

            # --- TH√äM √î NH·∫¨P GI√Å CH·ªêT & GI√Å TR·∫∫ EM ---
            st.write("")
            t_dict: Dict[str, Any] = dict(tour_info) if tour_info else {}
            c_final_p, c_child_p = st.columns(2)
            with c_final_p:
                # Gi√° ch·ªët tour - Text Input for dots formatting
                cur_final_price = float(t_dict.get('final_tour_price', 0) or 0)
                cur_final_price_str = "{:,.0f}".format(cur_final_price).replace(",", ".")
                final_tour_price_input = st.text_input("Gi√° ch·ªët tour (VND)", value=cur_final_price_str, disabled=is_disabled, help="Nh·∫≠p s·ªë ti·ªÅn (VD: 1.000.000)")
                try: final_tour_price_val = float(final_tour_price_input.replace('.', '').replace(',', ''))
                except: final_tour_price_val = 0.0

                # S·ªë l∆∞·ª£ng ng∆∞·ªùi l·ªõn
                cur_qty = float(t_dict.get('final_qty', 0))
                if cur_qty == 0: cur_qty = float(t_dict.get('guest_count', 1))
                final_qty_val = st.number_input("S·ªë l∆∞·ª£ng ng∆∞·ªùi l·ªõn", value=cur_qty, min_value=0.0, step=1.0, disabled=is_disabled)

            with c_child_p:
                # Gi√° tr·∫ª em - Text Input
                cur_child_price = float(t_dict.get('child_price', 0) or 0)
                cur_child_price_str = "{:,.0f}".format(cur_child_price).replace(",", ".")
                child_price_input = st.text_input("Gi√° tr·∫ª em (VND)", value=cur_child_price_str, disabled=is_disabled)
                try: child_price_val = float(child_price_input.replace('.', '').replace(',', ''))
                except: child_price_val = 0.0

                cur_child_qty = float(t_dict.get('child_qty', 0))
                child_qty_val = st.number_input("S·ªë l∆∞·ª£ng tr·∫ª em", value=cur_child_qty, min_value=0.0, step=1.0, disabled=is_disabled)
            
            total_final_manual = (final_tour_price_val * final_qty_val) + (child_price_val * child_qty_val)
            st.markdown(f"""<div style="background-color: #e8f5e9; padding: 15px; border-radius: 10px; margin-top: 10px; border: 1px solid #c8e6c9;"><div style="display:flex; justify-content:space-between; font-size: 1.3em; color: #2e7d32;"><span><b>T·ªîNG DOANH THU</b></span> <b>{format_vnd(total_final_manual)} VND</b></div></div>""", unsafe_allow_html=True)

            # --- EXPORT EXCEL ---
            st.write("")
            df_exp = st.session_state.est_df_temp.copy()
            
            # Chu·∫©n h√≥a d·ªØ li·ªáu s·ªë
            def clean_price_exp(x): # type: ignore
                if isinstance(x, str):
                    return float(x.replace('.', '').replace(',', '').replace(' VND', '').strip())
                return float(x) if x else 0.0
            
            df_exp['unit_price'] = df_exp['unit_price'].apply(clean_price_exp)
            df_exp['quantity'] = pd.to_numeric(df_exp['quantity'], errors='coerce').fillna(0)
            if 'times' not in df_exp.columns: df_exp['times'] = 1
            df_exp['times'] = pd.to_numeric(df_exp['times'], errors='coerce').fillna(1)
            
            # T√≠nh to√°n c√°c c·ªôt hi·ªÉn th·ªã gi·ªëng Web
            df_exp['total_amount'] = df_exp['quantity'] * df_exp['unit_price'] * df_exp['times']
            g_cnt = tour_info['guest_count'] if tour_info['guest_count'] else 1 # type: ignore
            df_exp['price_per_pax'] = df_exp['total_amount'] / g_cnt
            
            # Ch·ªçn v√† ƒë·ªïi t√™n c·ªôt
            df_exp = df_exp[['category', 'description', 'unit', 'unit_price', 'quantity', 'times', 'price_per_pax', 'total_amount']]
            df_exp.columns = ['H·∫°ng m·ª•c', 'Di·ªÖn gi·∫£i', 'ƒê∆°n v·ªã', 'ƒê∆°n gi√°', 'S·ªë l∆∞·ª£ng', 'S·ªë l·∫ßn', 'Gi√°/Pax', 'T·ªïng chi ph√≠']

            buffer = io.BytesIO()
            file_ext = "xlsx"
            mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            try:
                with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer: # type: ignore
                        # Start table at row 11 (index 10) to leave space for info
                        start_row = 10
                        df_exp.to_excel(writer, index=False, sheet_name='DuToan', startrow=start_row)
                        
                        # --- FORMATTING (N·∫øu d√πng xlsxwriter) ---
                        workbook: Any = writer.book
                        worksheet = writer.sheets['DuToan']
                        
                        # --- STYLES ---
                        company_name_fmt = workbook.add_format({'bold': True, 'font_size': 14, 'font_color': '#1B5E20', 'font_name': 'Times New Roman'})
                        company_info_fmt = workbook.add_format({'font_size': 10, 'italic': True, 'font_color': '#424242', 'font_name': 'Times New Roman'})
                        
                        title_fmt = workbook.add_format({'bold': True, 'font_size': 18, 'align': 'center', 'valign': 'vcenter', 'font_color': '#0D47A1', 'bg_color': '#E3F2FD', 'border': 1, 'font_name': 'Times New Roman'})
                        section_fmt = workbook.add_format({'bold': True, 'font_size': 11, 'font_color': '#E65100', 'underline': True, 'font_name': 'Times New Roman'})
                        
                        header_fmt = workbook.add_format({'bold': True, 'fg_color': '#2E7D32', 'font_color': 'white', 'border': 1, 'align': 'center', 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                        body_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'text_wrap': True, 'font_size': 10, 'font_name': 'Times New Roman'})
                        body_center_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'align': 'center', 'font_size': 10, 'font_name': 'Times New Roman'})
                        money_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'font_size': 10, 'font_name': 'Times New Roman'})
                        
                        # Summary Section Styles
                        sum_header_bg_fmt = workbook.add_format({'bold': True, 'bg_color': '#FFF3E0', 'border': 1, 'font_color': '#E65100', 'align': 'center', 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                        sum_label_fmt = workbook.add_format({'bold': True, 'align': 'left', 'border': 1, 'bg_color': '#FAFAFA', 'font_name': 'Times New Roman'})
                        sum_val_fmt = workbook.add_format({'num_format': '#,##0', 'align': 'right', 'border': 1, 'font_name': 'Times New Roman'})
                        sum_val_bold_fmt = workbook.add_format({'bold': True, 'num_format': '#,##0', 'align': 'right', 'border': 1, 'font_name': 'Times New Roman'})
                        sum_total_fmt = workbook.add_format({'bold': True, 'bg_color': '#C8E6C9', 'font_color': '#1B5E20', 'num_format': '#,##0', 'align': 'right', 'border': 1, 'font_size': 12, 'font_name': 'Times New Roman'})
                        
                        # --- 1. COMPANY INFO (Rows 0-3) ---
                        if comp['logo_b64_str']:
                            try:
                                logo_data = base64.b64decode(comp['logo_b64_str'])
                                image_stream = io.BytesIO(logo_data)
                                img = Image.open(image_stream)
                                w, h = img.size
                                scale = 60 / h if h > 0 else 0.5
                                image_stream.seek(0)
                                worksheet.insert_image('A1', 'logo.png', {'image_data': image_stream, 'x_scale': scale, 'y_scale': scale, 'x_offset': 5, 'y_offset': 5})
                            except: pass
                        
                        worksheet.write('B1', comp['name'], company_name_fmt)
                        worksheet.write('B2', f"ƒêC: {comp['address']}", company_info_fmt)
                        worksheet.write('B3', f"MST: {comp['phone']}", company_info_fmt)
                        
                        # --- 2. TOUR INFO (Rows 4-9) ---
                        worksheet.merge_range('A5:G5', "B·∫¢NG D·ª∞ TO√ÅN CHI PH√ç TOUR", title_fmt)
                        
                        # Info Data
                        t_info_dict = dict(tour_info) if tour_info else {}
                        t_name = t_info_dict.get('tour_name', '')
                        t_code = t_info_dict.get('tour_code', '')
                        t_sale = t_info_dict.get('sale_name', '')
                        t_start = t_info_dict.get('start_date', '')
                        t_end = t_info_dict.get('end_date', '')
                        t_cust = t_info_dict.get('customer_name', '')
                        t_phone = t_info_dict.get('customer_phone', '')
                        t_guest = t_info_dict.get('guest_count', 0)
                        
                        # Layout Info nicely
                        worksheet.write('A7', "T√™n ƒëo√†n:", sum_label_fmt)
                        worksheet.merge_range('B7:D7', t_name, sum_val_fmt)
                        worksheet.write('E7', "M√£ ƒëo√†n:", sum_label_fmt)
                        worksheet.merge_range('F7:G7', t_code, sum_val_fmt)
                        
                        worksheet.write('A8', "Kh√°ch h√†ng:", sum_label_fmt)
                        worksheet.merge_range('B8:D8', f"{t_cust} - {t_phone}", sum_val_fmt)
                        worksheet.write('E8', "Sales:", sum_label_fmt)
                        worksheet.merge_range('F8:G8', t_sale, sum_val_fmt)
                        
                        worksheet.write('A9', "Th·ªùi gian:", sum_label_fmt)
                        worksheet.merge_range('B9:D9', f"{t_start} - {t_end}", sum_val_fmt)
                        worksheet.write('E9', "S·ªë kh√°ch:", sum_label_fmt)
                        worksheet.merge_range('F9:G9', t_guest, sum_val_fmt)

                        # --- 3. TABLE HEADER & BODY ---
                        # Apply Header
                        for col_num, value in enumerate(df_exp.columns):
                            worksheet.write(start_row, col_num, value, header_fmt)
                        
                        # Apply Body
                        for row in range(len(df_exp)):
                            for col in range(len(df_exp.columns)):
                                val = df_exp.iloc[row, col]
                                # Cols: 0=Cat, 1=Desc, 2=Unit, 3=Price, 4=Qty, 5=PaxPrice, 6=Total
                                if col == 2: fmt = body_center_fmt # Unit centered
                                elif col in [3, 4, 5, 6, 7]: fmt = money_fmt # Money columns
                                else: fmt = body_fmt
                                
                                if pd.isna(val): val = ""
                                worksheet.write(row+start_row+1, col, val, fmt)
                        
                        # --- 4. SUMMARY SECTION ---
                        last_row = start_row + 1 + len(df_exp)
                        sum_row = last_row + 2
                        
                        # --- B·∫¢NG T√çNH GI√Å TH√ÄNH ---
                        worksheet.merge_range(sum_row, 0, sum_row, 3, "PH√ÇN T√çCH GI√Å TH√ÄNH & L·ª¢I NHU·∫¨N", sum_header_bg_fmt)
                        
                        # D√≤ng 1: T·ªïng chi ph√≠
                        worksheet.write(sum_row+1, 0, "1. T·ªïng chi ph√≠ d·ª± to√°n:", sum_label_fmt)
                        worksheet.merge_range(sum_row+1, 1, sum_row+1, 3, total_cost, sum_val_bold_fmt)
                        
                        # D√≤ng 2: L·ª£i nhu·∫≠n
                        worksheet.write(sum_row+2, 0, "2. L·ª£i nhu·∫≠n mong mu·ªën:", sum_label_fmt)
                        worksheet.write(sum_row+2, 1, f"{p_percent:g}%", body_center_fmt)
                        worksheet.merge_range(sum_row+2, 2, sum_row+2, 3, profit_amt, sum_val_fmt)
                        
                        # D√≤ng 3: Thu·∫ø
                        worksheet.write(sum_row+3, 0, "3. Thu·∫ø VAT:", sum_label_fmt)
                        worksheet.write(sum_row+3, 1, f"{t_percent:g}%", body_center_fmt)
                        worksheet.merge_range(sum_row+3, 2, sum_row+3, 3, tax_amt, sum_val_fmt)
                        
                        # D√≤ng 4: Gi√° t√≠nh to√°n
                        worksheet.write(sum_row+4, 0, "4. Gi√° b√°n t√≠nh to√°n:", sum_label_fmt)
                        worksheet.merge_range(sum_row+4, 1, sum_row+4, 3, final_price, sum_total_fmt)
                        
                        # --- B·∫¢NG CH·ªêT GI√Å B√ÅN ---
                        # ƒê·∫∑t b√™n ph·∫£i b·∫£ng gi√° th√†nh (C·ªôt E, F, G)
                        worksheet.merge_range(sum_row, 4, sum_row, 6, "B·∫¢NG CH·ªêT GI√Å B√ÅN TH·ª∞C T·∫æ", sum_header_bg_fmt)
                        
                        # Ng∆∞·ªùi l·ªõn
                        worksheet.write(sum_row+1, 4, "Ng∆∞·ªùi l·ªõn:", sum_label_fmt)
                        worksheet.write(sum_row+1, 5, final_qty_val, sum_val_fmt) # SL
                        worksheet.write(sum_row+1, 6, final_tour_price_val, sum_val_fmt) # Gi√°
                        
                        # Tr·∫ª em
                        worksheet.write(sum_row+2, 4, "Tr·∫ª em:", sum_label_fmt)
                        worksheet.write(sum_row+2, 5, child_qty_val, sum_val_fmt) # SL
                        worksheet.write(sum_row+2, 6, child_price_val, sum_val_fmt) # Gi√°
                        
                        # T·ªïng doanh thu
                        worksheet.write(sum_row+4, 4, "T·ªîNG DOANH THU:", sum_label_fmt)
                        worksheet.merge_range(sum_row+4, 5, sum_row+4, 6, total_final_manual, sum_total_fmt)

                        # Column Widths
                        worksheet.set_column('A:A', 25) # Category
                        worksheet.set_column('B:B', 40) # Desc
                        worksheet.set_column('C:C', 10) # Unit
                        worksheet.set_column('D:G', 18) # Numbers
            except Exception as e:
                # If xlsxwriter fails, fall back to a simple CSV export
                buffer.seek(0)
                buffer.truncate()
                df_exp.to_csv(buffer, index=False, encoding='utf-8-sig')
                file_ext = "csv"
                mime_type = "text/csv"
                st.error(f"‚ö†Ô∏è L·ªói khi t·∫°o file Excel: {e}. ƒê√£ chuy·ªÉn sang xu·∫•t file CSV.")
                st.info("üí° G·ª£i √Ω: N·∫øu b·∫°n v·ª´a c√†i th∆∞ vi·ªán, h√£y T·∫ÆT H·∫≤N ·ª©ng d·ª•ng (Ctrl+C t·∫°i terminal) v√† ch·∫°y l·∫°i l·ªánh `streamlit run app.py`.")

            clean_t_name = re.sub(r'[\\/*?:"<>|]', "", tour_info['tour_name'] if tour_info else "Tour") # type: ignore
            st.download_button(label=f"üì• T·∫£i B·∫£ng D·ª± To√°n ({file_ext.upper()})", data=buffer.getvalue(), file_name=f"DuToan_{clean_t_name}.{file_ext}", mime=mime_type, use_container_width=True)

            # --- N√∫t Ch·ªânh s·ª≠a / L∆∞u ---
            if st.session_state.est_edit_mode:
                if st.button("üíæ L∆∞u v√† chuy·ªÉn sang ph·∫ßn Danh s√°ch v√† d·ªãch v·ª•", type="primary", use_container_width=True):
                    # 1. Update Tour Meta
                    run_query("UPDATE tours SET est_profit_percent=?, est_tax_percent=?, final_tour_price=?, child_price=?, final_qty=?, child_qty=? WHERE id=?", (p_percent, t_percent, final_tour_price_val, child_price_val, final_qty_val, child_qty_val, tour_id), commit=True)
                    
                    # 2. Update Tour Items (X√≥a c≈© th√™m m·ªõi)
                    run_query("DELETE FROM tour_items WHERE tour_id=? AND item_type='EST'", (tour_id,), commit=True)

                    data_to_insert = []
                    query = """INSERT INTO tour_items (tour_id, item_type, category, description, unit, quantity, unit_price, total_amount, times)
                               VALUES (?, 'EST', ?, ?, ?, ?, ?, ?, ?)"""

                    for _, row in edited_est.iterrows():
                        if row['category'] or row['description']: # type: ignore
                            # X·ª≠ l√Ω d·ªØ li·ªáu
                            u_price = float(str(row['unit_price']).replace('.', '').replace(' VND', '').strip()) if row['unit_price'] else 0 # type: ignore
                            t_times = row.get('times', 1) # type: ignore
                            if pd.isna(t_times): t_times = 1
                            total_row = row['quantity'] * u_price * t_times # type: ignore
                            
                            # Th√™m v√†o danh s√°ch ch·ªù (ch∆∞a ghi ngay)
                            data_to_insert.append((
                                tour_id, 
                                row['category'], 
                                row['description'], 
                                row['unit'], 
                                row['quantity'],  # type: ignore
                                u_price, 
                                total_row, 
                                t_times
                            ))

                    # Ghi t·∫•t c·∫£ trong 1 l·∫ßn b·∫Øn
                    if data_to_insert:
                        run_query_many(query, data_to_insert)

                    if "est_df_temp" in st.session_state: del st.session_state.est_df_temp
                    st.session_state.est_edit_mode = False
                    st.success("ƒê√£ l∆∞u d·ª± to√°n th√†nh c√¥ng! H√£y chuy·ªÉn sang tab 'Danh s√°ch & D·ªãch v·ª•'.")
                    time.sleep(1); st.rerun()
            else:
                if st.button("‚úèÔ∏è Ch·ªânh s·ª≠a D·ª± to√°n", use_container_width=True):
                    st.session_state.est_edit_mode = True
                    st.rerun()

    # ---------------- TAB M·ªöI: DANH S√ÅCH & D·ªäCH V·ª§ ----------------
    with tab_list_srv:
        st.subheader("üìã Danh s√°ch ƒêo√†n & D·ªãch v·ª•")
        
        def clean_vnd_val(x):
            if isinstance(x, (int, float)): return float(x)
            if isinstance(x, str):
                return float(x.replace('.', '').replace(',', '').replace(' VND', '').strip()) if x.strip() else 0.0
            return 0.0

        def clean_df_cols(df, cols):
            for col in cols:
                if col in df.columns:
                    df[col] = df[col].apply(clean_vnd_val)
            return df

        selected_tour_ls_label = st.selectbox("Ch·ªçn ƒêo√†n:", list(tour_options.keys()) if tour_options else [], key="sel_tour_ls")
        
        if selected_tour_ls_label:
            tour_id_ls = tour_options[selected_tour_ls_label]
            tour_info_ls = next((t for t in all_tours if t['id'] == tour_id_ls), None)
            
            if tour_info_ls:
                tour_info_ls = dict(tour_info_ls)
                
                # --- SESSION STATE INIT FOR LIST & SERVICES (AUTO CALCULATION) ---
                if "current_tour_id_ls" not in st.session_state or st.session_state.current_tour_id_ls != tour_id_ls or "ls_incurred_temp" not in st.session_state:
                    st.session_state.current_tour_id_ls = tour_id_ls
                    
                    # Load Hotels
                    hotels = run_query("SELECT * FROM tour_hotels WHERE tour_id=?", (tour_id_ls,))
                    df_h = pd.DataFrame([dict(r) for r in hotels]) if hotels else pd.DataFrame(columns=['hotel_name', 'address', 'phone', 'total_rooms', 'room_type', 'total_amount', 'deposit'])
                    if df_h.empty: df_h = pd.DataFrame(columns=['hotel_name', 'address', 'phone', 'total_rooms', 'room_type', 'total_amount', 'deposit'])
                    for col in ['total_amount', 'deposit']:
                        if col not in df_h.columns: df_h[col] = 0.0
                    st.session_state.ls_hotels_temp = df_h[['hotel_name', 'address', 'phone', 'total_rooms', 'room_type', 'total_amount', 'deposit']]

                    # Load Restaurants
                    rests = run_query("SELECT * FROM tour_restaurants WHERE tour_id=?", (tour_id_ls,))
                    if not rests:
                        est_meals = run_query("SELECT description FROM tour_items WHERE tour_id=? AND item_type='EST' AND (category LIKE '%ƒÇn%' OR description LIKE '%B·ªØa%' OR description LIKE '%ƒÇn%')", (tour_id_ls,))
                        if est_meals:
                            df_r = pd.DataFrame([{'date': '', 'meal_name': m['description'], 'restaurant_name': '', 'address': '', 'phone': '', 'menu': '', 'total_amount': 0, 'deposit': 0} for m in est_meals])
                        else:
                            df_r = pd.DataFrame(columns=['date', 'meal_name', 'restaurant_name', 'address', 'phone', 'menu', 'total_amount', 'deposit'])
                    else:
                        df_r = pd.DataFrame([dict(r) for r in rests])
                    
                    for col in ['total_amount', 'deposit']:
                        if col not in df_r.columns: df_r[col] = 0.0
                    if 'date' not in df_r.columns: df_r['date'] = ''
                    st.session_state.ls_rests_temp = df_r[['date', 'meal_name', 'restaurant_name', 'address', 'phone', 'menu', 'total_amount', 'deposit']]

                    # Load Sightseeings
                    sights = run_query("SELECT * FROM tour_sightseeings WHERE tour_id=?", (tour_id_ls,))
                    df_s = pd.DataFrame([dict(r) for r in sights]) if sights else pd.DataFrame(columns=['date', 'name', 'address', 'quantity', 'total_amount', 'deposit', 'note'])
                    if df_s.empty: df_s = pd.DataFrame(columns=['date', 'name', 'address', 'quantity', 'total_amount', 'deposit', 'note'])
                    if 'date' not in df_s.columns: df_s['date'] = ''
                    if 'total_amount' not in df_s.columns: df_s['total_amount'] = 0.0
                    if 'deposit' not in df_s.columns: df_s['deposit'] = 0.0
                    st.session_state.ls_sight_temp = df_s[['date', 'name', 'address', 'quantity', 'total_amount', 'deposit', 'note']]

                    # Load Incurred Costs
                    incurred = run_query("SELECT * FROM tour_incurred_costs WHERE tour_id=?", (tour_id_ls,))
                    df_inc = pd.DataFrame([dict(r) for r in incurred]) if incurred else pd.DataFrame(columns=['name', 'unit', 'quantity', 'price', 'total_amount', 'deposit', 'note'])
                    if df_inc.empty: df_inc = pd.DataFrame(columns=['name', 'unit', 'quantity', 'price', 'total_amount', 'deposit', 'note'])
                    st.session_state.ls_incurred_temp = df_inc[['name', 'unit', 'quantity', 'price', 'total_amount', 'deposit', 'note']]

                # 1. DANH S√ÅCH ƒêO√ÄN
                st.markdown("##### 1. Danh s√°ch ƒëo√†n")
                
                # --- [NEW] TH√îNG TIN B√ÄN GIAO & ƒê√ìN TI·ªÑN ---
                with st.expander("üöå Th√¥ng Tin B√†n Giao & ƒê√≥n Ti·ªÖn (ƒêi·ªÅu h√†nh)", expanded=False):
                    with st.form(f"handover_form_{tour_id_ls}"):
                        c_h1, c_h2, c_h3 = st.columns(3)
                        pk_loc = c_h1.text_input("ƒêi·ªÉm ƒë√≥n", value=tour_info_ls.get('pickup_location', ''))
                        pk_time = c_h2.text_input("Th·ªùi gian ƒë√≥n", value=tour_info_ls.get('pickup_time', ''))
                        fl_code = c_h3.text_area("Chuy·∫øn bay/T√†u", value=tour_info_ls.get('flight_code', ''), height=68)
                        
                        c_d1, c_d2, c_d3, c_d4 = st.columns(4)
                        drv_name = c_d1.text_input("T√™n l√°i xe", value=tour_info_ls.get('driver_name', ''))
                        drv_phone = c_d2.text_input("SƒêT L√°i xe", value=tour_info_ls.get('driver_phone', ''))
                        car_plate = c_d3.text_input("Bi·ªÉn s·ªë xe", value=tour_info_ls.get('car_plate', ''))
                        car_type = c_d4.text_input("Lo·∫°i xe", value=tour_info_ls.get('car_type', ''))
                        
                        c_g1, c_g2 = st.columns(2)
                        gd_name = c_g1.text_input("T√™n HDV", value=tour_info_ls.get('guide_name', ''))
                        gd_phone = c_g2.text_input("SƒêT HDV", value=tour_info_ls.get('guide_phone', ''))

                        if st.form_submit_button("üíæ L∆∞u th√¥ng tin v·∫≠n h√†nh"):
                            run_query("""UPDATE tours SET 
                                pickup_location=?, pickup_time=?, flight_code=?, 
                                driver_name=?, driver_phone=?, car_plate=?, car_type=?, 
                                guide_name=?, guide_phone=? WHERE id=?""",
                                (pk_loc, pk_time, fl_code, drv_name, drv_phone, car_plate, car_type, gd_name, gd_phone, tour_id_ls), commit=True)
                            st.success("ƒê√£ c·∫≠p nh·∫≠t th√¥ng tin ƒëi·ªÅu h√†nh!"); time.sleep(0.5); st.rerun()

                    st.markdown("##### üìÖ L·ªãch tr√¨nh chi ti·∫øt")
                    try:
                        s_raw = tour_info_ls.get('start_date', '')
                        e_raw = tour_info_ls.get('end_date', '')
                        
                        def try_parse_date(d_str):
                            if not d_str: return None
                            for fmt in ('%d/%m/%Y', '%Y-%m-%d', '%d-%m-%Y'):
                                try: return datetime.strptime(str(d_str).strip(), fmt)
                                except: continue
                            return None

                        s_d = try_parse_date(s_raw)
                        e_d = try_parse_date(e_raw)

                        if s_d and e_d:
                            num_days = (e_d - s_d).days + 1
                            if num_days < 1: num_days = 1
                        else:
                            num_days = 1
                            s_d = datetime.now()
                    except: 
                        num_days = 1
                        s_d = datetime.now()
                    
                    exist_itin = run_query("SELECT * FROM tour_itineraries WHERE tour_id=? ORDER BY day_index", (tour_id_ls,))
                    itin_map = {r['day_index']: r['content'] for r in exist_itin} if exist_itin else {}
                    
                    itin_data = []
                    for i in range(num_days):
                        d_str = (s_d + pd.Timedelta(days=i)).strftime('%d/%m/%Y')
                        itin_data.append({
                            "day_label": f"Ng√†y {i+1} ({d_str})",
                            "content": itin_map.get(i, ""),
                            "day_index": i
                        })
                    
                    df_itin = pd.DataFrame(itin_data)
                    edited_itin = st.data_editor(
                        df_itin,
                        column_config={
                            "day_label": st.column_config.TextColumn("Ng√†y", disabled=True),
                            "content": st.column_config.TextColumn("N·ªôi dung l·ªãch tr√¨nh", width="large"),
                            "day_index": st.column_config.NumberColumn("Hidden", disabled=True)
                        },
                        column_order=("day_label", "content"),
                        use_container_width=True,
                        hide_index=True,
                        key=f"itin_ed_{tour_id_ls}"
                    )
                    
                    if st.button("üíæ L∆∞u l·ªãch tr√¨nh", key=f"save_itin_{tour_id_ls}"):
                        run_query("DELETE FROM tour_itineraries WHERE tour_id=?", (tour_id_ls,), commit=True)
                        data_itin = [(tour_id_ls, r['day_index'], r['content']) for _, r in edited_itin.iterrows()]
                        if data_itin:
                            run_query_many("INSERT INTO tour_itineraries (tour_id, day_index, content) VALUES (?,?,?)", data_itin)
                        st.success("ƒê√£ l∆∞u l·ªãch tr√¨nh!"); time.sleep(0.5); st.rerun()

                guests = run_query("SELECT * FROM tour_guests WHERE tour_id=?", (tour_id_ls,))
                df_guests = pd.DataFrame([dict(r) for r in guests]) if guests else pd.DataFrame(columns=['name', 'dob', 'hometown', 'cccd', 'type'])
                if df_guests.empty:
                    df_guests = pd.DataFrame(columns=['name', 'dob', 'hometown', 'cccd', 'type'])
                else:
                    df_guests = df_guests[['name', 'dob', 'hometown', 'cccd', 'type']]

                edited_guests = st.data_editor(
                    df_guests,
                    num_rows="dynamic",
                    key="guest_editor",
                    column_config={
                        "name": st.column_config.TextColumn("H·ªç v√† t√™n", required=True),
                        "dob": st.column_config.TextColumn("Ng√†y sinh"),
                        "hometown": st.column_config.TextColumn("Qu√™ qu√°n"),
                        "cccd": st.column_config.TextColumn("S·ªë CCCD"),
                        "type": st.column_config.SelectboxColumn("Ph√¢n lo·∫°i", options=["Kh√°ch", "N·ªôi b·ªô", "HDV"], required=True)
                    },
                    use_container_width=True
                )

                # 2. DANH S√ÅCH PH√íNG KH√ÅCH S·∫†N
                st.markdown("##### 2. Danh s√°ch ph√≤ng Kh√°ch s·∫°n")
                df_hotels = st.session_state.ls_hotels_temp.copy()
                
                # Calculate remaining for display
                df_hotels['total_amount'] = pd.to_numeric(df_hotels['total_amount'], errors='coerce').fillna(0)
                df_hotels['deposit'] = pd.to_numeric(df_hotels['deposit'], errors='coerce').fillna(0)
                df_hotels['remaining'] = df_hotels['total_amount'].fillna(0) - df_hotels['deposit'].fillna(0)

                # Format hi·ªÉn th·ªã ti·ªÅn t·ªá
                df_hotels['total_amount'] = df_hotels['total_amount'].apply(lambda x: format_vnd(x) + " VND")
                df_hotels['deposit'] = df_hotels['deposit'].apply(lambda x: format_vnd(x) + " VND")
                df_hotels['remaining'] = df_hotels['remaining'].apply(lambda x: format_vnd(x) + " VND")

                edited_hotels = st.data_editor(
                    df_hotels,
                    num_rows="dynamic",
                    key="hotel_editor",
                    column_config={
                        "hotel_name": st.column_config.TextColumn("T√™n Kh√°ch s·∫°n", required=True),
                        "address": "ƒê·ªãa ch·ªâ",
                        "phone": "SƒêT",
                        "total_rooms": st.column_config.TextColumn("T·ªïng s·ªë ph√≤ng"),
                        "room_type": st.column_config.TextColumn("Lo·∫°i ph√≤ng"),
                        "total_amount": st.column_config.TextColumn("T·ªïng ti·ªÅn"),
                        "deposit": st.column_config.TextColumn("ƒê√£ ·ª©ng/c·ªçc"),
                        "remaining": st.column_config.TextColumn("C√≤n l·∫°i (Guide tr·∫£)", disabled=True)
                    },
                    use_container_width=True
                )
                
                # L√†m s·∫°ch d·ªØ li·ªáu sau khi edit (Chuy·ªÉn v·ªÅ s·ªë)
                edited_hotels = clean_df_cols(edited_hotels, ['total_amount', 'deposit'])

                # Detect changes for Hotels
                cols_h = ['hotel_name', 'address', 'phone', 'total_rooms', 'room_type', 'total_amount', 'deposit']
                if not edited_hotels[cols_h].equals(st.session_state.ls_hotels_temp[cols_h]):
                    st.session_state.ls_hotels_temp = edited_hotels[cols_h]
                    st.rerun()

                # 3. MENU NH√Ä H√ÄNG
                st.markdown("##### 3. Menu nh√† h√†ng")
                df_rests = st.session_state.ls_rests_temp.copy()
                
                df_rests['total_amount'] = pd.to_numeric(df_rests['total_amount'], errors='coerce').fillna(0)
                df_rests['deposit'] = pd.to_numeric(df_rests['deposit'], errors='coerce').fillna(0)
                df_rests['remaining'] = df_rests['total_amount'].fillna(0) - df_rests['deposit'].fillna(0)

                # Format hi·ªÉn th·ªã ti·ªÅn t·ªá
                df_rests['total_amount'] = df_rests['total_amount'].apply(lambda x: format_vnd(x) + " VND")
                df_rests['deposit'] = df_rests['deposit'].apply(lambda x: format_vnd(x) + " VND")
                df_rests['remaining'] = df_rests['remaining'].apply(lambda x: format_vnd(x) + " VND")

                edited_rests = st.data_editor(
                    df_rests,
                    num_rows="dynamic",
                    key="rest_editor",
                    column_config={
                        "date": st.column_config.TextColumn("Ng√†y"),
                        "meal_name": st.column_config.TextColumn("B·ªØa ƒÉn (D·ª± to√°n)", required=True),
                        "restaurant_name": "T√™n nh√† h√†ng",
                        "address": "ƒê·ªãa ch·ªâ",
                        "phone": "SƒêT",
                        "menu": st.column_config.TextColumn("Th·ª±c ƒë∆°n", width="large"),
                        "total_amount": st.column_config.TextColumn("T·ªïng ti·ªÅn"),
                        "deposit": st.column_config.TextColumn("ƒê√£ ·ª©ng/c·ªçc"),
                        "remaining": st.column_config.TextColumn("C√≤n l·∫°i (Guide tr·∫£)", disabled=True)
                    },
                    column_order=("date", "meal_name", "restaurant_name", "address", "phone", "menu", "total_amount", "deposit", "remaining"),
                    use_container_width=True
                )
                
                # L√†m s·∫°ch d·ªØ li·ªáu sau khi edit
                edited_rests = clean_df_cols(edited_rests, ['total_amount', 'deposit'])

                cols_r = ['date', 'meal_name', 'restaurant_name', 'address', 'phone', 'menu', 'total_amount', 'deposit']
                if not edited_rests[cols_r].equals(st.session_state.ls_rests_temp[cols_r]):
                    st.session_state.ls_rests_temp = edited_rests[cols_r]
                    st.rerun()

                # 4. ƒêI·ªÇM THAM QUAN (M·ªöI)
                st.markdown("##### 4. ƒêi·ªÉm tham quan")
                df_sightseeings = st.session_state.ls_sight_temp.copy()
                
                df_sightseeings['total_amount'] = pd.to_numeric(df_sightseeings['total_amount'], errors='coerce').fillna(0)
                df_sightseeings['deposit'] = pd.to_numeric(df_sightseeings['deposit'], errors='coerce').fillna(0)
                df_sightseeings['remaining'] = df_sightseeings['total_amount'].fillna(0) - df_sightseeings['deposit'].fillna(0)

                # Format hi·ªÉn th·ªã ti·ªÅn t·ªá
                df_sightseeings['total_amount'] = df_sightseeings['total_amount'].apply(lambda x: format_vnd(x) + " VND")
                df_sightseeings['deposit'] = df_sightseeings['deposit'].apply(lambda x: format_vnd(x) + " VND")
                df_sightseeings['remaining'] = df_sightseeings['remaining'].apply(lambda x: format_vnd(x) + " VND")

                edited_sightseeings = st.data_editor(
                    df_sightseeings,
                    num_rows="dynamic",
                    key="sightseeing_editor",
                    column_config={
                        "date": st.column_config.TextColumn("Ng√†y"),
                        "name": st.column_config.TextColumn("T√™n ƒë·ªãa ƒëi·ªÉm", required=True),
                        "address": "ƒê·ªãa ch·ªâ",
                        "quantity": st.column_config.NumberColumn("S·ªë l∆∞·ª£ng", min_value=0),
                        "total_amount": st.column_config.TextColumn("T·ªïng ti·ªÅn"),
                        "deposit": st.column_config.TextColumn("ƒê√£ c·ªçc"),
                        "remaining": st.column_config.TextColumn("C√≤n l·∫°i", disabled=True),
                        "note": st.column_config.TextColumn("L∆∞u √Ω")
                    },
                    column_order=("date", "name", "address", "quantity", "total_amount", "deposit", "remaining", "note"),
                    use_container_width=True
                )
                
                # L√†m s·∫°ch d·ªØ li·ªáu sau khi edit
                edited_sightseeings = clean_df_cols(edited_sightseeings, ['total_amount', 'deposit'])

                cols_s = ['date', 'name', 'address', 'quantity', 'total_amount', 'deposit', 'note']
                if not edited_sightseeings[cols_s].equals(st.session_state.ls_sight_temp[cols_s]):
                    st.session_state.ls_sight_temp = edited_sightseeings[cols_s]
                    st.rerun()

                # 5. CHI PH√ç PH√ÅT SINH (ƒê√£ ƒë·ªïi th·ª© t·ª± l√™n tr√™n)
                st.divider()
                st.markdown("##### 5. Chi ph√≠ ph√°t sinh (N∆∞·ªõc, Sim, Banner...)")
                df_incurred = st.session_state.ls_incurred_temp.copy()
                
                # Clean numbers for display
                df_incurred['price'] = pd.to_numeric(df_incurred['price'], errors='coerce').fillna(0)
                df_incurred['quantity'] = pd.to_numeric(df_incurred['quantity'], errors='coerce').fillna(0)
                df_incurred['total_amount'] = df_incurred['price'] * df_incurred['quantity']
                df_incurred['deposit'] = pd.to_numeric(df_incurred['deposit'], errors='coerce').fillna(0)
                df_incurred['remaining'] = df_incurred['total_amount'] - df_incurred['deposit']

                # Format
                df_incurred['price'] = df_incurred['price'].apply(lambda x: format_vnd(x) + " VND").astype(str)
                df_incurred['total_amount'] = df_incurred['total_amount'].apply(lambda x: format_vnd(x) + " VND").astype(str)
                df_incurred['deposit'] = df_incurred['deposit'].apply(lambda x: format_vnd(x) + " VND").astype(str)
                df_incurred['remaining'] = df_incurred['remaining'].apply(lambda x: format_vnd(x) + " VND").astype(str)

                edited_incurred = st.data_editor(
                    df_incurred,
                    num_rows="dynamic",
                    key="incurred_editor",
                    column_config={
                        "name": st.column_config.TextColumn("T√™n chi ph√≠", required=True),
                        "unit": st.column_config.TextColumn("ƒêVT"),
                        "quantity": st.column_config.NumberColumn("S·ªë l∆∞·ª£ng", min_value=0),
                        "price": st.column_config.TextColumn("ƒê∆°n gi√°"),
                        "total_amount": st.column_config.TextColumn("Th√†nh ti·ªÅn", disabled=True),
                        "deposit": st.column_config.TextColumn("ƒê√£ ·ª©ng/c·ªçc"),
                        "remaining": st.column_config.TextColumn("C√≤n l·∫°i", disabled=True),
                        "note": st.column_config.TextColumn("Ghi ch√∫")
                    },
                    column_order=("name", "unit", "quantity", "price", "total_amount", "deposit", "remaining", "note"),
                    use_container_width=True
                )
                
                # Clean data back
                def clean_vnd_val_inc(x):
                    if isinstance(x, (int, float)): return float(x)
                    if isinstance(x, str):
                        return float(x.replace('.', '').replace(',', '').replace(' VND', '').strip()) if x.strip() else 0.0
                    return 0.0

                edited_incurred['price'] = edited_incurred['price'].apply(clean_vnd_val_inc)
                edited_incurred['deposit'] = edited_incurred['deposit'].apply(clean_vnd_val_inc)
                
                cols_inc = ['name', 'unit', 'quantity', 'price', 'total_amount', 'deposit', 'note']
                if not edited_incurred[cols_inc].equals(st.session_state.ls_incurred_temp[cols_inc]):
                     st.session_state.ls_incurred_temp = edited_incurred[cols_inc]
                     st.rerun()

                st.write("")
                # 6. CHECKLIST B√ÄN GIAO (ƒê√£ ƒë·ªïi th·ª© t·ª± xu·ªëng d∆∞·ªõi)
                st.markdown("##### 6. Checklist b√†n giao h·ªì s∆° HDV")
                checklist_items = ["Ch∆∞∆°ng tr√¨nh ƒë√≥ng m·ªôc", "Danh s√°ch ƒë√≥ng m·ªôc", "B·∫£o hi·ªÉm du l·ªãch", "Th·ª±c ƒë∆°n ƒë√≥ng m·ªôc", "V√© m√°y bay", "X√°c nh·∫≠n khu du l·ªãch/nh√† h√†ng (N·∫øu c√≥)", "H·ª£p ƒë·ªìng h∆∞·ªõng d·∫´n"]
                
                current_checklist = dict(tour_info_ls).get('handover_checklist', '')
                checked_items = current_checklist.split(',') if current_checklist else []
                
                cols_chk = st.columns(2)
                new_checked_list = []
                all_checked = True
                
                for i, item in enumerate(checklist_items):
                    is_checked = item in checked_items
                    if cols_chk[i % 2].checkbox(item, value=is_checked, key=f"chk_ho_{tour_id_ls}_{i}"):
                        new_checked_list.append(item)
                    else:
                        all_checked = False

                st.write("")
                st.markdown("##### T·∫°m ·ª©ng cho HDV")
                
                # [FIX] T·ª± ƒë·ªông t√≠nh t·ªïng ti·ªÅn c√≤n l·∫°i ƒë·ªÉ l√†m T·∫°m ·ª©ng
                def calc_rem_total(df):
                    if df.empty: return 0.0
                    def clean_val(x):
                        if isinstance(x, (int, float)): return float(x)
                        try: return float(str(x).replace('.', '').replace(' VND', '').strip())
                        except: return 0.0
                    
                    t = df['total_amount'].apply(clean_val)
                    d = df['deposit'].apply(clean_val)
                    return (t - d).sum()

                rem_h = calc_rem_total(st.session_state.ls_hotels_temp)
                rem_r = calc_rem_total(st.session_state.ls_rests_temp)
                rem_s = calc_rem_total(st.session_state.ls_sight_temp)
                
                # T√≠nh ri√™ng cho Incurred v√¨ c·∫ßn t√≠nh l·∫°i t·ª´ price * qty
                df_inc_c = st.session_state.ls_incurred_temp.copy()
                q_i = pd.to_numeric(df_inc_c['quantity'], errors='coerce').fillna(0)
                p_i = pd.to_numeric(df_inc_c['price'], errors='coerce').fillna(0)
                d_i = pd.to_numeric(df_inc_c['deposit'], errors='coerce').fillna(0)
                rem_i = ((q_i * p_i) - d_i).sum()
                
                total_rem_all = rem_h + rem_r + rem_s + rem_i
                
                tam_ung = float(total_rem_all)
                st.markdown(f"""<div style="background-color: #e8f5e9; padding: 15px; border-radius: 10px; margin-top: 10px; border: 1px solid #c8e6c9;"><div style="display:flex; justify-content:space-between; font-size: 1.3em; color: #2e7d32;"><span><b>T·∫†M ·ª®NG CHO HDV</b></span> <b>{format_vnd(tam_ung)} VND</b></div></div>""", unsafe_allow_html=True)
                
                st.write("")
                st.write("")
                c_save, c_export = st.columns([1, 2])
                
                if c_save.button("üíæ L∆∞u v√† ch·ªù quy·∫øt to√°n", type="primary", use_container_width=True):
                    if not all_checked:
                        st.error("‚õî B·∫°n ch∆∞a ho√†n th√†nh Checklist b√†n giao! Vui l√≤ng ki·ªÉm tra ƒë·ªß c√°c m·ª•c tr∆∞·ªõc khi l∆∞u.")
                    else:
                        # L∆∞u Danh s√°ch ƒëo√†n
                        run_query("DELETE FROM tour_guests WHERE tour_id=?", (tour_id_ls,), commit=True)
                        if not edited_guests.empty:
                            data_guests = [(tour_id_ls, r['name'], r['dob'], r['hometown'], r['cccd'], r['type']) for _, r in edited_guests.iterrows() if r['name']]
                            if data_guests: run_query_many("INSERT INTO tour_guests (tour_id, name, dob, hometown, cccd, type) VALUES (?,?,?,?,?,?)", data_guests)
                        
                        # L∆∞u Kh√°ch s·∫°n
                        run_query("DELETE FROM tour_hotels WHERE tour_id=?", (tour_id_ls,), commit=True)
                        if not edited_hotels.empty:
                            data_hotels = [(tour_id_ls, r['hotel_name'], r['address'], r['phone'], r['total_rooms'], r['room_type'], r['total_amount'], r['deposit']) for _, r in edited_hotels.iterrows() if r['hotel_name']]
                            if data_hotels: run_query_many("INSERT INTO tour_hotels (tour_id, hotel_name, address, phone, total_rooms, room_type, total_amount, deposit) VALUES (?,?,?,?,?,?,?,?)", data_hotels)

                        # L∆∞u Nh√† h√†ng
                        run_query("DELETE FROM tour_restaurants WHERE tour_id=?", (tour_id_ls,), commit=True)
                        if not edited_rests.empty:
                            data_rests = [(tour_id_ls, r['meal_name'], r['restaurant_name'], r['address'], r['phone'], r['menu'], r['total_amount'], r['deposit'], r['date']) for _, r in edited_rests.iterrows() if r['meal_name']]
                            if data_rests: run_query_many("INSERT INTO tour_restaurants (tour_id, meal_name, restaurant_name, address, phone, menu, total_amount, deposit, date) VALUES (?,?,?,?,?,?,?,?,?)", data_rests)

                        # L∆∞u ƒêi·ªÉm tham quan
                        run_query("DELETE FROM tour_sightseeings WHERE tour_id=?", (tour_id_ls,), commit=True)
                        if not edited_sightseeings.empty:
                            data_sightseeings = [(tour_id_ls, r['name'], r['address'], r['quantity'], r['note'], r['date'], r['total_amount'], r['deposit']) for _, r in edited_sightseeings.iterrows() if r['name']]
                            if data_sightseeings: run_query_many("INSERT INTO tour_sightseeings (tour_id, name, address, quantity, note, date, total_amount, deposit) VALUES (?,?,?,?,?,?,?,?)", data_sightseeings)

                        # L∆∞u Chi ph√≠ ph√°t sinh
                        run_query("DELETE FROM tour_incurred_costs WHERE tour_id=?", (tour_id_ls,), commit=True)
                        if not edited_incurred.empty:
                            data_inc = [(tour_id_ls, r['name'], r['unit'], r['quantity'], r['price'], r['quantity']*r['price'], r['deposit'], r['note']) for _, r in edited_incurred.iterrows() if r['name']]
                            if data_inc: run_query_many("INSERT INTO tour_incurred_costs (tour_id, name, unit, quantity, price, total_amount, deposit, note) VALUES (?,?,?,?,?,?,?,?)", data_inc)

                        # L∆∞u Checklist
                        checklist_str = ",".join(new_checked_list)
                        run_query("UPDATE tours SET handover_checklist=? WHERE id=?", (checklist_str, tour_id_ls), commit=True)
                        
                        st.success("‚úÖ ƒê√£ l∆∞u h·ªì s∆° v√† checklist th√†nh c√¥ng! Tour ƒëang ch·ªù quy·∫øt to√°n.")
                        time.sleep(1); st.rerun()
                
                with c_export:
                    # --- XU·∫§T FILE T·ªîNG H·ª¢P (B√ÄN GIAO + TH·ª∞C ƒê∆†N) ---
                    buffer_combined = io.BytesIO()
                    with pd.ExcelWriter(buffer_combined, engine='xlsxwriter') as writer:
                        workbook: Any = writer.book
                        # ws = workbook.add_worksheet("ThucDon")
                        
                        # Formats
                        fmt_comp = workbook.add_format({'bold': True, 'font_size': 12, 'font_color': '#1B5E20', 'font_name': 'Times New Roman'})
                        fmt_info = workbook.add_format({'font_size': 10, 'italic': True, 'font_name': 'Times New Roman'})
                        fmt_title = workbook.add_format({'bold': True, 'font_size': 16, 'align': 'center', 'valign': 'vcenter', 'font_color': '#E65100', 'border': 0, 'font_name': 'Times New Roman'})
                        fmt_header = workbook.add_format({'bold': True, 'bg_color': '#FFF3E0', 'border': 1, 'align': 'center', 'valign': 'vcenter', 'text_wrap': True, 'font_color': '#E65100', 'font_name': 'Times New Roman'})
                        fmt_text = workbook.add_format({'border': 1, 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                        # ==========================================
                        # SHEET 1: B√ÄN GIAO (BAN_GIAO_HDV)
                        # ==========================================
                        ws_bg = workbook.add_worksheet("BAN_GIAO_HDV")
                        
                        # --- FORMATS (B√ÄN GIAO) ---
                        fmt_title_bg = workbook.add_format({'bold': True, 'font_size': 16, 'align': 'center', 'valign': 'vcenter', 'font_color': '#0D47A1', 'border': 0, 'font_name': 'Times New Roman'})
                        fmt_comp_bg = workbook.add_format({'bold': True, 'font_size': 11, 'font_color': '#1B5E20', 'font_name': 'Times New Roman'})
                        fmt_header_bg = workbook.add_format({'bold': True, 'bg_color': '#E0F7FA', 'border': 1, 'align': 'center', 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                        fmt_label_bg = workbook.add_format({'bold': True, 'bg_color': '#F5F5F5', 'border': 1, 'align': 'left', 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                        fmt_text_bg = workbook.add_format({'border': 1, 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                        fmt_center_bg = workbook.add_format({'border': 1, 'align': 'center', 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                        fmt_section_bg = workbook.add_format({'bold': True, 'bg_color': '#FFF3E0', 'border': 1, 'font_color': '#E65100', 'align': 'left', 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                        money_fmt_bg = workbook.add_format({'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'font_name': 'Times New Roman'})

                        # Helper to safely parse float from potential strings
                        def safe_float_exp(x):
                            if isinstance(x, (int, float)): return float(x)
                            try: return float(str(x).replace('.', '').replace(',', '').replace(' VND', '').strip())
                            except: return 0.0

                        # --- DATA PREP ---
                        t = dict(tour_info_ls)
                        
                        # --- LAYOUT B√ÄN GIAO ---
                        ws_bg.merge_range('A1:F1', comp['name'].upper(), fmt_comp_bg)
                        ws_bg.merge_range('A2:F2', "PHI·∫æU B√ÄN GIAO TOUR / TOUR ORDER", fmt_title_bg)
                        
                        # SECTION A
                        row = 3
                        ws_bg.merge_range(row, 0, row, 5, "A. TH√îNG TIN ƒêO√ÄN", fmt_section_bg)
                        row += 1
                        ws_bg.write(row, 0, "T√™n ƒëo√†n:", fmt_label_bg)
                        ws_bg.merge_range(row, 1, row, 2, t.get('tour_name', ''), fmt_text_bg)
                        ws_bg.write(row, 3, "M√£ Tour:", fmt_label_bg)
                        ws_bg.merge_range(row, 4, row, 5, t.get('tour_code', ''), fmt_center_bg)
                        
                        row += 1
                        ws_bg.write(row, 0, "S·ªë l∆∞·ª£ng:", fmt_label_bg)
                        ws_bg.merge_range(row, 1, row, 2, f"{t.get('guest_count', 0)} kh√°ch", fmt_text_bg)
                        ws_bg.write(row, 3, "Th·ªùi gian:", fmt_label_bg)
                        ws_bg.merge_range(row, 4, row, 5, f"{t.get('start_date','')} - {t.get('end_date','')}", fmt_center_bg)
                        
                        row += 1
                        ws_bg.write(row, 0, "ƒêi·ªÉm ƒë√≥n:", fmt_label_bg)
                        ws_bg.write(row, 1, t.get('pickup_location', ''), fmt_text_bg)
                        ws_bg.write(row, 2, "Gi·ªù ƒë√≥n:", fmt_label_bg)
                        ws_bg.write(row, 3, t.get('pickup_time', ''), fmt_text_bg)
                        ws_bg.write(row, 4, "Chuy·∫øn bay:", fmt_label_bg)
                        ws_bg.write(row, 5, t.get('flight_code', ''), fmt_text_bg)

                        # SECTION B: L·ªäCH TR√åNH (D·ªúI T·ª™ D L√äN B V√Ä CHIA THEO NG√ÄY)
                        row += 2
                        ws_bg.merge_range(row, 0, row, 5, "B. L·ªäCH TR√åNH CHI TI·∫æT", fmt_section_bg)
                        row += 1
                            
                        # L·∫•y d·ªØ li·ªáu l·ªãch tr√¨nh t·ª´ DB
                        itins_xls = run_query("SELECT * FROM tour_itineraries WHERE tour_id=? ORDER BY day_index", (tour_id_ls,))
                        itin_map_xls = {r['day_index']: r['content'] for r in itins_xls} if itins_xls else {}

                        # T·ª± ƒë·ªông t·∫°o d√≤ng theo ng√†y
                        try:
                            s_date = datetime.strptime(t.get('start_date', ''), '%d/%m/%Y')
                            e_date = datetime.strptime(t.get('end_date', ''), '%d/%m/%Y')
                            delta = (e_date - s_date).days + 1
                            
                            if delta > 0:
                                for i in range(delta):
                                    curr_date = s_date + pd.Timedelta(days=i)
                                    date_str = curr_date.strftime('%d/%m')
                                    content_str = itin_map_xls.get(i, "")
                                    ws_bg.write(row, 0, f"Ng√†y {i+1} ({date_str})", fmt_label_bg)
                                    ws_bg.merge_range(row, 1, row, 5, content_str, fmt_text_bg)
                                    row += 1
                        except:
                            pass

                        # SECTION C: NH√ÇN S·ª∞ (D·ªúI T·ª™ B XU·ªêNG C)
                        row += 1
                        ws_bg.merge_range(row, 0, row, 5, "C. TH√îNG TIN NH√ÇN S·ª∞ & V·∫¨N CHUY·ªÇN", fmt_section_bg)
                        row += 1
                        headers_b = ["Vai tr√≤", "H·ªç v√† t√™n", "ƒêi·ªán tho·∫°i", "Ghi ch√∫ / Bi·ªÉn s·ªë", "", ""]
                        for i, h in enumerate(headers_b): 
                            if h: ws_bg.write(row, i, h, fmt_header_bg)
                        
                        row += 1
                        ws_bg.write(row, 0, "H∆∞·ªõng d·∫´n vi√™n", fmt_center_bg)
                        ws_bg.write(row, 1, t.get('guide_name', ''), fmt_text_bg)
                        ws_bg.write(row, 2, t.get('guide_phone', ''), fmt_center_bg)
                        ws_bg.write(row, 3, "", fmt_text_bg)
                        
                        row += 1
                        ws_bg.write(row, 0, "L√°i xe", fmt_center_bg)
                        ws_bg.write(row, 1, t.get('driver_name', ''), fmt_text_bg)
                        ws_bg.write(row, 2, t.get('driver_phone', ''), fmt_center_bg)
                        ws_bg.write(row, 3, f"{t.get('car_plate', '')} ({t.get('car_type', '')})", fmt_text_bg)
                        
                        row += 1
                        ws_bg.write(row, 0, "ƒêi·ªÅu h√†nh/Sale", fmt_center_bg)
                        ws_bg.write(row, 1, t.get('sale_name', ''), fmt_text_bg)
                        ws_bg.write(row, 2, "", fmt_center_bg)
                        ws_bg.write(row, 3, "", fmt_text_bg)

                        # SECTION C
                        row += 2
                        ws_bg.merge_range(row, 0, row, 5, "C. CHI TI·∫æT D·ªäCH V·ª§", fmt_section_bg)
                        
                        # 1. Kh√°ch s·∫°n
                        row += 1
                        ws_bg.merge_range(row, 0, row, 5, "1. L∆∞u tr√∫ (Kh√°ch s·∫°n)", fmt_label_bg)
                        row += 1
                        ws_bg.write(row, 0, "T√™n KS", fmt_header_bg)
                        ws_bg.write(row, 1, "Li√™n h·ªá", fmt_header_bg)
                        ws_bg.write(row, 2, "Ph√≤ng/Lo·∫°i", fmt_header_bg)
                        ws_bg.write(row, 3, "T·ªïng ti·ªÅn", fmt_header_bg)
                        ws_bg.write(row, 4, "ƒê√£ c·ªçc", fmt_header_bg)
                        ws_bg.write(row, 5, "C√≤n l·∫°i", fmt_header_bg)
                        
                        df_hotels_exp = st.session_state.ls_hotels_temp
                        if not df_hotels_exp.empty:
                            for _, h in df_hotels_exp.iterrows():
                                total = safe_float_exp(h.get('total_amount', 0))
                                dep = safe_float_exp(h.get('deposit', 0))
                                rem = total - dep
                                row += 1
                                ws_bg.write(row, 0, h['hotel_name'], fmt_text_bg)
                                ws_bg.write(row, 1, f"{h['address']}\n{h['phone']}", fmt_text_bg)
                                ws_bg.write(row, 2, f"{h['total_rooms']} ({h['room_type']})", fmt_center_bg)
                                ws_bg.write(row, 3, total, money_fmt_bg)
                                ws_bg.write(row, 4, dep, money_fmt_bg)
                                ws_bg.write(row, 5, rem, money_fmt_bg)
                        else:
                            row += 1; ws_bg.merge_range(row, 0, row, 5, "(Ch∆∞a c√≥ th√¥ng tin)", fmt_center_bg)

                        # 2. Nh√† h√†ng
                        row += 1
                        ws_bg.merge_range(row, 0, row, 5, "2. ·∫®m th·ª±c (Nh√† h√†ng)", fmt_label_bg)
                        row += 1
                        ws_bg.write(row, 0, "B·ªØa ƒÉn", fmt_header_bg)
                        ws_bg.write(row, 1, "Nh√† h√†ng", fmt_header_bg)
                        ws_bg.write(row, 2, "Li√™n h·ªá", fmt_header_bg)
                        ws_bg.write(row, 3, "T·ªïng ti·ªÅn", fmt_header_bg)
                        ws_bg.write(row, 4, "ƒê√£ c·ªçc", fmt_header_bg)
                        ws_bg.write(row, 5, "C√≤n l·∫°i", fmt_header_bg)
                        
                        df_rests_exp = st.session_state.ls_rests_temp
                        if not df_rests_exp.empty:
                            for _, r in df_rests_exp.iterrows():
                                total = safe_float_exp(r.get('total_amount', 0))
                                dep = safe_float_exp(r.get('deposit', 0))
                                rem = total - dep
                                row += 1
                                ws_bg.write(row, 0, r['meal_name'], fmt_text_bg)
                                ws_bg.write(row, 1, r['restaurant_name'], fmt_text_bg)
                                ws_bg.write(row, 2, f"{r['address']}\n{r['phone']}", fmt_text_bg)
                                ws_bg.write(row, 3, total, money_fmt_bg)
                                ws_bg.write(row, 4, dep, money_fmt_bg)
                                ws_bg.write(row, 5, rem, money_fmt_bg)
                        else:
                            row += 1; ws_bg.merge_range(row, 0, row, 5, "(Ch∆∞a c√≥ th√¥ng tin)", fmt_center_bg)

                        # 3. ƒêi·ªÉm tham quan
                        row += 1
                        ws_bg.merge_range(row, 0, row, 5, "3. ƒêi·ªÉm tham quan", fmt_label_bg)
                        row += 1
                        ws_bg.write(row, 0, "T√™n ƒë·ªãa ƒëi·ªÉm", fmt_header_bg)
                        ws_bg.merge_range(row, 1, row, 2, "ƒê·ªãa ch·ªâ", fmt_header_bg)
                        ws_bg.write(row, 3, "S·ªë l∆∞·ª£ng", fmt_header_bg)
                        ws_bg.merge_range(row, 4, row, 5, "L∆∞u √Ω", fmt_header_bg)
                        
                        df_sightseeings_exp = st.session_state.ls_sight_temp
                        if not df_sightseeings_exp.empty:
                            for _, s in df_sightseeings_exp.iterrows():
                                row += 1
                                ws_bg.write(row, 0, s['name'], fmt_text_bg)
                                ws_bg.merge_range(row, 1, row, 2, s['address'], fmt_text_bg)
                                ws_bg.write(row, 3, s['quantity'], fmt_center_bg)
                                ws_bg.merge_range(row, 4, row, 5, s['note'], fmt_text_bg)
                        else:
                            row += 1; ws_bg.merge_range(row, 0, row, 5, "(Ch∆∞a c√≥ th√¥ng tin)", fmt_center_bg)

                        # 4. Chi ph√≠ ph√°t sinh (M·ªöI)
                        row += 1
                        ws_bg.merge_range(row, 0, row, 5, "4. Chi ph√≠ ph√°t sinh (N∆∞·ªõc, Sim, Banner...)", fmt_label_bg)
                        row += 1
                        ws_bg.write(row, 0, "T√™n chi ph√≠", fmt_header_bg)
                        ws_bg.write(row, 1, "ƒêVT", fmt_header_bg)
                        ws_bg.write(row, 2, "S·ªë l∆∞·ª£ng", fmt_header_bg)
                        ws_bg.write(row, 3, "T·ªïng ti·ªÅn", fmt_header_bg)
                        ws_bg.write(row, 4, "ƒê√£ c·ªçc", fmt_header_bg)
                        ws_bg.write(row, 5, "C√≤n l·∫°i", fmt_header_bg)
                        
                        df_inc_exp = st.session_state.ls_incurred_temp
                        if not df_inc_exp.empty:
                            for _, inc in df_inc_exp.iterrows():
                                try:
                                    qty = safe_float_exp(inc.get('quantity', 0))
                                    price = safe_float_exp(inc.get('price', 0))
                                    total = qty * price
                                    dep = safe_float_exp(inc.get('deposit', 0))
                                    rem = total - dep
                                except: total=0; dep=0; rem=0; qty=0

                                row += 1
                                ws_bg.write(row, 0, inc['name'], fmt_text_bg)
                                ws_bg.write(row, 1, inc['unit'], fmt_center_bg)
                                ws_bg.write(row, 2, qty, fmt_center_bg)
                                ws_bg.write(row, 3, total, money_fmt_bg)
                                ws_bg.write(row, 4, dep, money_fmt_bg)
                                ws_bg.write(row, 5, rem, money_fmt_bg)
                        else:
                            row += 1; ws_bg.merge_range(row, 0, row, 5, "(Kh√¥ng c√≥)", fmt_center_bg)

                        # --- [FIX] T√çNH TO√ÅN T·ªîNG K·∫æT (L√†m s·∫°ch d·ªØ li·ªáu tr∆∞·ªõc khi t√≠nh) ---
                        def get_clean_sum(df, col_name):
                            if df.empty or col_name not in df.columns: return 0.0
                            def clean_val(x):
                                if isinstance(x, (int, float)): return float(x)
                                try: return float(str(x).replace('.', '').replace(' VND', '').strip())
                                except: return 0.0
                            return df[col_name].apply(clean_val).sum()

                        # 1. T√≠nh T·ªïng chi ph√≠ (Total Amount)
                        t_h = get_clean_sum(st.session_state.ls_hotels_temp, 'total_amount')
                        t_r = get_clean_sum(st.session_state.ls_rests_temp, 'total_amount')
                        t_s = get_clean_sum(st.session_state.ls_sight_temp, 'total_amount')
                        
                        # T√≠nh ri√™ng cho Incurred (v√¨ c·∫ßn nh√¢n quantity * price)
                        df_inc_calc = st.session_state.ls_incurred_temp.copy()
                        df_inc_calc['price'] = pd.to_numeric(df_inc_calc['price'], errors='coerce').fillna(0)
                        df_inc_calc['quantity'] = pd.to_numeric(df_inc_calc['quantity'], errors='coerce').fillna(0)
                        t_i = (df_inc_calc['price'] * df_inc_calc['quantity']).sum()

                        grand_total = t_h + t_r + t_s + t_i

                        # 2. T√≠nh ƒê√£ c·ªçc (Deposit)
                        d_h = get_clean_sum(st.session_state.ls_hotels_temp, 'deposit')
                        d_r = get_clean_sum(st.session_state.ls_rests_temp, 'deposit')
                        d_s = get_clean_sum(st.session_state.ls_sight_temp, 'deposit')
                        d_i = get_clean_sum(st.session_state.ls_incurred_temp, 'deposit')
                        
                        grand_deposit = d_h + d_r + d_s + d_i

                        # 3. C√≤n l·∫°i (HDV c·∫ßn thanh to√°n cho NCC)
                        grand_remaining = grand_total - grand_deposit
                        
                        # 4. Quy·∫øt to√°n (C√≤n l·∫°i - T·∫°m ·ª©ng)
                        # tam_ung ƒë√£ ƒë∆∞·ª£c t√≠nh ·ªü UI v√† truy·ªÅn v√†o ƒë√¢y
                        balance = grand_remaining - tam_ung

                        # SECTION D: T·ªîNG K·∫æT & T·∫†M ·ª®NG
                        row += 2
                        ws_bg.merge_range(row, 0, row, 5, "D. T·ªîNG K·∫æT KINH PH√ç", fmt_section_bg)
                        
                        # 1. T·ªïng chi ph√≠
                        row += 1
                        ws_bg.merge_range(row, 0, row, 3, "1. T·ªîNG CHI PH√ç TOUR:", fmt_label_bg)
                        ws_bg.merge_range(row, 4, row, 5, grand_total, money_fmt_bg)
                        
                        # 2. ƒê√£ c·ªçc (Ph√¢n r√£)
                        row += 1
                        ws_bg.merge_range(row, 0, row, 3, "2. ƒê√É C·ªåC / THANH TO√ÅN TR∆Ø·ªöC (C√îNG TY):", fmt_label_bg)
                        ws_bg.merge_range(row, 4, row, 5, grand_deposit, money_fmt_bg)
                        
                        # 3. C√≤n l·∫°i
                        row += 1
                        ws_bg.merge_range(row, 0, row, 3, "3. C√íN L·∫†I C·∫¶N THANH TO√ÅN (HDV CHI):", fmt_label_bg)
                        ws_bg.merge_range(row, 4, row, 5, grand_remaining, workbook.add_format({'bold': True, 'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'bg_color': '#FFF9C4', 'font_name': 'Times New Roman'}))

                        # [FIX] Chi ti·∫øt c√≤n l·∫°i (Thay v√¨ chi ti·∫øt c·ªçc)
                        r_h = t_h - d_h
                        r_r = t_r - d_r
                        r_s = t_s - d_s
                        r_i = t_i - d_i

                        row += 1
                        ws_bg.write(row, 0, "   - Kh√°ch s·∫°n:", fmt_text_bg)
                        ws_bg.merge_range(row, 1, row, 3, r_h, money_fmt_bg)
                        row += 1
                        ws_bg.write(row, 0, "   - Nh√† h√†ng:", fmt_text_bg)
                        ws_bg.merge_range(row, 1, row, 3, r_r, money_fmt_bg)
                        row += 1
                        ws_bg.write(row, 0, "   - Tham quan:", fmt_text_bg)
                        ws_bg.merge_range(row, 1, row, 3, r_s, money_fmt_bg)
                        row += 1
                        ws_bg.write(row, 0, "   - Ph√°t sinh:", fmt_text_bg)
                        ws_bg.merge_range(row, 1, row, 3, r_i, money_fmt_bg)

                        # 4. T·∫°m ·ª©ng cho HDV
                        row += 1
                        ws_bg.merge_range(row, 0, row, 3, "4. T·∫†M ·ª®NG CHO HDV:", fmt_label_bg)
                        ws_bg.merge_range(row, 4, row, 5, tam_ung, money_fmt_bg)

                        # 5. Quy·∫øt to√°n
                        row += 1
                        ws_bg.merge_range(row, 0, row, 3, "5. QUY·∫æT TO√ÅN (THU L·∫†I / CHI TH√äM):", fmt_label_bg)
                        ws_bg.merge_range(row, 4, row, 5, balance, workbook.add_format({'bold': True, 'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'font_color': '#D32F2F', 'font_size': 11, 'font_name': 'Times New Roman'}))

                        # FOOTER: CH·ªÆ K√ù
                        row += 3
                        fmt_sig_title = workbook.add_format({'bold': True, 'align': 'center', 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                        fmt_sig_name = workbook.add_format({'italic': True, 'align': 'center', 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                        
                        ws_bg.write(row, 0, "NG∆Ø·ªúI L·∫¨P PHI·∫æU", fmt_sig_title)
                        ws_bg.merge_range(row, 1, row, 2, "K·∫æ TO√ÅN", fmt_sig_title)
                        ws_bg.write(row, 3, "GI√ÅM ƒê·ªêC", fmt_sig_title)
                        ws_bg.merge_range(row, 4, row, 5, "H∆Ø·ªöNG D·∫™N VI√äN", fmt_sig_title)
                        
                        row += 1
                        ws_bg.write(row, 0, "(K√Ω, h·ªç t√™n)", fmt_sig_name)
                        ws_bg.merge_range(row, 1, row, 2, "(K√Ω, h·ªç t√™n)", fmt_sig_name)
                        ws_bg.write(row, 3, "(K√Ω, h·ªç t√™n)", fmt_sig_name)
                        ws_bg.merge_range(row, 4, row, 5, "(K√Ω, h·ªç t√™n)", fmt_sig_name)
                        
                        # Space for signature
                        row += 5
                        
                        # Names
                        ws_bg.write(row, 0, t.get('sale_name', ''), fmt_sig_title)
                        ws_bg.merge_range(row, 4, row, 5, t.get('guide_name', ''), fmt_sig_title)

                        ws_bg.set_column('A:A', 20)
                        ws_bg.set_column('B:F', 18)

                        # ==========================================
                        # SHEET: DANH S√ÅCH ƒêO√ÄN (DanhSachDoan)
                        # ==========================================
                        ws_ds = workbook.add_worksheet("DanhSachDoan")
                        
                        # 1. Company Info
                        ws_ds.merge_range('A1:F1', comp['name'].upper(), fmt_comp_bg)
                        ws_ds.merge_range('A2:F2', "DANH S√ÅCH ƒêO√ÄN / GUEST LIST", fmt_title_bg)
                        
                        # 2. Tour Info
                        ws_ds.write('A4', "T√™n ƒëo√†n:", fmt_label_bg)
                        ws_ds.merge_range('B4:C4', t.get('tour_name', ''), fmt_text_bg)
                        ws_ds.write('D4', "M√£ Tour:", fmt_label_bg)
                        ws_ds.merge_range('E4:F4', t.get('tour_code', ''), fmt_center_bg)
                        
                        ws_ds.write('A5', "Th·ªùi gian:", fmt_label_bg)
                        ws_ds.merge_range('B5:C5', f"{t.get('start_date','')} - {t.get('end_date','')}", fmt_center_bg)
                        ws_ds.write('D5', "S·ªë kh√°ch:", fmt_label_bg)
                        ws_ds.merge_range('E5:F5', f"{t.get('guest_count', 0)} kh√°ch", fmt_center_bg)

                        # 3. Table Header
                        row_ds = 7
                        headers_ds = ["STT", "H·ªç v√† t√™n", "Ng√†y sinh", "Qu√™ qu√°n", "S·ªë CCCD", "Ph√¢n lo·∫°i"]
                        for i, h in enumerate(headers_ds):
                            ws_ds.write(row_ds, i, h, fmt_header_bg)
                        
                        # 4. Data
                        if not edited_guests.empty:
                            for i, (idx, row_g) in enumerate(edited_guests.iterrows()):
                                row_ds += 1
                                ws_ds.write(row_ds, 0, i + 1, fmt_center_bg)
                                ws_ds.write(row_ds, 1, row_g.get('name', ''), fmt_text_bg)
                                ws_ds.write(row_ds, 2, row_g.get('dob', ''), fmt_center_bg)
                                ws_ds.write(row_ds, 3, row_g.get('hometown', ''), fmt_text_bg)
                                ws_ds.write(row_ds, 4, row_g.get('cccd', ''), fmt_center_bg)
                                ws_ds.write(row_ds, 5, row_g.get('type', ''), fmt_center_bg)
                        
                        # Column Widths
                        ws_ds.set_column('A:A', 5)
                        ws_ds.set_column('B:B', 25)
                        ws_ds.set_column('C:C', 15)
                        ws_ds.set_column('D:D', 20)
                        ws_ds.set_column('E:F', 15)

                        # ==========================================
                        # SHEET 2: TH·ª∞C ƒê∆†N (ThucDon)
                        # ==========================================
                        ws_menu = workbook.add_worksheet("ThucDon")
                        
                        # Formats (Menu)
                        fmt_comp_menu = workbook.add_format({'bold': True, 'font_size': 12, 'font_color': '#1B5E20', 'font_name': 'Times New Roman'})
                        fmt_info_menu = workbook.add_format({'font_size': 10, 'italic': True, 'font_name': 'Times New Roman'})
                        fmt_title_menu = workbook.add_format({'bold': True, 'font_size': 16, 'align': 'center', 'valign': 'vcenter', 'font_color': '#E65100', 'border': 0, 'font_name': 'Times New Roman'})
                        fmt_header_menu = workbook.add_format({'bold': True, 'bg_color': '#FFF3E0', 'border': 1, 'align': 'center', 'valign': 'vcenter', 'text_wrap': True, 'font_color': '#E65100', 'font_name': 'Times New Roman'})
                        fmt_text_menu = workbook.add_format({'border': 1, 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                        
                        # Company Info
                        comp_data = get_company_data()
                        ws_menu.write('A1', comp_data['name'], fmt_comp_menu)
                        ws_menu.write('A2', f"ƒêC: {comp_data['address']}", fmt_info_menu)
                        ws_menu.write('A3', f"MST: {comp_data['phone']}", fmt_info_menu)
                        
                        # Title
                        ws_menu.merge_range('A5:C5', f"DANH S√ÅCH TH·ª∞C ƒê∆†N TOUR: {tour_info_ls['tour_name']}", fmt_title_menu)
                        
                        # Table Header
                        headers_menu = ["Th√¥ng tin nh√† h√†ng", "B·ªØa ƒÉn / Th·ªùi gian", "Th·ª±c ƒë∆°n"]
                        for i, h in enumerate(headers_menu):
                            ws_menu.write(6, i, h, fmt_header_menu)
                            
                        # Data
                        row_menu = 7
                        if not df_rests.empty:
                            df_rests_exp = df_rests.fillna('')
                            for _, r in df_rests_exp.iterrows():
                                # G·ªôp th√¥ng tin: T√™n, ƒê·ªãa ch·ªâ, Li√™n h·ªá
                                info_parts = [str(r[k]) for k in ['restaurant_name', 'address', 'phone'] if str(r[k]).strip()]
                                info_str = "\n".join(info_parts)
                                
                                ws_menu.write(row_menu, 0, info_str, fmt_text_menu)
                                ws_menu.write(row_menu, 1, r['meal_name'], fmt_text_menu)
                                ws_menu.write(row_menu, 2, r['menu'], fmt_text_menu)
                                row_menu += 1
                        
                        # Column widths
                        ws_menu.set_column('A:A', 40) # Th√¥ng tin nh√† h√†ng
                        ws_menu.set_column('B:B', 25) # B·ªØa ƒÉn / Th·ªùi gian
                        ws_menu.set_column('C:C', 50) # Th·ª±c ƒë∆°n

                    st.download_button("üì• Xu·∫•t H·ªì S∆° B√†n Giao & Th·ª±c ƒê∆°n (Excel)", buffer_combined.getvalue(), f"HoSo_BanGiao_{tour_info_ls['tour_code']}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)

    # ---------------- TAB 2: QUY·∫æT TO√ÅN ----------------
    with tab_act:
        st.subheader("üí∏ Quy·∫øt To√°n ")
        
        selected_tour_act_label = st.selectbox("Ch·ªçn ƒêo√†n quy·∫øt to√°n:", list(tour_options.keys()) if tour_options else [], key="sel_tour_act")
        
        if selected_tour_act_label:
            tour_id_act = tour_options[selected_tour_act_label] # type: ignore
            tour_info_act = next((t for t in all_tours if t['id'] == tour_id_act), None)
            if not tour_info_act:
                st.error("Kh√¥ng t√¨m th·∫•y th√¥ng tin tour.")
                st.stop()
            assert tour_info_act is not None
            
            # --- L·∫•y D·ª± to√°n ƒë·ªÉ so s√°nh ---
            est_items = run_query("SELECT SUM(total_amount) as total FROM tour_items WHERE tour_id=? AND item_type='EST'", (tour_id_act,), fetch_one=True)
            # If the query returns a row and the 'total' is not None (SQL SUM can return NULL), use it. Otherwise, default to 0.
            est_total_cost = est_items['total'] if est_items and est_items['total'] is not None else 0
            # T√≠nh l·∫°i gi√° b√°n ch·ªët (D·ª±a tr√™n % ƒë√£ l∆∞u)
            p_pct = tour_info_act['est_profit_percent'] # type: ignore
            t_pct = tour_info_act['est_tax_percent'] # type: ignore
            est_profit_val = est_total_cost * (p_pct / 100)
            est_final_sale = (est_total_cost + est_profit_val) * (1 + t_pct/100)
            
            # [UPDATED] L·∫•y T·ªïng doanh thu t·ª´ b√™n D·ª± to√°n (Gi√° ch·ªët * SL)
            t_act_dict_calc = dict(tour_info_act)
            final_price_est = float(t_act_dict_calc.get('final_tour_price', 0) or 0)
            child_price_est = float(t_act_dict_calc.get('child_price', 0) or 0)
            final_qty_est = float(t_act_dict_calc.get('final_qty', 0) or 0)
            child_qty_est = float(t_act_dict_calc.get('child_qty', 0) or 0)
            if final_qty_est == 0: final_qty_est = float(t_act_dict_calc.get('guest_count', 1))
            total_revenue_est = (final_price_est * final_qty_est) + (child_price_est * child_qty_est)
            
            if total_revenue_est > 0:
                est_final_sale = total_revenue_est
            else:
                est_profit_val = est_total_cost * (p_pct / 100)
                est_final_sale = (est_total_cost + est_profit_val) * (1 + t_pct/100)
            
            st.info(f"T·ªîNG DOANH THU: {format_vnd(est_final_sale)} VND")

            # --- [UPDATED] PH√ÇN T√çCH CHI PH√ç ---
            st.divider()
            st.markdown("### üìä Ph√¢n t√≠ch Chi ph√≠")
            
            linked_docs = run_query("SELECT * FROM invoices WHERE cost_code=? AND status='active'", (tour_info_act['tour_code'],)) # type: ignore
            df_linked = pd.DataFrame([dict(r) for r in linked_docs]) if linked_docs else pd.DataFrame()
            
            total_unc = 0
            total_inv = 0
            df_unc = pd.DataFrame()
            df_inv = pd.DataFrame()

            if not df_linked.empty:
                unc_mask = df_linked['invoice_number'].astype(str).str.contains("UNC", case=False, na=False) # type: ignore
                df_unc = df_linked.loc[unc_mask]
                total_unc = df_unc['total_amount'].sum()
                
                inv_mask = (df_linked['type'] == 'IN') & (~unc_mask)
                df_inv = df_linked.loc[inv_mask]
                total_inv = df_inv['total_amount'].sum()

            c_unc_t, c_inv_t = st.columns(2)
            with c_unc_t:
                st.markdown(f"#### üí∏ 1. Chi ph√≠ UNC: {format_vnd(total_unc)}")
                if not df_unc.empty:
                    # [UPDATED] Format ti·ªÅn t·ªá Vi·ªát Nam c√≥ d·∫•u ch·∫•m v√† ch·ªØ VND
                    df_unc_show = df_unc.copy()
                    df_unc_show['total_show'] = df_unc_show['total_amount'].apply(lambda x: format_vnd(x) + " VND") # type: ignore
                    st.dataframe(df_unc_show[['date', 'invoice_number', 'memo', 'total_show']],
                                 column_config={
                                     "date": "Ng√†y", 
                                     "invoice_number": "S·ªë ch·ª©ng t·ª´", 
                                     "memo": "N·ªôi dung", 
                                     "total_show": "Th√†nh ti·ªÅn"
                                 },
                                 use_container_width=True, hide_index=True)
                else: st.caption("Ch∆∞a c√≥ UNC.")
            
            with c_inv_t:
                st.markdown(f"#### üìÑ 2. H√≥a ƒë∆°n ƒë·∫ßu v√†o: {format_vnd(total_inv)}")
                if not df_inv.empty:
                    # [UPDATED] Format ti·ªÅn t·ªá Vi·ªát Nam c√≥ d·∫•u ch·∫•m v√† ch·ªØ VND
                    df_inv_show = df_inv.copy()
                    df_inv_show['total_show'] = df_inv_show['total_amount'].apply(lambda x: format_vnd(x) + " VND") # type: ignore
                    st.dataframe(df_inv_show[['date', 'invoice_number', 'seller_name', 'total_show']], 
                                 column_config={"date": "Ng√†y", "invoice_number": "S·ªë h√≥a ƒë∆°n", "seller_name": "ƒê∆°n v·ªã b√°n", "total_show": "Th√†nh ti·ªÅn"}, 
                                 use_container_width=True, hide_index=True)
                else: st.caption("Ch∆∞a c√≥ h√≥a ƒë∆°n ƒë·∫ßu v√†o.")

            # [CODE M·ªöI] L·∫•y d·ªØ li·ªáu D·ª± to√°n ƒë·ªÉ so s√°nh
            est_items_ref = run_query("SELECT category, description, total_amount FROM tour_items WHERE tour_id=? AND item_type='EST'", (tour_id_act,))
            est_lookup = {}
            if est_items_ref:
                for r in est_items_ref:
                    key = (str(r['category']).strip().lower(), str(r['description']).strip().lower()) # type: ignore
                    est_lookup[key] = float(r['total_amount'] or 0) # type: ignore
            
            with st.expander("üëÄ B·∫£ng D·ª± To√°n (ƒê·ªÉ ƒë·ªëi chi·∫øu)", expanded=False):
                if est_items_ref:
                    df_est_ref = pd.DataFrame([dict(r) for r in est_items_ref])
                    df_est_ref['total_amount'] = df_est_ref['total_amount'].apply(lambda x: format_vnd(x)) # type: ignore
                    st.dataframe(df_est_ref, column_config={"category": "H·∫°ng m·ª•c", "description": "Di·ªÖn gi·∫£i", "total_amount": "D·ª± to√°n"}, use_container_width=True, hide_index=True)
                else: st.info("Ch∆∞a c√≥ d·ªØ li·ªáu d·ª± to√°n.")

            # --- Fetch Items (ACT) with Session State ---
            if "current_tour_id_act" not in st.session_state: st.session_state.current_tour_id_act = None
            if st.session_state.current_tour_id_act != tour_id_act:
                if "act_df_temp" in st.session_state: del st.session_state.act_df_temp
                st.session_state.current_tour_id_act = tour_id_act

            if "act_df_temp" not in st.session_state:
                act_items = run_query("SELECT * FROM tour_items WHERE tour_id=? AND item_type='ACT'", (tour_id_act,))
                if act_items:
                    df_act = pd.DataFrame([dict(r) for r in act_items])
                    if 'times' not in df_act.columns: df_act['times'] = 1.0
                    df_act = df_act[['category', 'description', 'unit', 'unit_price', 'quantity', 'times']]
                else:
                     # G·ª£i √Ω: N·∫øu ch∆∞a c√≥ item ACT, load item EST ƒë·ªÉ s·ª≠a cho nhanh
                     est_items_raw = run_query("SELECT * FROM tour_items WHERE tour_id=? AND item_type='EST'", (tour_id_act,))
                     if est_items_raw:
                         df_act = pd.DataFrame([dict(r) for r in est_items_raw])
                         if 'times' not in df_act.columns: df_act['times'] = 1.0
                         df_act = df_act[['category', 'description', 'unit', 'unit_price', 'quantity', 'times']]
                     else:
                         df_act = pd.DataFrame([{"category": "", "description": "", "unit": "", "quantity": 0, "unit_price": 0, "times": 1}])
                st.session_state.act_df_temp = df_act

            # Prepare Display Data
            df_act_display = st.session_state.act_df_temp.copy()
            guest_cnt_act = tour_info_act['guest_count'] if tour_info_act['guest_count'] else 1 # type: ignore
            
            # Calculate numeric totals
            # Ensure numeric types
            df_act_display['quantity'] = pd.to_numeric(df_act_display['quantity'], errors='coerce').fillna(0)
            df_act_display['unit_price'] = pd.to_numeric(df_act_display['unit_price'], errors='coerce').fillna(0)
            df_act_display['times'] = pd.to_numeric(df_act_display['times'], errors='coerce').fillna(1)

            # Formula: Total = Unit * Qty * Times
            df_act_display['total_val'] = df_act_display['quantity'] * df_act_display['unit_price'] * df_act_display['times']
            # Formula: Pax = Total / Guests
            df_act_display['price_per_pax'] = df_act_display['total_val'] / guest_cnt_act
            
            # Format strings
            df_act_display['price_per_pax'] = df_act_display['price_per_pax'].apply(lambda x: format_vnd(x) + " VND")
            df_act_display['total_display'] = df_act_display['total_val'].apply(lambda x: format_vnd(x) + " VND") # type: ignore
            df_act_display['unit_price'] = df_act_display['unit_price'].apply(lambda x: format_vnd(x) + " VND") # type: ignore

            # [CODE M·ªöI] T√≠nh to√°n so s√°nh (D·ª± to√°n vs Th·ª±c t·∫ø)
            def get_est_val(row): # type: ignore
                k = (str(row['category']).strip().lower(), str(row['description']).strip().lower()) # type: ignore
                return est_lookup.get(k, 0.0)
            
            df_act_display['est_val'] = df_act_display.apply(get_est_val, axis=1)
            df_act_display['diff_val'] = df_act_display['est_val'] - df_act_display['total_val']
            df_act_display['est_display'] = df_act_display['est_val'].apply(lambda x: format_vnd(x) + " VND")
            df_act_display['diff_display'] = df_act_display['diff_val'].apply(lambda x: format_vnd(x) + " VND")

            # --- LOGIC KH√ìA / DUY·ªÜT QUY·∫æT TO√ÅN ---
            req_act_status = tour_info_act['request_edit_act'] # type: ignore
            has_act_data = False
            check_act = run_query("SELECT id FROM tour_items WHERE tour_id=? AND item_type='ACT' LIMIT 1", (tour_id_act,))
            if check_act: has_act_data = True

            is_act_editable = False
            if current_user_role_tour in ['admin', 'admin_f1']:
                is_act_editable = True

            st.divider()
            st.markdown("#### ‚úçÔ∏è 3.Quy·∫øt to√°n")
            edited_act = st.data_editor(
                df_act_display,
                num_rows="dynamic",
                column_config={
                    "category": st.column_config.TextColumn("H·∫°ng m·ª•c chi ph√≠", required=False),
                    "description": st.column_config.TextColumn("Di·ªÖn gi·∫£i"),
                    "unit": st.column_config.TextColumn("ƒê∆°n v·ªã"),
                    "unit_price": st.column_config.TextColumn("ƒê∆°n gi√° (VND)", required=False),
                    "quantity": st.column_config.NumberColumn("S·ªë l∆∞·ª£ng", min_value=0),
                    "times": st.column_config.NumberColumn("S·ªë l·∫ßn", min_value=1),
                    "price_per_pax": st.column_config.TextColumn("Gi√°/Pax", disabled=True),
                    "total_display": st.column_config.TextColumn("Th·ª±c t·∫ø (VND)", disabled=True),
                    "est_display": st.column_config.TextColumn("D·ª± to√°n (VND)", disabled=True),
                    "diff_display": st.column_config.TextColumn("Ch√™nh l·ªách", disabled=True),
                    "total_val": st.column_config.NumberColumn("Hidden", disabled=True),
                    "est_val": st.column_config.NumberColumn("Hidden", disabled=True),
                    "diff_val": st.column_config.NumberColumn("Hidden", disabled=True),
                },
                disabled=not is_act_editable, # Kh√≥a n·∫øu kh√¥ng ƒë∆∞·ª£c ph√©p s·ª≠a
                column_order=("category", "description", "unit", "unit_price", "quantity", "times", "price_per_pax", "total_display", "est_display", "diff_display"),
                use_container_width=True,
                hide_index=True,
                key="editor_act"
            )
            
            # --- AUTO-UPDATE CALCULATION (ACTUAL) ---
            if is_act_editable:
                # T·ª± ƒë·ªông c·∫≠p nh·∫≠t khi d·ªØ li·ªáu thay ƒë·ªïi
                df_new_act = edited_act.copy()
                
                def clean_vnd_act_auto(x):
                    if isinstance(x, str):
                        return float(x.replace('.', '').replace(',', '').replace(' VND', '').strip())
                    return float(x) if x else 0.0
                
                df_new_act['unit_price'] = df_new_act['unit_price'].apply(clean_vnd_act_auto)
                df_new_act['quantity'] = pd.to_numeric(df_new_act['quantity'], errors='coerce').fillna(0)
                if 'times' not in df_new_act.columns: df_new_act['times'] = 1
                df_new_act['times'] = pd.to_numeric(df_new_act['times'], errors='coerce').fillna(1)
                
                # So s√°nh v·ªõi d·ªØ li·ªáu c≈©
                cols_check_act = ['category', 'description', 'unit', 'unit_price', 'quantity', 'times']
                df_old_act = st.session_state.act_df_temp.copy()
                if 'times' not in df_old_act.columns: df_old_act['times'] = 1
                
                # Reset index v√† fillna ƒë·ªÉ so s√°nh
                df_new_check_act = df_new_act[cols_check_act].reset_index(drop=True).fillna(0)
                df_old_check_act = df_old_act[cols_check_act].reset_index(drop=True).fillna(0)
                
                if len(df_new_check_act) != len(df_old_check_act) or not df_new_check_act.equals(df_old_check_act):
                    st.session_state.act_df_temp = df_new_act[cols_check_act]
                    st.rerun()

            act_total_cost = 0
            if not edited_act.empty:
                # Parse unit_price
                # [FIX] Handle case where a cell is None, which becomes the string 'None' after astype(str)
                cleaned_prices_act = edited_act['unit_price'].astype(str).str.replace('.', '', regex=False).str.replace(' VND', '', regex=False).str.strip()
                p_price_act = cleaned_prices_act.apply(lambda x: float(x) if x and x.lower() != 'none' else 0.0)
                # Ensure 'times' column exists and is numeric before accessing it
                # Use .get() with a default Series to handle cases where 'times' might be missing
                times_col_act = edited_act.get('times', pd.Series([1.0] * len(edited_act), index=edited_act.index)).fillna(1).astype(float) # type: ignore
                act_total_cost = (edited_act['quantity'] * p_price_act * times_col_act).sum()
            # T·ªîNG CHI PH√ç TH·ª∞C T·∫æ = H√≥a ƒë∆°n + Ph√°t sinh (Nh·∫≠p tay)
            final_act_cost = act_total_cost + total_inv

            # --- T·ªîNG K·∫æT QUY·∫æT TO√ÅN ---
            st.divider()
            st.markdown("### ‚öñÔ∏è T·ªïng k·∫øt & ƒê·ªëi chi·∫øu")
            
            c_sum1, c_sum2, c_sum3 = st.columns(3)
            c_sum1.metric("T·ªïng Chi ph√≠ (Hƒê + Ph√°t sinh)", format_vnd(final_act_cost), help="T·ªïng chi ph√≠ th·ª±c t·∫ø c·ªßa tour")
            c_sum2.metric("T·ªïng UNC (ƒê√£ thanh to√°n)", format_vnd(total_unc), help="T·ªïng ti·ªÅn ƒë√£ chi ra t·ª´ t√†i kho·∫£n")
            
            diff = total_unc - final_act_cost
            if diff == 0:
                c_sum3.success("‚úÖ ƒê√£ kh·ªõp (UNC = Chi ph√≠)")
            elif diff > 0:
                c_sum3.warning(f"‚ö†Ô∏è UNC d∆∞: {format_vnd(diff)}")
            else:
                c_sum3.error(f"‚ö†Ô∏è Thi·∫øu UNC: {format_vnd(abs(diff))}")
            
            # L·ª£i nhu·∫≠n = T·ªïng doanh thu (D·ª± to√°n) - T·ªïng chi
            final_profit = est_final_sale - final_act_cost
            
            st.markdown(f"""<div class="profit-summary-card">
                <h3>T·ªîNG DOANH THU - T·ªîNG CHI = L·ª¢I NHU·∫¨N</h3>
                <div class="formula">{format_vnd(est_final_sale)} - {format_vnd(final_act_cost)} = <span class="result">{format_vnd(final_profit)} VND</span></div>
            </div>
            """, unsafe_allow_html=True)

            # --- EXPORT EXCEL (ACT) ---
            st.write("")
            # Prepare Data for Export
            df_exp_act = edited_act.copy()
            if 'times' not in df_exp_act.columns: df_exp_act['times'] = 1
            df_exp_act['times'] = df_exp_act.get('times', pd.Series([1.0] * len(df_exp_act), index=df_exp_act.index)).fillna(1).astype(float)

            # Clean numbers

            def clean_num_act(x): # type: ignore
                if isinstance(x, str):
                    return float(x.replace('.', '').replace(',', '').replace(' VND', '').strip())
                return float(x) if x else 0.0
            
            df_exp_act['unit_price'] = df_exp_act['unit_price'].apply(clean_num_act)
            df_exp_act['quantity'] = pd.to_numeric(df_exp_act['quantity'], errors='coerce').fillna(0)
            df_exp_act['total_amount'] = df_exp_act['quantity'] * df_exp_act['unit_price'] * df_exp_act['times']
            df_exp_act['price_per_pax'] = df_exp_act['total_amount'] / guest_cnt_act
            
            # --- COMPARISON LOGIC ---
            # [CODE M·ªöI] S·ª≠ d·ª•ng l·∫°i est_lookup ƒë√£ t·∫°o ·ªü tr√™n ƒë·ªÉ t√≠nh c·ªôt D·ª± to√°n v√† Ch√™nh l·ªách cho Excel
            def get_est_val_exp(row): # type: ignore
                k = (str(row['category']).strip().lower(), str(row['description']).strip().lower()) # type: ignore
                return est_lookup.get(k, 0.0)

            df_exp_act['est_amount'] = df_exp_act.apply(get_est_val_exp, axis=1)
            df_exp_act['diff_amount'] = df_exp_act['est_amount'] - df_exp_act['total_amount'] # type: ignore
            
            def classify_item(row):
                if row['diff_amount'] < 0: return "V∆∞·ª£t chi"
                elif row['diff_amount'] > 0: return "Ti·∫øt ki·ªám"
                return ""

            df_exp_act['Ghi ch√∫'] = df_exp_act.apply(classify_item, axis=1)

            # Rename
            df_exp_act = df_exp_act.rename(columns={
                'category': 'H·∫°ng m·ª•c', 
                'description': 'Di·ªÖn gi·∫£i', 
                'unit': 'ƒê∆°n v·ªã', 
                'unit_price': 'ƒê∆°n gi√°', 
                'quantity': 'S·ªë l∆∞·ª£ng', 
                'times': 'S·ªë l·∫ßn',
                'price_per_pax': 'Gi√°/Pax',
                'total_amount': 'Th·ª±c t·∫ø',
                'est_amount': 'D·ª± to√°n',
                'diff_amount': 'Ch√™nh l·ªách'
            })
            
            # [REQUEST 1] B·ªè c·ªôt 'S·ªë l·∫ßn' -> Keep it
            cols_to_export = ['H·∫°ng m·ª•c', 'Di·ªÖn gi·∫£i', 'ƒê∆°n v·ªã', 'ƒê∆°n gi√°', 'S·ªë l∆∞·ª£ng', 'S·ªë l·∫ßn', 'Gi√°/Pax', 'D·ª± to√°n', 'Th·ª±c t·∫ø', 'Ch√™nh l·ªách', 'Ghi ch√∫']
            df_exp_act_filtered = df_exp_act[cols_to_export]

            # [REQUEST 2] T√°ch th√†nh 2 b·∫£ng: Chi ph√≠ trong d·ª± to√°n v√† chi ph√≠ ph√°t sinh
            df_in_est = df_exp_act_filtered[df_exp_act_filtered['D·ª± to√°n'] > 0].copy()
            df_extra_cost = df_exp_act_filtered[df_exp_act_filtered['D·ª± to√°n'] == 0].copy()

            buffer_act = io.BytesIO()
            file_ext_act = "xlsx"
            mime_type_act = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            
            try:
                with pd.ExcelWriter(buffer_act, engine='xlsxwriter') as writer:
                    workbook: Any = writer.book
                    worksheet = workbook.add_worksheet('QuyetToan')
                    
                    # Styles (Copied and adapted)
                    company_name_fmt = workbook.add_format({'bold': True, 'font_size': 14, 'font_color': '#D84315', 'font_name': 'Times New Roman'}) # Orange for Act
                    company_info_fmt = workbook.add_format({'font_size': 10, 'italic': True, 'font_color': '#424242', 'font_name': 'Times New Roman'})
                    title_fmt = workbook.add_format({'bold': True, 'font_size': 18, 'align': 'center', 'valign': 'vcenter', 'font_color': '#BF360C', 'bg_color': '#FBE9E7', 'border': 1, 'font_name': 'Times New Roman'})
                    
                    header_fmt = workbook.add_format({'bold': True, 'fg_color': '#D84315', 'font_color': 'white', 'border': 1, 'align': 'center', 'valign': 'vcenter', 'text_wrap': True, 'font_name': 'Times New Roman'})
                    body_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'text_wrap': True, 'font_size': 10, 'font_name': 'Times New Roman'})
                    body_center_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'align': 'center', 'font_size': 10, 'font_name': 'Times New Roman'})
                    money_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'font_size': 10, 'font_name': 'Times New Roman'})
                    
                    # Summary Styles
                    sum_header_bg_fmt = workbook.add_format({'bold': True, 'bg_color': '#FFF3E0', 'border': 1, 'font_color': '#E65100', 'align': 'center', 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                    sum_label_fmt = workbook.add_format({'bold': True, 'align': 'left', 'border': 1, 'bg_color': '#FAFAFA', 'font_name': 'Times New Roman'})
                    sum_val_fmt = workbook.add_format({'num_format': '#,##0', 'align': 'right', 'border': 1, 'font_name': 'Times New Roman'})
                    sum_val_bold_fmt = workbook.add_format({'bold': True, 'num_format': '#,##0', 'align': 'right', 'border': 1, 'font_name': 'Times New Roman'})
                    
                    # [CODE M·ªöI] Format m√†u ƒë·ªè cho d√≤ng √¢m
                    alert_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'text_wrap': True, 'font_size': 10, 'font_color': '#D32F2F', 'font_name': 'Times New Roman'})
                    alert_money_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'font_size': 10, 'font_color': '#D32F2F', 'font_name': 'Times New Roman'})

                    # [CODE M·ªöI] Format cho ti√™u ƒë·ªÅ c√°c b·∫£ng chi ph√≠
                    section_title_fmt = workbook.add_format({'bold': True, 'font_size': 12, 'font_color': '#004D40', 'bg_color': '#E0F2F1', 'border': 1, 'align': 'center', 'font_name': 'Times New Roman'})

                    # 1. Company Info
                    if comp['logo_b64_str']:
                        try:
                            logo_data = base64.b64decode(comp['logo_b64_str'])
                            image_stream = io.BytesIO(logo_data)
                            img = Image.open(image_stream)
                            w, h = img.size
                            scale = 60 / h if h > 0 else 0.5
                            image_stream.seek(0)
                            worksheet.insert_image('A1', 'logo.png', {'image_data': image_stream, 'x_scale': scale, 'y_scale': scale, 'x_offset': 5, 'y_offset': 5})
                        except: pass
                    
                    worksheet.write('B1', comp['name'], company_name_fmt)
                    worksheet.write('B2', f"ƒêC: {comp['address']}", company_info_fmt)
                    worksheet.write('B3', f"MST: {comp['phone']}", company_info_fmt)
                    
                    # 2. Tour Info
                    worksheet.merge_range('A5:I5', "B·∫¢NG QUY·∫æT TO√ÅN CHI PH√ç TOUR", title_fmt)
                    
                    t_info_dict = {k: tour_info_act[k] for k in tour_info_act.keys()}
                    worksheet.write('A7', "T√™n ƒëo√†n:", sum_label_fmt)
                    worksheet.merge_range('B7:D7', t_info_dict.get('tour_name',''), sum_val_fmt)
                    worksheet.write('E7', "M√£ ƒëo√†n:", sum_label_fmt)
                    worksheet.merge_range('F7:I7', t_info_dict.get('tour_code',''), sum_val_fmt)
                    
                    worksheet.write('A8', "Kh√°ch h√†ng:", sum_label_fmt)
                    worksheet.merge_range('B8:D8', f"{t_info_dict.get('customer_name','')} - {t_info_dict.get('customer_phone','')}", sum_val_fmt)
                    worksheet.write('E8', "Sales:", sum_label_fmt)
                    worksheet.merge_range('F8:I8', t_info_dict.get('sale_name',''), sum_val_fmt)
                    
                    worksheet.write('A9', "Th·ªùi gian:", sum_label_fmt)
                    worksheet.merge_range('B9:D9', f"{t_info_dict.get('start_date','')} - {t_info_dict.get('end_date','')}", sum_val_fmt)
                    worksheet.write('E9', "S·ªë kh√°ch:", sum_label_fmt)
                    worksheet.merge_range('F9:I9', t_info_dict.get('guest_count',0), sum_val_fmt)

                    # 3. Table Header & Body (MODIFIED)
                    current_row = 10 # B·∫Øt ƒë·∫ßu t·ª´ d√≤ng 11

                    # --- B·∫£ng 1: Chi ph√≠ trong d·ª± to√°n ---
                    if not df_in_est.empty:
                        worksheet.merge_range(current_row, 0, current_row, len(df_in_est.columns)-1, "CHI PH√ç TRONG D·ª∞ TO√ÅN", section_title_fmt)
                        current_row += 1
                        for col_num, value in enumerate(df_in_est.columns):
                            worksheet.write(current_row, col_num, value, header_fmt)
                        for row_idx in range(len(df_in_est)):
                            diff_val = df_in_est.iloc[row_idx, 7] # Ch√™nh l·ªách
                            is_negative = isinstance(diff_val, (int, float)) and diff_val < 0
                            for col_idx in range(len(df_in_est.columns)):
                                val = df_in_est.iloc[row_idx, col_idx]
                                if col_idx == 2: fmt = body_center_fmt
                                elif col_idx in [3, 4, 5, 6, 7, 8, 9]: fmt = money_fmt
                                else: fmt = body_fmt
                                if is_negative:
                                    if col_idx in [3, 4, 5, 6, 7, 8, 9]: fmt = alert_money_fmt
                                    else: fmt = alert_fmt
                                if pd.isna(val): val = ""
                                worksheet.write(current_row + 1 + row_idx, col_idx, val, fmt)
                        current_row += len(df_in_est) + 1

                    # Th√™m d√≤ng tr·ªëng
                    current_row += 1

                    # --- B·∫£ng 2: Chi ph√≠ ph√°t sinh ngo√†i d·ª± to√°n ---
                    if not df_extra_cost.empty:
                        worksheet.merge_range(current_row, 0, current_row, len(df_extra_cost.columns)-1, "CHI PH√ç PH√ÅT SINH NGO√ÄI D·ª∞ TO√ÅN", section_title_fmt)
                        current_row += 1
                        for col_num, value in enumerate(df_extra_cost.columns):
                            worksheet.write(current_row, col_num, value, header_fmt)
                        for row_idx in range(len(df_extra_cost)):
                            # Chi ph√≠ ph√°t sinh lu√¥n l√† √¢m (v∆∞·ª£t chi)
                            is_negative = True
                            for col_idx in range(len(df_extra_cost.columns)):
                                val = df_extra_cost.iloc[row_idx, col_idx]
                                if col_idx == 2: fmt = body_center_fmt
                                elif col_idx in [3, 4, 5, 6, 7, 8, 9]: fmt = money_fmt
                                else: fmt = body_fmt
                                if is_negative:
                                    if col_idx in [3, 4, 5, 6, 7, 8, 9]: fmt = alert_money_fmt
                                    else: fmt = alert_fmt
                                if pd.isna(val): val = ""
                                worksheet.write(current_row + 1 + row_idx, col_idx, val, fmt)
                        current_row += len(df_extra_cost) + 1
                    
                    # 4. Summary
                    sum_row = current_row + 1
                    
                    worksheet.merge_range(sum_row, 0, sum_row, 3, "T·ªîNG K·∫æT QUY·∫æT TO√ÅN", sum_header_bg_fmt)
                    
                    # [CODE M·ªöI] Hi·ªÉn th·ªã ƒë·∫ßy ƒë·ªß th√¥ng tin t√†i ch√≠nh
                    # 1. T·ªïng doanh thu
                    worksheet.write(sum_row+1, 0, "1. T·ªïng doanh thu:", sum_label_fmt)
                    worksheet.merge_range(sum_row+1, 1, sum_row+1, 3, est_final_sale, sum_val_bold_fmt)
                    
                    # 2. T·ªïng chi ph√≠ (B·∫£ng k√™ + H√≥a ƒë∆°n ngo√†i)
                    worksheet.write(sum_row+2, 0, "2. T·ªïng chi ph√≠ th·ª±c t·∫ø:", sum_label_fmt)
                    worksheet.merge_range(sum_row+2, 1, sum_row+2, 3, final_act_cost, sum_val_bold_fmt)
                    
                    # 3. L·ª£i nhu·∫≠n
                    worksheet.write(sum_row+3, 0, "3. L·ª£i nhu·∫≠n th·ª±c t·∫ø:", sum_label_fmt)
                    profit_fmt = workbook.add_format({'bold': True, 'num_format': '#,##0', 'align': 'right', 'border': 1, 'bg_color': '#C8E6C9', 'font_color': '#1B5E20', 'font_name': 'Times New Roman'})
                    worksheet.merge_range(sum_row+3, 1, sum_row+3, 3, final_profit, profit_fmt)
                    
                    # Note nh·ªè v·ªÅ chi ph√≠ ngo√†i
                    if total_inv > 0:
                        worksheet.write(sum_row+4, 0, f"(Bao g·ªìm {format_vnd(total_inv)} h√≥a ƒë∆°n ph√°t sinh ngo√†i b·∫£ng k√™)", workbook.add_format({'italic': True, 'font_size': 9, 'font_name': 'Times New Roman'}))
                    
                    # Column Widths
                    worksheet.set_column('A:A', 25)
                    worksheet.set_column('B:B', 40)
                    worksheet.set_column('C:C', 10)
                    worksheet.set_column('D:I', 15)

            except Exception as e:
                # If xlsxwriter fails, fall back to a simple CSV export
                st.error(f"‚ö†Ô∏è L·ªói khi t·∫°o file Excel: {e}. ƒê√£ chuy·ªÉn sang xu·∫•t file CSV.")
                buffer_act.seek(0)
                buffer_act.truncate()
                df_exp_act_filtered.to_csv(buffer_act, index=False, encoding='utf-8-sig')
                file_ext_act = "csv"
                mime_type_act = "text/csv"

            clean_t_name_act = re.sub(r'[\\/*?:"<>|]', "", tour_info_act['tour_name'] if tour_info_act else "Tour") # type: ignore
            st.download_button(label=f"üì• T·∫£i B·∫£ng Quy·∫øt To√°n ({file_ext_act.upper()})", data=buffer_act.getvalue(), file_name=f"QuyetToan_{clean_t_name_act}.{file_ext_act}", mime=mime_type_act, use_container_width=True)

            def save_act_logic():
                run_query("DELETE FROM tour_items WHERE tour_id=? AND item_type='ACT'", (tour_id_act,), commit=True)
                data_to_insert = []
                query = """INSERT INTO tour_items (tour_id, item_type, category, description, unit, quantity, unit_price, total_amount, times)
                           VALUES (?, 'ACT', ?, ?, ?, ?, ?, ?, ?)"""

                for _, row in edited_act.iterrows():
                    if row['category'] or row['description']: # type: ignore
                        u_price = float(str(row['unit_price']).replace('.', '').replace(' VND', '').strip()) if row['unit_price'] else 0 # type: ignore
                        # Handle times safely
                        t_times = row.get('times', 1) # type: ignore
                        if pd.isna(t_times): t_times = 1
                        total_row = row['quantity'] * u_price * t_times # type: ignore

                        data_to_insert.append((
                            tour_id_act,
                            row['category'],
                            row['description'],
                            row['unit'],
                            row['quantity'],
                            u_price, # type: ignore
                            total_row,
                            t_times
                        ))

                if data_to_insert:
                    run_query_many(query, data_to_insert)

            if is_act_editable:
                if st.button("üíæ L∆ØU QUY·∫æT TO√ÅN", type="primary", use_container_width=True):
                    save_act_logic()
                    st.success("ƒê√£ l∆∞u quy·∫øt to√°n!"); time.sleep(1); st.rerun()
            else:
                st.info("üîí Ch·ªâ Admin m·ªõi ƒë∆∞·ª£c ch·ªânh s·ª≠a quy·∫øt to√°n.")
            
            st.divider()
            if st.button("‚úÖ HO√ÄN TH√ÄNH TOUR (Chuy·ªÉn v√†o L·ªãch s·ª≠)", type="primary", use_container_width=True, key="complete_tour_btn"):
                run_query("UPDATE tours SET status='completed' WHERE id=?", (tour_id_act,), commit=True)
                st.success("ƒê√£ ho√†n th√†nh tour! Tour ƒë√£ ƒë∆∞·ª£c chuy·ªÉn sang tab L·ªãch s·ª≠.")
                time.sleep(1)

                st.rerun()

    # ---------------- TAB 4: L·ªäCH S·ª¨ TOUR ----------------
    with tab_hist:
        st.subheader("üìú L·ªãch s·ª≠ Tour ƒë√£ ho√†n th√†nh")
        completed_tours = [t for t in all_tours if t['status'] == 'completed']
        
        if completed_tours:
            df_hist = pd.DataFrame([dict(t) for t in completed_tours])
            st.dataframe(
                df_hist[['tour_code', 'tour_name', 'start_date', 'end_date', 'guest_count', 'sale_name']],
                column_config={
                    "tour_code": "M√£ Tour",
                    "tour_name": "T√™n Tour",
                    "start_date": "Ng√†y ƒëi",
                    "end_date": "Ng√†y v·ªÅ",
                    "guest_count": "S·ªë kh√°ch",
                    "sale_name": "Sales"
                },
                use_container_width=True,
                hide_index=True
            )
            
            st.divider()
            st.write("üõ†Ô∏è Thao t√°c:")
            hist_opts = {f"[{t['tour_code']}] {t['tour_name']}": t['id'] for t in completed_tours} # type: ignore
            sel_hist = st.selectbox("Ch·ªçn tour ƒë·ªÉ xem l·∫°i ho·∫∑c m·ªü l·∫°i:", list(hist_opts.keys()), key="sel_hist_tour")
            if sel_hist:
                tid_hist = hist_opts[sel_hist] # type: ignore
                if st.button("üîì M·ªü l·∫°i Tour (Chuy·ªÉn v·ªÅ ƒêang ch·∫°y)", key="reopen_tour_btn"):
                    run_query("UPDATE tours SET status='running' WHERE id=?", (tid_hist,), commit=True)
                    st.success("ƒê√£ m·ªü l·∫°i tour! Ki·ªÉm tra l·∫°i b√™n tab Quy·∫øt to√°n.")
                    time.sleep(1)
                    st.rerun()
        else:
            st.info("Ch∆∞a c√≥ tour n√†o trong l·ªãch s·ª≠.")

    # ---------------- TAB 3: T·ªîNG H·ª¢P L·ª¢I NHU·∫¨N ----------------
    with tab_rpt:
        st.subheader("üìà T·ªïng H·ª£p L·ª£i Nhu·∫≠n & Doanh S·ªë")
        
        # L·ªçc theo th·ªùi gian
        rpt_df = pd.DataFrame([dict(r) for r in all_tours])
        if not rpt_df.empty:
            rpt_df['dt'] = pd.to_datetime(rpt_df['start_date'], format='%d/%m/%Y', errors='coerce') # type: ignore
            rpt_df = rpt_df.dropna(subset=['dt'])
            
            rpt_df['Month'] = rpt_df['dt'].apply(lambda x: x.strftime('%m/%Y'))
            rpt_df['Quarter'] = rpt_df['dt'].apply(lambda x: f"Q{(x.month-1)//3+1}/{x.year}")
            rpt_df['Year'] = rpt_df['dt'].apply(lambda x: x.strftime('%Y'))
            
            # --- PRE-FETCH DATA FOR PERFORMANCE ---
            all_items = run_query("SELECT tour_id, item_type, total_amount FROM tour_items")
            items_map = {} 
            if all_items:
                for item in all_items:
                    tid = item['tour_id']
                    itype = item['item_type']
                    amt = item['total_amount'] or 0
                    if tid not in items_map: items_map[tid] = {'EST': 0, 'ACT': 0}
                    items_map[tid][itype] += amt
            
            # T√≠nh to√°n ch·ªâ s·ªë cho t·ª´ng tour
            results = []
            for _, t in rpt_df.iterrows():
                tid = t['id'] # type: ignore
                costs = items_map.get(tid, {'EST': 0, 'ACT': 0})
                est_cost = costs['EST']
                act_cost = costs['ACT']
                
                p_pct = t.get('est_profit_percent', 0) or 0
                t_pct = t.get('est_tax_percent', 0) or 0

                # T√≠nh doanh thu (∆Øu ti√™n gi√° ch·ªët tay)
                final_price_manual = float(t.get('final_tour_price', 0) or 0)
                child_price_manual = float(t.get('child_price', 0) or 0)
                final_qty = float(t.get('final_qty', 0) or 0)
                child_qty = float(t.get('child_qty', 0) or 0)
                if final_qty == 0: final_qty = float(t.get('guest_count', 1))
                
                manual_revenue = (final_price_manual * final_qty) + (child_price_manual * child_qty)
                
                if manual_revenue > 0:
                    final_sale = manual_revenue
                else:
                    profit_est_val = est_cost * (p_pct/100)
                    final_sale = (est_cost + profit_est_val) * (1 + t_pct/100)

                net_revenue = final_sale / (1 + t_pct/100) if (1 + t_pct/100) != 0 else final_sale
                
                real_profit = net_revenue - act_cost
                
                results.append({
                    **t.to_dict(),
                    "T√™n ƒêo√†n": t['tour_name'], # type: ignore
                    "Sales": t['sale_name'], # type: ignore
                    "Ng√†y ƒëi": t['start_date'], # type: ignore
                    "Doanh Thu Thu·∫ßn": net_revenue,
                    "Chi Ph√≠ TT": act_cost,
                    "L·ª£i Nhu·∫≠n TT": real_profit,
                })
            
            res_df = pd.DataFrame(results)

            # --- UI CONTROLS ---
            c_type, c_period, c_val = st.columns(3)
            report_type = c_type.selectbox("Lo·∫°i b√°o c√°o:", ["Theo Tour (Chi ti·∫øt)", "Theo Sales (T·ªïng h·ª£p)"])
            period_type = c_period.selectbox("Xem theo:", ["Th√°ng", "Qu√Ω", "NƒÉm"])
            
            period_options = []
            period_col = 'Month'
            if period_type == "Th√°ng":
                period_col = 'Month'
                period_options = sorted(res_df['Month'].unique(), reverse=True)
            elif period_type == "Qu√Ω":
                period_col = 'Quarter'
                period_options = sorted(res_df['Quarter'].unique(), reverse=True)
            else:
                period_col = 'Year'
                period_options = sorted(res_df['Year'].unique(), reverse=True)
            
            selected_period = c_val.selectbox("Ch·ªçn th·ªùi gian:", ["T·∫•t c·∫£"] + period_options)
            
            # Filter
            if selected_period != "T·∫•t c·∫£":
                res_df = res_df[res_df[period_col] == selected_period]
            
            if report_type == "Theo Tour (Chi ti·∫øt)":
                res_df['T·ª∑ su·∫•t LN'] = res_df.apply(lambda x: (x['L·ª£i Nhu·∫≠n TT']/x['Doanh Thu Thu·∫ßn']*100) if x['Doanh Thu Thu·∫ßn'] else 0, axis=1)
                
                c_sum1, c_sum2 = st.columns(2)
                c_sum1.metric("T·ªïng L·ª£i Nhu·∫≠n", format_vnd(res_df['L·ª£i Nhu·∫≠n TT'].sum()))
                c_sum2.metric("T·ªïng Doanh Thu", format_vnd(res_df['Doanh Thu Thu·∫ßn'].sum()))
                
                st.dataframe(
                    res_df[['T√™n ƒêo√†n', 'Sales', 'Ng√†y ƒëi', 'Doanh Thu Thu·∫ßn', 'Chi Ph√≠ TT', 'L·ª£i Nhu·∫≠n TT', 'T·ª∑ su·∫•t LN']],
                    column_config={
                        "Doanh Thu Thu·∫ßn": st.column_config.NumberColumn(format="%d VND"),
                        "Chi Ph√≠ TT": st.column_config.NumberColumn(format="%d VND"),
                        "L·ª£i Nhu·∫≠n TT": st.column_config.NumberColumn(format="%d VND"),
                        "T·ª∑ su·∫•t LN": st.column_config.NumberColumn(format="%.2f %%"),
                    },
                    use_container_width=True,
                    hide_index=True
                )
                
                # Chu·∫©n b·ªã d·ªØ li·ªáu xu·∫•t Excel
                df_export = res_df[['T√™n ƒêo√†n', 'Sales', 'Ng√†y ƒëi', 'Doanh Thu Thu·∫ßn', 'Chi Ph√≠ TT', 'L·ª£i Nhu·∫≠n TT', 'T·ª∑ su·∫•t LN']].copy()
                file_name_rpt = f"BaoCao_LoiNhuan_Tour_{selected_period.replace('/', '_')}.xlsx"
            else: # Theo Sales
                df_sales = res_df.groupby('Sales').agg({
                    'Doanh Thu Thu·∫ßn': 'sum',
                    'Chi Ph√≠ TT': 'sum',
                    'L·ª£i Nhu·∫≠n TT': 'sum',
                    'id': 'count'
                }).reset_index()
                df_sales.columns = ["Nh√¢n vi√™n Sales", "Doanh Thu Thu·∫ßn", "Chi Ph√≠ TT", "L·ª£i Nhu·∫≠n TT", "S·ªë Tour"]
                df_sales['T·ª∑ su·∫•t LN'] = df_sales.apply(lambda x: (x['L·ª£i Nhu·∫≠n TT']/x['Doanh Thu Thu·∫ßn']*100) if x['Doanh Thu Thu·∫ßn'] else 0, axis=1)
                df_sales = df_sales.sort_values('L·ª£i Nhu·∫≠n TT', ascending=False)
                
                st.markdown(f"##### üèÜ B·∫£ng x·∫øp h·∫°ng Sales ({selected_period})")
                if not df_sales.empty:
                    best = df_sales.iloc[0]
                    c1, c2, c3 = st.columns(3)
                    c1.metric("Top Sales", best['Nh√¢n vi√™n Sales'], delta=format_vnd(best['L·ª£i Nhu·∫≠n TT']))
                    c2.metric("T·ªïng Doanh S·ªë", format_vnd(df_sales['Doanh Thu Thu·∫ßn'].sum()))
                    c3.metric("T·ªïng L·ª£i Nhu·∫≠n", format_vnd(df_sales['L·ª£i Nhu·∫≠n TT'].sum()))
                    
                    st.bar_chart(df_sales.set_index("Nh√¢n vi√™n Sales")[['Doanh Thu Thu·∫ßn', 'L·ª£i Nhu·∫≠n TT']])
                
                st.dataframe(
                    df_sales,
                    column_config={
                        "Doanh Thu Thu·∫ßn": st.column_config.NumberColumn(format="%d VND"),
                        "Chi Ph√≠ TT": st.column_config.NumberColumn(format="%d VND"),
                        "L·ª£i Nhu·∫≠n TT": st.column_config.NumberColumn(format="%d VND"),
                        "T·ª∑ su·∫•t LN": st.column_config.NumberColumn(format="%.2f %%"),
                        "S·ªë Tour": st.column_config.NumberColumn(format="%d"),
                    },
                    use_container_width=True,
                    hide_index=True
                )
                
                # Chu·∫©n b·ªã d·ªØ li·ªáu xu·∫•t Excel
                df_export = df_sales.copy()
                file_name_rpt = f"BaoCao_DoanhSo_Sales_{selected_period.replace('/', '_')}.xlsx"

            # --- T√çNH NƒÇNG XU·∫§T EXCEL ---
            st.write("")
            buffer_rpt = io.BytesIO()
            with pd.ExcelWriter(buffer_rpt, engine='xlsxwriter') as writer:
                df_export.to_excel(writer, index=False, sheet_name='Report')
                workbook: Any = writer.book
                worksheet = writer.sheets['Report']
                
                # ƒê·ªãnh d·∫°ng
                header_fmt = workbook.add_format({'bold': True, 'fg_color': '#2E7D32', 'font_color': 'white', 'border': 1, 'align': 'center', 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                body_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'font_name': 'Times New Roman'})
                money_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'num_format': '#,##0', 'font_name': 'Times New Roman'})
                pct_fmt = workbook.add_format({'border': 1, 'valign': 'vcenter', 'num_format': '0.00"%"', 'font_name': 'Times New Roman'})
                
                # √Åp d·ª•ng ƒë·ªãnh d·∫°ng header
                for col_num, value in enumerate(df_export.columns):
                    worksheet.write(0, col_num, value, header_fmt)
                
                # √Åp d·ª•ng ƒë·ªãnh d·∫°ng body
                for row_idx in range(len(df_export)):
                    for col_idx in range(len(df_export.columns)):
                        val = df_export.iloc[row_idx, col_idx]
                        col_name = df_export.columns[col_idx]
                        
                        fmt = body_fmt
                        if col_name in ['Doanh Thu Thu·∫ßn', 'Chi Ph√≠ TT', 'L·ª£i Nhu·∫≠n TT']: fmt = money_fmt
                        elif col_name == 'T·ª∑ su·∫•t LN': fmt = pct_fmt
                        
                        if pd.isna(val): val = ""
                        worksheet.write(row_idx + 1, col_idx, val, fmt)
                
                worksheet.set_column('A:A', 25)
                worksheet.set_column('B:Z', 18)

            st.download_button("üì• Xu·∫•t b√°o c√°o Excel", buffer_rpt.getvalue(), file_name_rpt, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            st.info("Ch∆∞a c√≥ d·ªØ li·ªáu tour.")

def render_customer_management():
    st.title("ü§ù Qu·∫£n L√Ω Kh√°ch H√†ng")
    
    # L·∫•y th√¥ng tin user hi·ªán t·∫°i ƒë·ªÉ l·ªçc
    current_user_info_cust = st.session_state.get("user_info", {})
    current_user_name_cust = current_user_info_cust.get('name', 'N/A')
    current_user_role_cust = current_user_info_cust.get('role')
    
    tab_list, tab_add = st.tabs(["üìã Danh s√°ch kh√°ch h√†ng", "‚ûï Th√™m kh√°ch h√†ng"])
    
    with tab_add:
        with st.form("add_cust_form"):
            st.subheader("Th√™m kh√°ch h√†ng m·ªõi")
            c1, c2 = st.columns(2)
            name = c1.text_input("T√™n kh√°ch h√†ng (*)", placeholder="Nguy·ªÖn VƒÉn A")
            phone = c2.text_input("S·ªë ƒëi·ªán tho·∫°i", placeholder="090...")
            email = c1.text_input("Email", placeholder="abc@gmail.com")
            addr = c2.text_input("ƒê·ªãa ch·ªâ")
            note = st.text_area("Ghi ch√∫")
            
            if st.form_submit_button("L∆∞u kh√°ch h√†ng", type="primary"):
                if name:
                    add_row_to_table('customers', {
                        'name': name, 'phone': phone, 'email': email, 'address': addr, 'notes': note,
                        'created_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S"), 'sale_name': current_user_name_cust
                    })
                    st.success("ƒê√£ th√™m kh√°ch h√†ng m·ªõi!"); time.sleep(1); st.rerun()
                else:
                    st.warning("Vui l√≤ng nh·∫≠p t√™n kh√°ch h√†ng.")

    with tab_list:
        # Search bar
        search_term = st.text_input("üîç T√¨m ki·∫øm", placeholder="Nh·∫≠p t√™n ho·∫∑c s·ªë ƒëi·ªán tho·∫°i...")
        
        query = "SELECT * FROM customers"
        params = []

        # Base filter for sales role
        if current_user_role_cust == 'sale':
            query += " WHERE sale_name=?"
            params.append(current_user_name_cust)

        # Additional filter for search term
        if search_term:
            if "WHERE" in query:
                query += " AND (name LIKE ? OR phone LIKE ?)"
            else:
                query += " WHERE name LIKE ? OR phone LIKE ?"
            params.extend([f"%{search_term}%", f"%{search_term}%"])
        query += " ORDER BY id DESC"
        
        customers = run_query(query, tuple(params))
        
        if customers:
            # Display as dataframe for overview
            df_cust = pd.DataFrame([dict(r) for r in customers])
            st.dataframe(
                df_cust[['name', 'phone', 'email', 'address', 'notes']],
                column_config={
                    "name": "T√™n kh√°ch h√†ng",
                    "phone": "SƒêT",
                    "email": "Email",
                    "address": "ƒê·ªãa ch·ªâ",
                    "notes": "Ghi ch√∫"
                },
                use_container_width=True,
                hide_index=True
            )
            
            st.divider()
            st.markdown("##### üõ†Ô∏è Ch·ªânh s·ª≠a th√¥ng tin")
            
            cust_options = {f"{c['name']} - {c['phone']}": c['id'] for c in customers} # type: ignore
            selected_cust = st.selectbox("Ch·ªçn kh√°ch h√†ng ƒë·ªÉ s·ª≠a/x√≥a:", list(cust_options.keys()))
            
            if selected_cust:
                cid = cust_options[selected_cust] # type: ignore
                c_info = next((c for c in customers if c['id'] == cid), None)
                
                if c_info:
                    with st.form(f"edit_cust_{cid}"):
                        c1, c2 = st.columns(2)
                        n_name = c1.text_input("T√™n", value=c_info['name']) # type: ignore
                        n_phone = c2.text_input("SƒêT", value=c_info['phone']) # type: ignore
                        n_email = c1.text_input("Email", value=c_info['email']) # type: ignore
                        n_addr = c2.text_input("ƒê·ªãa ch·ªâ", value=c_info['address']) # type: ignore
                        n_note = st.text_area("Ghi ch√∫", value=c_info['notes']) # type: ignore
                        
                        c_save, c_del = st.columns(2)
                        if c_save.form_submit_button("üíæ C·∫≠p nh·∫≠t"):
                            run_query("UPDATE customers SET name=?, phone=?, email=?, address=?, notes=? WHERE id=?", 
                                      (n_name, n_phone, n_email, n_addr, n_note, cid), commit=True)
                            st.success("ƒê√£ c·∫≠p nh·∫≠t!"); time.sleep(0.5); st.rerun()
                        
                        if c_del.form_submit_button("üóëÔ∏è X√≥a kh√°ch h√†ng"):
                            run_query("DELETE FROM customers WHERE id=?", (cid,), commit=True)
                            st.success("ƒê√£ x√≥a!"); time.sleep(0.5); st.rerun()
        else:
            st.info("Ch∆∞a c√≥ kh√°ch h√†ng n√†o.")

def render_hr_management():
    st.title("üë• Qu·∫£n L√Ω Nh√¢n S·ª± & T√†i Kho·∫£n")
    
    if (st.session_state.user_info or {}).get('role') not in ['admin', 'admin_f1']:
        st.warning("‚õî Khu v·ª±c n√†y ch·ªâ d√†nh cho Admin ho·∫∑c Admin F1. Vui l√≤ng li√™n h·ªá qu·∫£n tr·ªã vi√™n.")
    else:
        tab_list, tab_req = st.tabs(["üìã Danh s√°ch t√†i kho·∫£n", "üìù Duy·ªát ƒëƒÉng k√Ω m·ªõi"])
        
        with tab_list:
            st.subheader("Danh s√°ch t√†i kho·∫£n h·ªá th·ªëng")
            
            # L·∫•y d·ªØ li·ªáu users
            users = run_query("SELECT id, username, role, status FROM users ORDER BY id ASC")
            if users:
                df_users = pd.DataFrame([dict(r) for r in users])
                original_df = df_users.copy()
                
                # X√°c ƒë·ªãnh c√°c quy·ªÅn c√≥ th·ªÉ g√°n
                role_options = ["admin", "admin_f1", "user", "sale", "accountant"]
                if (st.session_state.user_info or {}).get('role') == 'admin_f1':
                    role_options = ["admin_f1", "user", "sale", "accountant"] # Admin F1 kh√¥ng th·ªÉ t·∫°o admin ch√≠nh

                # Hi·ªÉn th·ªã b·∫£ng
                edited_df = st.data_editor(
                    df_users,
                    column_config={
                        "id": st.column_config.NumberColumn("ID", width="small", disabled=True),
                        "username": st.column_config.TextColumn("T√™n ƒëƒÉng nh·∫≠p", width="medium", disabled=True),
                        "role": st.column_config.SelectboxColumn("Quy·ªÅn h·∫°n", options=role_options, required=True, width="medium"),
                        "status": st.column_config.SelectboxColumn("Tr·∫°ng th√°i", options=["approved", "pending", "blocked"], required=True, width="medium")
                    },
                    use_container_width=True,
                    hide_index=True
                )
                
                if st.button("üíæ L∆∞u thay ƒë·ªïi ph√¢n quy·ªÅn", type="primary"):
                    if not original_df.equals(edited_df):
                        with st.spinner("ƒêang c·∫≠p nh·∫≠t..."):
                            current_user_role = (st.session_state.user_info or {}).get('role')
                            # Iterate through the edited dataframe
                            for index, row in edited_df.iterrows():
                                original_row = original_df.loc[index]# type: ignore
                                # Check if the row has changed
                                if not row.equals(original_row):
                                    user_id = row['id'] # type: ignore
                                    username = row['username'] # type: ignore
                                    new_role = row['role'] # type: ignore
                                    new_status = row['status'] # type: ignore
                                    original_role = original_row['role'] # type: ignore

                                    # Prevent changing the main admin
                                    if username == 'admin':
                                        st.warning("Kh√¥ng th·ªÉ thay ƒë·ªïi quy·ªÅn c·ªßa t√†i kho·∫£n 'admin' ch√≠nh.")
                                        continue
                                    
                                    # Prevent F1 from editing a full admin
                                    if current_user_role == 'admin_f1' and original_role == 'admin':
                                        st.warning(f"B·∫°n kh√¥ng c√≥ quy·ªÅn ch·ªânh s·ª≠a t√†i kho·∫£n admin '{username}'.")
                                        continue
                                    
                                    run_query(
                                        "UPDATE users SET role=?, status=? WHERE id=?",
                                        (new_role, new_status, user_id),
                                        commit=True
                                    )
                        st.success("ƒê√£ c·∫≠p nh·∫≠t th√†nh c√¥ng!")
                        time.sleep(1); st.rerun()
                    else:
                        st.toast("Kh√¥ng c√≥ thay ƒë·ªïi n√†o.")
                
                st.divider()
                st.markdown("##### üóëÔ∏è X√≥a t√†i kho·∫£n")
                # Lo·∫°i b·ªè admin ch√≠nh ra kh·ªèi danh s√°ch x√≥a ƒë·ªÉ tr√°nh l·ªói
                del_options = [u['username'] for u in users if u['username'] != 'admin'] # type: ignore
                user_to_del = st.selectbox("Ch·ªçn t√†i kho·∫£n c·∫ßn x√≥a:", del_options, key="sel_del_u")
                
                if st.button("X√°c nh·∫≠n x√≥a t√†i kho·∫£n", type="primary", key="btn_del_u"):
                    if user_to_del:
                        # Ki·ªÉm tra quy·ªÅn tr∆∞·ªõc khi x√≥a
                        user_to_del_info = run_query("SELECT role FROM users WHERE username=?", (user_to_del,), fetch_one=True)
                        current_user_role = (st.session_state.user_info or {}).get('role')

                        if current_user_role == 'admin_f1' and user_to_del_info and user_to_del_info['role'] == 'admin': # type: ignore
                            st.error(f"B·∫°n kh√¥ng c√≥ quy·ªÅn x√≥a t√†i kho·∫£n admin '{user_to_del}'.")
                        else:
                            run_query("DELETE FROM users WHERE username=?", (user_to_del,), commit=True)
                            st.success(f"ƒê√£ x√≥a t√†i kho·∫£n {user_to_del}!")
                            time.sleep(1); st.rerun()
            else:
                st.info("Ch∆∞a c√≥ t√†i kho·∫£n n√†o.")

        with tab_req:
            st.subheader("Y√™u c·∫ßu ƒëƒÉng k√Ω ch·ªù duy·ªát")
            pending = run_query("SELECT * FROM users WHERE status='pending'")
            if pending:
                for p in pending:
                    with st.container(border=True):
                        c1, c2, c3 = st.columns([2, 1, 1])
                        c1.write(f"User: **{p['username']}**") # type: ignore
                        if c2.button("‚úî Duy·ªát", key=f"hr_app_{p['id']}", use_container_width=True): # type: ignore
                            run_query("UPDATE users SET status='approved' WHERE id=?", (p['id'],), commit=True) # type: ignore
                            st.success("ƒê√£ duy·ªát!"); time.sleep(0.5); st.rerun()
                        if c3.button("‚úñ X√≥a", key=f"hr_del_{p['id']}", use_container_width=True): # type: ignore
                            run_query("DELETE FROM users WHERE id=?", (p['id'],), commit=True) # type: ignore
                            st.success("ƒê√£ x√≥a!"); time.sleep(0.5); st.rerun()
            else:
                st.info("Hi·ªán kh√¥ng c√≥ y√™u c·∫ßu n√†o.")

def render_search_module():
    st.title("üîç Tra c·ª©u th√¥ng tin h·ªá th·ªëng")
    
    # L·∫•y th√¥ng tin user hi·ªán t·∫°i ƒë·ªÉ l·ªçc
    current_user_info = st.session_state.get("user_info", {})
    current_user_name = current_user_info.get('name', 'N/A')
    current_user_role = current_user_info.get('role')

    query = st.text_input("Nh·∫≠p t·ª´ kh√≥a t√¨m ki·∫øm", placeholder="Nh·∫≠p M√£ Tour, S·ªë H√≥a ƒê∆°n, M√£ V√©, M√£ Chi Ph√≠, ho·∫∑c T√™n Kh√°ch...", help="H·ªá th·ªëng s·∫Ω t√¨m trong Tour, H√≥a ƒë∆°n, UNC v√† V√© m√°y bay")
        
    if query:
        st.divider()
        term = f"%{query.strip()}%"
        found_any = False
        
        # 1. T√åM TRONG TOUR
        tour_sql = "SELECT * FROM tours WHERE (tour_code LIKE ? OR tour_name LIKE ?)"
        tour_params = [term, term]
        if current_user_role == 'sale':
            tour_sql += " AND sale_name=?"
            tour_params.append(current_user_name)
            
        tours = run_query(tour_sql, tuple(tour_params))
        if tours:
            found_any = True
            st.subheader(f"üì¶ T√¨m th·∫•y {len(tours)} Tour")
            for t in tours:
                with st.expander(f"Tour: {t['tour_name']} (M√£: {t['tour_code']})", expanded=True):
                    c1, c2, c3 = st.columns(3) # type: ignore
                    c1.write(f"**Sales:** {t['sale_name']}") # type: ignore
                    c2.write(f"**Ng√†y:** {t['start_date']} - {t['end_date']}") # type: ignore
                    c3.write(f"**Kh√°ch:** {t['guest_count']}") # type: ignore
                    
                    est = run_query("SELECT SUM(total_amount) as t FROM tour_items WHERE tour_id=? AND item_type='EST'", (t['id'],), fetch_one=True) # type: ignore
                    act = run_query("SELECT SUM(total_amount) as t FROM tour_items WHERE tour_id=? AND item_type='ACT'", (t['id'],), fetch_one=True) # type: ignore
                    est_val = est['t'] if isinstance(est, sqlite3.Row) and est['t'] else 0 # type: ignore
                    act_val = act['t'] if isinstance(act, sqlite3.Row) and act['t'] else 0 # type: ignore
                    
                    st.info(f"üí∞ D·ª± to√°n: {format_vnd(est_val)} | üí∏ Quy·∫øt to√°n: {format_vnd(act_val)}")

        # 2. T√åM TRONG KH√ÅCH H√ÄNG (M·ªöI)
        cust_sql = "SELECT * FROM customers WHERE (name LIKE ? OR phone LIKE ?)"
        cust_params = [term, term]
        if current_user_role == 'sale':
            cust_sql += " AND sale_name=?"
            cust_params.append(current_user_name)
            
        custs = run_query(cust_sql, tuple(cust_params))
        if custs:
            found_any = True
            st.subheader(f"üë• T√¨m th·∫•y {len(custs)} Kh√°ch h√†ng")
            for c in custs:
                with st.expander(f"Kh√°ch h√†ng: {c['name']} - {c['phone']}", expanded=True):
                    st.write(f"**Email:** {c['email']}")
                    st.write(f"**ƒê·ªãa ch·ªâ:** {c['address']}")
                    st.write(f"**Ghi ch√∫:** {c['notes']}")

        # 3. T√åM TRONG H√ìA ƒê∆†N / UNC
        invs = run_query("SELECT * FROM invoices WHERE invoice_number LIKE ? OR cost_code LIKE ? OR memo LIKE ? ORDER BY date DESC", (term, term, term))
        if invs:
            found_any = True
            st.subheader(f"üí∞ T√¨m th·∫•y {len(invs)} H√≥a ƒë∆°n / UNC")
            
            for inv in invs:
                icon = "üí∏" if "UNC" in (inv['invoice_number'] or "") else "üìÑ"
                i_num = inv['invoice_number'] if inv['invoice_number'] else "(Kh√¥ng s·ªë)" # type: ignore
                label = f"{icon} {inv['date']} | {i_num} | {format_vnd(inv['total_amount'])} | {inv['memo']}" # type: ignore
                
                with st.expander(label):
                    c_info, c_file = st.columns([1, 1])
                    with c_info:
                        st.markdown(f"**B√™n b√°n:** {inv['seller_name']}") # type: ignore
                        st.markdown(f"**B√™n mua:** {inv['buyer_name']}") # type: ignore
                        st.markdown(f"**T·ªïng ti·ªÅn:** {format_vnd(inv['total_amount'])}") # type: ignore
                        st.markdown(f"**M√£ chi ph√≠:** `{inv['cost_code']}`") # type: ignore
                        st.caption(f"Tr·∫°ng th√°i: {inv['status']}") # type: ignore
                    
                    with c_file:
                        file_path = inv['file_path'] # type: ignore
                        if file_path and os.path.exists(file_path):
                            # The 'file_path' from the database is a Google Drive link, not a local path.
                            # The original code to check os.path.exists(file_path) and open it is incorrect.
                            # We should just provide the link.
                            st.link_button("üîó M·ªü file tr√™n Google Drive", file_path, use_container_width=True)

        if not found_any:
            st.warning("üì≠ Kh√¥ng t√¨m th·∫•y d·ªØ li·ªáu n√†o ph√π h·ª£p.")

def main():
    if not st.session_state.logged_in:
        render_login_page(comp)
        return

    module, menu = render_sidebar(comp)

    # --- HEADER CH√çNH ---
    l_html = f'<img src="data:image/png;base64,{comp["logo_b64_str"]}" class="company-logo-img">' if comp['logo_b64_str'] else ''
    st.markdown(f'''
    <div class="company-header-container">
        {l_html}
        <div class="company-info-text">
            <h1>{comp['name']}</h1>
            <p>üìç {comp['address']}</p>
            <p>MST: {comp['phone']}</p>
        </div>
    </div>
    ''', unsafe_allow_html=True)

    if module == "üè† Trang Ch·ªß":
        render_dashboard()
    elif module == "üìÖ L·ªãch Th√¥ng B√°o":
        render_notification_calendar()
    elif module == "üîñ Qu·∫£n L√Ω Booking":
        render_booking_management()
    elif module == "üí∞ Ki·ªÉm So√°t Chi Ph√≠":
        render_cost_control(menu)
    elif module == "üí≥ Qu·∫£n L√Ω C√¥ng N·ª£":
        render_debt_management()
    elif module == "üì¶ Qu·∫£n L√Ω Tour ":
        render_tour_management()
    elif module == "ü§ù Qu·∫£n L√Ω Kh√°ch H√†ng":
        render_customer_management()
    elif module == "üë• Qu·∫£n L√Ω Nh√¢n S·ª±":
        render_hr_management()
    elif module == "üîç Tra c·ª©u th√¥ng tin":
        render_search_module()

if __name__ == "__main__":
    main()
